from pathlib import Path

import pytest
from docx import Document

from app.core.config import settings
from app.services.document_processor import DocumentProcessor


def _url_to_storage_path(url: str) -> Path:
    prefix = f"{settings.static_url_prefix}/"
    assert url.startswith(prefix)
    relative = url[len(prefix) :]
    return settings.resolved_storage_dir / relative


def _find_repo_doc(name: str) -> Path:
    repo_root = Path(__file__).resolve().parents[2]
    matches = [
        path
        for path in repo_root.rglob(name)
        if ".git" not in path.parts and "__pycache__" not in path.parts
    ]
    assert matches, f"missing sample document: {name}"
    return matches[0]


def _create_pdf_with_pages(tmp_path: Path, page_count: int) -> Path:
    fitz = pytest.importorskip("fitz")
    pdf_path = tmp_path / f"sample_{page_count}.pdf"
    doc = fitz.open()
    try:
        for _ in range(page_count):
            doc.new_page(width=595, height=842)
        doc.save(str(pdf_path))
    finally:
        doc.close()
    return pdf_path


def test_split_pdf_preview_should_keep_full_pdf_when_no_answer_split(tmp_path):
    fitz = pytest.importorskip("fitz")
    source_pdf = _create_pdf_with_pages(tmp_path, page_count=4)
    processor = DocumentProcessor()

    main_pages, answer_pages = processor._split_pdf_preview(
        task_id=999101,
        source_pdf=source_pdf,
        answer_start_page=None,
    )

    assert len(main_pages) == 1
    assert answer_pages == []
    path = _url_to_storage_path(main_pages[0])
    assert path.exists()
    opened = fitz.open(str(path))
    try:
        assert opened.page_count == 4
    finally:
        opened.close()


def test_split_pdf_preview_should_keep_full_pdf_even_when_answer_split_detected(tmp_path):
    fitz = pytest.importorskip("fitz")
    source_pdf = _create_pdf_with_pages(tmp_path, page_count=5)
    processor = DocumentProcessor()

    main_pages, answer_pages = processor._split_pdf_preview(
        task_id=999102,
        source_pdf=source_pdf,
        answer_start_page=2,
    )

    assert len(main_pages) == 1
    assert answer_pages == []
    path = _url_to_storage_path(main_pages[0])
    assert path.exists()
    opened = fitz.open(str(path))
    try:
        assert opened.page_count == 5
    finally:
        opened.close()


def test_docx_preview_is_pdf(tmp_path):
    source = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("涓€銆侀€夋嫨棰橈紙10鍒嗭級")
    doc.add_paragraph("1. 绗竴棰橈紙2鍒嗭級")
    doc.add_paragraph("绛旀")
    doc.add_paragraph("A")
    doc.save(str(source))

    processor = DocumentProcessor()
    output = processor.process(
        task_id=999001,
        file_path=str(source),
        file_ext=".docx",
        max_level=8,
        second_level_mode="auto",
        answer_section_patterns=None,
        score_patterns=None,
    )

    result = output["result"]
    assert result["mainPages"], "mainPages should contain preview urls"
    assert result["mainPages"][0].endswith(".pdf")
    for url in result["mainPages"] + result["answerPages"]:
        assert url.endswith(".pdf")
        assert _url_to_storage_path(url).exists()


def test_create_adjusted_docx_for_preview_applies_layout_settings(tmp_path):
    source = tmp_path / "layout.docx"
    doc = Document()
    doc.add_paragraph("段落内容")
    doc.save(str(source))

    processor = DocumentProcessor()
    adjusted = processor._create_adjusted_docx_for_preview(
        source_docx=source,
        task_id=999002,
        layout_adjustments={
            "marginTopCm": 2.5,
            "marginLeftCm": 1.8,
            "paragraphLeftIndentCm": 0.8,
            "firstLineIndentCm": 0.6,
            "paragraphSpaceBeforePt": 12,
            "paragraphSpaceAfterPt": 6,
        },
    )

    adjusted_doc = Document(str(adjusted))
    section = adjusted_doc.sections[0]
    assert abs(section.top_margin.cm - 2.5) < 0.01
    assert abs(section.left_margin.cm - 1.8) < 0.01

    paragraph_fmt = adjusted_doc.paragraphs[0].paragraph_format
    assert paragraph_fmt.left_indent is not None
    assert paragraph_fmt.first_line_indent is not None
    assert paragraph_fmt.space_before is not None
    assert paragraph_fmt.space_after is not None
    assert abs(paragraph_fmt.left_indent.cm - 0.8) < 0.01
    assert abs(paragraph_fmt.first_line_indent.cm - 0.6) < 0.01
    assert abs(paragraph_fmt.space_before.pt - 12) < 0.01
    assert abs(paragraph_fmt.space_after.pt - 6) < 0.01


def test_process_docx_uses_adjusted_docx_for_preview_when_layout_settings_present(tmp_path, monkeypatch):
    source = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph("一、单选题（10分）")
    doc.add_paragraph("1. 题目（2分）")
    doc.save(str(source))

    processor = DocumentProcessor()
    fake_pdf = tmp_path / "preview.pdf"
    fake_pdf.write_bytes(b"%PDF-1.4\n%%EOF")

    captured = {"office_path": None}

    def fake_convert_office_to_pdf(office_path):
        captured["office_path"] = Path(office_path)
        return fake_pdf

    monkeypatch.setattr(processor, "_convert_office_to_pdf", fake_convert_office_to_pdf)
    monkeypatch.setattr(processor, "_extract_pdf_page_lines", lambda _path: [["一、单选题（10分）", "1. 题目（2分）"]])
    monkeypatch.setattr(processor, "_split_pdf_preview", lambda _task_id, _pdf, _answer_start: (["/storage/pages/main.pdf"], []))

    output = processor.process(
        task_id=999003,
        file_path=str(source),
        file_ext=".docx",
        max_level=8,
        second_level_mode="auto",
        answer_section_patterns=None,
        score_patterns=None,
        layout_adjustments={"marginTopCm": 1.2},
    )

    assert captured["office_path"] is not None
    assert captured["office_path"].suffix.lower() == ".docx"
    assert "_layout_" in captured["office_path"].stem
    assert output["result"]["mainPages"] == ["/storage/pages/main.pdf"]


def test_collect_docx_lines_keeps_auto_numbering(tmp_path):
    source = tmp_path / "auto_numbering.docx"
    doc = Document()
    doc.add_paragraph("Section")
    doc.add_paragraph("Question one", style="List Number")
    doc.add_paragraph("Question two", style="List Number")
    doc.save(str(source))

    processor = DocumentProcessor()
    lines = processor._collect_docx_lines(Document(str(source)))

    assert lines[1].startswith("1.")
    assert lines[2].startswith("2.")


def test_matching_question_should_keep_all_scored_unordered_children(tmp_path):
    source = tmp_path / "matching_question.docx"
    doc = Document()
    doc.add_paragraph("第三单元测试")
    doc.add_paragraph("一、单项选择题（每小题2分，共计40分）")
    doc.add_paragraph("1．我国的宰相制度开始和结束的时间是（ ）。（2分）")
    doc.add_paragraph("A．秦——明 B．秦——清 C．汉——明 D．汉——清")
    doc.add_paragraph("二、改一改（共计10分）")
    doc.add_paragraph("21．以下材料中有五处错误，找出并改正。")
    doc.add_paragraph("错误1：（ ） 改正：（ ）（2分）")
    doc.add_paragraph("三、连一连（每小题5分，共计10分）")
    doc.add_paragraph("22．将下列作者与作品进行正确连接。")
    doc.add_paragraph("罗贯中 《水浒传》 （1分）")
    doc.add_paragraph("吴承恩 《红楼梦》 （1分）")
    doc.add_paragraph("施耐庵 《三国演义》 （1分）")
    doc.add_paragraph("曹雪芹 《西游记》 （1分）")
    doc.add_paragraph("宋应星 《天工开物》 （1分）")
    doc.add_paragraph("23．将下列事件和人物进行正确连接。")
    doc.add_paragraph("朱元璋 组织雅克萨自卫反击战 （1分）")
    doc.add_paragraph("戚继光 从荷兰殖民者手中收复台湾 （1分）")
    doc.add_paragraph("郑成功 抗击东南沿海倭寇 （1分）")
    doc.add_paragraph("康熙帝 建立明朝 （1分）")
    doc.add_paragraph("顺治帝 册封五世达赖为“达赖喇嘛” （1分）")
    doc.save(str(source))

    processor = DocumentProcessor()
    output = processor.process(
        task_id=999004,
        file_path=str(source),
        file_ext=".docx",
        max_level=8,
        second_level_mode="auto",
        answer_section_patterns=None,
        score_patterns=None,
    )

    third_root = next(item for item in output["details"]["outlineItems"] if item.get("numbering") == "三")
    question_22 = next(item for item in third_root.get("children", []) if item.get("numbering") == "22")
    question_23 = next(item for item in third_root.get("children", []) if item.get("numbering") == "23")

    assert question_22["score"] == 5.0
    assert question_23["score"] == 5.0
    assert len(question_22["children"]) == 5
    assert len(question_23["children"]) == 5
    assert all(float(child.get("score") or 0) == 1.0 for child in question_22["children"])
    assert all(float(child.get("score") or 0) == 1.0 for child in question_23["children"])


def test_split_lines_by_answer_should_not_trigger_on_instruction_sentence():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "一、选择题（60分）选择题答案对应填写在题后的答题卡中。",
        "1. 下列说法正确的是（）",
        "A. 选项A",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines
    assert answer_lines == []


def test_split_lines_by_answer_should_split_on_repeated_title_followed_by_answer_block():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "第五单元达标检测卷",
        "时间：90分钟 满分：100分",
        "一、基础训练营(44分)",
        "1．给加点字选择正确读音，用“√”标出。(6分)",
        "五、习作百花园(25分)",
        "15．题目：__________里的新鲜事",
        "要求：1.把题目补充完整，紧扣“新鲜”来写。",
        "2．把事情的来龙去脉写清楚。",
        "第五单元达标检测卷",
        "一、1.bāo√ fèi√ mèn√ nǐ√ pēn√ ào√",
        "2．苔藓 帐篷 楼梯 裂缝 嘴唇 甘蔗 袖筒 草坪",
        "3．将 酱 降 胞 苞 包 砚 谚 验 宣 暄 喧",
    ]

    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)

    assert main_lines == lines[:8]
    assert answer_lines == lines[8:]


def test_repair_outline_extraction_artifacts_should_restore_statement_option_parent():
    processor = DocumentProcessor()
    outline_items = [
        {
            "lineNumber": 4,
            "level": 1,
            "numbering": "一",
            "title": "字词积累。(29分)",
            "rawText": "一、字词积累。(29分)",
            "blankText": "",
            "score": 29.0,
            "children": [
                {
                    "lineNumber": 5,
                    "level": 2,
                    "numbering": "1",
                    "title": "读句子，选择读音与字形完全正确的一项填空。(填序号)(9分)",
                    "rawText": "1.读句子，选择读音与字形完全正确的一项填空。(填序号)(9分) (1)这时候，小小的天窗是你唯一的( )。 (2)那( )的低语，是在和刚刚从雪被里伸出头来的麦苗谈心。",
                    "blankText": "____（9分） ____（0分）",
                    "score": 9.0,
                    "children": [
                        {
                            "lineNumber": 6,
                            "level": 3,
                            "numbering": "1",
                            "title": "这时候，小小的天窗是你唯一的( )。",
                            "rawText": "(1)这时候，小小的天窗是你唯一的( )。",
                            "blankText": "____（3分）",
                            "score": 3.0,
                            "children": [],
                        },
                        {
                            "lineNumber": 8,
                            "level": 3,
                            "numbering": "2",
                            "title": "那( )的低语，是在和刚刚从雪被里伸出头来的麦苗谈心。",
                            "rawText": "(2)那( )的低语，是在和刚刚从雪被里伸出头来的麦苗谈心。",
                            "blankText": "____（3分）",
                            "score": 3.0,
                            "children": [],
                        },
                    ],
                }
            ],
        }
    ]

    changed = processor._repair_outline_extraction_artifacts(outline_items)

    question = outline_items[0]["children"][0]
    assert changed is True
    assert question["rawText"] == "1.读句子，选择读音与字形完全正确的一项填空。(填序号)(9分)"
    assert question["blankText"] == ""


def test_repair_outline_extraction_artifacts_should_restore_leaf_blank_from_title():
    processor = DocumentProcessor()
    outline_items = [
        {
            "lineNumber": 22,
            "level": 1,
            "numbering": "二",
            "title": "句段运用。(9分)",
            "rawText": "二、句段运用。(9分)",
            "blankText": "",
            "score": 9.0,
            "children": [
                {
                    "lineNumber": 23,
                    "level": 2,
                    "numbering": "4",
                    "title": "按要求完成句子练习。(9分)",
                    "rawText": "4.按要求完成句子练习。(9分)",
                    "blankText": "",
                    "score": 9.0,
                    "children": [
                        {
                            "lineNumber": 24,
                            "level": 3,
                            "numbering": "1",
                            "title": "醉里吴音相媚好，白发谁家翁媪？(用自己的话说说句子的意思)________________________________________________________________________________",
                            "rawText": "(1)醉里吴音相媚好，白发谁家翁媪？(用自己的话说说句子的意思)",
                            "blankText": "",
                            "score": 3.0,
                            "children": [],
                        }
                    ],
                }
            ],
        }
    ]

    changed = processor._repair_outline_extraction_artifacts(outline_items)

    leaf = outline_items[0]["children"][0]["children"][0]
    assert changed is True
    assert leaf["blankText"] == "________________________________________________________________________________（3分）"


def test_repair_outline_extraction_artifacts_should_restore_leaf_raw_from_title():
    processor = DocumentProcessor()
    outline_items = [
        {
            "lineNumber": 21,
            "level": 1,
            "numbering": "一",
            "title": "基础训练营(44分)",
            "rawText": "一、基础训练营(44分)",
            "blankText": "",
            "score": 44.0,
            "children": [
                {
                    "lineNumber": 22,
                    "level": 2,
                    "numbering": "6",
                    "title": "按要求改写句子。(10分)",
                    "rawText": "6．按要求改写句子。(10分)",
                    "blankText": "",
                    "score": 10.0,
                    "children": [
                        {
                            "lineNumber": 25,
                            "level": 3,
                            "numbering": "4",
                            "title": "妈妈焖上米饭，说：“我什么时候说过要炖肉？”(改为转述句)_________________________________________________________",
                            "rawText": "(4)妈妈焖上米饭，说：“我什么时候说过要炖肉？”(改为转述句)",
                            "blankText": "",
                            "score": 2.0,
                            "children": [],
                        }
                    ],
                }
            ],
        }
    ]

    changed = processor._repair_outline_extraction_artifacts(outline_items)

    leaf = outline_items[0]["children"][0]["children"][0]
    assert changed is True
    assert leaf["rawText"].startswith("(4)")
    assert leaf["rawText"].endswith("_________________________________________________________")
    assert leaf["blankText"] == "_________________________________________________________（2分）"


def test_repair_outline_extraction_artifacts_should_restore_leaf_raw_when_title_has_more_blanks():
    processor = DocumentProcessor()
    outline_items = [
        {
            "lineNumber": 47,
            "level": 1,
            "numbering": "七",
            "title": "阅读。（29分）",
            "rawText": "七、阅读。（29分）",
            "blankText": "",
            "score": 29.0,
            "children": [
                {
                    "lineNumber": 48,
                    "level": 2,
                    "numbering": "2",
                    "title": "选文中画“_____”的句子采用了____________的说明方法。（2分）",
                    "rawText": "2.选文中画的句子采用了____________的说明方法。（2分）",
                    "blankText": "_____（1分） ____________（1分）",
                    "score": 2.0,
                    "children": [],
                }
            ],
        }
    ]

    changed = processor._repair_outline_extraction_artifacts(outline_items)

    leaf = outline_items[0]["children"][0]
    assert changed is True
    assert leaf["rawText"] == "2.选文中画“_____”的句子采用了____________的说明方法。（2分）"


def test_repair_misnested_restart_children_should_promote_to_siblings():
    processor = DocumentProcessor()
    outline_items = [
        {
            "lineNumber": 47,
            "level": 1,
            "numbering": "七",
            "title": "阅读。（29分）",
            "rawText": "七、阅读。（29分）",
            "blankText": "",
            "score": 29.0,
            "children": [
                {
                    "lineNumber": 56,
                    "level": 2,
                    "numbering": "4",
                    "title": "认真读文段，判断下面问题，对的打“√”，错的打“×”。（3分）",
                    "rawText": "4.认真读文段，判断下面问题，对的打“√”，错的打“×”。（3分）",
                    "blankText": "",
                    "score": 3.0,
                    "children": [
                        {
                            "lineNumber": 57,
                            "level": 3,
                            "numbering": "1",
                            "title": "因为科学家提出那么多设想……",
                            "rawText": "（1）因为科学家提出那么多设想……",
                            "blankText": "____（1分）",
                            "score": 1.0,
                            "children": [],
                        },
                        {
                            "lineNumber": 67,
                            "level": 3,
                            "numbering": "1",
                            "title": "写近义词。（2分）",
                            "rawText": "1.写近义词。（2分）",
                            "blankText": "",
                            "score": 2.0,
                            "children": [],
                        },
                        {
                            "lineNumber": 69,
                            "level": 3,
                            "numbering": "2",
                            "title": "写反义词。（2分）",
                            "rawText": "2.写反义词。（2分）",
                            "blankText": "",
                            "score": 2.0,
                            "children": [],
                        },
                        {
                            "lineNumber": 71,
                            "level": 3,
                            "numbering": "3",
                            "title": "根据下面意思在短文中找到相应的词语写下来。（2分）",
                            "rawText": "3.根据下面意思在短文中找到相应的词语写下来。（2分）",
                            "blankText": "",
                            "score": 2.0,
                            "children": [],
                        },
                    ],
                },
                {
                    "lineNumber": 74,
                    "level": 2,
                    "numbering": "4",
                    "title": "当大家都说鲁迅是天才时……",
                    "rawText": "4.当大家都说鲁迅是天才时……",
                    "blankText": "____（2分）",
                    "score": 2.0,
                    "children": [],
                },
            ],
        }
    ]

    changed = processor._repair_outline_extraction_artifacts(outline_items)

    children = outline_items[0]["children"]
    assert changed is True
    assert children[0]["numbering"] == "4"
    assert children[0]["children"][0]["numbering"] == "1"
    assert [child["numbering"] for child in children[1:]] == ["1", "2", "3", "4"]


def test_repair_outline_extraction_artifacts_should_raise_leaf_score_to_explicit_blank_total():
    processor = DocumentProcessor()
    outline_items = [
        {
            "lineNumber": 32,
            "level": 1,
            "numbering": "一",
            "title": "名著阅读（4分）",
            "rawText": "一、名著阅读（4分）",
            "blankText": "",
            "score": 4.0,
            "children": [
                {
                    "lineNumber": 34,
                    "level": 2,
                    "numbering": "1",
                    "title": "右图中的人物便是文中的“他”，结合选文判断，“他”是______________（填人名，1分），选自______________________（填作品，1分）。",
                    "rawText": "（1）右图中的人物便是文中的“他”，结合选文判断，“他”是______________（填人名，1分），选自______________________（填作品，1分）。",
                    "blankText": "______________（1分） ______________________（1分）",
                    "score": 1.0,
                    "children": [],
                }
            ],
        }
    ]

    changed = processor._repair_outline_extraction_artifacts(outline_items)

    leaf = outline_items[0]["children"][0]
    assert changed is True
    assert leaf["score"] == 2.0


def test_repair_overflow_arabic_children_should_promote_to_root_siblings():
    processor = DocumentProcessor()
    outline_items = [
        {
            "lineNumber": 1,
            "level": 1,
            "numbering": "6",
            "title": "默写。（8分）",
            "rawText": "6、默写。（8分）",
            "blankText": "",
            "score": 8.0,
            "_tokenType": "arabic",
            "children": [
                {
                    "lineNumber": 2,
                    "level": 2,
                    "numbering": "1",
                    "title": "（1）树树皆秋色",
                    "rawText": "（1）树树皆秋色，________。",
                    "blankText": "________（1分）",
                    "score": 1.0,
                    "_tokenType": "paren_arabic",
                    "children": [],
                },
                {
                    "lineNumber": 3,
                    "level": 2,
                    "numbering": "2",
                    "title": "（2）浮云游子意",
                    "rawText": "（2）________，落日故人情。",
                    "blankText": "________（1分）",
                    "score": 1.0,
                    "_tokenType": "paren_arabic",
                    "children": [],
                },
                {
                    "lineNumber": 4,
                    "level": 2,
                    "numbering": "7",
                    "title": "口语交际",
                    "rawText": "7、口语交际。（2分）",
                    "blankText": "",
                    "score": 2.0,
                    "_tokenType": "arabic",
                    "children": [],
                },
                {
                    "lineNumber": 5,
                    "level": 2,
                    "numbering": "8",
                    "title": "综合活动",
                    "rawText": "8、综合活动。（6分）",
                    "blankText": "",
                    "score": 6.0,
                    "_tokenType": "arabic",
                    "children": [],
                },
            ],
        }
    ]

    changed = processor._repair_overflow_arabic_children_to_root_siblings(outline_items)

    assert changed is True
    assert [item["numbering"] for item in outline_items] == ["6", "7", "8"]
    assert [child["numbering"] for child in outline_items[0]["children"]] == ["1", "2"]
    assert outline_items[1]["level"] == 1
    assert outline_items[2]["level"] == 1


def test_split_lines_by_answer_should_not_trigger_on_choice_option_lines_in_late_body():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "一、本卷（60分）在每小题给出的四个选项中，只有一项是最符合题目要求的。",
        "1．题目1",
        "A．选项A",
        "B．选项B",
        "C．选项C",
        "D．选项D",
        "2．题目2",
        "A．选项A",
        "B．选项B",
        "C．选项C",
        "D．选项D",
        "3．题目3",
        "A．选项A",
        "B．选项B",
        "C．选项C",
        "D．选项D",
        "21．下列灾害属于我国山区常见的一组是",
        "A．崩塌、滑坡、泥石流 B．滑坡、洪涝、干旱",
        "C．泥石流、寒潮、干旱 D．滑坡、泥石流、台风",
        "22．下列资源中，属于非可再生资源的",
        "A．矿产资源 B．土地资源",
        "C．森林资源 D．水资源",
        "二、本卷共（40分）",
        "31．（11分）阅读图文材料，回答问题。",
    ]

    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)

    assert main_lines == lines
    assert answer_lines == []


def test_tail_section_inline_answer_start_should_not_trigger_on_choice_option_line_from_real_doc():
    processor = DocumentProcessor()
    source = _find_repo_doc("515178e3cf914c6b8fe1432460b7884c.docx")
    lines = processor._collect_docx_lines(Document(str(source)))

    assert processor._looks_like_tail_section_inline_answer_start(lines, 87) is False


def test_split_lines_by_answer_should_trigger_on_answer_heading():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "一、选择题（40分）",
        "1. 第一题",
        "七年级地理月考答案",
        "1C 2B 3A",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines[:2]
    assert answer_lines == lines[2:]


def test_split_lines_by_answer_should_handle_inline_answer_marker_near_tail():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "一、选择题：（30分）",
        "1. 第一题",
        "2. 第二题",
        "[来源] 答案：1-5.CBBCA; 6-10.AABDC;",
        "11. A",
        "12. B",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines[:3] + ["[来源]"]
    assert answer_lines == ["答案：1-5.CBBCA; 6-10.AABDC;", "11. A", "12. B"]


def test_split_lines_by_answer_should_not_split_when_first_line_is_title_with_answer_word():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "南开教育八年级地理上册期末考试试题与答案",
        "一、单项选择题（10分）",
        "1．我国濒临的海洋，从北到南依次为（）",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines
    assert answer_lines == []


def test_split_lines_by_answer_with_loose_keyword_should_not_split_instruction_sentence():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(["答案"])
    lines = [
        "一、选择题（60分）选择题答案对应填写在题后的答题卡中。",
        "1. 下列说法正确的是（）",
        "A. 选项A",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines
    assert answer_lines == []


def test_split_lines_by_answer_with_loose_keyword_should_split_on_answer_heading():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(["答案"])
    lines = [
        "一、选择题（40分）",
        "1. 第一题",
        "参考答案与解析",
        "1C 2B 3A",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines[:2]
    assert answer_lines == lines[2:]


def test_split_lines_by_answer_should_split_on_answer_sheet_heading_even_without_answer_keyword():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(["答案"])
    lines = [
        "一、选择题（24分）",
        "1. 第一题",
        "初三物理12月份学情调研试卷（答题卷）（2015-12）",
        "13、__________、__________.",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines[:2]
    assert answer_lines == lines[2:]


def test_split_lines_by_answer_should_not_split_on_answer_sheet_instruction_sentence():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(["答案"])
    lines = [
        "一、选择题（60分）",
        "1. 下列说法正确的是（）",
        "请将答案填写在答题卡指定位置。",
        "A. 选项A",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines
    assert answer_lines == []


def test_split_lines_by_answer_should_split_on_tail_answer_card_hint():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "一、听力（20分）",
        "1. 第一题",
        "二、单项选择（15分）",
        "21. 第二十一题",
        "三、完形填空（15分）",
        "36. 第三十六题",
        "四、阅读理解（30分）",
        "51. 第五十一题",
        "五、词汇运用（10分）",
        "66. 第六十六题",
        "六、任务型阅读（10分）",
        "76. 第七十六题",
        "七、缺词填空（10分）",
        "86. 第八十六题",
        "八、书面表达（30分）",
        "96. 第九十六题",
        "友情提醒：请将答案填在答题卡上！",
        "第一部分 听对话，回答问题",
        "1. M: Hello.",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines[:16]
    assert answer_lines == lines[16:]


def test_split_lines_by_answer_should_split_on_answer_sheet_paper_heading():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(["答案"])
    lines = [
        "一、选择题（24分）",
        "1. 第一题",
        "七年级数学答题纸（2024）",
        "三、解答题（30分）",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines[:2]
    assert answer_lines == lines[2:]


def test_process_docx_fallbacks_to_pdf_lines_when_docx_package_is_invalid(tmp_path, monkeypatch):
    source = tmp_path / "broken.docx"
    source.write_text("not-a-real-docx", encoding="utf-8")

    processor = DocumentProcessor()
    fake_pdf = tmp_path / "fallback.pdf"
    fake_pdf.write_bytes(b"%PDF-1.4\n%%EOF")

    def fake_convert(_office_path):
        return fake_pdf

    def fake_extract_pdf_lines(_pdf_path):
        return [
            ["涓€銆侀€夋嫨棰橈紙10鍒嗭級", "1. 绗竴棰橈紙2鍒嗭級", "绛旀"],
            ["1A 2B"],
        ]

    def fake_split_pdf_preview(_task_id, _source_pdf, _answer_start_page):
        return ["/storage/pages/main.pdf"], ["/storage/pages/answer.pdf"]

    monkeypatch.setattr(processor, "_convert_office_to_pdf", fake_convert)
    monkeypatch.setattr(processor, "_extract_pdf_page_lines", fake_extract_pdf_lines)
    monkeypatch.setattr(processor, "_split_pdf_preview", fake_split_pdf_preview)

    output = processor.process(
        task_id=777001,
        file_path=str(source),
        file_ext=".docx",
        max_level=8,
        second_level_mode="auto",
        answer_section_patterns=None,
        score_patterns=None,
    )

    assert output["result"]["mainPages"] == ["/storage/pages/main.pdf"]
    assert output["result"]["answerPages"] == ["/storage/pages/answer.pdf"]
    assert output["details"]["outlineItems"]
    assert len(output["details"]["headerFooterItems"]) == 2


def test_process_docx_uses_pdf_lines_when_docx_paragraph_lines_empty(tmp_path, monkeypatch):
    source = tmp_path / "normal.docx"
    doc = Document()
    doc.add_paragraph("placeholder")
    doc.save(str(source))

    processor = DocumentProcessor()
    fake_pdf = tmp_path / "fallback2.pdf"
    fake_pdf.write_bytes(b"%PDF-1.4\n%%EOF")

    monkeypatch.setattr(processor, "_collect_docx_lines", lambda _doc: [])
    monkeypatch.setattr(processor, "_collect_docx_header_footer", lambda _doc: ([], []))
    monkeypatch.setattr(processor, "_convert_office_to_pdf", lambda _office_path: fake_pdf)
    monkeypatch.setattr(
        processor,
        "_extract_pdf_page_lines",
        lambda _pdf_path: [["涓€銆佸崟閫夐锛?0鍒嗭級", "1. 棰樼洰锛?鍒嗭級", "绛旀"], ["1A"]],
    )
    monkeypatch.setattr(
        processor,
        "_split_pdf_preview",
        lambda _task_id, _source_pdf, _answer_start_page: (["/storage/pages/main2.pdf"], ["/storage/pages/answer2.pdf"]),
    )

    output = processor.process(
        task_id=777002,
        file_path=str(source),
        file_ext=".docx",
        max_level=8,
        second_level_mode="auto",
        answer_section_patterns=None,
        score_patterns=None,
    )

    assert output["details"]["outlineItems"], "outline should come from fallback pdf lines"
    assert output["result"]["mainPages"] == ["/storage/pages/main2.pdf"]
    assert output["result"]["answerPages"] == ["/storage/pages/answer2.pdf"]


def test_process_docx_fallbacks_when_office_conversion_fails(tmp_path, monkeypatch):
    source = tmp_path / "broken2.docx"
    source.write_text("not-a-real-docx", encoding="utf-8")

    processor = DocumentProcessor()
    fake_pdf = tmp_path / "generated_fallback.pdf"
    fake_pdf.write_bytes(b"%PDF-1.4\n%%EOF")

    monkeypatch.setattr(
        processor,
        "_convert_office_to_pdf",
        lambda _office_path: (_ for _ in ()).throw(RuntimeError("conversion failed")),
    )
    monkeypatch.setattr(processor, "_read_fallback_text_lines", lambda _path: ["1. fallback title锛?鍒嗭級", "绛旀"])
    monkeypatch.setattr(processor, "_create_fallback_pdf", lambda _task_id, _lines, _tag="fallback_preview": fake_pdf)
    monkeypatch.setattr(processor, "_split_pdf_preview", lambda _task_id, _source_pdf, _answer_start_page: (["/storage/pages/fallback_main.pdf"], []))

    output = processor.process(
        task_id=777003,
        file_path=str(source),
        file_ext=".docx",
        max_level=8,
        second_level_mode="auto",
        answer_section_patterns=None,
        score_patterns=None,
    )

    assert output["result"]["mainPages"] == ["/storage/pages/fallback_main.pdf"]
    assert output["details"]["outlineItems"]


def test_process_docx_should_infer_answer_preview_split_from_line_level_split(tmp_path, monkeypatch):
    source = tmp_path / "docx_split_preview.docx"
    doc = Document()
    doc.add_paragraph("placeholder")
    doc.save(str(source))

    processor = DocumentProcessor()
    fake_pdf = tmp_path / "preview.pdf"
    fake_pdf.write_bytes(b"%PDF-1.4\n%%EOF")

    lines = [
        "一、单项选择题（69分）",
        "24. 题干内容（4分）",
        "25. 题干内容（8分）",
        "26. 题干内容（10分）",
        "27. 题干内容（9分）",
        "24. 0.1s 0.3m/s",
        "25. (1)低压直流电源 秒表 天平 交流电源",
        "26. (1) t=1s (2) x=10m (3) v=10m/s",
    ]
    page_lines = [lines[:5], lines[5:]]
    captured = {}

    def fake_split_pdf_preview(_task_id, _source_pdf, answer_start_page):
        captured["answer_start_page"] = answer_start_page
        if answer_start_page is None:
            return ["/storage/pages/main.pdf"], []
        return ["/storage/pages/main.pdf"], ["/storage/pages/answer.pdf"]

    monkeypatch.setattr(processor, "_collect_docx_lines", lambda _doc: lines)
    monkeypatch.setattr(processor, "_collect_docx_header_footer", lambda _doc: ([], []))
    monkeypatch.setattr(processor, "_convert_office_to_pdf", lambda _office_path: fake_pdf)
    monkeypatch.setattr(processor, "_extract_pdf_page_lines", lambda _pdf_path: page_lines)
    monkeypatch.setattr(processor, "_split_pdf_preview", fake_split_pdf_preview)

    output = processor.process(
        task_id=777004,
        file_path=str(source),
        file_ext=".docx",
        max_level=8,
        second_level_mode="auto",
        answer_section_patterns=None,
        score_patterns=None,
    )

    assert captured["answer_start_page"] == 1
    assert output["result"]["answerPages"] == ["/storage/pages/answer.pdf"]


def test_process_pdf_should_infer_answer_preview_split_from_line_level_split(tmp_path, monkeypatch):
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)

    source = tmp_path / "split_source.pdf"
    source.write_bytes(b"%PDF-1.4\n%%EOF")

    page_lines = [
        [
            "一、单项选择题（69分）",
            "24. 题干内容（4分）",
            "25. 题干内容（8分）",
            "26. 题干内容（10分）",
            "27. 题干内容（9分）",
        ],
        [
            "24. 0.1s 0.3m/s",
            "25. (1)低压直流电源 秒表 天平 交流电源",
            "26. (1) t=1s (2) x=10m (3) v=10m/s",
        ],
    ]
    captured = {}

    def fake_split_pdf_preview(_task_id, _source_pdf, answer_start_page):
        captured["answer_start_page"] = answer_start_page
        if answer_start_page is None:
            return ["/storage/pages/main.pdf"], []
        return ["/storage/pages/main.pdf"], ["/storage/pages/answer.pdf"]

    monkeypatch.setattr(processor, "_extract_pdf_page_lines", lambda _file_path: page_lines)
    monkeypatch.setattr(processor, "_split_pdf_preview", fake_split_pdf_preview)

    output = processor._process_pdf(
        task_id=777005,
        file_path=source,
        answer_patterns=patterns,
    )

    assert captured["answer_start_page"] == 1
    assert output[4] == ["/storage/pages/answer.pdf"]


def test_process_pdf_should_prefer_line_split_page_over_false_page_level_answer_match(tmp_path, monkeypatch):
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)

    source = tmp_path / "split_source_false_page_match.pdf"
    source.write_bytes(b"%PDF-1.4\n%%EOF")

    page_lines = [
        [
            "一、阅读理解（30分）",
            "31. 阅读短文，选择正确答案（ ）",
            "32. 第二题",
        ],
        [
            "参考答案",
            "31.A 32.C",
        ],
    ]
    captured = {}

    def fake_split_pdf_preview(_task_id, _source_pdf, answer_start_page):
        captured["answer_start_page"] = answer_start_page
        if answer_start_page is None:
            return ["/storage/pages/main.pdf"], []
        return ["/storage/pages/main.pdf"], ["/storage/pages/answer.pdf"]

    original_matches_answer = processor._matches_answer

    def fake_matches_answer(text, answer_patterns):
        # Simulate an over-greedy page-level match on page-1 full text.
        if "\n" in text and "阅读短文，选择正确答案" in text:
            return True
        return original_matches_answer(text, answer_patterns)

    monkeypatch.setattr(processor, "_matches_answer", fake_matches_answer)
    monkeypatch.setattr(processor, "_extract_pdf_page_lines", lambda _file_path: page_lines)
    monkeypatch.setattr(processor, "_split_pdf_preview", fake_split_pdf_preview)

    output = processor._process_pdf(
        task_id=777006,
        file_path=source,
        answer_patterns=patterns,
    )

    assert captured["answer_start_page"] == 1
    assert output[4] == ["/storage/pages/answer.pdf"]


def test_split_lines_by_answer_detects_tail_answer_key_without_heading():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "一、单项选择题（69分）",
        "24. 题干内容（4分）",
        "25. 题干内容（8分）",
        "26. 题干内容（10分）",
        "27. 题干内容（9分）",
        "24. 0.1s 0.3m/s",
        "25. (1)低压直流电源 秒表 天平 交流电源",
        "26. (1) t=1s (2) x=10m (3) v=10m/s",
        "27. (1) v=4m/s (2) N=300N (3) 80J",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines[:5]
    assert answer_lines == lines[5:]


def test_split_lines_by_answer_keeps_legit_section_number_restart():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "一、选择题（10分）",
        "1. 第一题（5分）",
        "2. 第二题（5分）",
        "二、填空题（10分）",
        "1. 填空一（5分）",
        "2. 填空二（5分）",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines
    assert answer_lines == []


def test_split_lines_by_answer_should_not_split_question_heading_with_correct_answer_phrase():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "一、听力（10分）",
        "1. 第一题",
        "Ⅱ. 听对话，选择正确答案（5分）",
        "11. What day is it today?",
        "A. Monday",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines
    assert answer_lines == []


def test_split_lines_by_answer_should_not_split_question_instruction_with_best_answer_phrase():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "三、阅读理解（30分）",
        "阅读短文，选择最佳答案。",
        "31. 第一题",
        "A. 选项A",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines
    assert answer_lines == []


def test_split_lines_by_answer_should_not_split_on_material_analysis_question_section_heading():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "一、选择题（14分）",
        "1. 第一题（1分）",
        "二、识图释图题（6分）",
        "15. 读图完成问题（6分）",
        "三、材料解析题（24分）",
        "16. 阅读材料，回答问题（12分）",
        "四、列举简答题（12分）",
        "18. 简答题（12分）",
        "参考答案",
        "一、选择题",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines[:8]
    assert answer_lines == lines[8:]


def test_split_lines_by_answer_should_split_on_tail_listening_material_heading():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "一、听力（20分）",
        "1. 第一题",
        "二、单项选择（15分）",
        "21. 第二十一题",
        "三、完形填空（15分）",
        "36. 第三十六题",
        "四、阅读理解（30分）",
        "51. 第五十一题",
        "五、词汇运用（10分）",
        "66. 第六十六题",
        "六、任务型阅读（10分）",
        "76. 第七十六题",
        "七、缺词填空（10分）",
        "86. 第八十六题",
        "八、书面表达（30分）",
        "96. 第九十六题",
        "一、听力材料",
        "1. M: Hello.",
        "2. W: Good morning.",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines[:16]
    assert answer_lines == lines[16:]


def test_split_lines_by_answer_should_split_tail_compact_chinese_answer_key_block():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "一、字词积累（29分）",
        "1. 第一题（3分）",
        "2. 第二题（3分）",
        "七、快乐习作（25分）",
        "13. 写作题（25分）",
        "第一单元综合素质评价",
        "一、1.(1)C (2)A (3)C",
        "2.(1)这里指屋子最高最上的部分。",
        "三、5.(1)× (2)× (3)√",
        "六、12. 示例：明天小区管道维修。",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines[:5]
    assert answer_lines == lines[5:]


def test_split_lines_by_answer_should_split_on_tail_section_heading_with_inline_answer_keys():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "1． 选择题（30分）",
        "1．第一题（2分）",
        "2．第二题（2分）",
        "三，实验题。（10分）",
        "25．实验题（5分）",
        "四，计算题。（20分）",
        "27．计算题（7分）",
        "一．选择题：CDBAA DBABD AACBD",
        "三：实验题",
        "25．答案略",
        "四：计算题",
        "27．答案略",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines[:7]
    assert answer_lines == lines[7:]


def test_split_lines_by_answer_should_not_split_before_scored_question_33_34_block():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "一、单项选择题（30分）",
        "30．由图可知，北京易发生的自然灾害是",
        "二、读图题（20分）",
        "31．读“中国空白政区图”，完成下列问题。（5分）",
        "32．读我国地形分布示意图，完成下列各题。（每空1分，共5分）",
        "33．读长江水系图，完成下列要求。",
        "（3）B是长江流域最长的支流________。",
        "34．读我国沿北纬36度附近地形剖面图，完成下列问题．（每空1分，共5分）",
        "八年级地理试卷答案",
        "1、 选择题：",
        "1-5   AADCC  6-10  CCCCC",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines[:8]
    assert answer_lines == lines[8:]

def test_split_lines_by_answer_should_not_split_on_scored_question_with_checkmark_symbols():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "四、阅读检阅台(21分)",
        "11．给加点字选择正确的读音，打\"√\"。(3分)",
        "撒(sā sǎ)谎 鸡爪(zhuǎ zhǎo) 骨折(zhé shé)",
        "12．从文中可以看出母亲是一个怎样的人？(3分)",
        "五、习作百花园(25分)",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines
    assert answer_lines == []


def test_split_lines_by_answer_should_not_split_tail_question_instruction_block_before_scored_next_section():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "八、课外阅读。（12分）",
        "1. 用“√”画出文中括号中恰当的词语。",
        "2. 在文中的括号里填上恰当的关联词。",
        "3. 把第一个画线句改为陈述句。",
        "_______________________________________________",
        "4. 给文章加个题目。",
        "九、习作。（30分）",
        "每当放假的时候……",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines
    assert answer_lines == []

def test_split_lines_by_answer_should_detect_spaced_answer_heading():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "三、作文。（40分）",
        "19、阅读下面材料，按要求作文。",
        "参 考 答 案",
        "一、1.答案:D",
        "2.答案:C",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines[:2]
    assert answer_lines == lines[2:]


def test_split_lines_by_answer_should_not_split_on_answer_invalid_instruction():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "2023年广州市初中学业水平考试语文",
        "注意事项：",
        "1.答题前，考生务必在答题卡上填写姓名。",
        "2.答案不能答在试卷上。",
        "3.不按以上要求作答的",
        "答案无效。",
        "第一部分积累与运用（共24分）",
        "一、（16分）",
        "1. 下列词语中，每对加点字的读音都相同的一项是（ ）",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines
    assert answer_lines == []


def test_split_lines_by_answer_should_ignore_first_line_title_with_answer_word_and_split_on_tail_heading():
    processor = DocumentProcessor()
    patterns = processor._compile_answer_patterns(None)
    lines = [
        "七年级地理试题与答案",
        "一、单项选择题（10分）",
        "1．第一题（2分）",
        "2．第二题（2分）",
        "参考答案",
        "1．B 2．C",
    ]
    main_lines, answer_lines = processor._split_lines_by_answer(lines, patterns)
    assert main_lines == lines[:4]
    assert answer_lines == lines[4:]


def test_repair_missing_root_heading_from_child_restart_should_split_duplicate_subquestion_runs():
    processor = DocumentProcessor()
    outline_items = [
        {
            "lineNumber": 29,
            "level": 1,
            "numbering": "5",
            "title": "观察下列图片：",
            "rawText": "5、观察下列图片：",
            "blankText": "",
            "score": 16.0,
            "children": [
                {"lineNumber": 32, "level": 2, "numbering": "1", "title": "第一问", "rawText": "（1）第一问（2分）", "blankText": "", "score": 2.0, "children": []},
                {"lineNumber": 33, "level": 2, "numbering": "2", "title": "第二问", "rawText": "（2）第二问（2分）", "blankText": "", "score": 2.0, "children": []},
                {"lineNumber": 34, "level": 2, "numbering": "3", "title": "第三问", "rawText": "（3）第三问（2分）", "blankText": "", "score": 2.0, "children": []},
                {"lineNumber": 38, "level": 2, "numbering": "1", "title": "第四问", "rawText": "（1）第四问（3分）", "blankText": "", "score": 3.0, "children": []},
                {"lineNumber": 39, "level": 2, "numbering": "2", "title": "第五问", "rawText": "（2）第五问（4分）", "blankText": "", "score": 4.0, "children": []},
                {"lineNumber": 42, "level": 2, "numbering": "3", "title": "第六问", "rawText": "（3）第六问（3分）", "blankText": "", "score": 3.0, "children": []},
            ],
        },
        {
            "lineNumber": 43,
            "level": 1,
            "numbering": "7",
            "title": "阅读下列材料：",
            "rawText": "7、阅读下列材料：",
            "blankText": "",
            "score": 6.0,
            "children": [],
        },
    ]

    changed = processor._repair_missing_root_heading_from_child_restart(outline_items)

    assert changed is True
    assert [item["numbering"] for item in outline_items] == ["5", "6", "7"]
    assert [child["numbering"] for child in outline_items[0]["children"]] == ["1", "2", "3"]
    assert [child["numbering"] for child in outline_items[1]["children"]] == ["1", "2", "3"]
    assert outline_items[0]["score"] == 6.0
    assert outline_items[1]["score"] == 10.0


def test_process_248_docx_should_not_extract_reference_paren_numbers_as_subquestions():
    processor = DocumentProcessor()
    source = _find_repo_doc("248.docx")

    output = processor.process(
        task_id=999248,
        file_path=str(source),
        file_ext=".docx",
        max_level=8,
        second_level_mode="auto",
        answer_section_patterns=None,
        score_patterns=None,
    )

    roots = output["details"]["outlineItems"]
    reading_root = next(item for item in roots if item.get("numbering") == "二")
    narrative_section = next(item for item in reading_root["children"] if item.get("numbering") == "三")

    assert [child.get("numbering") for child in narrative_section["children"]] == ["1", "2", "3", "4", "5", "6"]

    q2 = narrative_section["children"][1]
    assert q2["score"] == 2.0
    assert q2["children"] == []

    q6 = narrative_section["children"][5]
    assert q6["score"] == 3.0
    assert q6["children"] == []


def test_process_250_docx_should_not_extract_reading_article_title_as_child_node():
    processor = DocumentProcessor()
    source = _find_repo_doc("250.docx")

    output = processor.process(
        task_id=999250,
        file_path=str(source),
        file_ext=".docx",
        max_level=8,
        second_level_mode="auto",
        answer_section_patterns=None,
        score_patterns=None,
    )

    roots = output["details"]["outlineItems"]
    modern_reading_root = next(item for item in roots if item.get("numbering") == "二")

    assert [child.get("numbering") for child in modern_reading_root["children"]] == ["7", "8", "9", "10"]
    assert all(str(child.get("numbering") or "").strip() for child in modern_reading_root["children"])


def test_reference_continuation_node_should_reject_long_parenthesized_sentence():
    processor = DocumentProcessor()
    node = {
        "rawText": "（1）陪一个父亲，去八百里外的戒毒所，探视他在那里戒毒的儿子。戒毒结束已经一年了。"
    }
    assert processor._is_reference_continuation_node(node) is False


def test_reference_continuation_node_should_accept_score_only_paren_line():
    processor = DocumentProcessor()
    node = {"rawText": "（1分）"}
    assert processor._is_reference_continuation_node(node) is True
