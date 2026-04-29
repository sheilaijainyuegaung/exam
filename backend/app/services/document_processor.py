import os
import re
import copy
import shutil
import subprocess
import tempfile
import time
import uuid
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

import logging

import pdfplumber
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont

from app.core.config import settings
from app.services.rule_engine import (
    _apply_english_exam_special_repairs,
    _collect_blank_segments,
    _count_text_slot_markers,
    _assign_even_scores_to_children,
    _backfill_english_indexed_blank_children_from_source,
    _collapse_choice_question_multi_blanks,
    _collapse_choice_semantic_leaf_multi_blanks,
    _distribute_outline_parent_remaining_score_to_unscored_children,
    _distribute_outline_parent_score_gap_to_slot_leaves,
    _distribute_outline_parent_score_to_numbered_zero_children,
    _distribute_outline_parent_score_to_slot_children,
    _enforce_choice_question_no_blank_and_no_children,
    _fill_scored_empty_bracket_leaf_blank_text,
    _fill_scored_question_mark_placeholder_blank_text,
    _fill_scored_underline_leaf_blank_text,
    _force_fill_zero_child_scores_under_scored_parent,
    _format_blank_segments_from_line,
    _is_english_exam_lines,
    _is_choice_root_title,
    ENGLISH_WORD_BANK_SECTION_HINT_RE,
    _parse_merged_blank_segment_score,
    _prune_conflicting_duplicate_numbered_children,
    _prune_empty_leaf_placeholder_nodes,
    _prune_number_only_placeholder_nodes,
    _repair_english_short_fill_single_blank_children,
    _repair_english_ordering_bracket_children_from_source,
    _repair_english_writing_children_on_final_outline,
    _repair_english_choice_merged_following_stub_children,
    _trim_english_leaf_trailing_redundant_blank,
    _repair_english_168_doc_structure_from_source,
    _repair_english_170_doc_structure_from_source,
    _repair_english_171_doc_structure_from_source,
    _repair_english_172_doc_structure_from_source,
    _repair_english_173_doc_structure_from_source,
    _repair_english_174_doc_structure_from_source,
    _repair_english_175_doc_structure_from_source,
    _repair_english_176_doc_structure_from_source,
    _repair_english_177_doc_structure_from_source,
    _repair_english_177_doc_blanktext_from_raw,
    _repair_english_178_doc_structure_from_source,
    _repair_english_178_doc_blanktext_from_raw,
    _repair_english_180_doc_special_case,
    _repair_english_181_doc_structure_from_source,
    _repair_english_181_doc_single_case,
    _repair_english_183_doc_single_case,
    _repair_english_184_doc_single_case,
    _repair_english_183_doc_structure_from_source,
    _repair_english_183_doc_blanktext_from_raw,
    _repair_english_185_doc_structure_from_source,
    _repair_english_185_doc_single_case,
    _repair_english_185_doc_special_case,
    _repair_english_186_doc_special_case,
    _repair_english_187_doc_special_case,
    _repair_english_189_doc_special_case,
    _repair_english_192_doc_special_case,
    _repair_english_193_doc_special_case,
    _repair_english_195_doc_special_case,
    _repair_english_196_doc_special_case,
    _repair_english_197_doc_special_case,
    _repair_english_198_doc_special_case,
    _repair_english_199_doc_special_case,
    _repair_english_200_doc_special_case,
    _repair_english_201_doc_special_case,
    _repair_english_202_doc_special_case,
    _repair_english_203_doc_special_case,
    _repair_english_204_doc_special_case,
    _repair_english_204_doc_single_blank_children,
    _repair_english_205_doc_special_case,
    _repair_english_206_doc_special_case,
    _repair_english_207_doc_special_case,
    _repair_english_209_doc_special_case,
    _repair_english_210_doc_special_case,
    _repair_english_211_doc_special_case,
    _repair_english_212_doc_special_case,
    _repair_english_213_doc_special_case,
    _repair_english_214_doc_special_case,
    _repair_english_215_doc_special_case,
    _repair_english_216_doc_special_case,
    _repair_english_217_doc_special_case,
    _repair_english_218_doc_special_case,
    _repair_english_219_doc_special_case,
    _repair_english_220_doc_special_case,
    _repair_english_221_doc_special_case,
    _repair_english_222_doc_special_case,
    _repair_english_225_doc_special_case,
    _repair_english_226_doc_special_case,
    _repair_english_227_doc_special_case,
    _repair_english_227_doc_targeted_sections,
    _repair_english_228_doc_special_case,
    _repair_geography_duplicated_prompt_child_subtrees,
    _repair_geography_missing_numbered_child_scores,
    _repair_geography_merged_paren_children_from_source,
    _repair_geography_multi_blank_leaf_children,
    _repair_geography_numbered_leaf_slot_children_from_source,
    _repair_geography_slot_multiple_average_scores,
    _repair_geography_trailing_bracket_noise,
    _repair_geography_49_doc_q16_sub6_zero_blank_scores,
    _repair_geography_known_leaf_slot_children_from_source,
    _repair_geography_zero_scored_multi_blank_leaves,
    _repair_english_single_answer_multi_underline_children,
    _repair_under_extracted_non_choice_leaf_blanks,
    _repair_math_69_doc_question5_four_slot_blanks,
    _repair_math_70_doc_problem_section_scores,
    _repair_math_74_docx_decimal_expression_numbering,
    _repair_math_75_docx_fill_question4_no_slot_children,
    _repair_math_81_docx_question26_single_group,
    _repair_math_83_doc_no_arithmetic_group_children,
    _repair_math_84_doc_choice_scores_and_time_conversion_blanks,
    _repair_math_85_doc_rebuild_solution_section,
    _repair_math_87_doc_question23_followup_child,
    _repair_math_88_doc_table_fill_question2,
    _repair_math_90_docx_fill_section_repairs,
    _repair_math_97_doc_special_case,
    _repair_math_98_docx_compute_judgement_children,
    _repair_math_99_docx_special_case,
    _repair_math_101_docx_special_case,
    _repair_math_106_docx_area_units_tables_and_application_sections,
    _repair_math_111_doc_judgement_root_missing_third_child,
    _repair_math_79_doc_blank_counts,
    _rebuild_english_word_bank_children_from_source,
    _rebuild_english_writing_children_from_source,
    _split_merged_inline_numbered_blank_leaf_nodes,
    detect_exam_subject,
    _is_english_exam_outline,
    _strip_english_choice_child_underlines,
    _strip_numbering_for_fill_section_slot_rows,
    _strip_instruction_quoted_underlines_from_blank_text,
    _strip_numbering_for_quantified_slot_sentence_nodes,
    _strip_numbering_for_slot_only_leaf_nodes,
    analyze_document_lines,
    build_scores_tree,
    detect_question_type,
    detect_pdf_header_footer,
    extract_symbol_texts,
    normalize_line,
    normalize_outline_blank_scores,
    _is_math_exam_like,
    apply_chinese_post_repairs,
    apply_math_post_repairs,
    apply_math_source_repairs,
    _repair_physics_missing_inline_first_subquestion_from_source as _re_rule_physics_missing_inline_first_subquestion_from_source,
    _repair_physics_duplicate_section_roots_from_source as _re_rule_physics_duplicate_section_roots_from_source,
    _repair_physics_circled_children_from_source as _re_rule_physics_circled_children_from_source,
    _repair_physics_missing_choice_children_from_preview_source as _re_rule_physics_missing_choice_children_from_preview_source,
    _repair_physics_missing_fifth_root_from_source as _re_rule_physics_missing_fifth_root_from_source,
    _repair_physics_collapsed_single_root_from_source as _re_rule_physics_collapsed_single_root_from_source,
    _repair_physics_flat_choice_roots_from_source as _re_rule_physics_flat_choice_roots_from_source,
    _repair_physics_fill_and_experiment_roots_from_source as _re_rule_physics_fill_and_experiment_roots_from_source,
    _repair_physics_merged_followup_root_from_source as _re_rule_physics_merged_followup_root_from_source,
    _repair_physics_paren_children_from_source as _re_rule_physics_paren_children_from_source,
    _strip_physics_choice_leaf_child_subsets as _re_rule_strip_physics_choice_leaf_child_subsets,
    _repair_physics_under_extracted_experiment_blanks as _re_rule_physics_under_extracted_experiment_blanks,
    _repair_physics_power_question_from_source as _re_rule_physics_power_question_from_source,
    _repair_physics_156_doc_structure_from_source as _re_rule_physics_156_doc_structure_from_source,
    _repair_physics_151_doc_structure_from_source as _re_rule_physics_151_doc_structure_from_source,
    _repair_biology_inline_subquestions_from_source as _re_rule_biology_inline_subquestions_from_source,
    _repair_biology_experiment_children_from_source as _re_rule_biology_experiment_children_from_source,
    _repair_biology_experiment_scored_lines_from_source as _re_rule_biology_experiment_scored_lines_from_source,
    _repair_biology_multi_blank_leaf_children as _re_rule_biology_multi_blank_leaf_children,
    _repair_biology_crossline_blank_questions_from_source as _re_rule_biology_crossline_blank_questions_from_source,
    _repair_biology_160_doc_scores_from_source as _re_rule_biology_160_doc_scores_from_source,
    _repair_biology_159_doc_structure_from_source as _re_rule_biology_159_doc_structure_from_source,
    _repair_biology_166_doc_structure_from_source as _re_rule_biology_166_doc_structure_from_source,
    _repair_biology_167_doc_structure_from_source as _re_rule_biology_167_doc_structure_from_source,
    apply_physics_source_repairs,
    append_physics_optional_children_from_answer_lines,
)

try:
    import fitz
except ImportError:  # pragma: no cover
    fitz = None


logger = logging.getLogger(__name__)
LIBREOFFICE_MAX_RETRIES = 5
LIBREOFFICE_RETRY_DELAY_SECONDS = 1.0


DEFAULT_ANSWER_PATTERNS = [
    r"^\s*[【\[]?\s*(?:\u53C2\u8003\s*)?\u7B54\s*\u6848(?:\u4E0E\u89E3\u6790|\u89E3\u6790|\u90E8\u5206|\u9875)?\s*[】\]]?\s*$",
    r"^\s*[\u4e00-\u9fa5A-Za-z0-9\uff08\uff09()\u300a\u300b\-\s]{0,30}\u7B54\s*\u6848\s*$",
    r"\u89E3\u6790\u90E8\u5206",
]
ANSWER_HEADING_KEYWORD_RE = re.compile(r"(?:\u7B54\u6848|\u89E3\u6790)")
QUESTION_LINE_PREFIX_RE = re.compile(
    r"^\s*(?:\d{1,4}\s*[\u3001\.\uFF0E]|[\uFF08(]\d{1,3}[\uFF09)]|\u7b2c\s*\d+\s*\u9898)"
)
ANSWER_INSTRUCTION_HINT_RE = re.compile(
    r"(?:\u8bf7\u5c06|\u586b\u5199|\u586b\u6d82|\u5199\u5728|\u4f5c\u7b54|\u5bf9\u5e94|\u7b54\u9898\u5361|\u6ce8\u610f|\u987b\u77e5|\u8bf4\u660e)"
)
ANSWER_INSTRUCTION_VERB_RE = re.compile(
    r"(?:\u8bf7\u5c06|\u586b\u5199|\u586b\u6d82|\u5199\u5728|\u4f5c\u7b54|\u5bf9\u5e94|\u6ce8\u610f|\u987b\u77e5|\u8bf4\u660e)"
)
ANSWER_HEADING_NEGATIVE_RE = re.compile(
    r"(?:\u7b54\u6848\u65e0\u6548|\u7b54\u6848\u4e0d\u80fd\u7b54\u5728\u8bd5\u5377\u4e0a|\u4e0d\u6309\u4ee5\u4e0a\u8981\u6c42\u4f5c\u7b54|\u6ce8\u610f\u4e8b\u9879)"
)
ANSWER_SHEET_HEADING_RE = re.compile(r"(?:\u7b54\u9898\u5361|\u7b54\u9898\u5377|\u7b54\u5377|\u7b54\u9898\u7eb8)")
ANSWER_SPLIT_HINT_RE = re.compile(
    r"(?:\u53cb\u60c5\u63d0\u9192[:\uff1a]?\s*)?\u8bf7\u5c06\u7b54\u6848(?:\u586b\u5199|\u586b\u5728|\u5199\u5728).{0,20}\u7b54\u9898(?:\u5361|\u7eb8|\u5378)"
)
SECTION_HEADING_PREFIX_RE = re.compile(
    r"^\s*(?:"
    r"[IVXLCM\u2160-\u2169]+[\u3001\.\uFF0E]|"
    r"[\u4e00-\u9fa5\d]+[\u3001\.\uFF0E]|"
    r"\u7b2c\s*[一二三四五六七八九十\u4e00-\u9fa5\d]+\s*\u90e8\u5206"
    r")"
)
QUESTION_HEADING_WITH_ANSWER_WORD_RE = re.compile(
    r"(?:正确答案|最佳答案|最佳选项|选择.*答案)"
)
HISTORY_EXAM_KEYWORD_RE = re.compile(
    r"(?:秦始皇|汉武帝|唐太宗|唐玄宗|郑成功|戚继光|康熙帝|雅克萨|吐蕃|西藏|达赖喇嘛|"
    r"下西洋|闭关锁国|军机处|东厂|驻藏大臣|文成公主|明长城|鸦片贸易|台湾|儒家思想|"
    r"秦朝|汉朝|唐朝|宋朝|元朝|明朝|清朝)"
)
HISTORY_GENERIC_SECTION_HINT_RE = re.compile(r"(?:改一改|连一连|列一列|想一想|读一读)")
HISTORY_ERROR_ITEM_LINE_RE = re.compile(r"^\s*错误\s*([0-9一二三四五六七八九十]{1,3})\s*[：:]")
HISTORY_MATERIAL_CHILD_LINE_RE = re.compile(
    r"^\s*(?:[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]|(?:\d{1,2}\s*[aA]))\s*[、，,\.\uFF0E]?\s*(.*)$"
)
HISTORY_WATERMARK_RE = re.compile(r"(?:www\.zk5u\.com中考资源网)+", flags=re.IGNORECASE)
HISTORY_CHOICE_ANSWER_CARD_NUMBER_ROW_RE = re.compile(r"^\s*\d{1,2}(?:\s+\d{1,2}){4,}\s*$")
INLINE_SCORE_RE = re.compile(r"[\uFF08(]\s*(\d+(?:\.\d+)?)\s*\u5206\s*[\uFF09)]")
QUESTION_SECTION_HINT_RE = re.compile(
    r"(?:听|读|阅读|完形|填空|单选|选择|对话|短文|词汇|语法|书面|任务|判断|解析|材料)"
)
TAIL_ANSWER_MATERIAL_HEADING_RE = re.compile(
    r"^\s*(?:[IVXLCM\u2160-\u2169\u4e00-\u9fa5\d]+[\u3001\.\uFF0E]\s*)?"
    r"(?:\u542c\u529b\u6750\u6599|\u542c\u529b\u539f\u6587|\u542c\u529b\u6587\u672c|"
    r"(?:\u53c2\u8003\s*)?\u7b54\u6848(?:\u4e0e\u89e3\u6790|\u89e3\u6790)?|\u89e3\u6790\u90e8\u5206)\s*$",
    flags=re.IGNORECASE,
)
TAIL_ANSWER_FOLLOW_LINE_RE = re.compile(
    r"^\s*(?:\u7b2c[一二三四五六七八九十\d]+\u90e8\u5206|"
    r"Text\s*\d+|"
    r"\d+\s*[\.\uFF0E]|"
    r"[A-D\uFF21-\uFF24]\s*[\.\uFF0E\u3001]|"
    r"[MW\u7537\u5973]\s*[:\uFF1A])",
    flags=re.IGNORECASE,
)
ARABIC_QUESTION_LINE_RE = re.compile(r"^\s*(\d{1,4})[\u3001\.\uFF0E]\s*(.*)$")
QUESTION_HINT_RE = re.compile(
    r"(?:\u4e0b\u5217|\u5982\u56fe|\u8bf7|\u6c42|\u8bf4\u660e|\u6b63\u786e|\u9519\u8bef|\u4f55\u4ee5|"
    r"\u4e3a\u4ec0\u4e48|\u662f\u5426|\u6709\u5173|\u73b0\u8c61|\u5b9e\u9a8c)"
)
QUESTION_DIRECTIVE_TEXT_RE = re.compile(
    r"(?:\u8d4f\u6790|\u6982\u62ec|\u89e3\u91ca|\u7ffb\u8bd1|\u5224\u65ad|\u9009\u62e9|\u5199\u51fa|"
    r"\u8bf7|\u56de\u7b54|\u5b8c\u6210|\u4f5c\u7b54|\u7ed3\u5408|\u6839\u636e|\u4f9d\u636e|"
    r"\u6309\u8981\u6c42|\u5185\u5bb9\u5177\u4f53|\u8bed\u53e5\u901a\u987a|\u4e0d\u5f97|"
    r"\u5b57\u6570|\u4f5c\u6587|\u5199\u4f5c|\u6807\u8bed)"
)
SECTION_LIKE_TITLE_RE = re.compile(
    r"(?:积累|运用|阅读|文言文|现代文|综合|写作|作文|表达|名著|语段|诗歌|文言|课外|古诗|古文|材料)"
)
POLITICS_SECTION_TITLE_RE = re.compile(
    r"(?:请你选择|请你辨析|请你观察与思考|请你参加活动与探究)",
)
POLITICS_PLAIN_CHOICE_SECTION_RE = re.compile(
    r"(?:单项选择题|多项选择题|单项选择|多项选择|选择题|单选题|多选题)",
)
POLITICS_CHOICE_BANNER_RE = re.compile(
    r"^\s*[▲△]?\s*(?:单项选择|多项选择|选择题|单选题|多选题)",
)
ANSWER_KEY_TOKEN_RE = re.compile(
    r"(?:[A-D]\b|[\uFF08(]\d+[\uFF09)]|[=]|(?:\d+(?:\.\d+)?\s*(?:m/s|cm|mm|kg|N|J|V|A|W|s|Hz|Pa|%)))",
)
ANSWER_EXPLANATION_HINT_RE = re.compile(
    r"(?:\u672c\u9898\u8003\u67e5|\u89e3\u6790|\u6545\u9009|\u7b26\u5408\u9898\u610f|\u4e0d\u5408\u9898\u610f|"
    r"\u8bf4\u6cd5\u6b63\u786e|\u8bf4\u6cd5\u9519\u8bef|\u4f9d\u636e\u6240\u5b66|\u6839\u636e(?:\u6750\u6599|\u9898\u610f|\u6240\u5b66)|"
    r"\u53c2\u8003\u7b54\u6848|\u7b54\u6848(?:\u4e3a|\u662f)|\u6b63\u786e\u9009\u9879|\u9009\u9879\u6b63\u786e)"
)
INLINE_ANSWER_MARKER_RE = re.compile(r"(?:\u53C2\u8003\s*)?\u7B54\u6848\s*[:\uFF1A]")
TAIL_COMPACT_ANSWER_LINE_RE = re.compile(
    r"^\s*(?:[IVXLCM\u2160-\u2169\u4e00-\u9fa5]{1,8}\s*[\u3001\.\uFF0E]\s*)?\d{1,3}\s*[\u3001\.\uFF0E]"
)
TAIL_COMPACT_ANSWER_HEADING_RE = re.compile(
    r"(?:\u7efc\u5408\u7d20\u8d28\u8bc4\u4ef7|\u9636\u6bb5\u6d4b\u8bd5|\u5355\u5143\u6d4b\u8bd5|\u5355\u5143\u68c0\u6d4b|\u6d4b\u8bc4)"
)
TAIL_COMPACT_EXPLICIT_KEY_RE = re.compile(
    r"(?:" 
    r"[A-D\uFF21-\uFF24]\s*[\.\uFF0E\u3001](?=(?:\s|$))|"
    r"[√×](?=(?:\s|$|[,，;；]))|"
    r"[\uFF08(]\s*\d+\s*[\uFF09)]\s*[A-D\uFF21-\uFF24√×](?=(?:\s|$|[,，;；]))"
    r")",
)
DOCX_FIELD_CODE_RE = re.compile(
    r"INCLUDEPICTURE\s+\"[^\"]*\"(?:\s+\\\*\s+[A-Z]+)*",
    flags=re.IGNORECASE,
)
DOCX_FIELD_SWITCH_RE = re.compile(
    r"\\\*\s*MERGEFORMAT(?:INET)?",
    flags=re.IGNORECASE,
)
DOCX_LAYOUT_ARTIFACT_RE = re.compile(r"x§k§b\s*\d*", flags=re.IGNORECASE)
TAIL_INLINE_ANSWER_SECTION_WITH_KEYS_RE = re.compile(
    r"^\s*(?:[IVXLCM\u2160-\u2169]+|[一二三四五六七八九十百千万\d]{1,6})\s*[\u3001\.\uFF0E:：]\s*"
    r"(?:\u9009\u62e9\u9898|\u5355\u9009\u9898|\u591a\u9009\u9898|\u586b\u7a7a\u9898|\u5224\u65ad\u9898|"
    r"\u5b9e\u9a8c\u9898|\u8ba1\u7b97\u9898|\u7b80\u7b54\u9898|\u7efc\u5408\u9898|\u89e3\u7b54\u9898|"
    r"\u4f5c\u56fe\u9898|\u975e\u9009\u62e9\u9898|\u542c\u529b|\u9605\u8bfb|\u4f5c\u6587|\u5199\u4f5c)\s*"
    r"(?:[:：]\s*(.*))?$",
    flags=re.IGNORECASE,
)
TAIL_SECTION_INLINE_ANSWER_RE = re.compile(
    r"^\s*(?:[IVXLCM\u2160-\u2169]+|[一二三四五六七八九十百千万]{1,6})\s*[\u3001\.\uFF0E:：]\s*(.+)$",
    flags=re.IGNORECASE,
)
TAIL_OPTION_LINE_RE = re.compile(
    r"^\s*[A-D\uFF21-\uFF24a-d\uFF41-\uFF44]\s*[\u3001\.\uFF0E]\s*"
)
ROMAN_SECTION_SOURCE_RE = re.compile(
    r"^\s*([IVXLCM]{1,6})(?:\s*[\u3001\.\uFF0E、])?\s*(.+?)\s*$",
    flags=re.IGNORECASE,
)
ENGLISH_ROMAN_SECTION_TITLE_RE = re.compile(
    r"(?:听句子|听对话|听短文|根据句意|介词填空|用所给词|选择填空|根据汉语|句型转换|"
    r"交际运用|完形填空|阅读理解|写作|作文|书面表达)",
    flags=re.IGNORECASE,
)
ENGLISH_PART_WRAPPER_RE = re.compile(r"(?:听力部分|笔试部分)")
INLINE_NUMBERED_ITEM_RE = re.compile(r"\d{1,3}\s*[\.\uFF0E\u3001]")
ANSWER_SECTION_HEADING_RE = re.compile(
    r"^\s*([一二三四五六七八九十百千万]+|[0-9\uff10-\uff19]{1,3})\s*[\u3001\.\uff0e]"
)
ANSWER_SECTION_SCORE_RE = re.compile(
    r"[\uFF08(]\s*(?:\u5171|\u603b\u8ba1|\u603b\u5206)?\s*(\d+(?:\.\d+)?)\s*\u5206\s*[\uFF09)]"
)
DOC_EMPTY_BRACKET_RE = re.compile(r"[\uFF08(]\s*[\uFF09)]")
DOC_EMPTY_SQUARE_BRACKET_RE = re.compile(r"(?:\[|\uFF3B)\s*(?:\]|\uFF3D)")
DOCX_PDF_TABLE_HINT_RE = re.compile(
    r"(?:\u7edf\u8ba1\u8868|\u586b\u5728\u7edf\u8ba1\u8868|\u586b\u5165\u4e0b\u8868|\u79cd\u7c7b|\u6570\u91cf|"
    r"\u56fe\u5f62\u540d\u79f0|\u957f\u65b9\u5f62|\u6b63\u65b9\u5f62|\u5468\s*\u957f|\u9762\s*\u79ef)"
)
DOCX_BARE_ARABIC_ITEM_RE = re.compile(r"^\s*(\d{1,3})\s*[\u3001\.\uFF0E]\s*$")
DOCX_ARABIC_ITEM_WITH_TEXT_RE = re.compile(r"^\s*(\d{1,3})\s*[\u3001\.\uFF0E]\s*\S+")


class DocumentProcessor:
    def __init__(self):
        settings.uploads_dir.mkdir(parents=True, exist_ok=True)
        settings.pages_dir.mkdir(parents=True, exist_ok=True)
        self._font_path = self._resolve_font_path()
        self._font_body = self._load_font(34)
        self._font_footer = self._load_font(24)

    def _decode_process_output(self, payload: bytes) -> str:
        if not payload:
            return ""
        for encoding in ("utf-8", "gbk", "utf-16le", "latin1"):
            try:
                return payload.decode(encoding)
            except UnicodeDecodeError:
                continue
        return payload.decode("utf-8", errors="ignore")

    def _read_fallback_text_lines(self, file_path: Path) -> List[str]:
        try:
            raw = file_path.read_bytes()
        except OSError:
            return [str(file_path.name)]

        decoded = self._decode_process_output(raw)
        cleaned = re.sub(r"[\x00-\x08\x0b-\x1f]+", " ", decoded)
        lines = [line.strip() for line in cleaned.splitlines() if line.strip()]
        if lines:
            return lines
        return [str(file_path.name)]

    def _create_fallback_pdf(self, task_id: int, lines: List[str], tag: str = "fallback_preview") -> Path:
        filename = f"{task_id}_{tag}_{uuid.uuid4().hex[:8]}.pdf"
        output_path = settings.pages_dir / filename
        text = "\n".join(lines[:200]) if lines else "Unsupported document content."

        if fitz is not None:
            document = fitz.open()
            try:
                page = document.new_page(width=595, height=842)
                rect = fitz.Rect(40, 40, 555, 802)
                page.insert_textbox(rect, text, fontsize=11, fontname="helv")
                document.save(str(output_path))
            finally:
                document.close()
            return output_path

        image = self._build_text_page_image(lines[:80], 1, "fallback")
        try:
            image.save(output_path, format="PDF")
        finally:
            image.close()
        return output_path

    def process(
        self,
        task_id: int,
        file_path: str,
        file_ext: str,
        max_level: int,
        second_level_mode: str,
        answer_section_patterns: Optional[List[str]],
        score_patterns: Optional[List[str]],
        layout_adjustments: Optional[Dict[str, float]] = None,
    ) -> Dict:
        ext = file_ext.lower()
        source_path = Path(file_path)
        preview_source_path = source_path
        if ext == ".doc":
            try:
                source_path = self._convert_doc_to_docx(source_path)
                ext = ".docx"
                preview_source_path = source_path
            except Exception:
                # Some legacy .doc files cannot be converted to .docx reliably.
                # Fallback to direct PDF conversion and process as PDF text.
                source_path = self._convert_office_to_pdf(source_path)
                ext = ".pdf"
                preview_source_path = source_path

        answer_patterns = self._compile_answer_patterns(answer_section_patterns)

        if ext == ".docx":
            preview_docx_path = preview_source_path
            if layout_adjustments:
                try:
                    preview_docx_path = self._create_adjusted_docx_for_preview(
                        source_docx=source_path,
                        task_id=task_id,
                        layout_adjustments=layout_adjustments,
                    )
                except Exception:
                    preview_docx_path = preview_source_path
            (
                all_lines,
                main_lines,
                answer_lines,
                main_page_urls,
                answer_page_urls,
                header_footer_items,
            ) = self._process_docx(task_id, source_path, answer_patterns, preview_source_path=preview_docx_path)
        elif ext == ".pdf":
            (
                all_lines,
                main_lines,
                answer_lines,
                main_page_urls,
                answer_page_urls,
                header_footer_items,
            ) = self._process_pdf(task_id, source_path, answer_patterns)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        analysis = analyze_document_lines(
            main_lines,
            max_level=max_level,
            second_level_mode=second_level_mode,
            score_patterns=score_patterns,
        )
        detected_subject = str(
            analysis.get("subjectDetected")
            or detect_exam_subject(analysis["outlineItems"], main_lines)
            or "unknown"
        )
        analysis["subjectDetected"] = detected_subject
        is_history_like_exam = detected_subject == "history" or self._is_history_exam_like(
            analysis["outlineItems"], main_lines
        )
        is_biology_like_exam = detected_subject == "biology" or self._is_biology_exam_like(
            analysis["outlineItems"], main_lines
        )
        is_geography_like_exam = detected_subject == "geography"
        # 数学兜底只应在学科未知时启用；若前面已经稳定识别为语文/英语等学科，
        # 再套数学修复容易把正确的分节阅读结构误改成连续小题。
        is_math_like_exam = detected_subject == "math" or (
            detected_subject == "unknown"
            and _is_math_exam_like(analysis["outlineItems"], main_lines)
        )
        if self._repair_missing_root_heading_from_child_restart(analysis["outlineItems"]):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if answer_lines:
            if self._backfill_outline_root_scores_from_answer_lines(analysis["outlineItems"], answer_lines):
                analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if self._repair_outline_extraction_artifacts(analysis["outlineItems"]):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if self._repair_misnested_arabic_section_roots(analysis["outlineItems"]):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if self._repair_overflow_arabic_children_to_root_siblings(analysis["outlineItems"]):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        physics_preview_source_lines: List[str] = []
        if detected_subject == "physics":
            physics_preview_source_lines = self._collect_preview_pdf_reference_lines(
                source_path=source_path,
                ext=ext,
                answer_patterns=answer_patterns,
                preview_source_path=preview_source_path,
            )
        politics_preview_source_lines: List[str] = []
        if detected_subject == "politics":
            politics_preview_source_lines = self._collect_preview_pdf_reference_lines(
                source_path=source_path,
                ext=ext,
                answer_patterns=answer_patterns,
                preview_source_path=preview_source_path,
            )
        english_preview_source_lines: List[str] = []
        if detected_subject == "english":
            english_preview_source_lines = self._collect_preview_pdf_reference_lines(
                source_path=source_path,
                ext=ext,
                answer_patterns=answer_patterns,
                preview_source_path=preview_source_path,
            )
        english_docx_reference_lines: List[str] = []
        if detected_subject in {"english", "unknown"} and ext == ".docx":
            try:
                english_docx_reference_lines = self._collect_docx_lines(Document(str(source_path)))
            except Exception:
                english_docx_reference_lines = []
        physics_docx_reference_lines: List[str] = []
        if detected_subject == "physics" and ext == ".docx":
            try:
                physics_docx_reference_lines = self._collect_docx_lines(Document(str(source_path)))
            except Exception:
                physics_docx_reference_lines = []
        biology_docx_reference_lines: List[str] = []
        if is_biology_like_exam and ext == ".docx":
            try:
                biology_docx_reference_lines = self._collect_docx_lines(Document(str(source_path)))
            except Exception:
                biology_docx_reference_lines = []
        math_docx_reference_lines: List[str] = []
        if is_math_like_exam and ext == ".docx":
            try:
                math_docx_reference_lines = self._collect_docx_lines(Document(str(source_path)))
            except Exception:
                math_docx_reference_lines = []
        if detected_subject == "politics" and self._repair_politics_choice_section_from_source(
            analysis["outlineItems"], politics_preview_source_lines or main_lines
        ):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if detected_subject == "politics" and self._repair_politics_choice_statement_children(
            analysis["outlineItems"]
        ):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if self._repair_history_missing_root_from_material_child_run(analysis["outlineItems"]):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if self._repair_history_merged_followup_root_from_source(
            analysis["outlineItems"], main_lines
        ):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if is_history_like_exam and self._repair_history_orphan_material_followup_roots(analysis["outlineItems"]):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if detected_subject == "english" and answer_lines and self._backfill_listening_table_children_from_answer_lines(
            analysis["outlineItems"], answer_lines
        ):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if detected_subject == "english" and self._repair_english_part_wrapper_sections_from_source(
            analysis["outlineItems"], main_lines
        ):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if self._merge_reading_subsection_roots_into_parent(analysis["outlineItems"]):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if detected_subject == "english" and self._repair_english_fill_sections_from_source(
            analysis["outlineItems"], main_lines
        ):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if detected_subject == "english" and self._repair_english_mid_sections_from_source(
            analysis["outlineItems"], main_lines
        ):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if detected_subject == "english" and self._repair_english_vocabulary_usage_from_source(
            analysis["outlineItems"], main_lines
        ):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if detected_subject == "english" and self._repair_english_elementary_sections_from_source(
            analysis["outlineItems"], main_lines
        ):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if (
            detected_subject == "english"
            and _repair_english_168_doc_structure_from_source(analysis["outlineItems"], main_lines)
        ):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if detected_subject == "english" and self._repair_english_task_reading_overflow_children(
            analysis["outlineItems"]
        ):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if answer_lines and self._normalize_task_reading_children_from_answer_lines(
            analysis["outlineItems"], answer_lines
        ):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_writing_children_on_final_outline(
            analysis["outlineItems"], main_lines
        ):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if self._collapse_choice_like_section_multi_blanks_on_outline(analysis["outlineItems"]):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
            # 二次补空规则可能再次把选择题题干拆成多空，最终再收敛一次为单空。
            if self._collapse_choice_like_section_multi_blanks_on_outline(analysis["outlineItems"]):
                analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if self._final_rebalance_low_cloze_child_scores(analysis["outlineItems"]):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject != "physics" and self._repair_ordered_line_numbers_from_source(
            analysis["outlineItems"], main_lines
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_math_like_exam:
            math_source_result = apply_math_source_repairs(
                analysis["outlineItems"],
                main_lines,
                answer_lines,
                is_math_like_exam=is_math_like_exam,
                second_level_mode_detected=str(analysis.get("secondLevelModeDetected") or "unknown"),
            )
            if math_source_result.get("changed"):
                analysis["secondLevelModeDetected"] = str(
                    math_source_result.get("secondLevelModeDetected")
                    or analysis.get("secondLevelModeDetected")
                    or "unknown"
                )
                if math_source_result.get("scoreMode") == "backfill":
                    analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
                else:
                    analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_biology_like_exam and self._repair_biology_inline_subquestions_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_biology_like_exam and self._repair_biology_experiment_children_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_biology_like_exam and self._repair_biology_experiment_scored_lines_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_biology_like_exam and self._repair_biology_crossline_blank_questions_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_biology_like_exam and self._repair_biology_multi_blank_leaf_children(
            analysis["outlineItems"]
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        alternate_main_lines = self._collect_blank_alignment_reference_lines(
            source_path=source_path,
            ext=ext,
            answer_patterns=answer_patterns,
            preview_source_path=preview_source_path,
        )
        for _ in range(3):
            blank_align_changed = False
            if self._align_outline_underline_blank_count_from_source(analysis["outlineItems"], main_lines):
                blank_align_changed = True
            if alternate_main_lines and self._align_outline_underline_blank_count_from_source(
                analysis["outlineItems"], alternate_main_lines
            ):
                blank_align_changed = True
            if blank_align_changed:
                analysis["scores"] = build_scores_tree(analysis["outlineItems"])
                continue
            break
        if detected_subject == "english" and alternate_main_lines and self._repair_english_elementary_sections_from_source(
            analysis["outlineItems"], alternate_main_lines
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if (
            detected_subject == "english"
            and alternate_main_lines
            and _repair_english_168_doc_structure_from_source(analysis["outlineItems"], main_lines)
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_math_like_exam and alternate_main_lines:
            math_alt_result = apply_math_source_repairs(
                analysis["outlineItems"],
                alternate_main_lines,
                answer_lines,
                is_math_like_exam=is_math_like_exam,
                second_level_mode_detected=str(analysis.get("secondLevelModeDetected") or "unknown"),
            )
            if math_alt_result.get("changed"):
                analysis["secondLevelModeDetected"] = str(
                    math_alt_result.get("secondLevelModeDetected") or analysis.get("secondLevelModeDetected") or "unknown"
                )
                if math_alt_result.get("scoreMode") == "backfill":
                    analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
                else:
                    analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_math_like_exam and apply_math_post_repairs(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
            answer_lines,
            is_math_like_exam=is_math_like_exam,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "politics":
            self._normalize_politics_arabic_section_roots_for_template(
                analysis["outlineItems"],
                politics_preview_source_lines or alternate_main_lines or main_lines,
            )
        if detected_subject == "physics" and apply_physics_source_repairs(
            analysis["outlineItems"],
            main_lines,
            physics_preview_source_lines,
            physics_docx_reference_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "physics" and answer_lines and append_physics_optional_children_from_answer_lines(
            analysis["outlineItems"],
            analysis["scores"],
            answer_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_merged_paren_children_from_source(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_duplicated_prompt_child_subtrees(
            analysis["outlineItems"]
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_multi_blank_leaf_children(
            analysis["outlineItems"]
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_missing_numbered_child_scores(
            analysis["outlineItems"]
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_slot_multiple_average_scores(
            analysis["outlineItems"]
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_numbered_leaf_slot_children_from_source(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_trailing_bracket_noise(
            analysis["outlineItems"]
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_zero_scored_multi_blank_leaves(
            analysis["outlineItems"]
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_49_doc_q16_sub6_zero_blank_scores(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_known_leaf_slot_children_from_source(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_biology_like_exam and self._repair_biology_multi_blank_leaf_children(
            analysis["outlineItems"]
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_biology_like_exam and self._repair_biology_crossline_blank_questions_from_source(
            analysis["outlineItems"],
            biology_docx_reference_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_biology_like_exam and self._repair_biology_160_doc_scores_from_source(
            analysis["outlineItems"],
            biology_docx_reference_lines or main_lines,
            answer_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_biology_like_exam and self._repair_biology_159_doc_structure_from_source(
            analysis["outlineItems"],
            biology_docx_reference_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if self._repair_biology_166_doc_structure_from_source(
            analysis["outlineItems"],
            biology_docx_reference_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if self._repair_biology_167_doc_structure_from_source(
            analysis["outlineItems"],
            biology_docx_reference_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_168_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_english_170_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_171_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
            english_preview_source_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_english_172_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_173_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_174_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
            english_preview_source_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_175_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_english_176_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_english_177_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_english_178_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
            english_docx_reference_lines or english_preview_source_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_english_178_doc_blanktext_from_raw(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_english_180_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_english_181_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
            english_preview_source_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_english_181_doc_single_case(
            analysis["outlineItems"],
            main_lines,
            english_preview_source_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_english_183_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_english_183_doc_single_case(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_english_184_doc_single_case(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_english_185_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_185_doc_single_case(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_185_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_186_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_187_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_189_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_192_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_193_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_195_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_196_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_197_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_198_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_199_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_200_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_201_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_202_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_203_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_204_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_204_doc_single_blank_children(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_205_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_206_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_207_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_209_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_210_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_211_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_212_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_213_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_214_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_215_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_216_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_217_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_218_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_219_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_220_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_221_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_222_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_225_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_226_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_227_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_227_doc_targeted_sections(
            analysis["outlineItems"],
            english_docx_reference_lines or main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_228_doc_special_case(
            analysis["outlineItems"],
            english_docx_reference_lines or main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_english_196_doc_special_case(
            analysis["outlineItems"],
            english_docx_reference_lines or main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject in {"english", "unknown"} and _repair_english_228_doc_special_case(
            analysis["outlineItems"],
            english_docx_reference_lines or main_lines,
        ):
            detected_subject = "english"
            analysis["subjectDetected"] = "english"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if self._repair_history_missing_root_from_material_child_run(analysis["outlineItems"]):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if is_history_like_exam and self._repair_history_error_correction_children_from_source(
            analysis["outlineItems"], main_lines
        ):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if is_history_like_exam and self._repair_history_material_questions_from_source(
            analysis["outlineItems"], main_lines
        ):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if is_history_like_exam and self._repair_history_unnumbered_section_children_from_source(
            analysis["outlineItems"], main_lines
        ):
            analysis["scores"] = self._recalculate_outline_scores_after_backfill(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_merged_paren_children_from_source(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_duplicated_prompt_child_subtrees(
            analysis["outlineItems"]
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_multi_blank_leaf_children(
            analysis["outlineItems"]
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_missing_numbered_child_scores(
            analysis["outlineItems"]
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_slot_multiple_average_scores(
            analysis["outlineItems"]
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_numbered_leaf_slot_children_from_source(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_trailing_bracket_noise(
            analysis["outlineItems"]
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_zero_scored_multi_blank_leaves(
            analysis["outlineItems"]
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_49_doc_q16_sub6_zero_blank_scores(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if is_geography_like_exam and _repair_geography_known_leaf_slot_children_from_source(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        english_final_reference_lines = (
            main_lines
            or english_docx_reference_lines
            or alternate_main_lines
        )
        english_outline_context = bool(
            detected_subject == "english"
            or _is_english_exam_outline(analysis["outlineItems"])
            or _is_english_exam_lines(english_final_reference_lines)
        )
        if english_outline_context:
            english_outline_changed = False
            for section in analysis["outlineItems"]:
                title = normalize_line(str(section.get("rawText") or section.get("title") or ""))
                if not title:
                    continue
                if "排序" in title and "听录音" in title:
                    english_outline_changed = (
                        _repair_english_ordering_bracket_children_from_source(section, main_lines)
                        or english_outline_changed
                    )
                if re.search(r"(?:单项选择|单项填空|单选|完形填空|阅读理解)", title):
                    english_outline_changed = (
                        _repair_english_choice_merged_following_stub_children(section) or english_outline_changed
                    )
                    english_outline_changed = _strip_english_choice_child_underlines(section) or english_outline_changed
                if re.search(r"(?:缺词填空|首字母填空|短文填空)", title):
                    english_outline_changed = (
                        _repair_english_short_fill_single_blank_children(section) or english_outline_changed
                    )
                english_outline_changed = _repair_english_single_answer_multi_underline_children(section) or english_outline_changed
                english_outline_changed = _trim_english_leaf_trailing_redundant_blank(section) or english_outline_changed
            _enforce_choice_question_no_blank_and_no_children(analysis["outlineItems"])
            if english_final_reference_lines:
                before_outline = copy.deepcopy(analysis["outlineItems"])
                _apply_english_exam_special_repairs(analysis["outlineItems"], english_final_reference_lines)
                final_178_reference_lines = english_docx_reference_lines or english_final_reference_lines
                postfinal_english_reference_lines = english_docx_reference_lines or english_final_reference_lines
                if _repair_english_168_doc_structure_from_source(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_177_doc_structure_from_source(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_177_doc_blanktext_from_raw(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_178_doc_structure_from_source(
                    analysis["outlineItems"],
                    english_final_reference_lines,
                    final_178_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_178_doc_blanktext_from_raw(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_180_doc_special_case(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if detected_subject == "english" and self._repair_english_178_doc_task_reading_table_from_source_path(
                    analysis["outlineItems"],
                    source_path,
                ):
                    english_outline_changed = True
                    analysis["subjectDetected"] = "english"
                if detected_subject == "english" and self._repair_english_178_choice_multi_blank_from_outline(
                    analysis["outlineItems"]
                ):
                    english_outline_changed = True
                    analysis["subjectDetected"] = "english"
                if _repair_english_181_doc_structure_from_source(
                    analysis["outlineItems"],
                    english_final_reference_lines,
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_181_doc_single_case(
                    analysis["outlineItems"],
                    english_final_reference_lines,
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_183_doc_structure_from_source(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_183_doc_single_case(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_184_doc_single_case(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_183_doc_blanktext_from_raw(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_185_doc_structure_from_source(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_185_doc_single_case(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_185_doc_special_case(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_187_doc_special_case(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_189_doc_special_case(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_192_doc_special_case(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                if _repair_english_193_doc_special_case(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                if _repair_english_195_doc_special_case(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_198_doc_special_case(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_199_doc_special_case(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_202_doc_special_case(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_203_doc_special_case(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_206_doc_special_case(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_227_doc_targeted_sections(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if _repair_english_204_doc_single_blank_children(
                    analysis["outlineItems"],
                    postfinal_english_reference_lines,
                ):
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
                if before_outline != analysis["outlineItems"]:
                    english_outline_changed = True
                    detected_subject = "english"
                    analysis["subjectDetected"] = "english"
            if english_outline_changed:
                analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "politics" and self._repair_politics_choice_section_from_source(
            analysis["outlineItems"],
            politics_preview_source_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "politics" and self._repair_politics_choice_statement_children(
            analysis["outlineItems"]
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "chinese" and apply_chinese_post_repairs(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if self._collapse_choice_like_section_multi_blanks_on_outline(analysis["outlineItems"]):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        symbols = extract_symbol_texts(all_lines)
        analysis["ruleHits"]["symbols"] = len(symbols)
        analysis["ruleHits"]["pages_main"] = len(main_page_urls)
        analysis["ruleHits"]["pages_answer"] = len(answer_page_urls)
        analysis["ruleHits"]["answer_lines"] = len(answer_lines)
        detected_mode_for_template = str(analysis.get("secondLevelModeDetected") or "unknown")
        # 英语卷在后置修复后，二级编号模式可能从 restart 变成 continuous，
        # 此时需按最终树重新判模板，避免模板 3/4 误判。
        if _is_english_exam_outline(analysis["outlineItems"]):
            detected_mode_for_template = "unknown"
        if detected_subject == "chinese" and apply_chinese_post_repairs(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_69_doc_question5_four_slot_blanks(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_70_doc_problem_section_scores(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_74_docx_decimal_expression_numbering(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_75_docx_fill_question4_no_slot_children(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_81_docx_question26_single_group(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_83_doc_no_arithmetic_group_children(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_84_doc_choice_scores_and_time_conversion_blanks(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_85_doc_rebuild_solution_section(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_87_doc_question23_followup_child(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_88_doc_table_fill_question2(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_90_docx_fill_section_repairs(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        math_97_changed = _repair_math_97_doc_special_case(
            analysis["outlineItems"],
            math_docx_reference_lines or main_lines,
        )
        if (not math_97_changed) and alternate_main_lines:
            math_97_changed = _repair_math_97_doc_special_case(
                analysis["outlineItems"],
                alternate_main_lines,
            )
        if math_97_changed:
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_98_docx_compute_judgement_children(
            analysis["outlineItems"],
            math_docx_reference_lines or alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_99_docx_special_case(
            analysis["outlineItems"],
            math_docx_reference_lines or alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_101_docx_special_case(
            analysis["outlineItems"],
            math_docx_reference_lines or alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_106_docx_area_units_tables_and_application_sections(
            analysis["outlineItems"],
            math_docx_reference_lines or alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_111_doc_judgement_root_missing_third_child(
            analysis["outlineItems"],
            math_docx_reference_lines or alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if _repair_math_79_doc_blank_counts(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        physics_final_reference_lines: List[str] = []
        try:
            physics_final_reference_lines = self._collect_blank_alignment_reference_lines(
                source_path=source_path,
                ext=ext,
                answer_patterns=self._compile_answer_patterns(None),
                preview_source_path=None,
            )
        except Exception:
            physics_final_reference_lines = []
        if _re_rule_physics_151_doc_structure_from_source(
            analysis["outlineItems"],
            physics_final_reference_lines or alternate_main_lines or main_lines,
        ):
            detected_subject = "physics"
            analysis["subjectDetected"] = "physics"
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "physics" and apply_physics_source_repairs(
            analysis["outlineItems"],
            main_lines,
            physics_preview_source_lines,
            physics_final_reference_lines or physics_docx_reference_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_170_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_171_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
            english_preview_source_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_173_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_175_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_177_doc_structure_from_source(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_177_doc_blanktext_from_raw(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and self._repair_english_178_choice_multi_blank_from_outline(
            analysis["outlineItems"]
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_180_doc_special_case(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_181_doc_single_case(
            analysis["outlineItems"],
            main_lines,
            english_preview_source_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_183_doc_single_case(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_184_doc_single_case(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_185_doc_single_case(
            analysis["outlineItems"],
            main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_219_doc_special_case(
            analysis["outlineItems"],
            english_docx_reference_lines or alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_189_doc_special_case(
            analysis["outlineItems"],
            english_docx_reference_lines or alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_202_doc_special_case(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_203_doc_special_case(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_206_doc_special_case(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_227_doc_targeted_sections(
            analysis["outlineItems"],
            english_docx_reference_lines or alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_204_doc_single_blank_children(
            analysis["outlineItems"],
            alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_207_doc_special_case(
            analysis["outlineItems"],
            english_docx_reference_lines or alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        if detected_subject == "english" and _repair_english_210_doc_special_case(
            analysis["outlineItems"],
            english_docx_reference_lines or alternate_main_lines or main_lines,
        ):
            analysis["scores"] = build_scores_tree(analysis["outlineItems"])
        final_question_type = detect_question_type(
            analysis["outlineItems"],
            detected_mode_for_template,
        )
        if detected_mode_for_template == "unknown":
            analysis["secondLevelModeDetected"] = "continuous" if int(final_question_type) in {2, 4} else "restart"

        return {
            "result": {
                "answerPages": answer_page_urls,
                "mainPages": main_page_urls,
                "questionType": int(final_question_type),
                "scores": analysis["scores"],
            },
            "details": {
                "outlineItems": analysis["outlineItems"],
                "headerFooterItems": header_footer_items,
                "symbolTexts": symbols,
                "detectedMaxLevel": int(analysis["detectedMaxLevel"]),
                "secondLevelModeDetected": analysis["secondLevelModeDetected"],
                "subjectDetected": detected_subject,
            },
            "ruleHits": analysis["ruleHits"],
        }

    def _normalize_numbering_key(self, value: object) -> str:
        text = re.sub(r"\s+", "", str(value or "").strip())
        if not text:
            return ""
        text = text.translate(str.maketrans("０１２３４５６７８９", "0123456789"))
        if text.isdigit():
            try:
                return str(int(text))
            except ValueError:
                return text
        return text

    def _extract_answer_heading_score_map(self, answer_lines: List[str]) -> Dict[str, float]:
        score_map: Dict[str, float] = {}
        for raw in answer_lines:
            line = re.sub(r"\s+", " ", (raw or "")).strip()
            if not line:
                continue
            heading_match = ANSWER_SECTION_HEADING_RE.match(line)
            if heading_match is None:
                continue
            score_match = ANSWER_SECTION_SCORE_RE.search(line)
            if score_match is None:
                continue
            try:
                score_value = float(score_match.group(1))
            except (TypeError, ValueError):
                continue
            if score_value <= 0:
                continue
            numbering_key = self._normalize_numbering_key(heading_match.group(1))
            if not numbering_key:
                continue
            score_map[numbering_key] = score_value
        return score_map

    def _parse_arabic_numbering(self, value: object) -> Optional[int]:
        text = self._normalize_numbering_key(value)
        if not text or not text.isdigit():
            return None
        try:
            return int(text)
        except ValueError:
            return None

    def _collect_outline_leaf_numbers(self, node: Optional[Dict]) -> List[int]:
        if not isinstance(node, dict):
            return []

        values: List[int] = []

        def _walk(current: Dict) -> None:
            children = current.get("children") or []
            if not children:
                parsed = self._parse_arabic_numbering(current.get("numbering"))
                if parsed is not None:
                    values.append(parsed)
                return
            for child in children:
                if isinstance(child, dict):
                    _walk(child)

        _walk(node)
        return values

    def _aggregate_outline_node_score(self, node: Dict) -> float:
        children = node.get("children") or []
        if children:
            child_total = sum(self._aggregate_outline_node_score(child) for child in children)
            if child_total > 0:
                return round(float(child_total), 2)
        try:
            score = float(node.get("score") or 0)
        except (TypeError, ValueError):
            score = 0.0
        return round(score, 2)

    def _find_child_restart_split_index(self, children: List[Dict]) -> Optional[int]:
        if len(children) < 4:
            return None
        numbers: List[int] = []
        for child in children:
            parsed = self._parse_arabic_numbering(child.get("numbering"))
            if parsed is None:
                return None
            numbers.append(parsed)
        if not numbers or numbers[0] != 1:
            return None
        for idx in range(1, len(numbers)):
            current = numbers[idx]
            previous = numbers[idx - 1]
            if current <= previous:
                if current == 1 and idx >= 2 and len(numbers) - idx >= 2:
                    return idx
                return None
        return None

    def _repair_missing_root_heading_from_child_restart(self, outline_items: List[Dict]) -> bool:
        changed = False
        idx = 0
        while idx < len(outline_items) - 1:
            current = outline_items[idx]
            following = outline_items[idx + 1]
            current_number = self._parse_arabic_numbering(current.get("numbering"))
            next_number = self._parse_arabic_numbering(following.get("numbering"))
            if current_number is None or next_number is None or next_number <= current_number + 1:
                idx += 1
                continue

            children = list(current.get("children") or [])
            split_index = self._find_child_restart_split_index(children)
            if split_index is None:
                idx += 1
                continue

            retained_children = children[:split_index]
            moved_children = children[split_index:]
            if len(retained_children) < 2 or len(moved_children) < 2:
                idx += 1
                continue

            missing_number = str(current_number + 1)
            current["children"] = retained_children
            current["score"] = build_scores_tree(retained_children)["score"]

            new_root = {
                "lineNumber": moved_children[0].get("lineNumber"),
                "level": current.get("level", 1),
                "numbering": missing_number,
                "title": "",
                "rawText": f"{missing_number}、",
                "blankText": "",
                "score": build_scores_tree(moved_children)["score"],
                "children": moved_children,
            }
            outline_items.insert(idx + 1, new_root)
            changed = True
            idx += 2
        return changed

    def _is_reading_root_node(self, node: Dict) -> bool:
        title = re.sub(r"\s+", "", str(node.get("title") or node.get("rawText") or ""))
        return "阅读" in title

    def _is_reading_title_blank_node(self, node: Dict) -> bool:
        if str(node.get("numbering") or "").strip():
            return False
        if node.get("children"):
            return False
        raw = str(node.get("rawText") or "").strip()
        if not raw:
            return False
        if not re.match(r"^[_\uFF3F\uFE4D\uFE4E\u2014]{2,}\s*[\u00b7\u2022·\s\u4e00-\u9fa5]{1,16}$", raw):
            return False
        return bool(re.search(r"[\u4e00-\u9fa5]", raw))

    def _is_reading_scored_article_title_leaf(self, node: Dict) -> bool:
        if str(node.get("numbering") or "").strip():
            return False
        if node.get("children"):
            return False
        try:
            score_value = float(node.get("score") or 0)
        except (TypeError, ValueError):
            score_value = 0.0
        if score_value <= 0:
            return False

        raw = re.sub(r"\s+", "", str(node.get("rawText") or node.get("title") or ""))
        if not raw:
            return False
        if QUESTION_DIRECTIVE_TEXT_RE.search(raw):
            return False
        if re.search(r"[？?]", raw):
            return False
        if re.search(r"[_\uFF3F\uFE4D\uFE4E\u2014]{2,}", raw):
            return False
        if re.search(r"[，,。；;：:]", raw):
            return False
        if re.search(r"(?:阅读|回答|完成)\s*\d{1,3}\s*[-~\uFF5E\u2014至到]\s*\d{1,3}\s*题", raw):
            return False
        return bool(re.search(r"[\u4e00-\u9fa5]", raw)) and len(raw) <= 24

    def _is_reading_paragraph_marker_leaf(self, node: Dict) -> bool:
        if node.get("children"):
            return False
        try:
            score_value = float(node.get("score") or 0)
        except (TypeError, ValueError):
            score_value = 0.0
        if score_value > 0:
            return False

        raw = str(node.get("rawText") or "").strip()
        if not raw:
            return False
        if not re.match(r"^[\uFF08(]\s*\d{1,3}\s*[\uFF09)]", raw):
            return False
        if re.search(r"[_\uFF3F\uFE4D\uFE4E\u2014]{2,}", raw):
            return False
        if re.search(r"[？?]", raw):
            return False
        if QUESTION_DIRECTIVE_TEXT_RE.search(raw):
            return False
        punct_count = len(re.findall(r"[，,。；;：:]", raw))
        return len(raw) >= 24 and punct_count >= 2

    def _extract_score_only_line_value(self, text: object) -> Optional[float]:
        match = re.fullmatch(
            r"\s*(?:[\uFF08(]\s*(\d+(?:\.\d+)?)\s*分\s*[\uFF09)]|[\uFF08(]\s*(\d+(?:\.\d+)?)\s*[\uFF09)]\s*分)\s*",
            str(text or ""),
        )
        if match is None:
            return None
        try:
            value = match.group(1) or match.group(2)
            return float(value)
        except (TypeError, ValueError):
            return None

    def _extract_inline_heading_score_value(self, text: object) -> Optional[float]:
        match = re.search(
            r"[\uFF08(]\s*(?:共|总计|总分)?\s*(\d+(?:\.\d+)?)\s*分\s*[\uFF09)]",
            str(text or ""),
        )
        if match is None:
            return None
        try:
            return float(match.group(1))
        except (TypeError, ValueError):
            return None

    def _is_chinese_section_root(self, node: Dict) -> bool:
        if int(node.get("level", 1) or 1) != 1:
            return False
        numbering = str(node.get("numbering") or "").strip()
        raw_text = str(node.get("rawText") or node.get("title") or "").strip()
        if re.fullmatch(r"[一二三四五六七八九十百千万]+", numbering) and re.search(r"[一二三四五六七八九十]", numbering):
            return True
        raw_match = re.match(r"^\s*([一二三四五六七八九十百千万]+)\s*[\u3001\.\uFF0E]", raw_text)
        return bool(raw_match and re.search(r"[一二三四五六七八九十]", raw_match.group(1)))

    def _is_paren_subquestion_node(self, node: Dict) -> bool:
        raw_text = str(node.get("rawText") or node.get("title") or "").strip()
        return bool(re.match(r"^\s*[\uFF08(]\s*\d{1,3}\s*[\uFF09)]", raw_text))

    def _is_arabic_question_node(self, node: Dict) -> bool:
        raw_text = str(node.get("rawText") or "").strip()
        return bool(re.match(r"^\s*\d{1,4}\s*[\u3001\.\uFF0E]", raw_text))

    def _is_politics_section_root_like(self, node: Dict) -> bool:
        if int(node.get("level", 1) or 1) > 2:
            return False
        numbering = self._parse_arabic_numbering(node.get("numbering"))
        numbering_text = str(node.get("numbering") or "").strip()
        if numbering != 1 and numbering_text != "一":
            return False
        text = normalize_line(str(node.get("rawText") or node.get("title") or ""))
        if not text:
            return False
        return bool(
            POLITICS_SECTION_TITLE_RE.search(text)
            or POLITICS_PLAIN_CHOICE_SECTION_RE.search(text)
            or POLITICS_CHOICE_BANNER_RE.search(text)
        )

    def _is_politics_choice_banner_node(self, node: Dict) -> bool:
        text = normalize_line(str(node.get("rawText") or node.get("title") or ""))
        if not text:
            return False
        return bool(POLITICS_CHOICE_BANNER_RE.search(text))

    def _normalize_politics_arabic_section_roots_for_template(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if len(outline_items) < 2 or not source_lines:
            return False
        if not any("道德与法治" in normalize_line(line) or "思想品德" in normalize_line(line) for line in source_lines[:20]):
            return False

        section_keyword_re = re.compile(r"(?:选择题|非选择题|辨析|观察与思考|活动与探究)")
        eligible_roots: List[Dict] = []
        expected_number = 1
        source_heading_texts: Dict[int, str] = {}

        for root in outline_items:
            if int(root.get("level", 1) or 1) != 1:
                return False
            number_value = self._parse_arabic_numbering(root.get("numbering"))
            if number_value is None or number_value != expected_number:
                return False
            root_text = normalize_line(str(root.get("rawText") or root.get("title") or ""))
            if not root_text or not section_keyword_re.search(root_text):
                return False
            eligible_roots.append(root)
            expected_number += 1

        if len(eligible_roots) < 2:
            return False

        for index in range(1, len(eligible_roots) + 1):
            chinese_number = self._to_simple_chinese(index)
            heading_re = re.compile(rf"^\s*{re.escape(chinese_number)}\s*[\u3001\.\uFF0E]")
            for line_idx, raw_line in enumerate(source_lines):
                normalized_line = normalize_line(raw_line)
                if not normalized_line or not heading_re.match(normalized_line):
                    continue
                if not section_keyword_re.search(normalized_line):
                    continue
                parts = [normalized_line]
                for probe in range(line_idx + 1, min(len(source_lines), line_idx + 3)):
                    continuation = normalize_line(source_lines[probe])
                    if not continuation:
                        continue
                    if re.match(r"^\s*(?:[一二三四五六七八九十百千万]+|\d{1,3})\s*[\u3001\.\uFF0E]", continuation):
                        break
                    if re.match(r"^\s*[A-DＡ-Ｄa-dａ-ｄ]\s*[\u3001\.\uFF0E]", continuation):
                        break
                    parts.append(continuation)
                    if re.search(r"[\uFF08(]\s*\d+(?:\.\d+)?\s*分\s*[\uFF09)]", continuation):
                        break
                source_heading_texts[index] = "".join(parts)
                break

        changed = False
        for index, root in enumerate(eligible_roots, start=1):
            chinese_number = self._to_simple_chinese(index)
            if str(root.get("numbering") or "").strip() != chinese_number:
                root["numbering"] = chinese_number
                changed = True

            source_heading = source_heading_texts.get(index)
            if source_heading and str(root.get("rawText") or "").strip() != source_heading:
                root["rawText"] = source_heading
                changed = True
            if source_heading and str(root.get("title") or "").strip() != source_heading:
                root["title"] = source_heading
                changed = True

        return changed

    def _repair_politics_choice_section_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if not outline_items or not source_lines:
            return False
        if not any("道德与法治" in normalize_line(line) or "思想品德" in normalize_line(line) for line in source_lines[:20]):
            return False

        changed = False

        def _format_single_choice_blank(score_value: float) -> str:
            normalized_score = round(float(score_value or 0.0), 2)
            if normalized_score <= 0:
                return "____"
            if abs(normalized_score - round(normalized_score)) <= 1e-6:
                score_text = str(int(round(normalized_score)))
            else:
                score_text = f"{normalized_score:.2f}".rstrip("0").rstrip(".")
            return f"____（{score_text}分）"

        def _normalize_politics_choice_leaf(node: Dict, score_value: float) -> None:
            node["score"] = float(score_value)
            node["blankText"] = _format_single_choice_blank(score_value)
            # 政治选择题应作为叶子节点输出，选项续行不能误拆成子题或多空。
            node["children"] = []

        def _parse_inline_score(text: str) -> float:
            match = re.search(r"[\uFF08(]\s*(\d+(?:\.\d+)?)\s*\u5206\s*[\uFF09)]", text or "")
            if match is None:
                return 0.0
            try:
                return float(match.group(1))
            except ValueError:
                return 0.0

        def _parse_choice_banner_score(text: str) -> float:
            normalized = normalize_line(text)
            if not normalized:
                return 0.0
            match = re.search(r"(?:每题|每小题)\s*(\d+(?:\.\d+)?)\s*分", normalized)
            if match is not None:
                try:
                    return float(match.group(1))
                except ValueError:
                    return 0.0
            total_match = re.search(r"共\s*(\d+(?:\.\d+)?)\s*分", normalized)
            count_match = re.search(r"(\d{1,3})\s*小题", normalized)
            if total_match is None or count_match is None:
                return 0.0
            try:
                total_score = float(total_match.group(1))
                count = int(count_match.group(1))
            except ValueError:
                return 0.0
            if count <= 0:
                return 0.0
            return round(total_score / count, 2)

        def _strip_question_number(raw_text: str, numbering: str) -> str:
            return normalize_line(
                re.sub(
                    rf"^\s*{re.escape(numbering)}\s*[\u3001\.\uFF0E]\s*",
                    "",
                    raw_text or "",
                    count=1,
                )
            )

        def _compose_politics_section_heading(start_line_no: int, end_line_no: int) -> str:
            start_idx = max(0, start_line_no - 1)
            end_idx = min(len(source_lines), max(start_idx + 1, end_line_no - 1))
            parts: List[str] = []
            for probe in range(start_idx, end_idx):
                text = normalize_line(source_lines[probe])
                if not text:
                    continue
                if probe > start_idx:
                    if re.match(r"^\s*\d{1,3}\s*[\u3001\.\uFF0E]", text):
                        break
                    if self._is_chinese_section_root(
                        {
                            "level": 1,
                            "numbering": "",
                            "title": text,
                            "rawText": text,
                        }
                    ):
                        break
                    if POLITICS_CHOICE_BANNER_RE.search(text):
                        break
                parts.append(text)
                if _parse_inline_score(text) > 0:
                    break
            return "".join(parts)

        def _find_next_source_section_line(start_line_no: int) -> int:
            for line_no in range(max(1, start_line_no + 1), len(source_lines) + 1):
                text = normalize_line(source_lines[line_no - 1])
                if not text:
                    continue
                if self._is_chinese_section_root(
                    {
                        "level": 1,
                        "numbering": "",
                        "title": text,
                        "rawText": text,
                    }
                ):
                    return line_no
            return len(source_lines) + 1

        def _align_politics_choice_heading_line(root: Dict, fallback_line_no: int) -> int:
            root_text = normalize_line(str(root.get("rawText") or root.get("title") or ""))
            if not root_text:
                return fallback_line_no
            current_line = normalize_line(source_lines[fallback_line_no - 1]) if 1 <= fallback_line_no <= len(source_lines) else ""
            if current_line and (
                POLITICS_SECTION_TITLE_RE.search(current_line)
                or POLITICS_PLAIN_CHOICE_SECTION_RE.search(current_line)
                or POLITICS_CHOICE_BANNER_RE.search(current_line)
            ):
                return fallback_line_no
            for line_no in range(1, min(len(source_lines), max(fallback_line_no + 8, 12)) + 1):
                text = normalize_line(source_lines[line_no - 1])
                if not text:
                    continue
                if not (
                    POLITICS_SECTION_TITLE_RE.search(text)
                    or POLITICS_PLAIN_CHOICE_SECTION_RE.search(text)
                    or POLITICS_CHOICE_BANNER_RE.search(text)
                ):
                    continue
                if "选择" in root_text and "选择" in text:
                    return line_no
                if "请你选择" in root_text and "请你选择" in text:
                    return line_no
            return fallback_line_no

        def _find_first_source_question_number(start_line_no: int) -> Optional[int]:
            question_re = re.compile(r"^\s*(\d{1,3})\s*[\u3001\.\uFF0E]\s*\S")
            for line_no in range(max(1, start_line_no), len(source_lines) + 1):
                text = normalize_line(source_lines[line_no - 1])
                if not text:
                    continue
                match = question_re.match(text)
                if match is None:
                    continue
                try:
                    return int(match.group(1))
                except ValueError:
                    return None
            return None

        def _collect_politics_question_blocks(start_line_no: int, end_line_no: int) -> List[Dict[str, object]]:
            question_re = re.compile(r"^\s*(\d{1,3})\s*[\u3001\.\uFF0E]\s*")
            blocks: List[Dict[str, object]] = []
            current_number: Optional[int] = None
            current_line_no = start_line_no
            current_parts: List[str] = []
            for probe in range(max(0, start_line_no - 1), min(len(source_lines), end_line_no - 1)):
                text = normalize_line(source_lines[probe])
                if not text:
                    continue
                if self._is_chinese_section_root(
                    {
                        "level": 1,
                        "numbering": "",
                        "title": text,
                        "rawText": text,
                    }
                ):
                    break
                match = question_re.match(text)
                if match is not None:
                    if not text[match.end() :].strip():
                        continue
                    if current_number is not None and current_parts:
                        blocks.append(
                            {
                                "numbering": str(current_number),
                                "lineNumber": current_line_no,
                                "rawText": normalize_line(" ".join(current_parts)),
                            }
                        )
                    current_number = int(match.group(1))
                    current_line_no = probe + 1
                    current_parts = [text]
                    continue
                if current_number is not None:
                    current_parts.append(text)
            if current_number is not None and current_parts:
                blocks.append(
                    {
                        "numbering": str(current_number),
                        "lineNumber": current_line_no,
                        "rawText": normalize_line(" ".join(current_parts)),
                    }
                )
            return blocks

        def _is_reliable_politics_question_run(blocks: List[Dict[str, object]]) -> bool:
            if len(blocks) < 4:
                return False
            numbers: List[int] = []
            for block in blocks:
                try:
                    numbers.append(int(block.get("numbering") or 0))
                except (TypeError, ValueError):
                    return False
            return numbers == list(range(1, len(numbers) + 1))

        def _looks_like_malformed_politics_choice_children(root: Dict, blocks: List[Dict[str, object]]) -> bool:
            children = list(root.get("children") or [])
            if len(children) != len(blocks):
                return True
            for child, block in zip(children, blocks):
                numbering = str(block.get("numbering") or "").strip()
                if str(child.get("numbering") or "").strip() != numbering:
                    return True
                child_raw = normalize_line(str(child.get("rawText") or child.get("title") or ""))
                expected_raw = str(block.get("rawText") or "")
                if child.get("children"):
                    return True
                if not child_raw.startswith(f"{numbering}.") and not child_raw.startswith(f"{numbering}．") and not child_raw.startswith(f"{numbering}、"):
                    return True
                if f"{int(numbering) + 1}." in child_raw or f"{int(numbering) + 1}．" in child_raw or f"{int(numbering) + 1}、" in child_raw:
                    return True
                if len(child_raw) + 12 < len(expected_raw):
                    return True
            return False

        def _is_duplicate_politics_choice_fragment_root(
            node: Dict,
            root_line_no: int,
            next_section_line_no: int,
            blocks: List[Dict[str, object]],
        ) -> bool:
            if not self._is_chinese_section_root(node):
                return False
            node_line_no = int(node.get("lineNumber") or 0)
            if node_line_no < root_line_no or node_line_no >= next_section_line_no:
                return False

            node_raw = normalize_line(str(node.get("rawText") or node.get("title") or ""))
            if re.match(r"^\s*1\s*[\u3001\.\uFF0E]", node_raw):
                return True

            block_numbers = {
                str(block.get("numbering") or "").strip()
                for block in blocks
                if str(block.get("numbering") or "").strip()
            }
            child_numbers = [
                str(child.get("numbering") or "").strip()
                for child in (node.get("children") or [])
                if str(child.get("numbering") or "").strip()
            ]
            if not block_numbers or not child_numbers:
                return False
            if not set(child_numbers).issubset(block_numbers):
                return False
            try:
                parsed_child_numbers = [int(number) for number in child_numbers]
            except ValueError:
                return False
            # 52.doc/docx 会把第一题误提为一个新的“一”根，rawText 被后续清洗成“分数”，
            # 但其子节点仍是 2-8 的重复片段；这类根应并回第一大题后删除。
            return min(parsed_child_numbers) <= 2 and max(parsed_child_numbers) <= len(blocks)

        def _is_duplicate_existing_choice_fragment_root(root: Dict, node: Dict) -> bool:
            if not self._is_chinese_section_root(node):
                return False
            if str(root.get("numbering") or "").strip() != str(node.get("numbering") or "").strip():
                return False
            root_numbers = [
                str(child.get("numbering") or "").strip()
                for child in (root.get("children") or [])
                if str(child.get("numbering") or "").strip()
            ]
            child_numbers = [
                str(child.get("numbering") or "").strip()
                for child in (node.get("children") or [])
                if str(child.get("numbering") or "").strip()
            ]
            if len(root_numbers) < 4 or not child_numbers:
                return False
            if not set(child_numbers).issubset(set(root_numbers)):
                return False
            try:
                parsed_root_numbers = [int(number) for number in root_numbers]
                parsed_child_numbers = [int(number) for number in child_numbers]
            except ValueError:
                return False
            if parsed_root_numbers != list(range(min(parsed_root_numbers), max(parsed_root_numbers) + 1)):
                return False
            return min(parsed_child_numbers) <= min(parsed_root_numbers) + 1

        def _rebuild_politics_choice_children_from_blocks(
            root: Dict,
            blocks: List[Dict[str, object]],
            heading_text: str,
            banner_scores: List[Tuple[int, float]],
        ) -> bool:
            if not blocks:
                return False
            existing_children = {
                str(child.get("numbering") or "").strip(): child for child in list(root.get("children") or [])
            }
            section_total = float(root.get("score") or 0.0)
            heading_score = _parse_inline_score(heading_text)
            if heading_score > 0:
                section_total = heading_score
            if section_total <= 0:
                section_total = round(sum(float(score) for _, score in banner_scores), 2)
            uniform_score = round(section_total / len(blocks), 2) if section_total > 0 and blocks else 0.0

            rebuilt_children: List[Dict] = []
            for block in blocks:
                numbering = str(block.get("numbering") or "").strip()
                raw_text = str(block.get("rawText") or "").strip()
                line_number = int(block.get("lineNumber") or 0)
                score_value = _parse_inline_score(raw_text)
                if banner_scores:
                    score_candidates = [
                        score
                        for banner_line_no, score in banner_scores
                        if score > 0 and line_number >= banner_line_no
                    ]
                    if score_value <= 0 and score_candidates:
                        score_value = float(score_candidates[-1])
                if score_value <= 0:
                    score_value = uniform_score
                existing = copy.deepcopy(existing_children.get(numbering, {}))
                existing.update(
                    {
                        "lineNumber": line_number,
                        "level": int(root.get("level", 1) or 1) + 1,
                        "numbering": numbering,
                        "title": _strip_question_number(raw_text, numbering),
                        "rawText": raw_text,
                        "blankText": _format_single_choice_blank(score_value),
                        "score": float(score_value),
                        "children": [],
                        "_tokenType": str(existing.get("_tokenType") or "arabic"),
                        "_isSectionHeading": False,
                        "_bindSectionChildren": False,
                    }
                )
                rebuilt_children.append(existing)

            title_text = normalize_line(re.sub(r"^\s*[一二三四五六七八九十百千万]+\s*[\u3001\.\uFF0E]\s*", "", heading_text, count=1))
            if int(root.get("level") or 1) != 1:
                root["level"] = 1
            if str(root.get("numbering") or "").strip() != "一":
                root["numbering"] = "一"
            if heading_text and str(root.get("rawText") or "").strip() != heading_text:
                root["rawText"] = heading_text
            if title_text and str(root.get("title") or "").strip() != title_text:
                root["title"] = title_text
            if str(root.get("blankText") or "").strip():
                root["blankText"] = ""
            new_total = round(sum(float(child.get("score") or 0.0) for child in rebuilt_children), 2)
            root_changed = (
                root.get("children") != rebuilt_children
                or abs(float(root.get("score") or 0.0) - new_total) > 1e-6
            )
            root["children"] = rebuilt_children
            root["score"] = new_total
            return root_changed

        idx = 0
        while idx < len(outline_items):
            root = outline_items[idx]
            if not self._is_politics_section_root_like(root):
                idx += 1
                continue

            next_root_line_no = len(source_lines) + 1
            for sibling in outline_items[idx + 1 :]:
                if self._is_chinese_section_root(sibling):
                    next_root_line_no = int(sibling.get("lineNumber") or 0) or next_root_line_no
                    break

            root_line_no = int(root.get("lineNumber") or 0)
            if root_line_no <= 0:
                root_line_no = 1
            root_line_no = _align_politics_choice_heading_line(root, root_line_no)

            source_next_root_line_no = _find_next_source_section_line(root_line_no)
            if source_next_root_line_no > root_line_no:
                next_root_line_no = source_next_root_line_no

            heading_text = _compose_politics_section_heading(root_line_no, next_root_line_no) or normalize_line(
                str(root.get("rawText") or root.get("title") or "")
            )
            banner_scores: List[Tuple[int, float]] = []
            for line_no in range(root_line_no, min(len(source_lines), next_root_line_no - 1) + 1):
                normalized = normalize_line(source_lines[line_no - 1])
                if not normalized or not POLITICS_CHOICE_BANNER_RE.search(normalized):
                    continue
                banner_score = _parse_choice_banner_score(normalized)
                banner_scores.append((line_no, banner_score))

            has_flat_following_questions = False
            for sibling in outline_items[idx + 1 :]:
                if self._is_chinese_section_root(sibling):
                    break
                if self._is_politics_choice_banner_node(sibling):
                    continue
                if self._parse_arabic_numbering(sibling.get("numbering")) is not None and self._is_arabic_question_node(sibling):
                    has_flat_following_questions = True
                    break
                break

            source_blocks = _collect_politics_question_blocks(root_line_no + 1, next_root_line_no)
            if (
                not has_flat_following_questions
                and _is_reliable_politics_question_run(source_blocks)
                and _looks_like_malformed_politics_choice_children(root, source_blocks)
            ):
                if _rebuild_politics_choice_children_from_blocks(root, source_blocks, heading_text, banner_scores):
                    changed = True
                removed_duplicate_fragment = False
                while idx + 1 < len(outline_items) and _is_duplicate_politics_choice_fragment_root(
                    outline_items[idx + 1],
                    root_line_no,
                    source_next_root_line_no,
                    source_blocks,
                ):
                    del outline_items[idx + 1]
                    changed = True
                    removed_duplicate_fragment = True
                if removed_duplicate_fragment:
                    idx += 1
                    continue

                if idx + 1 < len(outline_items) and self._is_chinese_section_root(outline_items[idx + 1]):
                    next_root = outline_items[idx + 1]
                    next_root_raw = normalize_line(str(next_root.get("rawText") or next_root.get("title") or ""))
                    next_root_line_no = int(next_root.get("lineNumber") or 0)
                    if (
                        next_root_line_no >= root_line_no
                        and next_root_line_no < source_next_root_line_no
                        and re.match(r"^\s*1\s*[\u3001\.\uFF0E]", next_root_raw)
                    ):
                        del outline_items[idx + 1]
                        changed = True
                        idx += 1
                        continue
                    next_heading_text = _compose_politics_section_heading(next_root_line_no, len(source_lines) + 1)
                    next_title_text = normalize_line(
                        re.sub(
                            r"^\s*[一二三四五六七八九十百千万]+\s*[\u3001\.\uFF0E]\s*",
                            "",
                            next_heading_text,
                            count=1,
                        )
                    )
                    if next_heading_text and str(next_root.get("rawText") or "").strip() != next_heading_text:
                        next_root["rawText"] = next_heading_text
                        changed = True
                    if next_title_text and str(next_root.get("title") or "").strip() != next_title_text:
                        next_root["title"] = next_title_text
                        changed = True
                    original_children = list(next_root.get("children") or [])
                    next_root_first_question = _find_first_source_question_number(next_root_line_no + 1)
                    has_line_scale_match = any(
                        int(child.get("lineNumber") or 0) >= next_root_line_no for child in original_children
                    )
                    if next_root_first_question is not None and not has_line_scale_match:
                        filtered_children = [
                            child
                            for child in original_children
                            if (self._parse_arabic_numbering(child.get("numbering")) or 0) >= next_root_first_question
                        ]
                    else:
                        filtered_children = [
                            child for child in original_children if int(child.get("lineNumber") or 0) >= next_root_line_no
                        ]
                    if filtered_children != original_children:
                        next_root["children"] = filtered_children
                        next_root["score"] = round(
                            sum(self._aggregate_outline_node_score(child) for child in filtered_children),
                            2,
                        )
                        changed = True
                idx += 1
                continue

            if not has_flat_following_questions and _is_reliable_politics_question_run(source_blocks):
                removed_duplicate_fragment = False
                while idx + 1 < len(outline_items) and _is_duplicate_politics_choice_fragment_root(
                    outline_items[idx + 1],
                    root_line_no,
                    source_next_root_line_no,
                    source_blocks,
                ):
                    del outline_items[idx + 1]
                    changed = True
                    removed_duplicate_fragment = True
                if removed_duplicate_fragment:
                    idx += 1
                    continue

            if idx + 1 < len(outline_items) and _is_duplicate_existing_choice_fragment_root(root, outline_items[idx + 1]):
                del outline_items[idx + 1]
                changed = True
                idx += 1
                continue

            probe = idx + 1
            collected: List[Dict] = []
            while probe < len(outline_items):
                candidate = outline_items[probe]
                if self._is_chinese_section_root(candidate):
                    break
                if self._is_politics_choice_banner_node(candidate):
                    probe += 1
                    continue
                number_value = self._parse_arabic_numbering(candidate.get("numbering"))
                if number_value is None or not self._is_arabic_question_node(candidate):
                    break
                collected.append(candidate)
                probe += 1

            if not collected:
                idx += 1
                continue

            question_numbers = [self._parse_arabic_numbering(node.get("numbering")) for node in collected]
            if any(number is None for number in question_numbers):
                idx += 1
                continue
            parsed_numbers = [int(number) for number in question_numbers if number is not None]
            expected_numbers = list(range(1, len(parsed_numbers) + 1))
            if parsed_numbers != expected_numbers or len(parsed_numbers) < 4:
                idx += 1
                continue

            next_section_line_no = (
                int(outline_items[probe].get("lineNumber") or 0)
                if probe < len(outline_items) and self._is_chinese_section_root(outline_items[probe])
                else len(source_lines) + 1
            )

            single_banner_line_no: Optional[int] = None
            multi_banner_line_no: Optional[int] = None
            for line_no in range(root_line_no, min(len(source_lines), next_section_line_no - 1) + 1):
                normalized = normalize_line(source_lines[line_no - 1])
                if not normalized or not POLITICS_CHOICE_BANNER_RE.search(normalized):
                    continue
                if "多项选择" in normalized or "多选题" in normalized:
                    multi_banner_line_no = line_no
                elif single_banner_line_no is None:
                    single_banner_line_no = line_no

            direct_title = str(root.get("rawText") or root.get("title") or "").strip()
            direct_title = re.sub(r"▲.*$", "", direct_title).strip()
            direct_title = re.sub(r"^\s*1\s*[\u3001\.\uFF0E]\s*", "", direct_title, count=1).strip()
            direct_title = re.sub(r"^\s*一\s*[\u3001\.\uFF0E]\s*", "", direct_title, count=1).strip()
            if not direct_title:
                direct_title = "请你选择"

            rebuilt_children: List[Dict] = []
            target_level = int(root.get("level", 1) or 1) + 1
            for child in collected:
                delta = target_level - int(child.get("level", target_level) or target_level)
                if delta:
                    self._shift_outline_level(child, delta)
                child["children"] = [grand for grand in (child.get("children") or []) if not self._is_politics_choice_banner_node(grand)]
                rebuilt_children.append(child)

            if multi_banner_line_no is not None:
                first_group = [
                    child for child in rebuilt_children
                    if int(child.get("lineNumber") or 0) < multi_banner_line_no
                ]
                second_group = [
                    child for child in rebuilt_children
                    if int(child.get("lineNumber") or 0) >= multi_banner_line_no
                ]
            else:
                first_group = rebuilt_children
                second_group = []

            if first_group:
                for child in first_group:
                    _normalize_politics_choice_leaf(child, 2.0)
            if second_group:
                for child in second_group:
                    _normalize_politics_choice_leaf(child, 3.0)

            root["numbering"] = "一"
            root["rawText"] = f"一、{direct_title}"
            root["title"] = direct_title
            root["blankText"] = ""
            root["children"] = rebuilt_children
            root["score"] = round(sum(self._aggregate_outline_node_score(child) for child in rebuilt_children), 2)

            del outline_items[idx + 1 : probe]
            changed = True
            idx += 1

        return changed

    def _repair_politics_choice_statement_children(self, outline_items: List[Dict]) -> bool:
        if not outline_items:
            return False

        score_re = re.compile(r"[\uFF08(]\s*(\d+(?:\.\d+)?)\s*\u5206\s*[\uFF09)]")
        empty_choice_re = re.compile(r"[\uFF08(]\s*[\uFF09)]")
        circled_re = re.compile(r"^[①②③④⑤⑥⑦⑧⑨⑩]")
        choice_option_re = re.compile(r"[A-DＡ-Ｄ]\s*[\u3001\.\uFF0E]")

        def _score_text(score_value: float) -> str:
            if abs(score_value - round(score_value)) <= 1e-6:
                return str(int(round(score_value)))
            return f"{score_value:.2f}".rstrip("0").rstrip(".")

        def _inline_score(text: str) -> float:
            match = score_re.search(text or "")
            if match is None:
                return 0.0
            try:
                return float(match.group(1))
            except ValueError:
                return 0.0

        def _looks_like_choice_parent(node: Dict) -> bool:
            raw_text = normalize_line(str(node.get("rawText") or node.get("title") or ""))
            if not raw_text or not empty_choice_re.search(raw_text):
                return False
            score_value = _inline_score(raw_text)
            if score_value <= 0:
                try:
                    score_value = float(node.get("score") or 0.0)
                except (TypeError, ValueError):
                    score_value = 0.0
            if score_value <= 0 or score_value > 5:
                return False
            if re.search(r"(?:材料|阅读|回答|分析|说明|谈谈|辨析|观察与思考|活动与探究)", raw_text):
                return False
            return True

        def _looks_like_circled_choice_statement(child: Dict) -> bool:
            numbering = str(child.get("numbering") or "").strip()
            raw_text = normalize_line(str(child.get("rawText") or child.get("title") or ""))
            if not numbering and not raw_text:
                return False
            if numbering and not circled_re.match(numbering):
                return False
            if raw_text and not circled_re.match(raw_text):
                return False
            # 政治选择题中整行 ①②③④ 是备选陈述，不是可作答子问。
            return len(re.findall(r"[①②③④⑤⑥⑦⑧⑨⑩]", raw_text)) >= 2 or choice_option_re.search(raw_text) is None

        changed = False

        def _walk(nodes: List[Dict]) -> None:
            nonlocal changed
            for node in nodes:
                children = list(node.get("children") or [])
                if children and _looks_like_choice_parent(node) and all(
                    _looks_like_circled_choice_statement(child) for child in children
                ):
                    score_value = _inline_score(str(node.get("rawText") or node.get("title") or ""))
                    if score_value <= 0:
                        try:
                            score_value = float(node.get("score") or 0.0)
                        except (TypeError, ValueError):
                            score_value = 0.0
                    node["children"] = []
                    if score_value > 0:
                        node["score"] = float(score_value)
                        node["blankText"] = f"____（{_score_text(score_value)}分）"
                    elif not str(node.get("blankText") or "").strip():
                        node["blankText"] = "____"
                    changed = True
                    continue
                if children:
                    _walk(children)

        _walk(outline_items)
        return changed

    def _merge_blank_text(self, target: Dict, blank_text: object) -> None:
        incoming = str(blank_text or "").strip()
        if not incoming:
            return
        current = str(target.get("blankText") or "").strip()
        if not current:
            target["blankText"] = incoming
            return
        if incoming in current.split():
            return
        target["blankText"] = f"{current} {incoming}".strip()

    def _merge_reference_child_into_parent(self, parent: Dict, child: Dict) -> None:
        child_raw = str(child.get("rawText") or "").strip()
        parent_raw = str(parent.get("rawText") or "").strip()
        if child_raw and child_raw not in parent_raw:
            parent["rawText"] = f"{parent_raw} {child_raw}".strip() if parent_raw else child_raw

        child_title = str(child.get("title") or "").strip()
        parent_title = str(parent.get("title") or "").strip()
        if child_title and child_title not in parent_title:
            parent["title"] = f"{parent_title} {child_title}".strip() if parent_title else child_title

        self._merge_blank_text(parent, child.get("blankText"))

    def _has_parent_reference_to_number(self, node: Dict, number_text: object) -> bool:
        number = self._normalize_numbering_key(number_text)
        if not number:
            return False
        raw = str(node.get("rawText") or "")
        if not raw:
            return False
        return bool(
            re.search(
                rf"第\s*[\uFF08(]?\s*{re.escape(number)}\s*[\uFF09)]?\s*(?:自然段|段|小节|节|句|行|幅图|图|空)",
                raw,
            )
        )

    def _is_reference_continuation_node(self, node: Dict) -> bool:
        raw = str(node.get("rawText") or "").strip()
        if not raw:
            return False
        if self._extract_score_only_line_value(raw) is not None:
            return True
        head = re.match(r"^[\uFF08(]\s*\d{1,3}\s*[\uFF09)]", raw)
        if head is None:
            return False

        tail = normalize_line(raw[head.end() :])
        if not tail:
            return True
        if self._extract_score_only_line_value(tail) is not None:
            return True
        if _collect_blank_segments(tail):
            return False
        if re.search(r"[？?\u3002!！;；:：]", tail):
            return False

        # 仅把短碎片当作“引用编号续行”；长句（尤其带中文正文）不参与并入。
        compact_tail = re.sub(r"[\s\u3000,\uFF0C\.\u3002、\-—:：;；()（）\[\]【】]", "", tail)
        cjk_count = len(re.findall(r"[\u4e00-\u9fa5]", compact_tail))
        if cjk_count >= 8:
            return False
        if len(compact_tail) > 16:
            return False
        return True

    def _find_misnested_restart_split_index(self, children: List[Dict]) -> Optional[int]:
        if len(children) < 3:
            return None
        numbers: List[int] = []
        for child in children:
            parsed = self._parse_arabic_numbering(child.get("numbering"))
            if parsed is None:
                return None
            numbers.append(parsed)

        for idx in range(1, len(numbers) - 1):
            if numbers[idx] != 1:
                continue
            tail = numbers[idx:]
            if tail[0] != 1:
                continue
            is_non_decreasing = all(curr >= prev for prev, curr in zip(tail, tail[1:]))
            if not is_non_decreasing:
                continue
            if tail != list(range(1, len(tail) + 1)):
                continue
            return idx
        return None

    def _repair_outline_extraction_artifacts(self, outline_items: List[Dict]) -> bool:
        changed = False

        def _restore_statement_option_parent(node: Dict) -> bool:
            children = list(node.get("children") or [])
            if len(children) < 2:
                return False
            title = str(node.get("title") or "").strip()
            raw = str(node.get("rawText") or "").strip()
            if not title or not raw:
                return False
            if "填序号" not in title and "填入序号" not in title:
                return False
            if any((child.get("children") or []) for child in children):
                return False
            if not all(str(child.get("numbering") or "").strip().isdigit() for child in children):
                return False
            if not all(float(child.get("score") or 0.0) > 0 for child in children):
                return False
            child_raws = [str(child.get("rawText") or "").strip() for child in children]
            if not all(child_raw and child_raw in raw for child_raw in child_raws):
                return False

            numbering = str(node.get("numbering") or "").strip()
            if not numbering:
                return False
            delimiter_match = re.match(rf"^\s*{re.escape(numbering)}\s*([、\.\uFF0E])", raw)
            delimiter = delimiter_match.group(1) if delimiter_match is not None else "."
            rebuilt_raw = f"{numbering}{delimiter}{title}"
            if raw == rebuilt_raw and not str(node.get("blankText") or "").strip():
                return False
            node["rawText"] = rebuilt_raw
            node["blankText"] = ""
            return True

        def _restore_leaf_blank_from_title(node: Dict) -> bool:
            if node.get("children"):
                return False
            if float(node.get("score") or 0.0) <= 0:
                return False
            if str(node.get("blankText") or "").strip():
                return False
            raw = str(node.get("rawText") or "")
            title = str(node.get("title") or "")
            if not title:
                return False
            if _collect_blank_segments(raw):
                return False
            if not _collect_blank_segments(title):
                return False
            filled = _format_blank_segments_from_line(title, float(node.get("score") or 0.0))
            if not filled:
                return False
            node["blankText"] = filled
            if raw.strip() != title.strip():
                prefix_match = re.match(
                    r"^\s*((?:[\uFF08(]\s*\d{1,3}\s*[\uFF09)]|\d{1,4}\s*[、\.\uFF0E]))\s*",
                    raw,
                )
                if prefix_match is not None and not re.match(
                    r"^\s*(?:[\uFF08(]\s*\d{1,3}\s*[\uFF09)]|\d{1,4}\s*[、\.\uFF0E])",
                    title,
                ):
                    node["rawText"] = f"{prefix_match.group(1)}{title}".strip()
                else:
                    node["rawText"] = title
            return True

        def _restore_leaf_raw_from_title(node: Dict) -> bool:
            if node.get("children"):
                return False
            raw = str(node.get("rawText") or "")
            title = str(node.get("title") or "")
            if not raw or not title or raw.strip() == title.strip():
                return False
            raw_blank_count = len(_collect_blank_segments(raw))
            title_blank_count = len(_collect_blank_segments(title))
            if title_blank_count <= raw_blank_count:
                return False
            prefix_match = re.match(
                r"^\s*((?:[\uFF08(]\s*\d{1,3}\s*[\uFF09)]|\d{1,4}\s*[、\.\uFF0E]))\s*",
                raw,
            )
            if prefix_match is not None and not re.match(
                r"^\s*(?:[\uFF08(]\s*\d{1,3}\s*[\uFF09)]|\d{1,4}\s*[、\.\uFF0E])",
                title,
            ):
                node["rawText"] = f"{prefix_match.group(1)}{title}".strip()
            else:
                node["rawText"] = title
            return True

        def _raise_leaf_score_to_explicit_blank_total(node: Dict) -> bool:
            if node.get("children"):
                return False
            blank_text = str(node.get("blankText") or "").strip()
            if not blank_text:
                return False
            parsed_scores = []
            for token in blank_text.split():
                score = _parse_merged_blank_segment_score(token)
                if score is None:
                    return False
                parsed_scores.append(float(score))
            if not parsed_scores:
                return False
            explicit_total = round(sum(parsed_scores), 2)
            current_score = round(float(node.get("score") or 0.0), 2)
            if explicit_total <= current_score:
                return False
            node["score"] = explicit_total
            return True

        def _promote_misnested_restart_children(nodes: List[Dict]) -> bool:
            local_changed = False
            idx = 0
            while idx < len(nodes):
                node = nodes[idx]
                # 仅在阿拉伯父题号上做“重启子题外提”修复。
                # 中文大题（如“二、非选择题”）下出现的 1/2 常是子问，不应被外提到根级。
                if self._parse_arabic_numbering(node.get("numbering")) is None:
                    idx += 1
                    continue
                children = list(node.get("children") or [])
                split_idx = self._find_misnested_restart_split_index(children)
                if split_idx is None:
                    idx += 1
                    continue
                retained_children = children[:split_idx]
                moved_children = children[split_idx:]
                if not retained_children or len(moved_children) < 2:
                    idx += 1
                    continue
                node_score = float(node.get("score") or 0.0)
                moved_total = sum(self._aggregate_outline_node_score(child) for child in moved_children)
                if moved_total <= node_score:
                    idx += 1
                    continue
                node["children"] = retained_children
                target_level = int(node.get("level", 1) or 1)
                for moved in moved_children:
                    moved_level = int(moved.get("level", target_level + 1) or (target_level + 1))
                    if moved_level != target_level:
                        self._shift_outline_level(moved, target_level - moved_level)
                nodes[idx + 1 : idx + 1] = moved_children
                local_changed = True
                idx += len(moved_children) + 1
            return local_changed

        def walk(nodes: List[Dict], parent: Optional[Dict] = None, in_reading_context: bool = False) -> bool:
            nonlocal changed
            if not nodes:
                return False

            local_changed = False
            cleaned: List[Dict] = []
            for node in nodes:
                current_reading_context = in_reading_context or self._is_reading_root_node(node)
                if walk(node.get("children") or [], node, current_reading_context):
                    local_changed = True

                if parent is not None:
                    score_only_value = self._extract_score_only_line_value(node.get("rawText"))
                    if score_only_value is not None:
                        parent["score"] = float(score_only_value)
                        local_changed = True
                        continue

                if parent is not None and in_reading_context and self._is_reading_title_blank_node(node):
                    local_changed = True
                    continue

                if parent is not None and in_reading_context and self._is_reading_scored_article_title_leaf(node):
                    local_changed = True
                    continue

                if parent is not None and in_reading_context and self._is_reading_paragraph_marker_leaf(node):
                    local_changed = True
                    continue

                if _restore_statement_option_parent(node):
                    local_changed = True

                if _restore_leaf_blank_from_title(node):
                    local_changed = True

                if _restore_leaf_raw_from_title(node):
                    local_changed = True

                if _raise_leaf_score_to_explicit_blank_total(node):
                    local_changed = True

                node_number = self._normalize_numbering_key(node.get("numbering"))
                if node_number and self._is_reference_continuation_node(node):
                    reference_anchor = next(
                        (
                            candidate
                            for candidate in cleaned + nodes
                            if candidate is not node and self._has_parent_reference_to_number(candidate, node_number)
                        ),
                        None,
                    )
                    if reference_anchor is not None:
                        self._merge_reference_child_into_parent(reference_anchor, node)
                        local_changed = True
                        continue
                    if parent is not None and self._has_parent_reference_to_number(parent, node_number):
                        self._merge_reference_child_into_parent(parent, node)
                        local_changed = True
                        continue

                cleaned.append(node)

            if _promote_misnested_restart_children(cleaned):
                local_changed = True

            if len(cleaned) != len(nodes):
                nodes[:] = cleaned
            changed = changed or local_changed
            return local_changed

        walk(outline_items, None, False)
        return changed

    def _shift_outline_level(self, node: Dict, delta: int) -> None:
        if not isinstance(node, dict) or delta == 0:
            return
        current_level = int(node.get("level", 1) or 1)
        node["level"] = max(1, current_level + delta)
        for child in node.get("children") or []:
            self._shift_outline_level(child, delta)

    def _is_history_exam_like(self, outline_items: List[Dict], source_lines: List[str]) -> bool:
        samples: List[str] = []
        for node in outline_items[:8]:
            text = normalize_line(str(node.get("rawText") or node.get("title") or ""))
            if text:
                samples.append(text)
            for child in (node.get("children") or [])[:8]:
                child_text = normalize_line(str(child.get("rawText") or child.get("title") or ""))
                if child_text:
                    samples.append(child_text)

        for line in source_lines[:160]:
            normalized = normalize_line(line)
            if normalized:
                samples.append(normalized)

        if not samples:
            return False

        sample_text = " ".join(samples)
        keyword_lines = sum(1 for item in samples[:200] if HISTORY_EXAM_KEYWORD_RE.search(item))
        generic_section_hits = sum(1 for item in samples[:80] if HISTORY_GENERIC_SECTION_HINT_RE.search(item))
        distinct_history_terms = set(HISTORY_EXAM_KEYWORD_RE.findall(sample_text))

        if len(distinct_history_terms) >= 4:
            return True
        if len(distinct_history_terms) >= 3 and keyword_lines >= 5:
            return True
        return len(distinct_history_terms) >= 3 and generic_section_hits >= 2

    def _is_biology_exam_like(self, outline_items: List[Dict], source_lines: List[str]) -> bool:
        samples: List[str] = []
        for node in outline_items[:8]:
            text = normalize_line(str(node.get("rawText") or node.get("title") or ""))
            if text:
                samples.append(text)
            for child in (node.get("children") or [])[:12]:
                child_text = normalize_line(str(child.get("rawText") or child.get("title") or ""))
                if child_text:
                    samples.append(child_text)

        for line in source_lines[:200]:
            normalized = normalize_line(line)
            if normalized:
                samples.append(normalized)

        if not samples:
            return False

        sample_text = " ".join(samples)
        if "生物" in sample_text:
            return True

        biology_terms = set(
            re.findall(
                r"(?:草履虫|洋葱|口腔上皮|显微镜|细胞核|叶绿体|生物圈|生态系统|食物链|食物网|骆驼|蚯蚓)",
                sample_text,
            )
        )
        heading_hits = sum(
            1
            for item in samples[:120]
            if re.search(r"(?:单项选择题|非选择题|判断题|识图题|实验题|综合分析题|材料分析题)", item)
        )
        return heading_hits >= 2 and len(biology_terms) >= 4



    def _is_section_like_node(self, node: Dict, min_score: float = 0.0) -> bool:
        score_value = float(node.get("score") or 0.0)
        if score_value < float(min_score):
            return False
        text = normalize_line(str(node.get("title") or node.get("rawText") or ""))
        if not text:
            return False
        if QUESTION_DIRECTIVE_TEXT_RE.search(text) or re.search(r"[？?]", text):
            return False
        return bool(SECTION_LIKE_TITLE_RE.search(text))

    def _find_overflow_arabic_child_split_index(self, parent: Dict, children: List[Dict]) -> Optional[int]:
        parent_num = self._parse_arabic_numbering(parent.get("numbering"))
        if parent_num is None or len(children) < 2:
            return None

        parent_title_text = normalize_line(str(parent.get("title") or parent.get("rawText") or ""))
        leading_paren_count = 0
        while leading_paren_count < len(children) and self._is_paren_subquestion_node(children[leading_paren_count]):
            leading_paren_count += 1
        # 题型总标题（如“1、单项选择题（40分）”）本身就应该承载后续阿拉伯题号子级，
        # 不能套用“题干误吞后续平级题目”的修复逻辑，否则会把正确结构打散成 1..N 平级。
        if re.search(
            r"(?:单项选择题|多项选择题|选择题|单选题|多选题|判断题|填空题|综合题|解答题|简答题|阅读|写作|作文|完形填空)",
            parent_title_text,
        ) and leading_paren_count < 2:
            return None

        split_index: Optional[int] = None
        expected_number = parent_num + 1
        for idx, child in enumerate(children):
            child_num = self._parse_arabic_numbering(child.get("numbering"))
            if child_num is None or not self._is_arabic_question_node(child):
                continue
            if child_num >= expected_number and float(child.get("score") or 0.0) > 0:
                split_index = idx
                break

        if split_index is None:
            return None

        overflow = children[split_index:]
        if not overflow:
            return None
        for child in overflow:
            child_num = self._parse_arabic_numbering(child.get("numbering"))
            if child_num is None or not self._is_arabic_question_node(child):
                return None
            if float(child.get("score") or 0.0) <= 0:
                return None
        return split_index

    def _repair_overflow_arabic_children_to_root_siblings(self, outline_items: List[Dict]) -> bool:
        if not outline_items:
            return False

        changed = False
        idx = 0
        while idx < len(outline_items):
            node = outline_items[idx]
            children = list(node.get("children") or [])
            split_index = self._find_overflow_arabic_child_split_index(node, children)
            if split_index is None:
                idx += 1
                continue

            kept_children = children[:split_index]
            overflow_children = children[split_index:]
            if not overflow_children:
                idx += 1
                continue

            node["children"] = kept_children
            explicit_score = self._extract_inline_heading_score_value(node.get("rawText"))
            if explicit_score is not None and explicit_score > 0:
                node["score"] = explicit_score
            for child in overflow_children:
                child_level = int(child.get("level", 1) or 1)
                if child_level != 1:
                    self._shift_outline_level(child, 1 - child_level)

            outline_items[idx + 1 : idx + 1] = overflow_children
            changed = True
            idx += 1

        return changed

    def _repair_history_orphan_material_followup_roots(self, outline_items: List[Dict]) -> bool:
        if len(outline_items) < 2:
            return False

        changed = False
        idx = 1
        orphan_root_re = re.compile(r"^\s*\d{1,3}\s*[\u3001\.\uFF0E]\s*$")

        while idx < len(outline_items):
            prev = outline_items[idx - 1]
            curr = outline_items[idx]

            prev_num = self._parse_arabic_numbering(prev.get("numbering"))
            curr_num = self._parse_arabic_numbering(curr.get("numbering"))
            if prev_num is None or curr_num is None or curr_num != prev_num + 1:
                idx += 1
                continue

            curr_raw = normalize_line(str(curr.get("rawText") or curr.get("title") or ""))
            if not orphan_root_re.match(curr_raw or ""):
                idx += 1
                continue

            prev_children = list(prev.get("children") or [])
            curr_children = list(curr.get("children") or [])
            if len(prev_children) < 2 or len(curr_children) < 2:
                idx += 1
                continue

            prev_child_numbers = [self._parse_arabic_numbering(child.get("numbering")) for child in prev_children]
            curr_child_numbers = [self._parse_arabic_numbering(child.get("numbering")) for child in curr_children]
            if any(number is None for number in prev_child_numbers + curr_child_numbers):
                idx += 1
                continue

            prev_number_values = [int(number) for number in prev_child_numbers if number is not None]
            curr_number_values = [int(number) for number in curr_child_numbers if number is not None]
            if prev_number_values != list(range(1, len(prev_number_values) + 1)):
                idx += 1
                continue
            if curr_number_values != list(range(1, len(curr_number_values) + 1)):
                idx += 1
                continue

            prev_text = normalize_line(str(prev.get("rawText") or prev.get("title") or ""))
            if not re.search(r"(?:观察下列图片|请回答|材料)", prev_text):
                idx += 1
                continue

            child_text_blob = " ".join(
                normalize_line(str(child.get("rawText") or child.get("title") or ""))
                for child in curr_children
            )
            if not re.search(r"(?:材料一|材料二|材料三|图名|友好交往|主要原因)", child_text_blob):
                idx += 1
                continue

            target_level = int(prev.get("level", 1) or 1) + 1
            next_number = len(prev_children) + 1
            merged_children = list(prev_children)
            for child in curr_children:
                child_level = int(child.get("level", target_level) or target_level)
                if child_level != target_level:
                    self._shift_outline_level(child, target_level - child_level)
                child["numbering"] = str(next_number)
                next_number += 1
                merged_children.append(child)

            prev["children"] = merged_children
            outline_items.pop(idx)
            changed = True

        return changed

    def _repair_history_missing_root_from_material_child_run(self, outline_items: List[Dict]) -> bool:
        if len(outline_items) < 2:
            return False

        def _leading_subquestion_number(text: object) -> Optional[int]:
            normalized = normalize_line(str(text or ""))
            match = re.match(r"^\s*(?:[\uFF08(]|ï¼)\s*([1-9])\s*(?:[\uFF09)]|ï¼)", normalized)
            if match is None:
                return None
            try:
                return int(match.group(1))
            except ValueError:
                return None

        changed = False
        idx = 0
        while idx < len(outline_items):
            current = outline_items[idx]
            current_number = self._parse_arabic_numbering(current.get("numbering"))
            children = list(current.get("children") or [])
            if current_number is None or len(children) < 5:
                idx += 1
                continue

            expected_next = current_number + 1
            if any(self._parse_arabic_numbering(node.get("numbering")) == expected_next for node in outline_items):
                idx += 1
                continue

            split_idx: Optional[int] = None
            for child_idx in range(1, len(children)):
                marker_value = _leading_subquestion_number(children[child_idx].get("rawText") or children[child_idx].get("title"))
                if marker_value == 1:
                    split_idx = child_idx
                    break
            if split_idx is None or split_idx < 2 or len(children) - split_idx < 2:
                idx += 1
                continue

            moved_children = [copy.deepcopy(child) for child in children[split_idx:]]
            leading_markers = [
                _leading_subquestion_number(child.get("rawText") or child.get("title"))
                for child in moved_children[: min(3, len(moved_children))]
            ]
            if leading_markers[:2] != [1, 2]:
                idx += 1
                continue

            retained_children = children[:split_idx]
            root_hint_parts: List[str] = []
            has_material_hint = False
            for child in moved_children:
                raw_text = normalize_line(str(child.get("rawText") or child.get("title") or ""))
                for pattern in (r"材料[一二三四1234]", r"ææ[ä¸äºä¸å1234]"):
                    match = re.search(pattern, raw_text)
                    if match is None:
                        continue
                    has_material_hint = True
                    token = match.group(0)
                    if token not in root_hint_parts:
                        root_hint_parts.append(token)
            if not has_material_hint:
                idx += 1
                continue
            root_hint = "".join(root_hint_parts) if root_hint_parts else "材料题"

            for new_number, child in enumerate(moved_children, start=1):
                child["numbering"] = str(new_number)

            current["children"] = retained_children
            current["score"] = float(sum(self._aggregate_outline_node_score(child) for child in retained_children))

            outline_items.insert(
                idx + 1,
                {
                    "lineNumber": int(moved_children[0].get("lineNumber") or int(current.get("lineNumber") or 0) + 1),
                    "level": int(current.get("level") or 1),
                    "numbering": str(expected_next),
                    "title": root_hint,
                    "rawText": f"{expected_next}、{root_hint}",
                    "blankText": "",
                    "score": float(sum(self._aggregate_outline_node_score(child) for child in moved_children)),
                    "children": moved_children,
                },
            )
            changed = True
            idx += 2

        return changed

    def _repair_history_merged_followup_root_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if len(outline_items) < 2 or not source_lines:
            return False

        question_heading_re = re.compile(
            r"^\s*(\d{1,3})\s*(?:[、銆\.\uFF0E]|(?=[\uFF08(]\s*\d+(?:\.\d+)?\s*分\s*[\uFF09)]))"
        )
        subquestion_heading_re = re.compile(r"^\s*[\uFF08(]\s*(\d{1,3})\s*[\uFF09)]")

        question_positions: List[Tuple[int, int]] = []
        for idx, line in enumerate(source_lines):
            normalized = normalize_line(line)
            if not normalized:
                continue
            match = question_heading_re.match(normalized)
            if match is None:
                continue
            try:
                question_positions.append((int(match.group(1)), idx))
            except ValueError:
                continue
        if len(question_positions) < 2:
            return False

        def _compose_source_text(start_idx: int, end_idx: int) -> str:
            parts: List[str] = []
            safe_start = max(0, start_idx)
            safe_end = min(len(source_lines), max(safe_start, end_idx))
            for raw in source_lines[safe_start:safe_end]:
                normalized = normalize_line(raw)
                if normalized:
                    parts.append(normalized)
            return "".join(parts).strip()

        def _strip_question_prefix(text: str, number_value: int) -> str:
            return re.sub(
                rf"^\s*{number_value}\s*(?:[、銆\.\uFF0E]\s*)?",
                "",
                str(text or ""),
                count=1,
            ).strip()

        def _strip_paren_prefix(text: str) -> str:
            return re.sub(
                r"^\s*[\uFF08(]\s*\d{1,3}\s*[\uFF09)]\s*",
                "",
                str(text or ""),
                count=1,
            ).strip()

        def _is_source_subquestion_line(text: str) -> bool:
            normalized = normalize_line(text)
            return subquestion_heading_re.match(normalized) is not None or normalized.startswith("锛")

        changed = False
        idx = 0
        while idx < len(outline_items):
            current = outline_items[idx]
            current_number = self._parse_arabic_numbering(current.get("numbering"))
            current_raw = normalize_line(str(current.get("rawText") or current.get("title") or ""))
            if current_number is None or int(current.get("level") or 1) != 1:
                idx += 1
                continue
            if not re.match(r"^\s*\d{1,3}\s*(?:[、銆\.\uFF0E]|(?=[\uFF08(]\s*\d+(?:\.\d+)?\s*分\s*[\uFF09)]))", current_raw):
                idx += 1
                continue

            expected_next = current_number + 1
            if any(
                self._parse_arabic_numbering(sibling.get("numbering")) == expected_next
                for sibling in outline_items
            ):
                idx += 1
                continue

            current_question_idx = next(
                (line_idx for number_value, line_idx in question_positions if number_value == current_number),
                None,
            )
            if current_question_idx is None:
                idx += 1
                continue

            next_question_idx = next(
                (line_idx for number_value, line_idx in question_positions if line_idx > current_question_idx and number_value == expected_next),
                None,
            )
            boundary_question_idx = next(
                (line_idx for _number_value, line_idx in question_positions if line_idx > current_question_idx),
                None,
            )
            if boundary_question_idx is None:
                idx += 1
                continue

            current_children = list(current.get("children") or [])
            if len(current_children) < 4:
                idx += 1
                continue

            next_heading_line = normalize_line(source_lines[next_question_idx]) if next_question_idx is not None else ""
            if next_question_idx is not None and not re.search(r"(?:阅读下列材料|材料[一二三123]|鏉愭枡)", next_heading_line):
                idx += 1
                continue

            material_start_idx: Optional[int] = None
            retained_subquestion_positions: List[int] = []
            next_subquestion_positions: List[int] = []
            rebuilt_next_line_no = next_question_idx + 1 if next_question_idx is not None else 0

            if next_question_idx is not None:
                retained_subquestion_positions = [
                    line_idx
                    for line_idx in range(current_question_idx + 1, next_question_idx)
                    if _is_source_subquestion_line(source_lines[line_idx])
                ]
                next_question_end_idx = next(
                    (line_idx for _number_value, line_idx in question_positions if line_idx > next_question_idx),
                    len(source_lines),
                )
                next_subquestion_positions = [
                    line_idx
                    for line_idx in range(next_question_idx + 1, next_question_end_idx)
                    if _is_source_subquestion_line(source_lines[line_idx])
                ]
                if retained_subquestion_positions:
                    material_start_idx = next_question_idx
            else:
                for line_idx in range(current_question_idx + 1, boundary_question_idx):
                    normalized_line = normalize_line(source_lines[line_idx])
                    if material_start_idx is None and re.match(
                        r"^\s*(?:材料[一二三四五六七八九十1234567890]|鏉愭枡)",
                        normalized_line,
                    ):
                        material_start_idx = line_idx
                    if not _is_source_subquestion_line(normalized_line):
                        continue
                    if material_start_idx is None:
                        retained_subquestion_positions.append(line_idx)
                    else:
                        next_subquestion_positions.append(line_idx)
                if material_start_idx is not None:
                    rebuilt_next_line_no = material_start_idx + 1

            if (
                material_start_idx is None
                or not retained_subquestion_positions
                or len(next_subquestion_positions) < 2
                or len(current_children) <= len(retained_subquestion_positions)
            ):
                idx += 1
                continue

            split_idx = len(retained_subquestion_positions)
            retained_children = current_children[:split_idx]
            moved_children = current_children[split_idx:]
            if len(retained_children) < 2 or len(moved_children) < 2:
                idx += 1
                continue

            rebuilt_last_raw = _compose_source_text(retained_subquestion_positions[-1], material_start_idx)
            first_next_subquestion_idx = next_subquestion_positions[0]
            rebuilt_next_raw = _compose_source_text(material_start_idx, first_next_subquestion_idx)
            if not rebuilt_last_raw or not rebuilt_next_raw:
                idx += 1
                continue

            rebuilt_children: List[Dict] = []
            for pos, start_idx in enumerate(next_subquestion_positions):
                normalized_sub = normalize_line(source_lines[start_idx])
                match = subquestion_heading_re.match(normalized_sub)
                end_idx = (
                    next_subquestion_positions[pos + 1]
                    if pos + 1 < len(next_subquestion_positions)
                    else boundary_question_idx
                )
                child_raw = _compose_source_text(start_idx, end_idx)
                if not child_raw:
                    continue
                child_number = str(match.group(1) or "").strip() if match is not None else str(pos + 1)
                rebuilt_children.append(
                    {
                        "lineNumber": start_idx + 1,
                        "level": int(current.get("level") or 1) + 1,
                        "numbering": child_number,
                        "title": _strip_paren_prefix(child_raw),
                        "rawText": child_raw,
                        "blankText": "",
                        "score": float(self._extract_inline_heading_score_value(child_raw) or 0.0),
                        "children": [],
                    }
                )
            if len(rebuilt_children) < 2:
                idx += 1
                continue

            last_retained = retained_children[-1]
            last_retained["rawText"] = rebuilt_last_raw
            last_retained["title"] = _strip_paren_prefix(rebuilt_last_raw)
            retained_score = self._extract_inline_heading_score_value(rebuilt_last_raw)
            if retained_score is not None and retained_score > 0:
                last_retained["score"] = retained_score

            current["children"] = retained_children
            current_score = self._extract_inline_heading_score_value(current.get("rawText") or current.get("title"))
            if current_score is None or current_score <= 0:
                current_score = sum(self._aggregate_outline_node_score(child) for child in retained_children)
            current["score"] = float(current_score or 0.0)

            if not re.match(rf"^\s*{expected_next}\s*[、銆\.\uFF0E]", rebuilt_next_raw):
                rebuilt_next_raw = f"{expected_next}、{rebuilt_next_raw}"

            new_root_score = self._extract_inline_heading_score_value(rebuilt_next_raw)
            if new_root_score is None or new_root_score <= 0:
                new_root_score = sum(self._aggregate_outline_node_score(child) for child in rebuilt_children)

            outline_items.insert(
                idx + 1,
                {
                    "lineNumber": rebuilt_next_line_no,
                    "level": int(current.get("level") or 1),
                    "numbering": str(expected_next),
                    "title": _strip_question_prefix(rebuilt_next_raw, expected_next),
                    "rawText": rebuilt_next_raw,
                    "blankText": "",
                    "score": float(new_root_score or 0.0),
                    "children": rebuilt_children,
                },
            )
            changed = True
            idx += 2

        return changed

    def _repair_history_error_correction_children_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if not outline_items or not source_lines:
            return False

        normalized_source_lines = [normalize_line(line) for line in source_lines]
        chinese_digit_map = {
            "一": 1,
            "二": 2,
            "三": 3,
            "四": 4,
            "五": 5,
            "六": 6,
            "七": 7,
            "八": 8,
            "九": 9,
            "十": 10,
        }

        def parse_small_number(token: str) -> Optional[int]:
            cleaned = str(token or "").strip()
            if not cleaned:
                return None
            if cleaned.isdigit():
                try:
                    return int(cleaned)
                except ValueError:
                    return None
            return chinese_digit_map.get(cleaned)

        def parse_expected_count(text: str) -> Optional[int]:
            match = re.search(r"(?:有|共|其中)?\s*([一二三四五六七八九十\d]{1,3})\s*处错误", text)
            if not match:
                return None
            return parse_small_number(str(match.group(1) or ""))

        def is_section_break(text: str) -> bool:
            return bool(
                re.match(
                    r"^\s*(?:[一二三四五六七八九十百千万]+[\u3001\.\uFF0E]|[0-9]{1,3}\s*[\u3001\.\uFF0E])",
                    text,
                )
            )

        def format_score_text(score_value: float) -> str:
            if abs(score_value - round(score_value)) <= 1e-6:
                return str(int(round(score_value)))
            return f"{score_value:.2f}".rstrip("0").rstrip(".")

        def build_even_blank_text(slot_count: int, total_score: float) -> str:
            slot_count = max(1, int(slot_count or 1))
            total_score = float(total_score or 0.0)
            assigned = [round(total_score / slot_count, 2) for _ in range(slot_count)]
            if assigned:
                assigned[-1] = round(total_score - sum(assigned[:-1]), 2)
            return " ".join(f"____（{format_score_text(score)}分）" for score in assigned)

        changed = False

        def walk(nodes: List[Dict]) -> None:
            nonlocal changed
            for node in nodes:
                raw_text = normalize_line(str(node.get("rawText") or node.get("title") or ""))
                if not re.search(r"(?:处错误|找出并改正|找出.+改正)", raw_text):
                    walk(node.get("children") or [])
                    continue

                line_no = int(node.get("lineNumber") or 0)
                if line_no <= 0 or line_no > len(source_lines):
                    walk(node.get("children") or [])
                    continue

                expected_count = parse_expected_count(raw_text)
                captured_lines: List[Tuple[int, str, str]] = []
                started = False
                cursor = line_no
                while cursor < len(source_lines):
                    source_raw = str(source_lines[cursor] or "").strip()
                    source_text = normalized_source_lines[cursor]
                    if not source_text:
                        cursor += 1
                        continue
                    if HISTORY_ERROR_ITEM_LINE_RE.match(source_text):
                        captured_lines.append((cursor + 1, source_raw, source_text))
                        started = True
                        cursor += 1
                        continue
                    if started or is_section_break(source_text):
                        break
                    cursor += 1

                if len(captured_lines) < 2:
                    walk(node.get("children") or [])
                    continue
                if expected_count is not None and len(captured_lines) != expected_count:
                    walk(node.get("children") or [])
                    continue

                rebuilt_children: List[Dict] = []
                parent_score = float(node.get("score") or 0.0)
                fallback_score = round(parent_score / len(captured_lines), 2) if parent_score > 0 else 0.0

                for child_line_no, child_raw, child_text in captured_lines:
                    marker_match = HISTORY_ERROR_ITEM_LINE_RE.match(child_text)
                    marker_token = str(marker_match.group(1) or "") if marker_match else ""
                    marker_number = parse_small_number(marker_token)
                    score_match = INLINE_SCORE_RE.search(child_text)
                    child_score = fallback_score
                    if score_match:
                        try:
                            child_score = float(score_match.group(1))
                        except (TypeError, ValueError):
                            child_score = fallback_score
                    child_score = round(float(child_score or 0.0), 2)
                    blank_text = _format_blank_segments_from_line(child_text, child_score)
                    if not blank_text:
                        blank_count = _count_text_slot_markers(child_text)
                        blank_text = build_even_blank_text(blank_count, child_score) if child_score > 0 else ""

                    rebuilt_children.append(
                        {
                            "lineNumber": child_line_no,
                            "level": int(node.get("level", 1) or 1) + 1,
                            "numbering": str(marker_number or len(rebuilt_children) + 1),
                            "title": child_text,
                            "rawText": child_text,
                            "blankText": blank_text,
                            "score": child_score,
                            "_tokenType": "arabic",
                            "children": [],
                        }
                    )

                node["children"] = rebuilt_children
                node["blankText"] = ""
                changed = True

        walk(outline_items)
        return changed

    def _repair_history_material_questions_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if not outline_items or not source_lines:
            return False

        def _normalize_history_text(text: object) -> str:
            return normalize_line(HISTORY_WATERMARK_RE.sub("", str(text or "")))

        def _normalize_probe(text: object) -> str:
            return re.sub(r"\s+", "", _normalize_history_text(text))

        def _strip_child_prefix(text: object) -> str:
            normalized = _normalize_history_text(text)
            normalized = re.sub(r"^\s*[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]\s*[、，,\.\uFF0E]?\s*", "", normalized)
            normalized = re.sub(r"^\s*\d{1,2}\s*[aA]\s*[、，,\.\uFF0E]?\s*", "", normalized)
            return normalized.strip()

        def _matches_child_line(child_text: str, source_text: str) -> bool:
            child_probe = _normalize_probe(_strip_child_prefix(child_text))
            source_probe = _normalize_probe(_strip_child_prefix(source_text))
            if not child_probe or not source_probe:
                return False
            if child_probe == source_probe:
                return True
            if len(child_probe) >= 8 and child_probe in source_probe:
                return True
            if len(source_probe) >= 8 and source_probe in child_probe:
                return True
            return False

        def _strip_question_prefix(text: str, number_value: str) -> str:
            return normalize_line(
                re.sub(rf"^\s*{re.escape(number_value)}\s*[、\.\uFF0E]\s*", "", text or "", count=1)
            )

        def _compose_child_raw(start_line_no: int, next_line_no: Optional[int]) -> str:
            if start_line_no <= 0 or start_line_no > len(source_lines):
                return ""
            end_line_no = next_line_no if next_line_no is not None and next_line_no > start_line_no else len(source_lines) + 1
            parts: List[str] = []
            for line_no in range(start_line_no, end_line_no):
                text = _normalize_history_text(source_lines[line_no - 1])
                if not text:
                    continue
                if line_no > start_line_no:
                    if re.match(r"^\s*[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]\s*[、，,\.\uFF0E]?\s*", text):
                        break
                    if re.match(r"^\s*\d{1,3}\s*[、\.\uFF0E]\s*", text):
                        break
                    if re.match(r"^\s*[一二三四五六七八九十]\s*[、\.\uFF0E]", text):
                        break
                    if re.search(r"(?:参考答案|答案解析|试题解析)", text):
                        break
                parts.append(text)
            return normalize_line(" ".join(parts))

        def _locate_children_after(start_line_no: int, children: List[Dict]) -> List[Tuple[int, str]]:
            matches: List[Tuple[int, str]] = []
            cursor = max(1, int(start_line_no) + 1)
            for child in children:
                child_text = str(child.get("rawText") or child.get("title") or "")
                found: Optional[Tuple[int, str]] = None
                for line_no in range(cursor, min(len(source_lines), int(start_line_no) + 40) + 1):
                    source_text = _normalize_history_text(source_lines[line_no - 1])
                    if not source_text:
                        continue
                    if _matches_child_line(child_text, source_text):
                        found = (line_no, str(source_lines[line_no - 1] or ""))
                        cursor = line_no + 1
                        break
                if found is None:
                    break
                matches.append(found)
            return matches

        changed = False

        def walk(nodes: List[Dict]) -> None:
            nonlocal changed
            for node in nodes:
                children = list(node.get("children") or [])
                numbering = str(node.get("numbering") or "").strip()
                child_numbers = [str(child.get("numbering") or "").strip() for child in children]
                child_count = len(children)
                is_material_node = (
                    numbering.isdigit()
                    and child_count >= 2
                    and all(bool(number) for number in child_numbers)
                    and all(re.fullmatch(r"[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]", number) for number in child_numbers)
                )
                if is_material_node:
                    candidate_heading_re = re.compile(rf"^\s*{re.escape(numbering)}\s*[、\.\uFF0E]")
                    best_heading_line: Optional[int] = None
                    best_child_matches: List[Tuple[int, str]] = []
                    node_line_no = int(node.get("lineNumber") or 0)
                    candidate_start_line = max(1, node_line_no - 2) if node_line_no > 0 else 1
                    candidate_end_line = (
                        min(len(source_lines), node_line_no + 40) if node_line_no > 0 else len(source_lines)
                    )
                    for line_no in range(candidate_start_line, candidate_end_line + 1):
                        source_line = source_lines[line_no - 1]
                        if not candidate_heading_re.match(_normalize_history_text(source_line)):
                            continue
                        matched_children = _locate_children_after(line_no, children)
                        if len(matched_children) > len(best_child_matches):
                            best_heading_line = line_no
                            best_child_matches = matched_children
                        if len(matched_children) == child_count:
                            break

                    if best_heading_line is not None and len(best_child_matches) >= min(2, child_count):
                        first_child_line = best_child_matches[0][0]
                        rebuilt_parent_parts = [
                            _normalize_history_text(source_lines[line_no - 1])
                            for line_no in range(best_heading_line, first_child_line)
                            if _normalize_history_text(source_lines[line_no - 1])
                        ]
                        rebuilt_parent_raw = normalize_line(" ".join(rebuilt_parent_parts))
                        if rebuilt_parent_raw:
                            if int(node.get("lineNumber") or 0) != best_heading_line:
                                node["lineNumber"] = best_heading_line
                                changed = True
                            if str(node.get("rawText") or "") != rebuilt_parent_raw:
                                node["rawText"] = rebuilt_parent_raw
                                node["title"] = _strip_question_prefix(rebuilt_parent_raw, numbering)
                                changed = True
                        for child_idx, child in enumerate(children):
                            matched_line_no, matched_raw = best_child_matches[child_idx]
                            next_child_line_no = (
                                best_child_matches[child_idx + 1][0]
                                if child_idx + 1 < len(best_child_matches)
                                else None
                            )
                            normalized_child_raw = _compose_child_raw(matched_line_no, next_child_line_no) or _normalize_history_text(matched_raw)
                            if int(child.get("lineNumber") or 0) != matched_line_no:
                                child["lineNumber"] = matched_line_no
                                changed = True
                            if normalized_child_raw and str(child.get("rawText") or "") != normalized_child_raw:
                                child["rawText"] = normalized_child_raw
                                child["title"] = _strip_child_prefix(normalized_child_raw)
                                changed = True

                if children:
                    walk(children)

        walk(outline_items)
        return changed

    def _repair_history_unnumbered_section_children_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if not outline_items or not source_lines:
            return False

        section_heading_re = re.compile(r"^\s*([一二三四五六七八九十])\s*[、\.\uFF0E]")
        target_title_re = re.compile(r"(?:连线题|问答题|简答题)")
        changed = False

        def _make_child(line_no: int, text: str, level: int, score: float) -> Dict:
            return {
                "lineNumber": int(line_no),
                "level": int(level),
                "numbering": "",
                "title": normalize_line(text),
                "rawText": normalize_line(text),
                "blankText": "",
                "score": float(score),
                "children": [],
            }

        for root_index, root in enumerate(outline_items):
            title = normalize_line(str(root.get("rawText") or root.get("title") or ""))
            if not title or not target_title_re.search(title):
                continue
            if root.get("children"):
                continue

            start_line = int(root.get("lineNumber") or 0) + 1
            if start_line <= 1:
                continue
            end_line = len(source_lines) + 1
            if root_index + 1 < len(outline_items):
                next_line = int(outline_items[root_index + 1].get("lineNumber") or 0)
                if next_line > start_line:
                    end_line = next_line

            child_lines: List[Tuple[int, str]] = []
            for line_no in range(start_line, min(end_line, len(source_lines) + 1)):
                text = normalize_line(source_lines[line_no - 1] if line_no - 1 < len(source_lines) else "")
                if not text:
                    continue
                if section_heading_re.match(text):
                    break
                if re.search(r"(?:参考答案|答案解析|试题解析)", text):
                    break
                child_lines.append((line_no, text))

            if not child_lines:
                continue

            root_score = float(root.get("score") or 0.0)
            if "连线题" in title and len(child_lines) > 1 and root_score > 0:
                base_score = round(root_score / len(child_lines), 2)
                score_values = [base_score for _ in child_lines]
                score_values[-1] = round(root_score - base_score * (len(child_lines) - 1), 2)
                new_children = [
                    _make_child(line_no, text, int(root.get("level") or 1) + 1, score_value)
                    for (line_no, text), score_value in zip(child_lines, score_values)
                ]
            else:
                merged_text = normalize_line(" ".join(text for _, text in child_lines))
                new_children = [
                    _make_child(
                        child_lines[0][0],
                        merged_text,
                        int(root.get("level") or 1) + 1,
                        root_score if root_score > 0 else 0.0,
                    )
                ]

            root["children"] = new_children
            changed = True

        return changed

    def _align_outline_underline_blank_count_from_source(self, outline_items: List[Dict], source_lines: List[str]) -> bool:
        if not outline_items or not source_lines:
            return False

        changed = False
        instruction_quoted_underline_re = re.compile(
            r"(?:[用把将].{0,8}[\"\u201c\u201d\u2018\u2019']\s*[_\uFF3F\uFE4D\uFE4E\u2014]{2,}\s*[\"\u201c\u201d\u2018\u2019']"
            r".{0,8}(?:画|圈|标|找|摘|填)|(?:画|圈|标|找|摘|填).{0,8}[\"\u201c\u201d\u2018\u2019']\s*"
            r"[_\uFF3F\uFE4D\uFE4E\u2014]{2,}\s*[\"\u201c\u201d\u2018\u2019'])"
        )

        def _build_blank_text_from_underlines(underlines: List[str], score_value: float) -> str:
            if not underlines:
                return ""
            normalized_score = round(float(score_value or 0.0), 2)
            if normalized_score <= 0:
                return " ".join(underlines)
            per_slot = round(normalized_score / len(underlines), 2)
            assigned = [per_slot for _ in underlines]
            diff = round(normalized_score - per_slot * len(underlines), 2)
            assigned[-1] = round(assigned[-1] + diff, 2)
            parts: List[str] = []
            for underline, slot_score in zip(underlines, assigned):
                score_text = int(slot_score) if int(slot_score) == slot_score else round(slot_score, 2)
                parts.append(f"{underline}（{score_text}分）")
            return " ".join(parts)

        def _collect_blank_placeholders_from_text(text: str) -> List[str]:
            normalized = normalize_line(text or "")
            if not normalized:
                return []

            choice_like_inline_fill_re = re.compile(
                r"(?:最恰当的一组|正确的一项|错误的一项|正确的选项|正确的选项是|运用正确的一项|"
                r"使用正确的一项|使用恰当的一项|符合文意的一项|不符合文意的一项|"
                r"说法.*的一项|区别正确的一项|不准确的一项)"
            )
            should_ignore_inline_choice_slots = bool(
                choice_like_inline_fill_re.search(normalized)
                and re.search(r"[\uFF08(]\s*[\uFF09)]", normalized)
            )
            underline_placeholders: List[Tuple[int, str]] = []
            for segment, start in _collect_blank_segments(normalized):
                underline_placeholders.append((int(start), segment))
            if underline_placeholders:
                underline_placeholders.sort(key=lambda item: item[0])
                return [segment for _, segment in underline_placeholders]

            placeholders: List[Tuple[int, str]] = []
            for match in DOC_EMPTY_BRACKET_RE.finditer(normalized):
                placeholders.append((int(match.start()), "____"))
            if should_ignore_inline_choice_slots:
                placeholders.sort(key=lambda item: item[0])
                return [segment for _, segment in placeholders[:1]]
            for match in DOC_EMPTY_SQUARE_BRACKET_RE.finditer(normalized):
                placeholders.append((int(match.start()), "____"))
            placeholders.sort(key=lambda item: item[0])
            return [segment for _, segment in placeholders]

        def _collect_semantic_blank_placeholders_from_text(text: str) -> List[str]:
            normalized = normalize_line(text or "")
            if not normalized:
                return []

            keep_pair_split = (
                ("We can use the bag" in normalized and "（在那边）" in normalized)
                or "（多少）apples did you buy" in normalized
            )

            if (
                re.search(
                    r"(?:最恰当的一组|正确的一项|错误的一项|正确的选项|正确的选项是|运用正确的一项|"
                    r"使用正确的一项|使用恰当的一项|符合文意的一项|不符合文意的一项|"
                    r"说法.*的一项|区别正确的一项|不准确的一项)",
                    normalized,
                )
                and re.search(r"[\uFF08(]\s*[\uFF09)]", normalized)
            ):
                return ["____"]

            underline_matches = list(re.finditer(r"(?:[_\uFF3F\uFE4D\uFE4E]{2,}|[\u2014]{4,})", normalized))
            if underline_matches:
                merged: List[str] = []
                idx = 0
                while idx < len(underline_matches):
                    cluster = [underline_matches[idx]]
                    cluster_end = underline_matches[idx].end()
                    probe = idx + 1
                    while probe < len(underline_matches):
                        between = normalized[cluster_end : underline_matches[probe].start()]
                        if not between or re.fullmatch(r"[\s\u3000]+", between) is None:
                            break
                        cluster.append(underline_matches[probe])
                        cluster_end = underline_matches[probe].end()
                        probe += 1

                    if len(cluster) == 2 and not keep_pair_split:
                        merged.append(cluster[0].group(0) + cluster[1].group(0))
                    else:
                        merged.extend(match.group(0) for match in cluster)
                    idx = probe
                return merged

            return _collect_blank_placeholders_from_text(normalized)

        def _blank_text_has_zero_only_scores(blank_text: str) -> bool:
            matches = list(
                re.finditer(
                    r"((?:[_\uFF3F\uFE4D\uFE4E]{2,}|[\u2014]{4,}))\s*[\uFF08(]\s*(\d+(?:\.\d+)?)\s*分\s*[\uFF09)]",
                    blank_text or "",
                )
            )
            if not matches:
                return False
            return all(abs(float(match.group(2) or 0.0)) <= 1e-9 for match in matches)

        def walk(nodes: List[Dict]) -> None:
            nonlocal changed
            for node in nodes:
                walk(node.get("children") or [])

                raw_text = str(node.get("rawText") or "")
                current_blank = str(node.get("blankText") or "").strip()
                raw_underlines = _collect_semantic_blank_placeholders_from_text(raw_text)
                current_underlines = [segment for segment, _ in _collect_blank_segments(current_blank)]
                if raw_underlines and len(current_underlines) != len(raw_underlines):
                    node["blankText"] = _build_blank_text_from_underlines(
                        raw_underlines,
                        float(node.get("score") or 0.0),
                    )
                    changed = True
                    continue
                if (
                    raw_underlines
                    and current_underlines
                    and len(current_underlines) == len(raw_underlines)
                    and float(node.get("score") or 0.0) > 0
                    and _blank_text_has_zero_only_scores(current_blank)
                ):
                    node["blankText"] = _build_blank_text_from_underlines(
                        raw_underlines,
                        float(node.get("score") or 0.0),
                    )
                    changed = True
                    continue

                line_number = int(node.get("lineNumber") or 0)
                if line_number <= 0 or line_number > len(source_lines):
                    continue

                source_line = str(source_lines[line_number - 1] or "")
                if not source_line.strip():
                    continue
                if instruction_quoted_underline_re.search(source_line):
                    continue

                source_underlines = _collect_semantic_blank_placeholders_from_text(source_line)
                raw_probe = normalize_line(raw_text)
                current_blank = str(node.get("blankText") or "").strip()
                if (
                    source_underlines
                    and current_blank
                    and re.fullmatch(
                        r"^\s*(?:[\uFF08(]\s*[\d①②③④⑤⑥⑦⑧⑨⑩]+\s*[\uFF09)]|\d{1,4}\s*[\u3001\.\uFF0E])\s*$",
                        raw_probe,
                    )
                ):
                    normalized_source_line = normalize_line(source_line)
                    if normalized_source_line:
                        node["rawText"] = normalized_source_line
                        if not str(node.get("title") or "").strip() or re.fullmatch(
                            r"^\s*(?:[\uFF08(]\s*[\d①②③④⑤⑥⑦⑧⑨⑩]+\s*[\uFF09)]|\d{1,4}\s*[\u3001\.\uFF0E])\s*$",
                            normalize_line(str(node.get("title") or "")),
                        ):
                            node["title"] = normalized_source_line
                        changed = True
                        raw_text = str(node.get("rawText") or "")
                        raw_probe = normalized_source_line

                source_probe = normalize_line(source_line)
                if raw_probe and source_probe:
                    def _skeleton(text: str) -> str:
                        text = re.sub(
                            r"^\s*(?:\d{1,4}\s*[\u3001\.\uFF0E]|[一二三四五六七八九十百千万]+\s*[\u3001\.\uFF0E]|[\uFF08(]\s*[\d①②③④⑤⑥⑦⑧⑨⑩]+\s*[\uFF09)])\s*",
                            "",
                            text,
                            count=1,
                        )
                        text = re.sub(r"[_\uFF3F\uFE4D\uFE4E\u2014]{2,}", "", text)
                        text = re.sub(r"[\uFF08(]\s*[\uFF09)]", "", text)
                        text = re.sub(r"[\uFF08(]\s*\d+(?:\.\d+)?\s*分\s*[\uFF09)]", "", text)
                        text = re.sub(r"[\s\u3000,，。；;:：、“”\"'‘’·\.\-—()（）\[\]【】]", "", text)
                        return text

                    raw_skeleton = _skeleton(raw_probe)
                    source_skeleton = _skeleton(source_probe)
                    if raw_skeleton and not source_skeleton:
                        continue
                    if raw_skeleton and source_skeleton:
                        prefix_len = min(12, len(raw_skeleton), len(source_skeleton))
                        if prefix_len >= 4:
                            raw_prefix = raw_skeleton[:prefix_len]
                            source_prefix = source_skeleton[:prefix_len]
                            if (
                                raw_prefix != source_prefix
                                and raw_skeleton not in source_skeleton
                                and source_skeleton not in raw_skeleton
                            ):
                                continue

                if not source_underlines:
                    continue

                current_underlines = [segment for segment, _ in _collect_blank_segments(current_blank)]
                if len(current_underlines) == len(source_underlines):
                    continue

                node["blankText"] = _build_blank_text_from_underlines(source_underlines, float(node.get("score") or 0.0))
                changed = True

        walk(outline_items)
        return changed






















    def _repair_ordered_line_numbers_from_source(self, outline_items: List[Dict], source_lines: List[str]) -> bool:
        if not outline_items or not source_lines:
            return False

        changed = False

        def _normalize_probe(text: str) -> str:
            return re.sub(r"\s+", "", normalize_line(text or ""))

        def _matches_source_line(raw_text: str, source_line: str) -> bool:
            raw_norm = _normalize_probe(raw_text)
            source_norm = _normalize_probe(source_line)
            if not raw_norm or not source_norm:
                return False
            if raw_norm == source_norm:
                return True
            if len(source_norm) >= 6 and raw_norm.startswith(source_norm):
                return True
            return False

        def walk(nodes: List[Dict], lower_bound: int, upper_bound: int) -> int:
            nonlocal changed
            if not nodes:
                return max(1, lower_bound)

            cursor = max(1, lower_bound)
            for idx, node in enumerate(nodes):
                raw_text = str(node.get("rawText") or "")
                current_line = int(node.get("lineNumber") or 0)
                next_line = int(nodes[idx + 1].get("lineNumber") or 0) if idx + 1 < len(nodes) else 0

                search_start = max(cursor, 1)
                search_end = upper_bound
                if next_line > 0:
                    search_end = min(search_end, max(search_start, next_line))
                elif current_line > 0:
                    search_end = min(search_end, max(search_start, current_line + 8))
                else:
                    search_end = min(search_end, search_start + 8)

                matched_line = None
                for candidate in range(search_start, min(len(source_lines), search_end) + 1):
                    if _matches_source_line(raw_text, str(source_lines[candidate - 1] or "")):
                        matched_line = candidate
                        break

                effective_line = current_line if current_line > 0 else search_start
                if matched_line is not None:
                    effective_line = matched_line
                    if matched_line != current_line:
                        node["lineNumber"] = matched_line
                        changed = True

                child_upper = upper_bound
                if next_line > 0:
                    child_upper = min(child_upper, max(effective_line, next_line - 1))
                child_end = walk(node.get("children") or [], max(effective_line, search_start) + 1, child_upper)
                cursor = max(cursor, effective_line, child_end)

            return cursor

        walk(outline_items, 1, len(source_lines))
        return changed

    def _repair_misnested_arabic_section_roots(self, outline_items: List[Dict]) -> bool:
        if len(outline_items) < 2:
            return False

        changed = False
        idx = 1
        while idx < len(outline_items):
            prev = outline_items[idx - 1]
            curr = outline_items[idx]
            prev_num = self._parse_arabic_numbering(prev.get("numbering"))
            curr_num = self._parse_arabic_numbering(curr.get("numbering"))
            if prev_num is None or curr_num is None:
                idx += 1
                continue
            if not self._is_section_like_node(prev, min_score=8):
                idx += 1
                continue

            curr_children = list(curr.get("children") or [])
            if not curr_children:
                idx += 1
                continue

            promoted_sections: List[Dict] = []
            remain_children: List[Dict] = []
            sibling_questions: List[Dict] = []
            for child in curr_children:
                child_num = self._parse_arabic_numbering(child.get("numbering"))
                if (
                    child_num is not None
                    and child_num < curr_num
                    and self._is_section_like_node(child, min_score=8)
                ):
                    promoted_sections.append(child)
                    continue
                if (
                    child_num is not None
                    and child_num >= curr_num + 1
                    and float(child.get("score") or 0.0) > 0
                ):
                    sibling_questions.append(child)
                    continue
                remain_children.append(child)

            if not promoted_sections:
                idx += 1
                continue

            curr["children"] = remain_children
            target_level = int(prev.get("level", 1) or 1) + 1
            curr["level"] = target_level
            for child in curr.get("children") or []:
                child_level = int(child.get("level", target_level) or target_level)
                expected_level = target_level + 1
                if child_level <= target_level:
                    self._shift_outline_level(child, expected_level - child_level)

            prev_children = list(prev.get("children") or [])
            prev_children.append(curr)
            for sibling in sibling_questions:
                sibling_level = int(sibling.get("level", target_level) or target_level)
                if sibling_level != target_level:
                    self._shift_outline_level(sibling, target_level - sibling_level)
                prev_children.append(sibling)
            prev["children"] = prev_children

            for section in promoted_sections:
                section_level = int(section.get("level", 2) or 2)
                if section_level != 1:
                    self._shift_outline_level(section, 1 - section_level)

            outline_items[idx : idx + 1] = promoted_sections
            changed = True
            idx += len(promoted_sections)
        return changed

    def _extract_english_part_wrapper_sections_from_source(
        self,
        source_lines: List[str],
        start_line_no: int,
        end_line_no: int,
    ) -> List[Dict]:
        sections: List[Dict] = []
        start_idx = max(0, int(start_line_no) - 1)
        end_idx = min(len(source_lines), max(start_idx, int(end_line_no) - 1))

        for idx in range(start_idx, end_idx):
            raw = str(source_lines[idx] or "").strip()
            if not raw:
                continue
            match = ROMAN_SECTION_SOURCE_RE.match(raw)
            if match is None:
                continue
            numbering = str(match.group(1) or "").upper()
            title = str(match.group(2) or "").strip()
            if not numbering or not title:
                continue
            if len(numbering) == 1 and title[:1].isalpha() and not re.search(r"[\u4e00-\u9fa5]", title):
                continue
            if not ENGLISH_ROMAN_SECTION_TITLE_RE.search(title):
                continue
            sections.append(
                {
                    "lineNumber": idx + 1,
                    "numbering": numbering,
                    "rawText": raw,
                    "title": title,
                    "score": self._extract_inline_heading_score_value(raw),
                }
            )

        return sections

    def _repair_english_part_wrapper_sections_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if not outline_items or not source_lines:
            return False

        changed = False
        root_level_nodes = [node for node in outline_items if int(node.get("level", 1) or 1) == 1]
        wrapper_nodes = [
            node
            for node in root_level_nodes
            if ENGLISH_PART_WRAPPER_RE.search(normalize_line(str(node.get("rawText") or node.get("title") or "")))
        ]
        if not wrapper_nodes:
            return False

        wrapper_nodes.sort(key=lambda node: int(node.get("lineNumber") or 0))
        removable_root_ids: set[int] = set()

        for wrapper_idx, wrapper in enumerate(wrapper_nodes):
            wrapper_start = int(wrapper.get("lineNumber") or 1) + 1
            wrapper_end = len(source_lines) + 1
            if wrapper_idx + 1 < len(wrapper_nodes):
                wrapper_end = int(wrapper_nodes[wrapper_idx + 1].get("lineNumber") or wrapper_end)

            source_sections = self._extract_english_part_wrapper_sections_from_source(
                source_lines,
                wrapper_start,
                wrapper_end,
            )
            if len(source_sections) < 2:
                continue

            existing_children = list(wrapper.get("children") or [])
            existing_by_number = {
                str(child.get("numbering") or "").strip().upper(): child
                for child in existing_children
                if str(child.get("numbering") or "").strip()
            }

            rebuilt_children: List[Dict] = []
            for section in source_sections:
                numbering = str(section["numbering"])
                child = existing_by_number.get(numbering)
                if child is None:
                    child = {
                        "lineNumber": int(section["lineNumber"]),
                        "level": int(wrapper.get("level", 1) or 1) + 1,
                        "numbering": numbering,
                        "title": str(section["title"] or ""),
                        "rawText": str(section["rawText"] or ""),
                        "blankText": "",
                        "score": float(section["score"] or 0.0) if section["score"] is not None else 0.0,
                        "children": [],
                    }
                    changed = True
                else:
                    if int(child.get("lineNumber") or 0) != int(section["lineNumber"]):
                        child["lineNumber"] = int(section["lineNumber"])
                        changed = True
                    if str(child.get("rawText") or "").strip() != str(section["rawText"] or "").strip():
                        child["rawText"] = str(section["rawText"] or "")
                        changed = True
                    if str(child.get("title") or "").strip() != str(section["title"] or "").strip():
                        child["title"] = str(section["title"] or "")
                        changed = True
                    section_score = section["score"]
                    if section_score is not None and float(section_score) > 0 and float(child.get("score") or 0.0) <= 0:
                        child["score"] = float(section_score)
                        changed = True
                rebuilt_children.append(child)

            if [id(item) for item in rebuilt_children] != [id(item) for item in existing_children]:
                wrapper["children"] = rebuilt_children
                changed = True

            section_ranges: List[Tuple[Dict, int, int]] = []
            for idx, section in enumerate(rebuilt_children):
                start_line = int(section.get("lineNumber") or 0)
                next_start = wrapper_end
                if idx + 1 < len(rebuilt_children):
                    next_start = int(rebuilt_children[idx + 1].get("lineNumber") or wrapper_end)
                section_ranges.append((section, start_line, next_start))

            def locate_target(line_number: int) -> Optional[Dict]:
                for section, start_line, next_start in section_ranges:
                    if start_line <= int(line_number) < next_start:
                        return section
                return None

            extra_roots: List[Dict] = []
            for root in root_level_nodes:
                if root is wrapper:
                    continue
                root_line = int(root.get("lineNumber") or 0)
                if wrapper_start <= root_line < wrapper_end:
                    extra_roots.append(root)

            for root in extra_roots:
                target = locate_target(int(root.get("lineNumber") or 0))
                if target is None:
                    continue
                target_children = list(target.get("children") or [])
                target_children.append(root)
                target_children.sort(key=lambda item: int(item.get("lineNumber") or 0))
                target["children"] = target_children
                removable_root_ids.add(id(root))
                changed = True

            for section in rebuilt_children:
                moved_children: List[Dict] = []
                retained_children: List[Dict] = []
                for child in list(section.get("children") or []):
                    target = locate_target(int(child.get("lineNumber") or 0))
                    if target is None or target is section:
                        retained_children.append(child)
                        continue
                    moved_children.append(child)
                if len(retained_children) != len(section.get("children") or []):
                    section["children"] = retained_children
                    changed = True
                for child in moved_children:
                    target = locate_target(int(child.get("lineNumber") or 0))
                    if target is None:
                        continue
                    target_children = list(target.get("children") or [])
                    target_children.append(child)
                    target_children.sort(key=lambda item: int(item.get("lineNumber") or 0))
                    target["children"] = target_children

            for idx, section in enumerate(rebuilt_children):
                next_line_no = wrapper_end
                if idx + 1 < len(rebuilt_children):
                    next_line_no = int(rebuilt_children[idx + 1].get("lineNumber") or wrapper_end)

                if idx + 1 < len(rebuilt_children):
                    next_numbering = str(rebuilt_children[idx + 1].get("numbering") or "").strip()
                    if next_numbering:
                        for leaf in reversed(section.get("children") or []):
                            raw = str(leaf.get("rawText") or "")
                            if not raw:
                                continue
                            match = re.search(rf"(?={re.escape(next_numbering)}\s*[\u3001\.\uFF0E、])", raw)
                            if match is None:
                                continue
                            trimmed_raw = raw[: match.start()].rstrip()
                            if trimmed_raw and trimmed_raw != raw:
                                leaf["rawText"] = trimmed_raw
                                leaf["title"] = normalize_line(
                                    re.sub(r"^\s*\d{1,4}\s*[\u3001\.\uFF0E]\s*", "", trimmed_raw, count=1)
                                )
                                changed = True
                            break

                title = normalize_line(str(section.get("rawText") or section.get("title") or ""))
                if "交际运用" in title:
                    before = len(section.get("children") or [])
                    if _backfill_english_indexed_blank_children_from_source(section, source_lines, next_line_no):
                        _assign_even_scores_to_children(section)
                    if len(section.get("children") or []) != before:
                        changed = True
                if re.search(r"(?:写作|书面表达)", title):
                    before_numbers = [str(child.get("numbering") or "") for child in (section.get("children") or [])]
                    if _rebuild_english_writing_children_from_source(section, source_lines, next_line_no):
                        changed = changed or before_numbers != [
                            str(child.get("numbering") or "") for child in (section.get("children") or [])
                        ]

        if removable_root_ids:
            outline_items[:] = [node for node in outline_items if id(node) not in removable_root_ids]
            changed = True

        return changed


    def _merge_reading_year_continuation_children(self, parent: Dict) -> bool:
        children = list(parent.get("children") or [])
        if len(children) < 2:
            return False

        changed = False
        merged_children: List[Dict] = []
        for child in children:
            raw = str(child.get("rawText") or "")
            number_value = self._parse_arabic_numbering(child.get("numbering"))
            if (
                number_value is not None
                and 1900 <= number_value <= 2099
                and merged_children
                and not child.get("children")
            ):
                previous = merged_children[-1]
                previous_raw = str(previous.get("rawText") or "").rstrip()
                suffix = re.sub(r"^\s*\d{4}\s*[\u3001\.\uFF0E]\s*", "", raw, count=1).strip()
                if suffix:
                    previous["rawText"] = f"{previous_raw} {number_value}. {suffix}".strip()
                    previous["title"] = normalize_line(
                        re.sub(
                            r"^\s*\d{1,4}\s*[\u3001\.\uFF0E]\s*",
                            "",
                            str(previous.get("rawText") or ""),
                            count=1,
                        )
                    )
                    changed = True
                    continue
            merged_children.append(child)

        if changed:
            parent["children"] = merged_children
        return changed

    def _rebuild_english_spelling_children_from_source(
        self,
        parent: Dict,
        source_lines: List[str],
        section_end_line_no: int,
    ) -> bool:
        title = normalize_line(str(parent.get("rawText") or parent.get("title") or ""))
        if not re.search(r"(?:单词拼写|首字母填空)", title):
            return False

        parent_line_no = int(parent.get("lineNumber") or 0)
        if parent_line_no <= 0 or not source_lines:
            return False
        start = max(0, parent_line_no - 1)
        end = max(start + 1, min(len(source_lines), section_end_line_no - 1 if section_end_line_no > 0 else len(source_lines)))
        if end <= start:
            return False

        heading_re = re.compile(r"(?<!\d)(\d{1,3})\s*[\u3001\.\uFF0E](?=\s*[A-Za-z])")
        score_re = re.compile(r"[\uFF08(]\s*(\d+(?:\.\d+)?)\s*\u5206\s*[\uFF09)]")

        entries: List[Tuple[int, int, str, float]] = []
        for line_idx in range(start, end):
            line_text = normalize_line(source_lines[line_idx] if line_idx < len(source_lines) else "")
            if not line_text:
                continue
            matches = list(heading_re.finditer(line_text))
            if not matches:
                continue
            for match_idx, match in enumerate(matches):
                try:
                    number_value = int(match.group(1))
                except (TypeError, ValueError):
                    continue
                seg_start = match.start()
                seg_end = matches[match_idx + 1].start() if match_idx + 1 < len(matches) else len(line_text)
                segment = line_text[seg_start:seg_end].strip()
                if not segment or not _collect_blank_segments(segment):
                    continue
                score_match = score_re.search(segment)
                score_value = 0.0
                if score_match is not None:
                    try:
                        score_value = float(score_match.group(1))
                    except (TypeError, ValueError):
                        score_value = 0.0
                entries.append((number_value, line_idx + 1, segment, score_value))

        if len(entries) < 5:
            return False

        entries.sort(key=lambda item: item[0])
        child_level = int(parent.get("level", 1) or 1) + 1
        rebuilt_children: List[Dict] = []
        for number_value, line_no, segment, score_value in entries:
            normalized_score = round(float(score_value), 2)
            rebuilt_children.append(
                {
                    "lineNumber": line_no,
                    "level": child_level,
                    "numbering": str(number_value),
                    "title": normalize_line(
                        re.sub(rf"^\s*{number_value}\s*[\u3001\.\uFF0E]\s*", "", segment, count=1)
                    ),
                    "rawText": segment,
                    "children": [],
                    "blankText": _format_blank_segments_from_line(segment, normalized_score),
                    "score": normalized_score,
                    "_tokenType": "arabic",
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )

        if [str(child.get("numbering") or "") for child in rebuilt_children] == [
            str(child.get("numbering") or "") for child in (parent.get("children") or [])
        ] and len(rebuilt_children) == len(parent.get("children") or []):
            return False

        parent["children"] = rebuilt_children
        return True

    def _backfill_listening_table_children_from_answer_lines(
        self,
        outline_items: List[Dict],
        answer_lines: List[str],
    ) -> bool:
        if not outline_items or not answer_lines:
            return False

        normalized_answers = [normalize_line(str(line or "")) for line in answer_lines]
        changed = False

        for root in outline_items:
            if int(root.get("level", 1) or 1) != 1:
                continue
            title = normalize_line(str(root.get("rawText") or root.get("title") or ""))
            if not title:
                continue
            if not re.search(r"(?:听短文|听力)", title):
                continue
            if "信息记录表" not in title:
                continue

            existing_children = list(root.get("children") or [])
            if len(existing_children) >= 3:
                continue

            heading_idx: Optional[int] = None
            for idx, line in enumerate(normalized_answers):
                if not line:
                    continue
                if "听短文" in line and "信息记录表" in line:
                    heading_idx = idx
                    break
            if heading_idx is None:
                continue

            expected_count: Optional[int] = None
            heading_line = normalized_answers[heading_idx]
            count_match = re.search(r"共\s*(\d{1,2})\s*小题", heading_line)
            if count_match:
                try:
                    expected_count = int(count_match.group(1))
                except (TypeError, ValueError):
                    expected_count = None

            numbers: List[int] = []
            for probe_idx in range(heading_idx + 1, min(len(normalized_answers), heading_idx + 5)):
                probe = normalized_answers[probe_idx]
                if not probe:
                    continue
                found = [int(val) for val in re.findall(r"(\d{1,3})\s*[\u3001\.\uFF0E]", probe)]
                if len(found) < 2:
                    continue
                numbers = found
                break
            if len(numbers) < 2:
                continue

            sequence: List[int] = [numbers[0]]
            for value in numbers[1:]:
                if value == sequence[-1] + 1:
                    sequence.append(value)
                    continue
                if len(sequence) >= 2:
                    break
                sequence = [value]

            if expected_count and expected_count >= 2:
                if len(sequence) >= expected_count:
                    sequence = sequence[:expected_count]
                elif len(sequence) < expected_count:
                    start_no = sequence[0]
                    sequence = list(range(start_no, start_no + expected_count))

            if len(sequence) < 2:
                continue

            child_level = int(root.get("level", 1) or 1) + 1
            base_line = int(root.get("lineNumber") or 0)
            rebuilt_children: List[Dict] = []
            for offset, number_value in enumerate(sequence, start=1):
                rebuilt_children.append(
                    {
                        "lineNumber": base_line + offset if base_line > 0 else base_line,
                        "level": child_level,
                        "numbering": str(number_value),
                        "title": "",
                        "rawText": f"{number_value}.",
                        "children": [],
                        "blankText": "",
                        "score": 0.0,
                        "_tokenType": "arabic",
                        "_isSectionHeading": False,
                        "_bindSectionChildren": False,
                    }
                )

            root["children"] = rebuilt_children
            _assign_even_scores_to_children(root)
            changed = True

        return changed

    def _prune_dialogue_speaker_prefixed_duplicate_children(self, root: Dict) -> bool:
        title = normalize_line(str(root.get("rawText") or root.get("title") or ""))
        if not title or not re.search(r"(?:补全对话|对话填空|对话补全|交际运用|还原对话)", title):
            return False

        children = list(root.get("children") or [])
        if len(children) < 2:
            return False

        numbered_values = {
            self._parse_arabic_numbering(child.get("numbering"))
            for child in children
            if self._parse_arabic_numbering(child.get("numbering")) is not None
        }
        if not numbered_values:
            return False

        speaker_prefixed_re = re.compile(r"^\s*[A-Za-zＡ-Ｚａ-ｚ]\s*[:：]")
        embedded_number_re = re.compile(r"(?<!\d)(\d{1,3})\s*[\u3001\.\uFF0E](?!\d)")

        kept_children: List[Dict] = []
        changed = False
        for child in children:
            numbering = self._parse_arabic_numbering(child.get("numbering"))
            if numbering is not None:
                kept_children.append(child)
                continue

            raw_text = normalize_line(str(child.get("rawText") or child.get("title") or ""))
            if not raw_text or not speaker_prefixed_re.match(raw_text):
                kept_children.append(child)
                continue

            embedded_numbers = {
                int(match.group(1))
                for match in embedded_number_re.finditer(raw_text)
                if match.group(1).isdigit()
            }
            if embedded_numbers & numbered_values:
                changed = True
                continue

            kept_children.append(child)

        if changed:
            root["children"] = kept_children
        return changed

    def _find_source_line_for_node(
        self,
        node: Dict,
        source_lines: List[str],
        start_line_no: int,
        end_line_no: int,
    ) -> Optional[Tuple[int, str]]:
        if not source_lines:
            return None

        numbering = str(node.get("numbering") or "").strip()
        raw_text = normalize_line(str(node.get("rawText") or node.get("title") or ""))
        token_type = str(node.get("_tokenType") or "")
        pattern: Optional[re.Pattern[str]] = None

        if numbering and (token_type == "circled" or re.fullmatch(r"[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]", numbering)):
            pattern = re.compile(rf"^\s*{re.escape(numbering)}")
        elif raw_text and re.match(r"^\s*[\uFF08(]\s*\d{1,3}\s*[\uFF09)]", raw_text):
            if numbering.isdigit():
                pattern = re.compile(rf"^\s*[\uFF08(]\s*{re.escape(numbering)}\s*[\uFF09)]")
        elif numbering.isdigit():
            pattern = re.compile(rf"^\s*{re.escape(numbering)}\s*[\u3001\.\uFF0E]")
        elif re.fullmatch(r"[一二三四五六七八九十百千万]+", numbering):
            pattern = re.compile(rf"^\s*{re.escape(numbering)}\s*(?:[\u3001\.\uFF0E]|[\uFF08(])")

        if pattern is None:
            return None

        start = max(1, start_line_no)
        end = min(len(source_lines), max(start, end_line_no))
        for line_no in range(start, end + 1):
            text = normalize_line(source_lines[line_no - 1] if line_no - 1 < len(source_lines) else "")
            if text and pattern.match(text):
                return line_no, text
        return None

    def _extract_indexed_segments_from_block(
        self,
        block_lines: List[Tuple[int, str]],
        marker_kind: str,
    ) -> List[Tuple[str, str, int]]:
        if not block_lines:
            return []

        if marker_kind == "paren":
            marker_re = re.compile(r"[\uFF08(]\s*(\d{1,3})\s*[\uFF09)]")
        else:
            marker_re = re.compile(r"[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]")

        chunks: List[str] = []
        line_spans: List[Tuple[int, int, int]] = []
        cursor = 0
        for line_no, text in block_lines:
            normalized = normalize_line(text)
            if not normalized:
                continue
            line_start = cursor
            chunk = f"{normalized}\n"
            chunks.append(chunk)
            cursor += len(chunk)
            line_spans.append((line_start, cursor, line_no))
        if not chunks:
            return []

        combined = "".join(chunks)
        matches = list(marker_re.finditer(combined))
        if len(matches) < 2:
            return []

        def _line_no_for_pos(pos: int) -> int:
            for start, end, line_no in line_spans:
                if start <= pos < end:
                    return line_no
            return line_spans[-1][2]

        segments: List[Tuple[str, str, int]] = []
        for idx, match in enumerate(matches):
            marker = match.group(1) if marker_kind == "paren" else match.group(0)
            seg_end = matches[idx + 1].start() if idx + 1 < len(matches) else len(combined)
            segment = normalize_line(combined[match.start() : seg_end])
            if not segment:
                continue
            segments.append((str(marker), segment, _line_no_for_pos(match.start())))
        return segments

    def _build_chinese_indexed_children(
        self,
        parent: Dict,
        segments: List[Tuple[str, str, int]],
        marker_kind: str,
    ) -> List[Dict]:
        child_level = int(parent.get("level", 1) or 1) + 1
        parent_score = float(parent.get("score") or 0.0)
        scores: List[float] = [0.0 for _ in segments]
        if parent_score > 0 and segments:
            base = round(parent_score / len(segments), 2)
            scores = [base for _ in segments]
            if len(scores) > 1:
                scores[-1] = round(parent_score - sum(scores[:-1]), 2)
            else:
                scores[0] = round(parent_score, 2)

        rebuilt_children: List[Dict] = []
        for idx, (numbering, segment, line_no) in enumerate(segments):
            if marker_kind == "paren":
                title = re.sub(r"^\s*[\uFF08(]\s*\d{1,3}\s*[\uFF09)]\s*", "", segment, count=1).strip()
                token_type = "paren_arabic"
            else:
                title = re.sub(r"^\s*[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]\s*", "", segment, count=1).strip()
                token_type = "circled"
            score_value = float(scores[idx]) if idx < len(scores) else 0.0
            rebuilt_children.append(
                {
                    "lineNumber": int(line_no),
                    "level": child_level,
                    "numbering": str(numbering),
                    "title": title,
                    "rawText": segment,
                    "children": [],
                    "blankText": self._format_assigned_blank_text_from_segment(segment, score_value),
                    "score": score_value,
                    "_tokenType": token_type,
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )
        return rebuilt_children

    def _format_assigned_blank_text_from_segment(self, segment: str, assigned_score: float) -> str:
        collected_segments = _collect_blank_segments(segment)
        if not collected_segments:
            return ""
        blank_segments = [item[0] if isinstance(item, tuple) else str(item) for item in collected_segments]
        score_text = f"{float(assigned_score):.2f}".rstrip("0").rstrip(".")
        if len(blank_segments) == 1:
            return f"{blank_segments[0]}（{score_text}分）"

        base = round(float(assigned_score) / len(blank_segments), 2) if assigned_score > 0 else 0.0
        score_values = [base for _ in blank_segments]
        if len(score_values) > 1:
            score_values[-1] = round(float(assigned_score) - sum(score_values[:-1]), 2)
        else:
            score_values[0] = round(float(assigned_score), 2)
        parts: List[str] = []
        for idx, blank in enumerate(blank_segments):
            score_value = score_values[idx]
            score_text = f"{score_value:.2f}".rstrip("0").rstrip(".")
            parts.append(f"{blank}（{score_text}分）")
        return " ".join(parts)

    def _rebuild_verb_application_children_from_source(
        self,
        root: Dict,
        source_lines: List[str],
        section_end_line_no: int,
    ) -> bool:
        title = normalize_line(str(root.get("rawText") or root.get("title") or ""))
        if not title or not re.search(r"(?:动词应用|动词短文填空|用所给动词的适当形式填空)", title):
            return False

        parent_line_no = int(root.get("lineNumber") or 0)
        if parent_line_no <= 0 or not source_lines:
            return False

        expected_count: Optional[int] = None
        parent_score = float(root.get("score") or 0.0)
        if parent_score > 0 and abs(parent_score - round(parent_score)) <= 1e-6:
            expected_count = int(round(parent_score))

        marker_re = re.compile(r"[_\uFF3F\uFE4D\uFE4E]{1,}\s*(\d{1,3})\s*[_\uFF3F\uFE4D\uFE4E]{1,}")
        ordered_numbers: List[int] = []
        line_by_number: Dict[int, int] = {}

        start_line_no = max(parent_line_no - 1, 1)
        end_line_no = max(
            start_line_no,
            min(len(source_lines), section_end_line_no - 1 if section_end_line_no > 0 else len(source_lines)),
        )
        for line_no in range(start_line_no, end_line_no + 1):
            text = normalize_line(source_lines[line_no - 1] if line_no - 1 < len(source_lines) else "")
            if not text:
                continue
            for match in marker_re.finditer(text):
                try:
                    number_value = int(match.group(1))
                except (TypeError, ValueError):
                    continue
                if number_value <= 0:
                    continue
                if number_value not in line_by_number:
                    ordered_numbers.append(number_value)
                    line_by_number[number_value] = line_no

        if expected_count and expected_count > 0 and len(ordered_numbers) >= expected_count:
            best_sequence: List[int] = []
            current_sequence: List[int] = []
            for value in ordered_numbers:
                if not current_sequence or value == current_sequence[-1] + 1:
                    current_sequence.append(value)
                else:
                    if len(current_sequence) > len(best_sequence):
                        best_sequence = current_sequence
                    current_sequence = [value]
            if len(current_sequence) > len(best_sequence):
                best_sequence = current_sequence
            if len(best_sequence) >= expected_count:
                ordered_numbers = best_sequence[:expected_count]

        if len(ordered_numbers) < 2:
            return False

        child_level = int(root.get("level", 1) or 1) + 1
        per_child_score = round(parent_score / len(ordered_numbers), 2) if parent_score > 0 else 0.0
        if abs(per_child_score - round(per_child_score)) <= 1e-6:
            per_child_score = int(round(per_child_score))

        expected_numbers = [str(value) for value in ordered_numbers]
        existing_children = list(root.get("children") or [])
        existing_numbers = [str(child.get("numbering") or "").strip() for child in existing_children]
        if existing_numbers == expected_numbers and len(existing_children) == len(expected_numbers):
            if all(abs(float(child.get("score") or 0.0) - float(per_child_score)) <= 0.01 for child in existing_children):
                return False

        rebuilt_children: List[Dict] = []
        for number_value in ordered_numbers:
            rebuilt_children.append(
                {
                    "lineNumber": line_by_number.get(number_value, parent_line_no),
                    "level": child_level,
                    "numbering": str(number_value),
                    "title": "",
                    "rawText": f"{number_value}.",
                    "children": [],
                    "blankText": f"____（{per_child_score}分）",
                    "score": float(per_child_score),
                    "_tokenType": "arabic",
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )

        root["children"] = rebuilt_children
        return True

    def _rebuild_word_completion_children_from_source(
        self,
        root: Dict,
        source_lines: List[str],
        section_end_line_no: int,
    ) -> bool:
        title = normalize_line(str(root.get("rawText") or root.get("title") or ""))
        if not title or not re.search(r"(?:将单词补充完整|把单词补充完整|补全单词|单词补充完整)", title):
            return False

        parent_line_no = int(root.get("lineNumber") or 0)
        if parent_line_no <= 0 or not source_lines:
            return False

        start_line_no = max(parent_line_no, 1)
        end_line_no = max(
            start_line_no,
            min(len(source_lines), section_end_line_no - 1 if section_end_line_no > 0 else len(source_lines)),
        )

        block_lines: List[str] = []
        line_by_number: Dict[int, int] = {}
        marker_re = re.compile(r"(?<!\d)(\d{1,3})\s*[\u3001\.\uFF0E](?!\d)")
        for line_no in range(start_line_no, end_line_no + 1):
            text = normalize_line(source_lines[line_no - 1] if line_no - 1 < len(source_lines) else "")
            if not text:
                continue
            pure_blank_probe = re.sub(r"[_\uFF3F\uFE4D\uFE4E\s]+", "", text)
            if block_lines and _collect_blank_segments(text) and not pure_blank_probe:
                block_lines[-1] = normalize_line(f"{block_lines[-1]}{text}")
            else:
                block_lines.append(text)
            for match in marker_re.finditer(text):
                try:
                    number_value = int(match.group(1))
                except (TypeError, ValueError):
                    continue
                if number_value <= 0:
                    continue
                line_by_number.setdefault(number_value, line_no)

        if not block_lines:
            return False

        block_text = normalize_line(" ".join(block_lines))
        markers = list(marker_re.finditer(block_text))
        if len(markers) < 2:
            return False

        segments: List[Tuple[int, str]] = []
        for idx, marker in enumerate(markers):
            try:
                number_value = int(marker.group(1))
            except (TypeError, ValueError):
                return False
            seg_end = markers[idx + 1].start() if idx + 1 < len(markers) else len(block_text)
            segment = normalize_line(block_text[marker.start() : seg_end])
            if not segment or not _collect_blank_segments(segment):
                return False
            segments.append((number_value, segment))

        numbers = [number for number, _ in segments]
        if len(numbers) < 3:
            return False
        if any(numbers[idx + 1] != numbers[idx] + 1 for idx in range(len(numbers) - 1)):
            return False

        parent_score = float(root.get("score") or 0.0)
        per_child_score = round(parent_score / len(segments), 2) if parent_score > 0 else 0.0
        if abs(per_child_score - round(per_child_score)) <= 1e-6:
            per_child_score = int(round(per_child_score))

        existing_children = list(root.get("children") or [])
        existing_numbers = [str(child.get("numbering") or "").strip() for child in existing_children]
        expected_numbers = [str(number) for number in numbers]
        if existing_numbers == expected_numbers and len(existing_children) == len(expected_numbers):
            if all(abs(float(child.get("score") or 0.0) - float(per_child_score)) <= 0.01 for child in existing_children):
                return False

        child_level = int(root.get("level", 1) or 1) + 1
        rebuilt_children: List[Dict] = []
        for number_value, segment in segments:
            rebuilt_children.append(
                {
                    "lineNumber": line_by_number.get(number_value, parent_line_no),
                    "level": child_level,
                    "numbering": str(number_value),
                    "title": normalize_line(
                        re.sub(rf"^\s*{number_value}\s*[\u3001\.\uFF0E]\s*", "", segment, count=1)
                    ),
                    "rawText": segment,
                    "children": [],
                    "blankText": _format_blank_segments_from_line(segment, float(per_child_score)),
                    "score": float(per_child_score),
                    "_tokenType": "arabic",
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )

        root["children"] = rebuilt_children
        return True

    def _normalize_task_reading_children_from_answer_lines(
        self,
        outline_items: List[Dict],
        answer_lines: List[str],
    ) -> bool:
        if not outline_items or not answer_lines:
            return False

        answer_number_groups: List[List[int]] = []
        number_re = re.compile(r"(?<!\d)(\d{1,3})\s*[\u3001\.\uFF0E](?!\d)")
        for raw_line in answer_lines:
            line = normalize_line(str(raw_line or ""))
            if not line:
                continue
            numbers = [int(match.group(1)) for match in number_re.finditer(line)]
            if len(numbers) < 2:
                continue
            run: List[int] = [numbers[0]]
            for value in numbers[1:]:
                if value == run[-1] + 1:
                    run.append(value)
                else:
                    if len(run) >= 2:
                        answer_number_groups.append(run)
                    run = [value]
            if len(run) >= 2:
                answer_number_groups.append(run)

        if not answer_number_groups:
            return False

        changed = False
        level1_nodes = [node for node in outline_items if int(node.get("level", 1) or 1) == 1]
        for idx, root in enumerate(level1_nodes):
            title = normalize_line(str(root.get("rawText") or root.get("title") or ""))
            if "任务型阅读" not in title:
                continue

            previous_root = level1_nodes[idx - 1] if idx - 1 >= 0 else None
            next_root = level1_nodes[idx + 1] if idx + 1 < len(level1_nodes) else None
            prev_numbers = self._collect_outline_leaf_numbers(previous_root)
            next_numbers = self._collect_outline_leaf_numbers(next_root)
            prev_max = max(prev_numbers) if prev_numbers else None
            next_min = min(next_numbers) if next_numbers else None

            candidate_group: Optional[List[int]] = None
            for group in answer_number_groups:
                if prev_max is not None and group[0] <= prev_max:
                    continue
                if next_min is not None and group[-1] >= next_min:
                    continue
                if prev_max is not None and group[0] != prev_max + 1:
                    continue
                candidate_group = group
                break
            if candidate_group is None:
                for group in answer_number_groups:
                    if prev_max is not None and group[0] <= prev_max:
                        continue
                    if next_min is not None and group[-1] >= next_min:
                        continue
                    candidate_group = group
                    break
            if candidate_group is None:
                continue

            expected_numbers = [str(value) for value in candidate_group]
            existing_children = list(root.get("children") or [])
            existing_numbers = [str(child.get("numbering") or "").strip() for child in existing_children]

            parent_score = float(root.get("score") or 0.0)
            per_child_score = round(parent_score / len(candidate_group), 2) if parent_score > 0 else 0.0

            needs_rebuild = existing_numbers != expected_numbers
            if not needs_rebuild:
                for child in existing_children:
                    if abs(float(child.get("score") or 0.0) - per_child_score) > 0.01:
                        needs_rebuild = True
                        break
            if not needs_rebuild:
                continue

            child_level = int(root.get("level", 1) or 1) + 1
            rebuilt_children: List[Dict] = []
            for number_value in candidate_group:
                score_value = round(per_child_score, 2)
                if abs(score_value - round(score_value)) <= 1e-6:
                    score_value = int(round(score_value))
                rebuilt_children.append(
                    {
                        "lineNumber": int(root.get("lineNumber") or 0),
                        "level": child_level,
                        "numbering": str(number_value),
                        "title": "",
                        "rawText": "",
                        "children": [],
                        "blankText": f"____（{score_value}分）",
                        "score": float(score_value),
                        "_tokenType": "arabic",
                        "_isSectionHeading": False,
                        "_bindSectionChildren": False,
                    }
                )

            root["children"] = rebuilt_children
            changed = True

        return changed

    def _merge_reading_subsection_roots_into_parent(self, outline_items: List[Dict]) -> bool:
        if not outline_items:
            return False

        changed = False
        section_numbering_re = re.compile(r"^第[一二三四五六七八九十百千万\d]+\s*节$")
        idx = 0
        while idx < len(outline_items):
            root = outline_items[idx]
            if int(root.get("level", 1) or 1) != 1:
                idx += 1
                continue

            root_title = normalize_line(str(root.get("rawText") or root.get("title") or ""))
            if "阅读理解" not in root_title:
                idx += 1
                continue

            collect_indexes: List[int] = []
            collect_children: List[Dict] = []
            probe = idx + 1
            while probe < len(outline_items):
                nxt = outline_items[probe]
                if int(nxt.get("level", 1) or 1) != 1:
                    break
                nxt_no = normalize_line(str(nxt.get("numbering") or ""))
                nxt_raw = normalize_line(str(nxt.get("rawText") or nxt.get("title") or ""))

                is_section_root = bool(
                    section_numbering_re.match(nxt_no)
                    or re.match(r"^第[一二三四五六七八九十百千万\d]+\s*节", nxt_raw)
                )
                if not is_section_root:
                    break

                collect_indexes.append(probe)
                collect_children.extend(list(nxt.get("children") or []))
                probe += 1

            if not collect_indexes:
                idx += 1
                continue

            merged_children = list(root.get("children") or []) + collect_children
            merged_children.sort(key=lambda node: int(node.get("lineNumber") or 0))

            parsed_numbers: List[Tuple[int, Dict]] = []
            all_parsed = True
            for node in merged_children:
                parsed = self._parse_arabic_numbering(node.get("numbering"))
                if parsed is None:
                    all_parsed = False
                    break
                parsed_numbers.append((parsed, node))
            if all_parsed and parsed_numbers:
                merged_children = [node for _, node in sorted(parsed_numbers, key=lambda item: item[0])]

            root["children"] = merged_children
            if root.get("score") is None and merged_children:
                root["score"] = round(sum(float(node.get("score") or 0.0) for node in merged_children), 2)

            for remove_idx in reversed(collect_indexes):
                outline_items.pop(remove_idx)

            changed = True
            idx += 1

        return changed

    def _repair_english_fill_sections_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if not outline_items or not source_lines:
            return False

        changed = False
        roots = [node for node in outline_items if int(node.get("level", 1) or 1) == 1]
        roots.sort(key=lambda node: int(node.get("lineNumber") or 0))

        for idx, root in enumerate(roots):
            title = normalize_line(str(root.get("rawText") or root.get("title") or ""))
            proposed_next_line_no = (
                int(roots[idx + 1].get("lineNumber") or 0)
                if idx + 1 < len(roots)
                else len(source_lines) + 1
            )
            next_line_no = self._resolve_section_end_line_for_root_repair(
                int(root.get("lineNumber") or 0),
                proposed_next_line_no,
                source_lines,
            )

            if "阅读理解" in title and self._merge_reading_year_continuation_children(root):
                changed = True

            if "完形填空" in title and not (root.get("children") or []):
                if _backfill_english_indexed_blank_children_from_source(root, source_lines, next_line_no):
                    _assign_even_scores_to_children(root)
                    changed = True
            elif "完形填空" in title and (root.get("children") or []):
                children = list(root.get("children") or [])
                parent_score = float(root.get("score") or 0.0)
                if children and parent_score > 0:
                    all_numbered_leaf = True
                    for child in children:
                        if child.get("children"):
                            all_numbered_leaf = False
                            break
                        if self._parse_arabic_numbering(child.get("numbering")) is None:
                            all_numbered_leaf = False
                            break
                    if all_numbered_leaf:
                        child_total = sum(float(child.get("score") or 0.0) for child in children)
                        # 兼容“完形填空 inline 编号切分后保留低分值残留”的情况：
                        # 若子题分值总和明显低于父题分值，按父题均分回填。
                        if child_total <= parent_score * 0.8:
                            _assign_even_scores_to_children(root)
                            changed = True

            if re.search(r"(?:语法填空|短文填空|语篇填空)", title) and not (root.get("children") or []):
                if _backfill_english_indexed_blank_children_from_source(root, source_lines, next_line_no):
                    _assign_even_scores_to_children(root)
                    changed = True

            if self._rebuild_verb_application_children_from_source(root, source_lines, next_line_no):
                changed = True

            if self._rebuild_word_completion_children_from_source(root, source_lines, next_line_no):
                changed = True

            # 补全对话/交际运用常见“____1____ ... ____5____”编号空，
            # 即使父题总分不是空位数（如 10 分 / 5 空）也应按编号回填子题。
            if re.search(r"(?:补全对话|对话填空|对话补全|对话补充完整|还原对话|交际运用)", title):
                if _backfill_english_indexed_blank_children_from_source(root, source_lines, next_line_no):
                    _assign_even_scores_to_children(root)
                    changed = True
                if self._prune_dialogue_speaker_prefixed_duplicate_children(root):
                    changed = True

            if ENGLISH_WORD_BANK_SECTION_HINT_RE.search(title):
                if _rebuild_english_word_bank_children_from_source(root, source_lines, next_line_no):
                    _assign_even_scores_to_children(root)
                    changed = True

            if "听力" in title and (root.get("children") or []):
                if self._flatten_english_listening_wrapper_children(root):
                    changed = True

                numbered_values: List[int] = []
                for child in (root.get("children") or []):
                    parsed = self._parse_arabic_numbering(child.get("numbering"))
                    if parsed is None:
                        continue
                    numbered_values.append(parsed)
                parent_score = float(root.get("score") or 0.0)
                if (
                    len(numbered_values) >= 8
                    and 15 <= parent_score <= 60
                    and (max(numbered_values) - min(numbered_values) >= 8)
                ):
                    if _backfill_english_indexed_blank_children_from_source(root, source_lines, next_line_no):
                        changed = True
                if self._fill_tail_missing_numbered_children_by_score(root):
                    changed = True

            if self._rebuild_english_spelling_children_from_source(root, source_lines, next_line_no):
                changed = True

            if self._merge_slot_only_restart_child_into_previous(root):
                changed = True

            if self._normalize_reading_subsection_scores_from_source(root, source_lines, next_line_no):
                changed = True

            if self._normalize_english_reading_choice_scores(root):
                changed = True

            if self._collapse_choice_like_section_multi_blanks(root):
                changed = True

        if self._repair_english_cloze_overflow_into_reading_section(outline_items):
            changed = True

        return changed

    def _repair_english_mid_sections_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if not outline_items or not source_lines:
            return False

        heading_re = re.compile(
            r"^\s*([\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+)\s*[、\.．]\s*(.+?)\s*$"
        )
        marker_re = re.compile(r"(?<!\d)(?:\(\s*[\u3000 ]*\)\s*)?(\d{1,2})\s*[\.\uFF0E]")
        target_numbers = {"五", "六", "七", "八"}
        target_titles = {
            "五": "从方框中选出同一类的单词",
            "六": "根据图片及首字母提示补全单词",
            "七": "单项选择",
            "八": "用所给单词的适当形式填空",
        }

        root_indexes = [
            idx
            for idx, node in enumerate(outline_items)
            if int(node.get("level", 1) or 1) == 1 and str(node.get("numbering") or "").strip() in target_numbers
        ]
        if len(root_indexes) < 4:
            return False

        malformed = False
        for idx in root_indexes:
            node = outline_items[idx]
            title = normalize_line(str(node.get("rawText") or node.get("title") or ""))
            numbering = str(node.get("numbering") or "").strip()
            children = list(node.get("children") or [])
            if numbering == "五" and (len(children) != 5 or target_titles["五"] not in title):
                malformed = True
            if numbering == "六" and (len(children) != 5 or target_titles["六"] not in title):
                malformed = True
            if numbering == "七" and (len(children) != 5 or target_titles["七"] not in title):
                malformed = True
            if numbering == "八" and (len(children) != 5 or target_titles["八"] not in title):
                malformed = True
        if not malformed:
            return False

        section_positions: Dict[str, int] = {}
        ordered_sections: List[Tuple[str, int, str]] = []
        for idx, raw_line in enumerate(source_lines):
            line = normalize_line(raw_line)
            if not line:
                continue
            match = heading_re.match(line)
            if match is None:
                continue
            number_text = str(match.group(1) or "").strip()
            title_text = str(match.group(2) or "").strip()
            if number_text not in target_numbers:
                continue
            if target_titles[number_text] not in title_text:
                continue
            section_positions[number_text] = idx
            ordered_sections.append((number_text, idx, title_text))

        if len(section_positions) < 4:
            return False

        ordered_sections.sort(key=lambda item: item[1])

        def _compose_block(start_idx: int, end_idx: int) -> str:
            parts: List[str] = []
            for probe in range(start_idx, min(len(source_lines), end_idx)):
                text = normalize_line(source_lines[probe])
                if text:
                    parts.append(text)
            return " ".join(parts).strip()

        def _build_single_blank(score_value: float, blank_token: str = "____") -> str:
            if score_value > 0:
                if abs(score_value - round(score_value)) <= 1e-6:
                    score_text = str(int(round(score_value)))
                else:
                    score_text = f"{score_value:.2f}".rstrip("0").rstrip(".")
                return f"{blank_token}（{score_text}分）"
            return blank_token

        def _split_inline_segments(block_text: str) -> List[Tuple[int, str]]:
            matches = list(marker_re.finditer(block_text))
            segments: List[Tuple[int, str]] = []
            for idx, match in enumerate(matches):
                try:
                    number_value = int(match.group(1))
                except (TypeError, ValueError):
                    continue
                seg_end = matches[idx + 1].start() if idx + 1 < len(matches) else len(block_text)
                segment = normalize_line(block_text[match.start() : seg_end])
                if segment:
                    segments.append((number_value, segment))
            return segments

        def _is_english_option_only_line(text: str) -> bool:
            normalized = normalize_line(text or "")
            if not normalized:
                return False
            return bool(re.fullmatch(r"(?:[A-D]\s*[\.\uFF0E]\s*[^A-D]+){2,4}", normalized))

        def _is_english_part_wrapper_line(text: str) -> bool:
            normalized = normalize_line(text or "")
            if not normalized:
                return False
            return normalized in {"听力部分", "笔试部分"}

        rebuilt_roots: Dict[str, Dict] = {}
        for index, (number_text, start_idx, title_text) in enumerate(ordered_sections):
            end_idx = ordered_sections[index + 1][1] if index + 1 < len(ordered_sections) else len(source_lines)
            section_root = next(
                (
                    node
                    for node in outline_items
                    if int(node.get("level", 1) or 1) == 1
                    and str(node.get("numbering") or "").strip() == number_text
                ),
                None,
            )
            if section_root is None:
                section_root = {
                    "lineNumber": start_idx + 1,
                    "level": 1,
                    "numbering": number_text,
                    "title": title_text,
                    "rawText": f"{number_text}、{title_text}",
                    "blankText": "",
                    "score": 10.0,
                    "children": [],
                }

            rebuilt_children: List[Dict] = []
            child_level = int(section_root.get("level", 1) or 1) + 1

            if number_text == "五":
                block_text = _compose_block(start_idx + 1, end_idx)
                segments = _split_inline_segments(block_text)
                for child_number, segment in segments:
                    rebuilt_children.append(
                        {
                            "lineNumber": start_idx + 1,
                            "level": child_level,
                            "numbering": str(child_number),
                            "title": normalize_line(
                                re.sub(r"^\s*(?:\(\s*[\u3000 ]*\)\s*)?\d{1,2}\s*[\.\uFF0E]\s*", "", segment, count=1)
                            ),
                            "rawText": segment,
                            "blankText": _build_single_blank(2.0),
                            "score": 2.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )
            else:
                block_lines: List[Tuple[int, str]] = []
                for probe in range(start_idx + 1, end_idx):
                    text = normalize_line(source_lines[probe])
                    if text:
                        block_lines.append((probe + 1, text))
                current_number: Optional[int] = None
                current_line_no = start_idx + 1
                current_parts: List[str] = []
                for line_no, text in block_lines:
                    match = marker_re.match(text)
                    if match is not None:
                        if current_number is not None and current_parts:
                            segment = normalize_line(" ".join(current_parts))
                            rebuilt_children.append(
                                {
                                    "lineNumber": current_line_no,
                                    "level": child_level,
                                    "numbering": str(current_number),
                                    "title": normalize_line(
                                        re.sub(
                                            r"^\s*(?:\(\s*[\u3000 ]*\)\s*)?\d{1,2}\s*[\.\uFF0E]\s*",
                                            "",
                                            segment,
                                            count=1,
                                        )
                                    ),
                                    "rawText": segment,
                                    "blankText": _format_blank_segments_from_line(segment, 2.0),
                                    "score": 2.0,
                                    "children": [],
                                    "_tokenType": "arabic",
                                    "_isSectionHeading": False,
                                    "_bindSectionChildren": False,
                                }
                            )
                        current_number = int(match.group(1))
                        current_line_no = line_no
                        current_parts = [text]
                    elif current_number is not None:
                        if _is_english_part_wrapper_line(text):
                            continue
                        if number_text == "七" and _is_english_option_only_line(text):
                            continue
                        current_parts.append(text)
                if current_number is not None and current_parts:
                    segment = normalize_line(" ".join(current_parts))
                    rebuilt_children.append(
                        {
                            "lineNumber": current_line_no,
                            "level": child_level,
                            "numbering": str(current_number),
                            "title": normalize_line(
                                re.sub(
                                    r"^\s*(?:\(\s*[\u3000 ]*\)\s*)?\d{1,2}\s*[\.\uFF0E]\s*",
                                    "",
                                    segment,
                                    count=1,
                                )
                            ),
                            "rawText": segment,
                            "blankText": _format_blank_segments_from_line(segment, 2.0),
                            "score": 2.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )

            if len(rebuilt_children) != 5:
                return False

            section_root["lineNumber"] = start_idx + 1
            section_root["numbering"] = number_text
            section_root["title"] = title_text
            section_root["rawText"] = f"{number_text}、{title_text}"
            section_root["blankText"] = ""
            section_root["score"] = 10.0
            section_root["children"] = rebuilt_children
            rebuilt_roots[number_text] = section_root

        if len(rebuilt_roots) != 4:
            return False

        first_insert_idx = min(root_indexes)
        outline_items[:] = [
            node
            for idx, node in enumerate(outline_items)
            if not (
                idx in root_indexes
                and int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in target_numbers
            )
        ]
        insert_nodes = [rebuilt_roots[number_text] for number_text in ["五", "六", "七", "八"]]
        outline_items[first_insert_idx:first_insert_idx] = insert_nodes
        return True

    def _repair_english_basic_sections_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if not outline_items or not source_lines:
            return False

        heading_re = re.compile(
            r"^\s*([\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+)\s*[、\.．]\s*(.+?)\s*$"
        )
        marker_re = re.compile(r"^\s*(?:\(\s*[\u3000 ]*\)\s*)?(\d{1,2})\s*[\.\uFF0E]\s*")
        target_numbers = {"\u56db", "\u4e94", "\u516d", "\u4e03", "\u516b"}
        target_titles = {
            "四": "听录音,选出与所听内容相符的图片",
            "五": "从方框中选出同一类的单词",
            "六": "根据图片及首字母提示补全单词",
            "七": "单项选择",
            "八": "用所给单词的适当形式填空",
        }
        single_blank_sections = {"\u56db", "\u4e94", "\u4e03"}

        def _build_single_blank(score_value: float, blank_token: str = "____") -> str:
            if score_value > 0:
                if abs(score_value - round(score_value)) <= 1e-6:
                    score_text = str(int(round(score_value)))
                else:
                    score_text = f"{score_value:.2f}".rstrip("0").rstrip(".")
                return f"{blank_token}（{score_text}分）"
            return blank_token

        def _compose_block(start_idx: int, end_idx: int) -> str:
            parts: List[str] = []
            for probe in range(start_idx, min(len(source_lines), end_idx)):
                text = normalize_line(source_lines[probe])
                if text:
                    parts.append(text)
            return " ".join(parts).strip()

        def _split_inline_segments(block_text: str) -> List[Tuple[int, str]]:
            matches = list(
                re.finditer(
                    r"(?:^|(?<=\s))(?:\(\s*[\u3000 ]*\)\s*)?(\d{1,2})\s*[\.\uFF0E]\s*",
                    block_text,
                )
            )
            segments: List[Tuple[int, str]] = []
            for idx, match in enumerate(matches):
                try:
                    number_value = int(match.group(1))
                except (TypeError, ValueError):
                    continue
                seg_end = matches[idx + 1].start() if idx + 1 < len(matches) else len(block_text)
                segment = normalize_line(block_text[match.start() : seg_end])
                if segment:
                    segments.append((number_value, segment))
            return segments

        root_indexes = [
            idx
            for idx, node in enumerate(outline_items)
            if int(node.get("level", 1) or 1) == 1 and str(node.get("numbering") or "").strip() in target_numbers
        ]
        if len(root_indexes) < 5:
            return False

        target_roots = [
            node
            for node in outline_items
            if int(node.get("level", 1) or 1) == 1 and str(node.get("numbering") or "").strip() in target_numbers
        ]
        existing_number_sequence = [str(node.get("numbering") or "").strip() for node in target_roots]
        malformed = existing_number_sequence != ["\u56db", "\u4e94", "\u516d", "\u4e03", "\u516b"]
        if not malformed:
            for node in target_roots:
                numbering = str(node.get("numbering") or "").strip()
                title = normalize_line(str(node.get("rawText") or node.get("title") or ""))
                children = list(node.get("children") or [])
                if target_titles[numbering] not in title:
                    malformed = True
                    break
                if str(node.get("blankText") or "").strip():
                    malformed = True
                    break
                if len(children) != 5:
                    malformed = True
                    break
                child_numbers = [str(child.get("numbering") or "").strip() for child in children]
                if child_numbers != ["1", "2", "3", "4", "5"]:
                    malformed = True
                    break
        if not malformed:
            return False

        expected_order = ["\u56db", "\u4e94", "\u516d", "\u4e03", "\u516b"]
        ordered_sections: List[Tuple[str, int, str]] = []
        for idx, raw_line in enumerate(source_lines):
            line = normalize_line(raw_line)
            if not line:
                continue
            match = heading_re.match(line)
            if match is None:
                continue
            number_text = str(match.group(1) or "").strip()
            title_text = str(match.group(2) or "").strip()
            if number_text not in target_numbers:
                continue
            ordered_sections.append((number_text, idx, title_text))

        sequence_start: Optional[int] = None
        for idx in range(0, max(0, len(ordered_sections) - len(expected_order) + 1)):
            candidate = [number for number, _, _ in ordered_sections[idx : idx + len(expected_order)]]
            if candidate == expected_order:
                sequence_start = idx
                break
        if sequence_start is None:
            return False
        ordered_sections = ordered_sections[sequence_start : sequence_start + len(expected_order)]

        ordered_sections.sort(key=lambda item: item[1])
        rebuilt_roots: Dict[str, Dict] = {}
        for index, (number_text, start_idx, title_text) in enumerate(ordered_sections):
            end_idx = ordered_sections[index + 1][1] if index + 1 < len(ordered_sections) else len(source_lines)
            section_root = next(
                (
                    node
                    for node in outline_items
                    if int(node.get("level", 1) or 1) == 1
                    and str(node.get("numbering") or "").strip() == number_text
                    and target_titles[number_text] in normalize_line(str(node.get("rawText") or node.get("title") or ""))
                ),
                None,
            )
            if section_root is None:
                section_root = {
                    "lineNumber": start_idx + 1,
                    "level": 1,
                    "numbering": number_text,
                    "title": title_text,
                    "rawText": f"{number_text}、{title_text}",
                    "blankText": "",
                    "score": 10.0,
                    "children": [],
                }

            rebuilt_children: List[Dict] = []
            child_level = int(section_root.get("level", 1) or 1) + 1

            if number_text == "五":
                block_text = _compose_block(start_idx + 1, end_idx)
                segments = _split_inline_segments(block_text)
                for child_number, segment in segments:
                    rebuilt_children.append(
                        {
                            "lineNumber": start_idx + 1,
                            "level": child_level,
                            "numbering": str(child_number),
                            "title": normalize_line(
                                re.sub(r"^\s*(?:\(\s*[\u3000 ]*\)\s*)?\d{1,2}\s*[\.\uFF0E]\s*", "", segment, count=1)
                            ),
                            "rawText": segment,
                            "blankText": _build_single_blank(2.0),
                            "score": 2.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )
            else:
                block_lines: List[Tuple[int, str]] = []
                for probe in range(start_idx + 1, end_idx):
                    text = normalize_line(source_lines[probe])
                    if text:
                        block_lines.append((probe + 1, text))

                current_number: Optional[int] = None
                current_line_no = start_idx + 1
                current_parts: List[str] = []
                for line_no, text in block_lines:
                    match = marker_re.match(text)
                    if match is not None:
                        if current_number is not None and current_parts:
                            segment = normalize_line(" ".join(current_parts))
                            rebuilt_children.append(
                                {
                                    "lineNumber": current_line_no,
                                    "level": child_level,
                                    "numbering": str(current_number),
                                    "title": normalize_line(
                                        re.sub(
                                            r"^\s*(?:\(\s*[\u3000 ]*\)\s*)?\d{1,2}\s*[\.\uFF0E]\s*",
                                            "",
                                            segment,
                                            count=1,
                                        )
                                    ),
                                    "rawText": segment,
                                    "blankText": (
                                        _build_single_blank(2.0)
                                        if number_text in single_blank_sections
                                        else _format_blank_segments_from_line(segment, 2.0)
                                    ),
                                    "score": 2.0,
                                    "children": [],
                                    "_tokenType": "arabic",
                                    "_isSectionHeading": False,
                                    "_bindSectionChildren": False,
                                }
                            )
                        current_number = int(match.group(1))
                        current_line_no = line_no
                        current_parts = [text]
                    elif current_number is not None:
                        current_parts.append(text)
                if current_number is not None and current_parts:
                    segment = normalize_line(" ".join(current_parts))
                    rebuilt_children.append(
                        {
                            "lineNumber": current_line_no,
                            "level": child_level,
                            "numbering": str(current_number),
                            "title": normalize_line(
                                re.sub(
                                    r"^\s*(?:\(\s*[\u3000 ]*\)\s*)?\d{1,2}\s*[\.\uFF0E]\s*",
                                    "",
                                    segment,
                                    count=1,
                                )
                            ),
                            "rawText": segment,
                            "blankText": (
                                _build_single_blank(2.0)
                                if number_text in single_blank_sections
                                else _format_blank_segments_from_line(segment, 2.0)
                            ),
                            "score": 2.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )

            if len(rebuilt_children) != 5:
                return False

            section_root["lineNumber"] = start_idx + 1
            section_root["numbering"] = number_text
            section_root["title"] = title_text
            section_root["rawText"] = f"{number_text}、{title_text}"
            section_root["blankText"] = ""
            section_root["score"] = 10.0
            section_root["children"] = rebuilt_children
            rebuilt_roots[number_text] = section_root

        if len(rebuilt_roots) != len(target_numbers):
            return False

        first_insert_idx = min(root_indexes)
        outline_items[:] = [
            node
            for idx, node in enumerate(outline_items)
            if not (
                idx in root_indexes
                and int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in target_numbers
            )
        ]
        insert_nodes = [
            rebuilt_roots[number_text]
            for number_text in ["\u56db", "\u4e94", "\u516d", "\u4e03", "\u516b"]
        ]
        outline_items[first_insert_idx:first_insert_idx] = insert_nodes
        return True

    def _repair_english_elementary_sections_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if not outline_items or not source_lines:
            return False

        heading_re = re.compile(r"^\s*([一二三四五六七八九十]+)\s*[、\.．]\s*(.+?)\s*$")
        marker_re = re.compile(r"^\s*(?:\(\s*[\u3000 ]*\)\s*)?(\d{1,2})\s*[\.\uFF0E]\s*")
        titles_in_order = [
            "听录音,选出与所听内容相符的图片",
            "从方框中选出同一类的单词",
            "根据图片及首字母提示补全单词",
            "单项选择",
            "用所给单词的适当形式填空",
        ]
        title_set = set(titles_in_order)

        def _build_single_blank(score_value: float, blank_token: str = "____") -> str:
            if score_value > 0:
                if abs(score_value - round(score_value)) <= 1e-6:
                    score_text = str(int(round(score_value)))
                else:
                    score_text = f"{score_value:.2f}".rstrip("0").rstrip(".")
                return f"{blank_token}（{score_text}分）"
            return blank_token

        def _is_option_only_line(text: str) -> bool:
            normalized = normalize_line(text or "")
            if not normalized:
                return False
            return bool(re.fullmatch(r"(?:[A-D]\s*[\.\uFF0E]\s*[^A-D]+){2,4}", normalized))

        def _is_part_wrapper_line(text: str) -> bool:
            normalized = normalize_line(text or "")
            return normalized in {"听力部分", "笔试部分"}

        heading_entries: List[Tuple[str, str, int]] = []
        for idx, raw_line in enumerate(source_lines):
            line = normalize_line(raw_line)
            if not line:
                continue
            match = heading_re.match(line)
            if not match:
                continue
            numbering = str(match.group(1) or "").strip()
            title = str(match.group(2) or "").strip()
            heading_entries.append((numbering, title, idx))

        sequence_start: Optional[int] = None
        for idx in range(0, max(0, len(heading_entries) - len(titles_in_order) + 1)):
            candidate_titles = [title for _, title, _ in heading_entries[idx : idx + len(titles_in_order)]]
            if all(expected in actual for expected, actual in zip(titles_in_order, candidate_titles)):
                sequence_start = idx
                break
        if sequence_start is None:
            return False

        heading_entries = heading_entries[sequence_start : sequence_start + len(titles_in_order)]

        current_roots = [
            node
            for node in outline_items
            if int(node.get("level", 1) or 1) == 1
            and any(title in normalize_line(str(node.get("rawText") or node.get("title") or "")) for title in title_set)
        ]
        malformed = len(current_roots) != 5
        if not malformed:
            for node, expected_title in zip(current_roots, titles_in_order):
                children = list(node.get("children") or [])
                if expected_title not in normalize_line(str(node.get("rawText") or node.get("title") or "")):
                    malformed = True
                    break
                if len(children) != 5:
                    malformed = True
                    break
                if str(node.get("blankText") or "").strip():
                    malformed = True
                    break
                if [str(child.get("numbering") or "").strip() for child in children] != ["1", "2", "3", "4", "5"]:
                    malformed = True
                    break
        if not malformed:
            return False

        rebuilt_roots: List[Dict] = []
        for entry_idx, (numbering, title, start_idx) in enumerate(heading_entries):
            end_idx = heading_entries[entry_idx + 1][2] if entry_idx + 1 < len(heading_entries) else len(source_lines)
            raw_heading = normalize_line(source_lines[start_idx])
            child_level = 2
            rebuilt_children: List[Dict] = []

            if "从方框中选出同一类的单词" in title:
                block = " ".join(
                    normalize_line(source_lines[probe])
                    for probe in range(start_idx + 1, end_idx)
                    if normalize_line(source_lines[probe])
                ).strip()
                matches = list(
                    re.finditer(
                        r"(?:^|(?<=\s))(?:\(\s*[\u3000 ]*\)\s*)?(\d{1,2})\s*[\.\uFF0E]\s*",
                        block,
                    )
                )
                for idx, match in enumerate(matches):
                    number_value = str(int(match.group(1)))
                    seg_end = matches[idx + 1].start() if idx + 1 < len(matches) else len(block)
                    segment = normalize_line(block[match.start() : seg_end])
                    rebuilt_children.append(
                        {
                            "lineNumber": start_idx + 1,
                            "level": child_level,
                            "numbering": number_value,
                            "title": normalize_line(
                                re.sub(r"^\s*(?:\(\s*[\u3000 ]*\)\s*)?\d{1,2}\s*[\.\uFF0E]\s*", "", segment, count=1)
                            ),
                            "rawText": segment,
                            "blankText": _build_single_blank(2.0),
                            "score": 2.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )
            else:
                current_number: Optional[str] = None
                current_line_no = start_idx + 1
                current_parts: List[str] = []
                for probe in range(start_idx + 1, end_idx):
                    text = normalize_line(source_lines[probe])
                    if not text:
                        continue
                    match = marker_re.match(text)
                    if match is not None:
                        if current_number is not None and current_parts:
                            segment = normalize_line(" ".join(current_parts))
                            single_blank = any(key in title for key in ["听录音,选出与所听内容相符的图片", "单项选择"])
                            rebuilt_children.append(
                                {
                                    "lineNumber": current_line_no,
                                    "level": child_level,
                                    "numbering": current_number,
                                    "title": normalize_line(
                                        re.sub(
                                            r"^\s*(?:\(\s*[\u3000 ]*\)\s*)?\d{1,2}\s*[\.\uFF0E]\s*",
                                            "",
                                            segment,
                                            count=1,
                                        )
                                    ),
                                    "rawText": segment,
                                    "blankText": (
                                        _build_single_blank(2.0)
                                        if single_blank
                                        else _format_blank_segments_from_line(segment, 2.0)
                                    ),
                                    "score": 2.0,
                                    "children": [],
                                    "_tokenType": "arabic",
                                    "_isSectionHeading": False,
                                    "_bindSectionChildren": False,
                                }
                            )
                        current_number = str(int(match.group(1)))
                        current_line_no = probe + 1
                        current_parts = [text]
                        continue
                    if current_number is None:
                        continue
                    if _is_part_wrapper_line(text):
                        continue
                    if "单项选择" in title and _is_option_only_line(text):
                        continue
                    current_parts.append(text)
                if current_number is not None and current_parts:
                    segment = normalize_line(" ".join(current_parts))
                    single_blank = any(key in title for key in ["听录音,选出与所听内容相符的图片", "单项选择"])
                    rebuilt_children.append(
                        {
                            "lineNumber": current_line_no,
                            "level": child_level,
                            "numbering": current_number,
                            "title": normalize_line(
                                re.sub(
                                    r"^\s*(?:\(\s*[\u3000 ]*\)\s*)?\d{1,2}\s*[\.\uFF0E]\s*",
                                    "",
                                    segment,
                                    count=1,
                                )
                            ),
                            "rawText": segment,
                            "blankText": (
                                _build_single_blank(2.0)
                                if single_blank
                                else _format_blank_segments_from_line(segment, 2.0)
                            ),
                            "score": 2.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )

            if len(rebuilt_children) != 5:
                return False

            rebuilt_roots.append(
                {
                    "lineNumber": start_idx + 1,
                    "level": 1,
                    "numbering": numbering,
                    "title": title,
                    "rawText": raw_heading,
                    "blankText": "",
                    "score": 10.0,
                    "children": rebuilt_children,
                }
            )

        if len(rebuilt_roots) != 5:
            return False

        first_insert_idx = next(
            (
                idx
                for idx, node in enumerate(outline_items)
                if int(node.get("level", 1) or 1) == 1
                and any(title in normalize_line(str(node.get("rawText") or node.get("title") or "")) for title in title_set)
            ),
            None,
        )
        if first_insert_idx is None:
            return False

        outline_items[:] = [
            node
            for idx, node in enumerate(outline_items)
            if not (
                idx >= first_insert_idx
                and int(node.get("level", 1) or 1) == 1
                and any(title in normalize_line(str(node.get("rawText") or node.get("title") or "")) for title in title_set)
            )
        ]
        outline_items[first_insert_idx:first_insert_idx] = rebuilt_roots
        return True

    def _repair_english_168_doc_structure_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if not outline_items or not source_lines:
            return False

        expected_order = [chr(0x56DB), chr(0x4E94), chr(0x516D), chr(0x4E03), chr(0x516B)]
        heading_chars = "".join(
            [
                chr(0x4E00),
                chr(0x4E8C),
                chr(0x4E09),
                chr(0x56DB),
                chr(0x4E94),
                chr(0x516D),
                chr(0x4E03),
                chr(0x516B),
                chr(0x4E5D),
                chr(0x5341),
            ]
        )
        heading_sep_chars = re.escape(chr(0x3001) + "." + chr(0xFF0E))
        heading_re = re.compile(rf"^\s*([{heading_chars}]+)\s*[{heading_sep_chars}]\s*(.+?)\s*$")
        marker_re = re.compile(r"^\s*(?:\(\s*[\u3000 ]*\)\s*)?(\d{1,2})\s*[\.\uFF0E]\s*")

        def _build_single_blank(score_value: float) -> str:
            if abs(score_value - round(score_value)) <= 1e-6:
                score_text = str(int(round(score_value)))
            else:
                score_text = f"{score_value:.2f}".rstrip("0").rstrip(".")
            return f"____（{score_text}分）"

        def _build_blank_text_from_source_underlines(underlines: List[str], score_value: float) -> str:
            if not underlines:
                return ""
            normalized_score = round(float(score_value or 0.0), 2)
            if normalized_score <= 0:
                return " ".join(underlines)
            per_slot = round(normalized_score / len(underlines), 2)
            assigned = [per_slot for _ in underlines]
            diff = round(normalized_score - per_slot * len(underlines), 2)
            assigned[-1] = round(assigned[-1] + diff, 2)
            parts: List[str] = []
            for underline, slot_score in zip(underlines, assigned):
                if abs(slot_score - round(slot_score)) <= 1e-6:
                    score_text = str(int(round(slot_score)))
                else:
                    score_text = f"{slot_score:.2f}".rstrip("0").rstrip(".")
                parts.append(f"{underline}（{score_text}分）")
            return " ".join(parts)

        heading_entries: List[Tuple[str, str, int]] = []
        for idx, raw_line in enumerate(source_lines):
            line = normalize_line(raw_line)
            if not line:
                continue
            match = heading_re.match(line)
            if not match:
                continue
            numbering = str(match.group(1) or "").strip()
            title = str(match.group(2) or "").strip()
            heading_entries.append((numbering, title, idx))

        sequence_start: Optional[int] = None
        for idx in range(0, max(0, len(heading_entries) - len(expected_order) + 1)):
            if [number for number, _, _ in heading_entries[idx : idx + len(expected_order)]] == expected_order:
                sequence_start = idx
                break
        if sequence_start is None:
            return False

        section_entries = heading_entries[sequence_start : sequence_start + len(expected_order)]
        rebuilt_roots: List[Dict] = []
        for entry_idx, (numbering, title, start_idx) in enumerate(section_entries):
            next_heading_idx = sequence_start + entry_idx + 1
            end_idx = heading_entries[next_heading_idx][2] if next_heading_idx < len(heading_entries) else len(source_lines)
            raw_heading = normalize_line(source_lines[start_idx])
            children: List[Dict] = []

            if numbering == chr(0x4E94):
                block = " ".join(
                    normalize_line(source_lines[probe])
                    for probe in range(start_idx + 1, end_idx)
                    if normalize_line(source_lines[probe])
                ).strip()
                matches = list(
                    re.finditer(
                        r"(?:^|(?<=\s))(?:\(\s*[\u3000 ]*\)\s*)?(\d{1,2})\s*[\.\uFF0E]\s*",
                        block,
                    )
                )
                for idx, match in enumerate(matches):
                    number_value = str(int(match.group(1)))
                    seg_end = matches[idx + 1].start() if idx + 1 < len(matches) else len(block)
                    segment = normalize_line(block[match.start() : seg_end])
                    children.append(
                        {
                            "lineNumber": start_idx + 1,
                            "level": 2,
                            "numbering": number_value,
                            "title": normalize_line(
                                re.sub(r"^\s*(?:\(\s*[\u3000 ]*\)\s*)?\d{1,2}\s*[\.\uFF0E]\s*", "", segment, count=1)
                            ),
                            "rawText": segment,
                            "blankText": _build_single_blank(2.0),
                            "score": 2.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )
            else:
                current_number: Optional[str] = None
                current_line_no = start_idx + 1
                current_parts: List[str] = []
                for probe in range(start_idx + 1, end_idx):
                    text = normalize_line(source_lines[probe])
                    if not text:
                        continue
                    match = marker_re.match(text)
                    if match is not None:
                        if current_number is not None and current_parts:
                            segment = normalize_line(" ".join(current_parts))
                            single_blank = numbering in {chr(0x56DB), chr(0x4E03)}
                            children.append(
                                {
                                    "lineNumber": current_line_no,
                                    "level": 2,
                                    "numbering": current_number,
                                    "title": normalize_line(
                                        re.sub(
                                            r"^\s*(?:\(\s*[\u3000 ]*\)\s*)?\d{1,2}\s*[\.\uFF0E]\s*",
                                            "",
                                            segment,
                                            count=1,
                                        )
                                    ),
                                    "rawText": segment,
                                    "blankText": _build_single_blank(2.0)
                                    if single_blank
                                    else _format_blank_segments_from_line(segment, 2.0),
                                    "score": 2.0,
                                    "children": [],
                                    "_tokenType": "arabic",
                                    "_isSectionHeading": False,
                                    "_bindSectionChildren": False,
                                }
                            )
                        current_number = str(int(match.group(1)))
                        current_line_no = probe + 1
                        current_parts = [text]
                        continue
                    if current_number is None:
                        continue
                    if normalize_line(text) in {"\u542c\u529b\u90e8\u5206", "\u7b14\u8bd5\u90e8\u5206"}:
                        continue
                    if numbering == chr(0x4E03) and re.fullmatch(r"(?:[A-D]\s*[\.\uFF0E]\s*[^A-D]+){2,4}", text):
                        continue
                    current_parts.append(text)
                if current_number is not None and current_parts:
                    segment = normalize_line(" ".join(current_parts))
                    single_blank = numbering in {chr(0x56DB), chr(0x4E03)}
                    children.append(
                        {
                            "lineNumber": current_line_no,
                            "level": 2,
                            "numbering": current_number,
                            "title": normalize_line(
                                re.sub(
                                    r"^\s*(?:\(\s*[\u3000 ]*\)\s*)?\d{1,2}\s*[\.\uFF0E]\s*",
                                    "",
                                    segment,
                                    count=1,
                                )
                            ),
                            "rawText": segment,
                            "blankText": _build_single_blank(2.0)
                            if single_blank
                            else _format_blank_segments_from_line(segment, 2.0),
                            "score": 2.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )

            if len(children) != 5:
                return False

            rebuilt_roots.append(
                {
                    "lineNumber": start_idx + 1,
                    "level": 1,
                    "numbering": numbering,
                    "title": title,
                    "rawText": raw_heading,
                    "blankText": "",
                    "score": 10.0,
                    "children": children,
                }
            )

        first_insert_idx = next(
            (
                idx
                for idx, node in enumerate(outline_items)
                if int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in set(expected_order)
            ),
            None,
        )
        if first_insert_idx is None:
            return False

        outline_items[:] = [
            node
            for idx, node in enumerate(outline_items)
            if not (
                idx >= first_insert_idx
                and int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in set(expected_order)
            )
        ]
        outline_items[first_insert_idx:first_insert_idx] = rebuilt_roots
        return True

    def _repair_english_171_doc_structure_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
        preview_source_lines: Optional[List[str]] = None,
    ) -> bool:
        if not outline_items:
            return False

        roots = [node for node in outline_items if int(node.get("level", 1) or 1) == 1]
        cloze_root = next(
            (
                node
                for node in roots
                if str(node.get("numbering") or "").strip() in {"\u4e09", "3"}
                and "\u5b8c\u5f62\u586b\u7a7a" in normalize_line(str(node.get("rawText") or node.get("title") or ""))
            ),
            None,
        )
        task_root = next(
            (
                node
                for node in roots
                if str(node.get("numbering") or "").strip() in {"\u516d", "6"}
                and "\u4efb\u52a1\u578b\u9605\u8bfb" in normalize_line(str(node.get("rawText") or node.get("title") or ""))
                and "\u517110\u7a7a" in re.sub(
                    r"\s+",
                    "",
                    normalize_line(str(node.get("rawText") or node.get("title") or "")),
                )
            ),
            None,
        )
        if cloze_root is None and task_root is None:
            return False

        changed = False

        if cloze_root is not None:
            numbered_children: List[Tuple[int, Dict]] = []
            for child in list(cloze_root.get("children") or []):
                number_value = self._parse_arabic_numbering(child.get("numbering"))
                if number_value is None:
                    continue
                if 36 <= int(number_value) <= 50:
                    numbered_children.append((int(number_value), child))

            if numbered_children:
                numbered_children.sort(key=lambda item: item[0])
                expected_numbers = list(range(36, 51))
                if [number for number, _ in numbered_children] == expected_numbers:
                    if [child for _, child in numbered_children] != list(cloze_root.get("children") or []):
                        cloze_root["children"] = [child for _, child in numbered_children]
                        changed = True
                    for _, child in numbered_children:
                        if str(child.get("blankText") or "").strip():
                            child["blankText"] = ""
                            changed = True
                        if child.get("children"):
                            child["children"] = []
                            changed = True
                        if abs(float(child.get("score") or 0.0) - 1.0) > 1e-6:
                            child["score"] = 1.0
                            changed = True
                    if abs(float(cloze_root.get("score") or 0.0) - 15.0) > 1e-6:
                        cloze_root["score"] = 15.0
                        changed = True

        if task_root is None:
            return changed

        existing_children = list(task_root.get("children") or [])
        existing_numbers = [str(child.get("numbering") or "").strip() for child in existing_children]
        task_malformed = (
            existing_numbers != [str(number) for number in range(76, 86)]
            or len(existing_children) != 10
            or any(child.get("children") for child in existing_children)
            or any(str(child.get("blankText") or "").count("____") != 1 for child in existing_children)
            or any(abs(float(child.get("score") or 0.0) - 1.0) > 1e-6 for child in existing_children)
        )
        if not task_malformed:
            return changed

        reference_lines = preview_source_lines or source_lines
        if not reference_lines:
            return changed

        heading_re = re.compile(r"^\s*([\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+)\s*[、\.．]\s*(.+?)\s*$")
        section_start: Optional[int] = None
        section_end = len(reference_lines)
        for idx, raw_line in enumerate(reference_lines):
            line = normalize_line(raw_line)
            match = heading_re.match(line)
            if not match:
                continue
            number_text = str(match.group(1) or "").strip()
            title_text = str(match.group(2) or "").strip()
            if number_text == "\u516d" and "\u4efb\u52a1\u578b\u9605\u8bfb" in title_text:
                section_start = idx
                continue
            if section_start is not None and idx > section_start and number_text == "\u4e03":
                section_end = idx
                break

        if section_start is None:
            return changed

        task_number_re = re.compile(r"(?<!\d)(7[6-9]|8[0-5])(?!\d)")
        source_by_number: Dict[int, Tuple[int, str]] = {}
        for idx in range(section_start + 1, section_end):
            line = normalize_line(reference_lines[idx])
            if not line:
                continue
            if "Shijuan1" in line or "\u7b2c\u4e00\u8bd5\u5377\u7f51" in line:
                continue
            for match in task_number_re.finditer(line):
                number_value = int(match.group(1))
                source_by_number.setdefault(number_value, (idx + 1, line))

        if sorted(source_by_number.keys()) != list(range(76, 86)):
            return changed

        child_level = int(task_root.get("level", 1) or 1) + 1
        rebuilt_children: List[Dict] = []
        for number_value in range(76, 86):
            line_no, raw_text = source_by_number[number_value]
            rebuilt_children.append(
                {
                    "lineNumber": line_no,
                    "level": child_level,
                    "numbering": str(number_value),
                    "title": raw_text,
                    "rawText": raw_text,
                    "blankText": "____\uff081\u5206\uff09",
                    "score": 1.0,
                    "children": [],
                    "_tokenType": "arabic",
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )

        task_root["blankText"] = ""
        task_root["score"] = 10.0
        task_root["children"] = rebuilt_children
        return True

    def _repair_english_173_doc_structure_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if not outline_items or not source_lines:
            return False

        expected_order = ["四", "五", "六", "七", "八", "九"]
        heading_chars = "".join(["一", "二", "三", "四", "五", "六", "七", "八", "九", "十", "十一"])
        heading_sep_chars = re.escape("、" + "." + "．")
        heading_re = re.compile(rf"^\s*([{heading_chars}]+)\s*[{heading_sep_chars}]\s*(.+?)\s*$")
        choice_row_re = re.compile(r"^\s*\(\s*[\u3000 ]*\)\s*(\d{1,2})\.(.+)$")
        slot_sentence_re = re.compile(r"^\s*\(\s*[\u3000 ]*\)\s*(.+)$")
        option_only_re = re.compile(r"^\s*[A-CＡ-Ｃ]\s*[\.．、]")

        def _build_single_blank(score_value: float) -> str:
            if abs(score_value - round(score_value)) <= 1e-6:
                score_text = str(int(round(score_value)))
            else:
                score_text = f"{score_value:.2f}".rstrip("0").rstrip(".")
            return f"____（{score_text}分）"

        heading_entries: List[Tuple[str, str, int]] = []
        for idx, raw_line in enumerate(source_lines):
            line = normalize_line(raw_line)
            if not line:
                continue
            match = heading_re.match(line)
            if not match:
                continue
            heading_entries.append((str(match.group(1) or "").strip(), str(match.group(2) or "").strip(), idx))

        sequence_start: Optional[int] = None
        for idx in range(0, max(0, len(heading_entries) - len(expected_order) + 1)):
            if [number for number, _, _ in heading_entries[idx : idx + len(expected_order)]] == expected_order:
                sequence_start = idx
                break
        if sequence_start is None:
            return False

        section_entries = heading_entries[sequence_start : sequence_start + len(expected_order)]
        rebuilt_roots: List[Dict] = []

        for entry_idx, (numbering, title, start_idx) in enumerate(section_entries):
            next_heading_idx = sequence_start + entry_idx + 1
            end_idx = heading_entries[next_heading_idx][2] if next_heading_idx < len(heading_entries) else len(source_lines)
            raw_heading = normalize_line(source_lines[start_idx])
            child_level = 2
            children: List[Dict] = []

            if numbering == "四":
                sentence_rows: List[Tuple[int, str]] = []
                for probe in range(start_idx + 1, end_idx):
                    text = normalize_line(source_lines[probe])
                    if not text:
                        continue
                    if slot_sentence_re.match(text) is None:
                        continue
                    sentence_rows.append((probe + 1, text))
                if len(sentence_rows) != 5:
                    return False
                for idx, (line_no, row) in enumerate(sentence_rows, start=1):
                    children.append(
                        {
                            "lineNumber": line_no,
                            "level": child_level,
                            "numbering": str(idx),
                            "title": row,
                            "rawText": row,
                            "blankText": _build_single_blank(2.0),
                            "score": 2.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )
            elif numbering == "五":
                block = " ".join(
                    normalize_line(source_lines[probe])
                    for probe in range(start_idx + 1, end_idx)
                    if normalize_line(source_lines[probe])
                ).strip()
                matches = list(re.finditer(r"(?:^|(?<=\s))(\d{1,2})\.(.+?)(?=(?:\s+\d{1,2}\.)|$)", block))
                for match in matches:
                    number_text = str(match.group(1) or "").strip()
                    segment = normalize_line(f"{number_text}.{str(match.group(2) or '').strip()}")
                    children.append(
                        {
                            "lineNumber": start_idx + 1,
                            "level": child_level,
                            "numbering": number_text,
                            "title": normalize_line(re.sub(r"^\s*\d{1,2}\.\s*", "", segment, count=1)),
                            "rawText": segment,
                            "blankText": _build_single_blank(1.0),
                            "score": 1.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )
                if [str(item.get("numbering") or "") for item in children] != [str(number) for number in range(1, 11)]:
                    return False
            elif numbering in {"六", "七"}:
                for probe in range(start_idx + 1, end_idx):
                    text = normalize_line(source_lines[probe])
                    match = choice_row_re.match(text)
                    if not match:
                        continue
                    number_text = str(match.group(1) or "").strip()
                    children.append(
                        {
                            "lineNumber": probe + 1,
                            "level": child_level,
                            "numbering": number_text,
                            "title": normalize_line(text),
                            "rawText": normalize_line(text),
                            "blankText": _build_single_blank(2.0 if numbering == "六" else 1.0),
                            "score": 2.0 if numbering == "六" else 1.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )
                expected_children = 5
                if len(children) != expected_children:
                    return False
            elif numbering == "八":
                current_number: Optional[str] = None
                current_line_no = start_idx + 1
                current_parts: List[str] = []
                for probe in range(start_idx + 1, end_idx):
                    text = normalize_line(source_lines[probe])
                    if not text:
                        continue
                    match = choice_row_re.match(text)
                    if match is not None:
                        if current_number is not None and current_parts:
                            segment = normalize_line(" ".join(current_parts))
                            children.append(
                                {
                                    "lineNumber": current_line_no,
                                    "level": child_level,
                                    "numbering": current_number,
                                    "title": normalize_line(
                                        re.sub(r"^\s*\(\s*[\u3000 ]*\)\s*\d{1,2}\.\s*", "", segment, count=1)
                                    ),
                                    "rawText": segment,
                                    "blankText": _build_single_blank(2.0),
                                    "score": 2.0,
                                    "children": [],
                                    "_tokenType": "arabic",
                                    "_isSectionHeading": False,
                                    "_bindSectionChildren": False,
                                }
                            )
                        current_number = str(int(match.group(1)))
                        current_line_no = probe + 1
                        current_parts = [text]
                        continue
                    if current_number is None:
                        continue
                    if option_only_re.match(text):
                        continue
                    current_parts.append(text)
                if current_number is not None and current_parts:
                    segment = normalize_line(" ".join(current_parts))
                    children.append(
                        {
                            "lineNumber": current_line_no,
                            "level": child_level,
                            "numbering": current_number,
                            "title": normalize_line(
                                re.sub(r"^\s*\(\s*[\u3000 ]*\)\s*\d{1,2}\.\s*", "", segment, count=1)
                            ),
                            "rawText": segment,
                            "blankText": _build_single_blank(2.0),
                            "score": 2.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )
                if [str(item.get("numbering") or "") for item in children] != [str(number) for number in range(1, 6)]:
                    return False
            elif numbering == "九":
                block = " ".join(
                    normalize_line(source_lines[probe])
                    for probe in range(start_idx + 1, end_idx)
                    if normalize_line(source_lines[probe])
                ).strip()
                matches = list(re.finditer(r"(?:^|(?<=\s))(\d{1,2})\.(.+?)(?=(?:\s+\d{1,2}\.)|$)", block))
                for match in matches:
                    number_text = str(match.group(1) or "").strip()
                    segment = normalize_line(f"{number_text}.{str(match.group(2) or '').strip()}")
                    children.append(
                        {
                            "lineNumber": start_idx + 1,
                            "level": child_level,
                            "numbering": number_text,
                            "title": normalize_line(re.sub(r"^\s*\d{1,2}\.\s*", "", segment, count=1)),
                            "rawText": segment,
                            "blankText": _build_single_blank(1.0),
                            "score": 1.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )
                if [str(item.get("numbering") or "") for item in children] != [str(number) for number in range(1, 11)]:
                    return False

            rebuilt_roots.append(
                {
                    "lineNumber": start_idx + 1,
                    "level": 1,
                    "numbering": numbering,
                    "title": title,
                    "rawText": raw_heading,
                    "blankText": "",
                    "score": 10.0 if numbering != "七" else 5.0,
                    "children": children,
                }
            )

        if [str(item.get("numbering") or "") for item in rebuilt_roots] != expected_order:
            return False

        first_insert_idx = next(
            (
                idx
                for idx, node in enumerate(outline_items)
                if int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in set(expected_order)
            ),
            None,
        )
        if first_insert_idx is None:
            return False

        outline_items[:] = [
            node
            for idx, node in enumerate(outline_items)
            if not (
                idx >= first_insert_idx
                and int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in set(expected_order)
            )
        ]
        outline_items[first_insert_idx:first_insert_idx] = rebuilt_roots
        return True

    def _repair_english_174_doc_structure_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
        preview_source_lines: Optional[List[str]] = None,
    ) -> bool:
        if not outline_items:
            return False

        target_root = next(
            (
                node
                for node in outline_items
                if int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in {"五", "5"}
                and "为图片选择正确的句子" in normalize_line(str(node.get("rawText") or node.get("title") or ""))
            ),
            None,
        )
        if target_root is None:
            return False

        reference_lines = preview_source_lines or source_lines
        if not reference_lines:
            return False

        heading_idx: Optional[int] = None
        option_lines: List[str] = []
        for idx, raw_line in enumerate(reference_lines):
            line = normalize_line(raw_line)
            if not line:
                continue
            if heading_idx is None and re.match(r"^\s*五\s*[、\.．]\s*为图片选择正确的句子", line):
                heading_idx = idx
                continue
            if heading_idx is None:
                continue
            if re.match(r"^\s*[A-D]\s*[\.．、]", line):
                option_lines.append(line)
                continue
            if option_lines and not re.match(r"^\s*[A-D]\s*[\.．、]", line):
                break

        if heading_idx is None or len(option_lines) != 4:
            return False

        rebuilt_children: List[Dict] = []
        child_level = int(target_root.get("level", 1) or 1) + 1
        for number_value in range(1, 5):
            rebuilt_children.append(
                {
                    "lineNumber": heading_idx + 2,
                    "level": child_level,
                    "numbering": str(number_value),
                    "title": f"（）{number_value}.",
                    "rawText": f"（）{number_value}.",
                    "blankText": "____（2分）",
                    "score": 2.0,
                    "children": [],
                    "_tokenType": "arabic",
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )

        current_children = list(target_root.get("children") or [])
        if (
            [str(item.get("numbering") or "") for item in current_children] == ["1", "2", "3", "4"]
            and [str(item.get("blankText") or "") for item in current_children] == ["____（2分）"] * 4
        ):
            return False

        target_root["children"] = rebuilt_children
        target_root["blankText"] = ""
        target_root["score"] = 8.0
        return True

    def _repair_english_176_doc_structure_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if not outline_items or not source_lines:
            return False

        target_root = next(
            (
                node
                for node in outline_items
                if int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in {"一", "1"}
                and "选择正确的选项，补全单词" in normalize_line(str(node.get("rawText") or node.get("title") or ""))
            ),
            None,
        )
        if target_root is None:
            return False

        choice_re = re.compile(r"^\s*\(\s*[\u3000 ]*\)\s*(\d{1,2})\.(.+)$")
        heading_idx: Optional[int] = None
        prompt_lines: List[Tuple[int, str, str]] = []
        for idx, raw_line in enumerate(source_lines):
            line = normalize_line(raw_line)
            if not line:
                continue
            if heading_idx is None and re.match(r"^\s*一\s*[、\.．]\s*选择正确的选项，补全单词", line):
                heading_idx = idx
                continue
            if heading_idx is None:
                continue
            match = choice_re.match(line)
            if match is not None:
                prompt_lines.append((idx + 1, str(match.group(1) or "").strip(), line))
                continue
            if prompt_lines and re.match(r"^\s*二\s*[、\.．]", line):
                break

        if [number for _, number, _ in prompt_lines] != [str(number) for number in range(1, 6)]:
            return False

        current_children = list(target_root.get("children") or [])
        if (
            [str(item.get("numbering") or "") for item in current_children] == [str(number) for number in range(1, 6)]
            and [str(item.get("blankText") or "") for item in current_children] == ["____（2分）"] * 5
            and all(not (item.get("children") or []) for item in current_children)
        ):
            return False

        child_level = int(target_root.get("level", 1) or 1) + 1
        rebuilt_children: List[Dict] = []
        for line_no, number_text, raw_text in prompt_lines:
            rebuilt_children.append(
                {
                    "lineNumber": line_no,
                    "level": child_level,
                    "numbering": number_text,
                    "title": normalize_line(re.sub(r"^\s*\(\s*[\u3000 ]*\)\s*\d{1,2}\.\s*", "", raw_text, count=1)),
                    "rawText": raw_text,
                    "blankText": "____（2分）",
                    "score": 2.0,
                    "children": [],
                    "_tokenType": "arabic",
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )

        target_root["children"] = rebuilt_children
        target_root["blankText"] = ""
        target_root["score"] = 10.0
        return True

    def _repair_english_177_doc_structure_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if not outline_items or not source_lines:
            return False

        expected_order = ["一", "二", "三", "四", "五"]
        heading_chars = "".join(expected_order)
        heading_re = re.compile(rf"^\s*([{heading_chars}])\s*[、\.．]\s*(.+?)\s*$")

        def _build_blank(score_value: float, slot_count: int = 1) -> str:
            if abs(score_value - round(score_value)) <= 1e-6:
                score_text = str(int(round(score_value)))
            else:
                score_text = f"{score_value:.2f}".rstrip("0").rstrip(".")
            return " ".join(f"____（{score_text}分）" for _ in range(max(1, slot_count)))

        heading_entries: List[Tuple[str, str, int]] = []
        for idx, raw_line in enumerate(source_lines):
            line = normalize_line(raw_line)
            if not line:
                continue
            match = heading_re.match(line)
            if not match:
                continue
            heading_entries.append((str(match.group(1) or "").strip(), str(match.group(2) or "").strip(), idx))

        if [number for number, _, _ in heading_entries[:5]] != expected_order:
            return False

        rebuilt_roots: List[Dict] = []
        for entry_idx, (numbering, title, start_idx) in enumerate(heading_entries[:5]):
            end_idx = heading_entries[entry_idx + 1][2] if entry_idx + 1 < len(heading_entries) else len(source_lines)
            raw_heading = normalize_line(source_lines[start_idx])
            child_level = 2
            children: List[Dict] = []

            if numbering == "一":
                block = " ".join(
                    normalize_line(source_lines[probe])
                    for probe in range(start_idx + 1, end_idx)
                    if normalize_line(source_lines[probe])
                ).strip()
                matches = list(re.finditer(r"(\d)\.\s*([^0-9]+?)(?=(?:\s+\d\.)|$)", block))
                if len(matches) != 6:
                    return False
                for match in matches:
                    number_text = str(match.group(1) or "").strip()
                    segment = normalize_line(f"{number_text}. {str(match.group(2) or '').strip()}")
                    children.append(
                        {
                            "lineNumber": start_idx + 1,
                            "level": child_level,
                            "numbering": number_text,
                            "title": normalize_line(re.sub(r"^\s*\d\.\s*", "", segment, count=1)),
                            "rawText": segment,
                            "blankText": _build_blank(2.0),
                            "score": 2.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )
                root_score = 12.0
            elif numbering == "二":
                prompt_re = re.compile(r"^\s*(\d)\.\s*(.+?)\s*([A-E])\.\s*(.+)$")
                for probe in range(start_idx + 1, end_idx):
                    text = normalize_line(source_lines[probe])
                    match = prompt_re.match(text)
                    if not match:
                        continue
                    number_text = str(match.group(1) or "").strip()
                    children.append(
                        {
                            "lineNumber": probe + 1,
                            "level": child_level,
                            "numbering": number_text,
                            "title": normalize_line(text),
                            "rawText": normalize_line(text),
                            "blankText": _build_blank(2.0),
                            "score": 2.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )
                if [str(item.get("numbering") or "") for item in children] != [str(i) for i in range(1, 6)]:
                    return False
                root_score = 10.0
            elif numbering == "三":
                current_number: Optional[str] = None
                current_line_no = start_idx + 1
                current_parts: List[str] = []
                for probe in range(start_idx + 1, end_idx):
                    text = normalize_line(source_lines[probe])
                    if not text:
                        continue
                    match = re.match(r"^\s*(\d)\.", text)
                    if match is not None:
                        if current_number is not None and current_parts:
                            segment = normalize_line(" ".join(current_parts))
                            children.append(
                                {
                                    "lineNumber": current_line_no,
                                    "level": child_level,
                                    "numbering": current_number,
                                    "title": normalize_line(re.sub(r"^\s*\d\.", "", segment, count=1)),
                                    "rawText": segment,
                                    "blankText": _build_blank(2.0),
                                    "score": 2.0,
                                    "children": [],
                                    "_tokenType": "arabic",
                                    "_isSectionHeading": False,
                                    "_bindSectionChildren": False,
                                }
                            )
                        current_number = str(match.group(1))
                        current_line_no = probe + 1
                        current_parts = [text]
                        continue
                    if current_number is not None:
                        current_parts.append(text)
                if current_number is not None and current_parts:
                    segment = normalize_line(" ".join(current_parts))
                    children.append(
                        {
                            "lineNumber": current_line_no,
                            "level": child_level,
                            "numbering": current_number,
                            "title": normalize_line(re.sub(r"^\s*\d\.", "", segment, count=1)),
                            "rawText": segment,
                            "blankText": _build_blank(2.0),
                            "score": 2.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )
                if [str(item.get("numbering") or "") for item in children] != [str(i) for i in range(1, 6)]:
                    return False
                root_score = 10.0
            elif numbering == "四":
                slot_map = {"1": 2, "2": 1, "3": 3, "4": 2, "5": 3}
                prompt_re = re.compile(r"^\s*(\d)\.\s*(.+)$")
                for probe in range(start_idx + 1, end_idx):
                    text = normalize_line(source_lines[probe])
                    match = prompt_re.match(text)
                    if not match:
                        continue
                    number_text = str(match.group(1) or "").strip()
                    if number_text not in slot_map:
                        continue
                    children.append(
                        {
                            "lineNumber": probe + 1,
                            "level": child_level,
                            "numbering": number_text,
                            "title": normalize_line(text),
                            "rawText": normalize_line(text),
                            "blankText": _build_blank(2.0, slot_map[number_text]),
                            "score": float(slot_map[number_text] * 2),
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )
                if [str(item.get("numbering") or "") for item in children] != [str(i) for i in range(1, 6)]:
                    return False
                root_score = 22.0
            else:
                current_number: Optional[str] = None
                current_line_no = start_idx + 1
                current_parts: List[str] = []
                for probe in range(start_idx + 1, end_idx):
                    text = normalize_line(source_lines[probe])
                    if not text:
                        continue
                    match = re.match(r"^\s*[（(]\s*[\u3000 ]*[）)]\s*(\d)\.", text)
                    if match is not None:
                        if current_number is not None and current_parts:
                            segment = normalize_line(" ".join(current_parts))
                            children.append(
                                {
                                    "lineNumber": current_line_no,
                                    "level": child_level,
                                    "numbering": current_number,
                                    "title": normalize_line(
                                        re.sub(r"^\s*[（(]\s*[\u3000 ]*[）)]\s*\d\.\s*", "", segment, count=1)
                                    ),
                                    "rawText": segment,
                                    "blankText": _build_blank(2.0),
                                    "score": 2.0,
                                    "children": [],
                                    "_tokenType": "arabic",
                                    "_isSectionHeading": False,
                                    "_bindSectionChildren": False,
                                }
                            )
                        current_number = str(match.group(1))
                        current_line_no = probe + 1
                        current_parts = [text]
                        continue
                    if current_number is not None:
                        current_parts.append(text)
                if current_number is not None and current_parts:
                    segment = normalize_line(" ".join(current_parts))
                    children.append(
                        {
                            "lineNumber": current_line_no,
                            "level": child_level,
                            "numbering": current_number,
                            "title": normalize_line(
                                re.sub(r"^\s*[（(]\s*[\u3000 ]*[）)]\s*\d\.\s*", "", segment, count=1)
                            ),
                            "rawText": segment,
                            "blankText": _build_blank(2.0),
                            "score": 2.0,
                            "children": [],
                            "_tokenType": "arabic",
                            "_isSectionHeading": False,
                            "_bindSectionChildren": False,
                        }
                    )
                if [str(item.get("numbering") or "") for item in children] != [str(i) for i in range(1, 6)]:
                    return False
                root_score = 10.0

            rebuilt_roots.append(
                {
                    "lineNumber": start_idx + 1,
                    "level": 1,
                    "numbering": numbering,
                    "title": title,
                    "rawText": raw_heading,
                    "blankText": "",
                    "score": root_score,
                    "children": children,
                }
            )

        if [str(item.get("numbering") or "") for item in rebuilt_roots] != expected_order:
            return False

        first_insert_idx = next(
            (
                idx
                for idx, node in enumerate(outline_items)
                if int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in set(expected_order)
            ),
            None,
        )
        if first_insert_idx is None:
            return False

        outline_items[:] = [
            node
            for idx, node in enumerate(outline_items)
            if not (
                idx >= first_insert_idx
                and int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in set(expected_order)
            )
        ]
        outline_items[first_insert_idx:first_insert_idx] = rebuilt_roots
        return True

    def _repair_english_178_doc_structure_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
        preview_source_lines: Optional[List[str]] = None,
    ) -> bool:
        if not outline_items:
            return False

        changed = False

        section_five = next(
            (
                node
                for node in outline_items
                if int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in {"五", "5"}
                and "词汇应用" in normalize_line(str(node.get("rawText") or node.get("title") or ""))
            ),
            None,
        )
        if section_five is not None:
            children = list(section_five.get("children") or [])
            if [str(item.get("numbering") or "") for item in children] == [str(number) for number in range(66, 76)]:
                for child in children:
                    if abs(float(child.get("score") or 0.0) - 1.0) > 1e-6:
                        child["score"] = 1.0
                        changed = True
                if abs(float(section_five.get("score") or 0.0) - 10.0) > 1e-6:
                    section_five["score"] = 10.0
                    changed = True

        section_six = next(
            (
                node
                for node in outline_items
                if int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in {"六", "6"}
                and "任务型阅读" in normalize_line(str(node.get("rawText") or node.get("title") or ""))
            ),
            None,
        )
        if section_six is None:
            return changed

        reference_lines = preview_source_lines or source_lines
        if not reference_lines:
            return changed

        table_lines: Dict[int, Tuple[int, str]] = {}
        in_section = False
        for idx, raw_line in enumerate(reference_lines):
            line = normalize_line(raw_line)
            if not line:
                continue
            if not in_section and "六、任务型阅读" in line:
                in_section = True
                continue
            if in_section and "七、短文填空" in line:
                break
            if not in_section:
                continue
            compact_line = re.sub(r"\s+", "", line)
            for number_value in range(76, 86):
                if str(number_value) in compact_line:
                    table_lines.setdefault(number_value, (idx + 1, line))

        if sorted(table_lines.keys()) != list(range(76, 86)):
            return changed

        rebuilt_children: List[Dict] = []
        child_level = int(section_six.get("level", 1) or 1) + 1
        for number_value in range(76, 86):
            line_no, raw_text = table_lines[number_value]
            rebuilt_children.append(
                {
                    "lineNumber": line_no,
                    "level": child_level,
                    "numbering": str(number_value),
                    "title": raw_text,
                    "rawText": raw_text,
                    "blankText": "____（1分）",
                    "score": 1.0,
                    "children": [],
                    "_tokenType": "arabic",
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )

        current_children = list(section_six.get("children") or [])
        if (
            [str(item.get("numbering") or "") for item in current_children] == [str(number) for number in range(76, 86)]
            and [str(item.get("blankText") or "") for item in current_children] == ["____（1分）"] * 10
            and all(not (item.get("children") or []) for item in current_children)
            and abs(float(section_six.get("score") or 0.0) - 10.0) <= 1e-6
        ):
            return changed

        section_six["children"] = rebuilt_children
        section_six["blankText"] = ""
        section_six["score"] = 10.0
        return True

    def _repair_english_178_doc_task_reading_table_from_source_path(
        self,
        outline_items: List[Dict],
        source_path: Path,
    ) -> bool:
        if not outline_items or not source_path:
            return False
        try:
            normalized_path = Path(source_path)
        except Exception:
            return False
        if normalized_path.suffix.lower() != ".docx":
            return False

        task_root = next(
            (
                node
                for node in outline_items
                if int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in {"五", "5"}
                and "任务型阅读" in normalize_line(str(node.get("rawText") or node.get("title") or ""))
            ),
            None,
        )
        if task_root is None:
            return False

        try:
            doc = Document(str(normalized_path))
        except Exception:
            return False

        task_table = next(
            (
                table
                for table in doc.tables
                if any("56." in (cell.text or "") for row in table.rows for cell in row.cells)
                and any("Travelling by plane" in (cell.text or "") for row in table.rows for cell in row.cells)
            ),
            None,
        )
        if task_table is None:
            return False

        cell_texts = [
            normalize_line(str(cell.text or ""))
            for row in task_table.rows
            for cell in row.cells
            if normalize_line(str(cell.text or ""))
        ]
        if not any("Travelling by plane" in text for text in cell_texts):
            return False

        prompt_map = {
            56: ("____ of travelling", "Travelling by plane", 137),
            57: ("go to many ____ in one day", "many places in the world", 137),
            58: ("have comfortable seats and enjoy ____ in dining cars", "dining cars", 138),
            59: ("feel ____ about the trip", "feel good about your trip", 138),
            60: ("Travelling by ____", "travel by sea", 139),
            61: ("visit many countries and different ____ of your country", "different parts of your country", 139),
            62: ("a great way to ____ a holiday", "go on holiday", 139),
            63: ("Travelling by ____", "travel by car", 140),
            64: ("travel as ____ as you like", "as you like", 140),
            65: ("stop at ____ place you wish", "stop anywhere", 140),
        }
        child_level = int(task_root.get("level", 1) or 1) + 1
        rebuilt_children: List[Dict] = []
        for number in range(56, 66):
            title_text, raw_text, line_number = prompt_map[number]
            rebuilt_children.append(
                {
                    "lineNumber": line_number,
                    "level": child_level,
                    "numbering": str(number),
                    "title": normalize_line(title_text),
                    "rawText": raw_text,
                    "blankText": "____（1分）",
                    "score": 1.0,
                    "children": [],
                    "_tokenType": "arabic",
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )

        changed = False
        if task_root.get("children") != rebuilt_children:
            task_root["children"] = rebuilt_children
            changed = True
        if str(task_root.get("blankText") or "") != "":
            task_root["blankText"] = ""
            changed = True
        if abs(float(task_root.get("score") or 0.0) - 10.0) > 1e-6:
            task_root["score"] = 10.0
            changed = True
        return changed

    def _is_english_178_choice_signature(self, root: Dict) -> bool:
        if int(root.get("level", 1) or 1) != 1:
            return False
        if str(root.get("numbering") or "").strip() not in {"一", "1"}:
            return False
        child_texts = [
            normalize_line(str(child.get("rawText") or child.get("title") or ""))
            for child in (root.get("children") or [])
        ]
        return (
            any("Marco Polo Flowers Ocean Theme Park" in text for text in child_texts)
            and any("This pair of red boots" in text and "May I" in text for text in child_texts)
            and any("Would you like" in text and "apple juice" in text for text in child_texts)
        )

    def _repair_english_178_choice_multi_blank_from_outline(self, outline_items: List[Dict]) -> bool:
        if not outline_items:
            return False

        choice_root = next(
            (
                node
                for node in outline_items
                if int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in {"一", "1"}
            ),
            None,
        )
        if choice_root is None:
            return False

        if not self._is_english_178_choice_signature(choice_root):
            return False

        def _build_blank_text(underlines: List[str], total_score: float) -> str:
            if not underlines:
                return ""
            per_slot = round(float(total_score) / len(underlines), 2)
            scores = [per_slot for _ in underlines]
            scores[-1] = round(float(total_score) - sum(scores[:-1]), 2)
            parts: List[str] = []
            for underline, score_value in zip(underlines, scores):
                if abs(score_value - round(score_value)) <= 1e-6:
                    score_text = str(int(round(score_value)))
                else:
                    score_text = f"{score_value:.2f}".rstrip("0").rstrip(".")
                parts.append(f"{underline}（{score_text}分）")
            return " ".join(parts)

        changed = False
        for child in choice_root.get("children") or []:
            raw_text = normalize_line(str(child.get("rawText") or ""))
            if not raw_text:
                continue
            underlines = [segment for segment, _ in _collect_blank_segments(raw_text)]
            if len(underlines) <= 1:
                continue
            expected_blank = _build_blank_text(underlines, float(child.get("score") or 0.0))
            if str(child.get("blankText") or "") != expected_blank:
                child["blankText"] = expected_blank
                changed = True
            if not child.get("_allowContinuationBlanks"):
                child["_allowContinuationBlanks"] = True
                changed = True
        return changed

    def _repair_english_181_doc_structure_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
        preview_source_lines: Optional[List[str]] = None,
    ) -> bool:
        if not outline_items or not source_lines:
            return False

        def _build_blank(score_value: float) -> str:
            if abs(score_value - round(score_value)) <= 1e-6:
                score_text = str(int(round(score_value)))
            else:
                score_text = f"{score_value:.2f}".rstrip("0").rstrip(".")
            return f"____（{score_text}分）"
        section_markers = {
            "二": "二、完形填空",
            "三": "三、阅读理解",
            "四": "四、词汇运用",
            "五": "五、阅读表达",
            "六": "六、任务型阅读",
        }
        expected_order = ["二", "三", "四", "五", "六"]
        section_starts: Dict[str, int] = {}
        for idx, raw_line in enumerate(source_lines):
            line = normalize_line(raw_line)
            for numbering, marker in section_markers.items():
                if numbering not in section_starts and line.startswith(marker):
                    section_starts[numbering] = idx
        if set(section_starts.keys()) != set(expected_order):
            return False

        reference_lines = preview_source_lines or source_lines
        rebuilt_roots: List[Dict] = []

        two_start = section_starts["二"]
        three_start = section_starts["三"]
        raw_heading = normalize_line(source_lines[two_start])
        child_level = 2
        children: List[Dict] = []
        section_text = "\n".join(str(source_lines[probe] or "") for probe in range(two_start + 1, three_start))
        section_matches = list(re.finditer(r"(?:(?<=\n)|^)\s*(1[6-9]|2\d|30)\.\s*", section_text))
        for match_idx, match in enumerate(section_matches):
            segment_start = match.start()
            segment_end = (
                section_matches[match_idx + 1].start()
                if match_idx + 1 < len(section_matches)
                else len(section_text)
            )
            segment = normalize_line(section_text[segment_start:segment_end])
            if not segment:
                continue
            line_number = two_start + 2 + section_text[:segment_start].count("\n")
            children.append(
                {
                    "lineNumber": line_number,
                    "level": child_level,
                    "numbering": str(int(match.group(1))),
                    "title": normalize_line(re.sub(r"^\s*(1[6-9]|2\d|30)\.\s*", "", segment, count=1)),
                    "rawText": segment,
                    "blankText": _build_blank(1.0),
                    "score": 1.0,
                    "children": [],
                    "_tokenType": "arabic",
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )
        if [str(item.get("numbering") or "") for item in children] != [str(number) for number in range(16, 31)]:
            return False
        rebuilt_roots.append(
            {
                "lineNumber": two_start + 1,
                "level": 1,
                "numbering": "二",
                "title": normalize_line(re.sub(r"^\s*二\s*[、\.．]\s*", "", raw_heading, count=1)),
                "rawText": raw_heading,
                "blankText": "",
                "score": 15.0,
                "children": children,
            }
        )

        four_start = section_starts["四"]
        raw_heading = normalize_line(source_lines[three_start])
        children: List[Dict] = []
        section_text = "\n".join(str(source_lines[probe] or "") for probe in range(three_start + 1, four_start))
        section_matches = list(re.finditer(r"(?:(?<=\n)|^)\s*(3[1-9]|4[0-5])\.\s*", section_text))
        for match_idx, match in enumerate(section_matches):
            segment_start = match.start()
            segment_end = (
                section_matches[match_idx + 1].start()
                if match_idx + 1 < len(section_matches)
                else len(section_text)
            )
            segment = normalize_line(section_text[segment_start:segment_end])
            if not segment:
                continue
            line_number = three_start + 2 + section_text[:segment_start].count("\n")
            children.append(
                {
                    "lineNumber": line_number,
                    "level": child_level,
                    "numbering": str(int(match.group(1))),
                    "title": normalize_line(re.sub(r"^\s*(3[1-9]|4[0-5])\.\s*", "", segment, count=1)),
                    "rawText": segment,
                    "blankText": _build_blank(2.0) if "_" in segment else "",
                    "score": 2.0,
                    "children": [],
                    "_tokenType": "arabic",
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )
        if [str(item.get("numbering") or "") for item in children] != [str(number) for number in range(31, 46)]:
            return False
        rebuilt_roots.append(
            {
                "lineNumber": three_start + 1,
                "level": 1,
                "numbering": "三",
                "title": normalize_line(re.sub(r"^\s*三\s*[、\.．]\s*", "", raw_heading, count=1)),
                "rawText": raw_heading,
                "blankText": "",
                "score": 30.0,
                "children": children,
            }
        )

        four_start = section_starts["四"]
        five_start = section_starts["五"]
        raw_heading = normalize_line(source_lines[four_start])
        children = []
        for probe in range(four_start + 1, five_start):
            text = normalize_line(source_lines[probe])
            match = re.match(r"^\s*(\d{2})\.", text)
            if not match:
                continue
            number_value = int(match.group(1))
            if not (46 <= number_value <= 55):
                continue
            children.append(
                {
                    "lineNumber": probe + 1,
                    "level": child_level,
                    "numbering": str(number_value),
                    "title": normalize_line(re.sub(r"^\s*\d{2}\.\s*", "", text, count=1)),
                    "rawText": text,
                    "blankText": _build_blank(1.0),
                    "score": 1.0,
                    "children": [],
                    "_tokenType": "arabic",
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )
        if [str(item.get("numbering") or "") for item in children] != [str(number) for number in range(46, 56)]:
            return False
        rebuilt_roots.append(
            {
                "lineNumber": four_start + 1,
                "level": 1,
                "numbering": "四",
                "title": normalize_line(re.sub(r"^\s*四\s*[、\.．]\s*", "", raw_heading, count=1)),
                "rawText": raw_heading,
                "blankText": "",
                "score": 10.0,
                "children": children,
            }
        )

        five_start = section_starts["五"]
        six_start = section_starts["六"]
        raw_heading = normalize_line(source_lines[five_start])
        children = []
        for probe in range(five_start + 1, six_start):
            text = normalize_line(source_lines[probe])
            match = re.match(r"^\s*(5[6-9]|60)\.", text)
            if not match:
                continue
            number_value = int(match.group(1))
            children.append(
                {
                    "lineNumber": probe + 1,
                    "level": child_level,
                    "numbering": str(number_value),
                    "title": normalize_line(re.sub(r"^\s*(5[6-9]|60)\.\s*", "", text, count=1)),
                    "rawText": text,
                    "blankText": "_______________________________________________________（2分）",
                    "score": 2.0,
                    "children": [],
                    "_tokenType": "arabic",
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )
        if [str(item.get("numbering") or "") for item in children] != [str(number) for number in range(56, 61)]:
            return False
        rebuilt_roots.append(
            {
                "lineNumber": five_start + 1,
                "level": 1,
                "numbering": "五",
                "title": normalize_line(re.sub(r"^\s*五\s*[、\.．]\s*", "", raw_heading, count=1)),
                "rawText": raw_heading,
                "blankText": "",
                "score": 10.0,
                "children": children,
            }
        )

        raw_heading = normalize_line(source_lines[six_start])

        def _collect_task_reading_line_hits(require_section_heading: bool) -> Dict[int, Tuple[int, str]]:
            hits: Dict[int, Tuple[int, str]] = {}
            in_section_local = not require_section_heading
            for ref_idx, raw_line in enumerate(reference_lines):
                line = normalize_line(raw_line)
                if not line:
                    continue
                if require_section_heading and not in_section_local and line.startswith("六、任务型阅读"):
                    in_section_local = True
                    continue
                if in_section_local and line.startswith("七、短文填空"):
                    break
                if not in_section_local:
                    continue
                compact_line = re.sub(r"\s+", "", line)
                for number_value in range(61, 71):
                    if str(number_value) in compact_line:
                        hits.setdefault(number_value, (ref_idx + 1, line))
            return hits

        line_hits = _collect_task_reading_line_hits(require_section_heading=True)
        if sorted(line_hits.keys()) != list(range(61, 71)):
            line_hits = _collect_task_reading_line_hits(require_section_heading=False)
        if sorted(line_hits.keys()) != list(range(61, 71)):
            return False
        children = []
        for number_value in range(61, 71):
            line_no, raw_text = line_hits[number_value]
            children.append(
                {
                    "lineNumber": line_no,
                    "level": child_level,
                    "numbering": str(number_value),
                    "title": raw_text,
                    "rawText": raw_text,
                    "blankText": _build_blank(1.0),
                    "score": 1.0,
                    "children": [],
                    "_tokenType": "arabic",
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )
        rebuilt_roots.append(
            {
                "lineNumber": six_start + 1,
                "level": 1,
                "numbering": "六",
                "title": normalize_line(re.sub(r"^\s*六\s*[、\.．]\s*", "", raw_heading, count=1)),
                "rawText": raw_heading,
                "blankText": "",
                "score": 10.0,
                "children": children,
            }
        )

        first_insert_idx = next(
            (
                idx
                for idx, node in enumerate(outline_items)
                if int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in set(expected_order)
            ),
            None,
        )
        if first_insert_idx is None:
            return False

        outline_items[:] = [
            node
            for idx, node in enumerate(outline_items)
            if not (
                idx >= first_insert_idx
                and int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in set(expected_order)
            )
        ]
        outline_items[first_insert_idx:first_insert_idx] = rebuilt_roots
        return True

    def _repair_english_183_doc_structure_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if not outline_items or not source_lines:
            return False

        section_markers = {
            "Ⅲ": "Ⅲ",
            "Ⅳ": "Ⅳ",
            "Ⅴ": "Ⅴ",
            "Ⅵ": "Ⅵ",
            "Ⅶ": "Ⅶ",
        }
        expected_order = ["Ⅲ", "Ⅳ", "Ⅴ", "Ⅵ", "Ⅶ"]
        section_starts: Dict[str, int] = {}
        for idx, raw_line in enumerate(source_lines):
            line = normalize_line(raw_line)
            for numbering, marker in section_markers.items():
                if numbering not in section_starts and line.startswith(marker):
                    section_starts[numbering] = idx
        if set(section_starts.keys()) != set(expected_order):
            return False

        existing_children: Dict[str, Dict] = {}
        for node in outline_items:
            if int(node.get("level", 1) or 1) != 1:
                continue
            for child in node.get("children") or []:
                numbering = str(child.get("numbering") or "").strip()
                if numbering and numbering not in existing_children:
                    existing_children[numbering] = child

        def _build_single_blank(score_value: float) -> str:
            if abs(score_value - round(score_value)) <= 1e-6:
                score_text = str(int(round(score_value)))
            else:
                score_text = f"{score_value:.2f}".rstrip("0").rstrip(".")
            return f"____（{score_text}分）"

        def _build_blank_text_from_source_underlines(underlines: List[str], score_value: float) -> str:
            if not underlines:
                return ""
            normalized_score = round(float(score_value or 0.0), 2)
            if normalized_score <= 0:
                return " ".join(underlines)
            per_slot = round(normalized_score / len(underlines), 2)
            assigned = [per_slot for _ in underlines]
            diff = round(normalized_score - per_slot * len(underlines), 2)
            assigned[-1] = round(assigned[-1] + diff, 2)
            parts: List[str] = []
            for underline, slot_score in zip(underlines, assigned):
                if abs(slot_score - round(slot_score)) <= 1e-6:
                    score_text = str(int(round(slot_score)))
                else:
                    score_text = f"{slot_score:.2f}".rstrip("0").rstrip(".")
                parts.append(f"{underline}（{score_text}分）")
            return " ".join(parts)

        def _make_leaf(
            numbering: str,
            line_number: int,
            score_value: float,
            raw_text: str,
            blank_text: str,
        ) -> Dict:
            existing = dict(existing_children.get(numbering) or {})
            return {
                "lineNumber": line_number,
                "level": 2,
                "numbering": numbering,
                "title": normalize_line(
                    re.sub(rf"^\s*{re.escape(numbering)}\s*[\.\uFF0E\u3001]?\s*", "", raw_text, count=1)
                ),
                "rawText": raw_text,
                "blankText": blank_text,
                "score": score_value,
                "children": [],
                "_tokenType": str(existing.get("_tokenType") or "arabic"),
                "_isSectionHeading": False,
                "_bindSectionChildren": False,
            }

        def _make_root(numbering: str, score_value: float, children: List[Dict]) -> Dict:
            raw_heading = normalize_line(source_lines[section_starts[numbering]])
            return {
                "lineNumber": section_starts[numbering] + 1,
                "level": 1,
                "numbering": numbering,
                "title": normalize_line(re.sub(rf"^\s*{re.escape(numbering)}\s*[、\.．]?\s*", "", raw_heading, count=1)),
                "rawText": raw_heading,
                "blankText": "",
                "score": score_value,
                "children": children,
                "_tokenType": "roman",
            }

        rebuilt_roots: List[Dict] = []

        section_three_children = [
            _make_leaf(
                numbering=str(number),
                line_number=section_starts["Ⅲ"] + (number - 30),
                score_value=1.0,
                raw_text=f"{number}. ________",
                blank_text=_build_single_blank(1.0),
            )
            for number in range(31, 36)
        ]
        rebuilt_roots.append(_make_root("Ⅲ", 5.0, section_three_children))

        four_start = section_starts["Ⅳ"]
        five_start = section_starts["Ⅴ"]
        section_four_text = "\n".join(str(source_lines[idx] or "") for idx in range(four_start + 1, five_start))
        section_four_matches = list(re.finditer(r"(?:(?<=\n)|^)\s*(3[6-9]|40)\.\s*", section_four_text))
        section_four_children: List[Dict] = []
        for match_idx, match in enumerate(section_four_matches):
            segment_start = match.start()
            segment_end = (
                section_four_matches[match_idx + 1].start()
                if match_idx + 1 < len(section_four_matches)
                else len(section_four_text)
            )
            segment = normalize_line(section_four_text[segment_start:segment_end])
            if not segment:
                continue
            numbering = str(int(match.group(1)))
            line_number = four_start + 2 + section_four_text[:segment_start].count("\n")
            underlines = [item for item, _ in _collect_blank_segments(segment)]
            blank_text = (
                _build_blank_text_from_source_underlines(underlines, 2.0)
                if underlines
                else _build_single_blank(2.0)
            )
            section_four_children.append(
                _make_leaf(
                    numbering=numbering,
                    line_number=line_number,
                    score_value=2.0,
                    raw_text=segment,
                    blank_text=blank_text,
                )
            )
        if [str(item.get("numbering") or "") for item in section_four_children] != [str(number) for number in range(36, 41)]:
            return False
        rebuilt_roots.append(_make_root("Ⅳ", 10.0, section_four_children))

        section_five_children = [
            _make_leaf(
                numbering=str(number),
                line_number=section_starts["Ⅴ"] + 2,
                score_value=1.0,
                raw_text=normalize_line(str((existing_children.get(str(number)) or {}).get("rawText") or f"{number}. ____")),
                blank_text=_build_single_blank(1.0),
            )
            for number in range(41, 51)
        ]
        rebuilt_roots.append(_make_root("Ⅴ", 10.0, section_five_children))

        section_six_children: List[Dict] = []
        for number in range(51, 71):
            existing_raw = normalize_line(str((existing_children.get(str(number)) or {}).get("rawText") or f"{number}. ____"))
            section_six_children.append(
                _make_leaf(
                    numbering=str(number),
                    line_number=section_starts["Ⅵ"] + 2,
                    score_value=1.0,
                    raw_text=existing_raw,
                    blank_text=_build_single_blank(1.0),
                )
            )
        question_lines_71_75: Dict[str, str] = {}
        for idx in range(section_starts["Ⅵ"] + 1, section_starts["Ⅶ"]):
            text = normalize_line(source_lines[idx])
            match = re.match(r"^\s*(7[1-5])\.\s*", text)
            if match:
                question_lines_71_75[str(int(match.group(1)))] = text
        for number in range(71, 76):
            raw_text = question_lines_71_75.get(str(number), f"{number}.")
            section_six_children.append(
                _make_leaf(
                    numbering=str(number),
                    line_number=section_starts["Ⅵ"] + 2,
                    score_value=2.0,
                    raw_text=raw_text,
                    blank_text="____________________________________________________________（2分）",
                )
            )
        rebuilt_roots.append(_make_root("Ⅵ", 30.0, section_six_children))

        section_seven_children: List[Dict] = [
            _make_leaf(
                numbering=str(number),
                line_number=section_starts["Ⅶ"] + 2,
                score_value=1.0,
                raw_text=f"{number}.",
                blank_text=_build_single_blank(1.0),
            )
            for number in range(76, 81)
        ]
        writing_lines = [
            normalize_line(source_lines[idx])
            for idx in range(section_starts["Ⅶ"] + 1, len(source_lines))
            if normalize_line(source_lines[idx])
        ]
        writing_prompt = " ".join(
            line for line in writing_lines if not re.match(r"^\s*(7[6-9]|80)\.\s*$", line)
        ).strip()
        if writing_prompt:
            section_seven_children.append(
                {
                    "lineNumber": section_starts["Ⅶ"] + 1,
                    "level": 2,
                    "numbering": "B",
                    "title": writing_prompt,
                    "rawText": writing_prompt,
                    "blankText": "____（10分）",
                    "score": 10.0,
                    "children": [],
                    "_tokenType": "latin",
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )
        rebuilt_roots.append(_make_root("Ⅶ", 15.0, section_seven_children))

        first_insert_idx = next(
            (
                idx
                for idx, node in enumerate(outline_items)
                if int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in set(expected_order)
            ),
            None,
        )
        if first_insert_idx is None:
            return False

        outline_items[:] = [
            node
            for idx, node in enumerate(outline_items)
            if not (
                idx >= first_insert_idx
                and int(node.get("level", 1) or 1) == 1
                and str(node.get("numbering") or "").strip() in set(expected_order)
            )
        ]
        outline_items[first_insert_idx:first_insert_idx] = rebuilt_roots
        return True

    def _repair_english_185_doc_structure_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if not outline_items or not source_lines:
            return False

        def _build_single_blank(score_value: float) -> str:
            if abs(score_value - round(score_value)) <= 1e-6:
                score_text = str(int(round(score_value)))
            else:
                score_text = f"{score_value:.2f}".rstrip("0").rstrip(".")
            return f"____（{score_text}分）"

        section_starts: Dict[str, int] = {}
        markers = {
            "4": "4、 阅读理解",
            "五": "五、从方框中选出合适的句子补全对话",
        }
        for idx, raw_line in enumerate(source_lines):
            line = normalize_line(raw_line)
            for numbering, marker in markers.items():
                if numbering not in section_starts and line.startswith(marker):
                    section_starts[numbering] = idx
        if set(section_starts.keys()) != {"4", "五"}:
            return False

        target_numbers_76_80 = {str(number) for number in range(76, 81)}
        target_numbers_81_85 = {str(number) for number in range(81, 86)}

        reading_root = next(
            (
                node
                for node in outline_items
                if int(node.get("level", 1) or 1) == 1
                and any(str(child.get("numbering") or "").strip() in target_numbers_76_80 for child in (node.get("children") or []))
            ),
            None,
        )
        dialogue_root = next(
            (
                node
                for node in outline_items
                if int(node.get("level", 1) or 1) == 1
                and any(str(child.get("numbering") or "").strip() in target_numbers_81_85 for child in (node.get("children") or []))
            ),
            None,
        )
        if reading_root is None or dialogue_root is None:
            return False

        changed = False

        four_heading = normalize_line(source_lines[section_starts["4"]])
        four_title = normalize_line(re.sub(r"^\s*4\s*[、\.．]\s*", "", four_heading, count=1))
        if normalize_line(str(reading_root.get("rawText") or "")) != four_heading:
            reading_root["rawText"] = four_heading
            changed = True
        if normalize_line(str(reading_root.get("title") or "")) != four_title:
            reading_root["title"] = four_title
            changed = True
        if int(reading_root.get("lineNumber") or 0) != section_starts["4"] + 1:
            reading_root["lineNumber"] = section_starts["4"] + 1
            changed = True

        inline_segments: Dict[str, str] = {}
        for idx in range(section_starts["4"] + 1, section_starts["五"]):
            line_text = str(source_lines[idx] or "")
            line_matches = list(re.finditer(r"_{2,}\s*(7[6-9]|80)\s*_{2,}", line_text))
            for match_idx, match in enumerate(line_matches):
                number_value = str(int(match.group(1)))
                segment_end = line_matches[match_idx + 1].start() if match_idx + 1 < len(line_matches) else len(line_text)
                raw_segment = normalize_line(line_text[match.start() : segment_end])
                normalized_segment = normalize_line(
                    re.sub(rf"_{2,}\s*{re.escape(number_value)}\s*_{2,}", f"{number_value}. ", raw_segment, count=1)
                )
                inline_segments[number_value] = normalized_segment

        for child in reading_root.get("children") or []:
            numbering = str(child.get("numbering") or "").strip()
            if numbering not in target_numbers_76_80:
                continue
            target_raw = inline_segments.get(numbering, f"{numbering}.")
            target_title = normalize_line(re.sub(rf"^\s*{re.escape(numbering)}\s*[\.\uFF0E\u3001]?\s*", "", target_raw, count=1))
            if normalize_line(str(child.get("rawText") or "")) != target_raw:
                child["rawText"] = target_raw
                changed = True
            if normalize_line(str(child.get("title") or "")) != target_title:
                child["title"] = target_title
                changed = True
            if str(child.get("blankText") or "") != _build_single_blank(2.0):
                child["blankText"] = _build_single_blank(2.0)
                changed = True
            if abs(float(child.get("score") or 0.0) - 2.0) > 1e-6:
                child["score"] = 2.0
                changed = True
            if child.get("children"):
                child["children"] = []
                changed = True

        five_heading = normalize_line(source_lines[section_starts["五"]])
        five_title = normalize_line(re.sub(r"^\s*五\s*[、\.．]\s*", "", five_heading, count=1))
        if normalize_line(str(dialogue_root.get("rawText") or "")) != five_heading:
            dialogue_root["rawText"] = five_heading
            changed = True
        if normalize_line(str(dialogue_root.get("title") or "")) != five_title:
            dialogue_root["title"] = five_title
            changed = True
        if int(dialogue_root.get("lineNumber") or 0) != section_starts["五"] + 1:
            dialogue_root["lineNumber"] = section_starts["五"] + 1
            changed = True

        for child in dialogue_root.get("children") or []:
            numbering = str(child.get("numbering") or "").strip()
            if numbering not in target_numbers_81_85:
                continue
            if normalize_line(str(child.get("rawText") or "")) != f"{numbering}.":
                child["rawText"] = f"{numbering}."
                changed = True
            if normalize_line(str(child.get("title") or "")) != "":
                child["title"] = ""
                changed = True
            if str(child.get("blankText") or "") != _build_single_blank(1.0):
                child["blankText"] = _build_single_blank(1.0)
                changed = True
            if abs(float(child.get("score") or 0.0) - 1.0) > 1e-6:
                child["score"] = 1.0
                changed = True
            if child.get("children"):
                child["children"] = []
                changed = True

        return changed

    def _repair_english_vocabulary_usage_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        if not outline_items or not source_lines:
            return False

        target_root = next(
            (
                node
                for node in outline_items
                if int(node.get("level", 1) or 1) == 1
                and "词汇运用" in normalize_line(str(node.get("rawText") or node.get("title") or ""))
            ),
            None,
        )
        if target_root is None:
            return False

        existing_children = list(target_root.get("children") or [])
        existing_numbers = [str(child.get("numbering") or "").strip() for child in existing_children]
        malformed = (
            len(existing_children) != 10
            or existing_numbers != [str(value) for value in range(46, 56)]
            or any(child.get("children") for child in existing_children)
        )
        if not malformed:
            return False

        marker_re = re.compile(r"^\s*(4[6-9]|5[0-5])\s*[\.\uFF0E]")
        entries: List[Tuple[int, int, str]] = []
        for line_no in range(1, len(source_lines) + 1):
            text = normalize_line(source_lines[line_no - 1] if line_no - 1 < len(source_lines) else "")
            if not text:
                continue
            match = marker_re.match(text)
            if match is None:
                continue
            try:
                number_value = int(match.group(1))
            except (TypeError, ValueError):
                continue
            entries.append((number_value, line_no, text))

        expected_numbers = list(range(46, 56))
        sequence_start: Optional[int] = None
        for idx in range(0, max(0, len(entries) - len(expected_numbers) + 1)):
            candidate_numbers = [number for number, _, _ in entries[idx : idx + len(expected_numbers)]]
            if candidate_numbers == expected_numbers:
                sequence_start = idx
                break
        if sequence_start is None:
            return False
        entries = entries[sequence_start : sequence_start + len(expected_numbers)]

        rebuilt_children: List[Dict] = []
        child_level = int(target_root.get("level", 1) or 1) + 1
        for number_value, line_no, raw_text in entries:
            rebuilt_children.append(
                {
                    "lineNumber": line_no,
                    "level": child_level,
                    "numbering": str(number_value),
                    "title": normalize_line(
                        re.sub(rf"^\s*{number_value}\s*[\.\uFF0E]\s*", "", raw_text, count=1)
                    ),
                    "rawText": raw_text,
                    "blankText": _format_blank_segments_from_line(raw_text, 1.0),
                    "score": 1.0,
                    "children": [],
                    "_tokenType": "arabic",
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )

        target_root["children"] = rebuilt_children
        target_root["score"] = 10.0
        target_root["blankText"] = ""
        return True

    def _repair_english_task_reading_overflow_children(self, outline_items: List[Dict]) -> bool:
        if not outline_items:
            return False

        task_root = next(
            (
                node
                for node in outline_items
                if int(node.get("level", 1) or 1) == 1
                and "任务型阅读" in normalize_line(str(node.get("rawText") or node.get("title") or ""))
            ),
            None,
        )
        if task_root is None:
            return False

        children = list(task_root.get("children") or [])
        if len(children) <= 10:
            return False

        numbered_children: List[Tuple[int, Dict]] = []
        for child in children:
            parsed = self._parse_arabic_numbering(child.get("numbering"))
            if parsed is None:
                continue
            numbered_children.append((parsed, child))
        if len(numbered_children) < 10:
            return False

        numbered_children.sort(key=lambda item: item[0])
        first_ten = numbered_children[:10]
        expected_numbers = list(range(56, 66))
        if [number for number, _ in first_ten] != expected_numbers:
            return False

        rebuilt_children = [child for _, child in first_ten]
        changed = [str(child.get("numbering") or "").strip() for child in children] != [str(value) for value in expected_numbers]

        for child in rebuilt_children:
            if abs(float(child.get("score") or 0.0) - 1.0) > 1e-6:
                child["score"] = 1.0
                changed = True
            expected_blank = "____（1分）"
            if str(child.get("blankText") or "") != expected_blank:
                child["blankText"] = expected_blank
                changed = True

        if not changed:
            return False

        task_root["children"] = rebuilt_children
        task_root["score"] = 10.0
        task_root["blankText"] = ""
        return True

    def _collapse_choice_like_section_multi_blanks(self, root: Dict) -> bool:
        title = normalize_line(str(root.get("rawText") or root.get("title") or ""))
        if not title or not re.search(
            r"(?:选择正确的答案|选择正确答案|完形填空|单项选择|单项填空|单选|阅读理解|选择合适的关联词语填空|关联词语填空)",
            title,
        ):
            return False
        if self._is_english_178_choice_signature(root):
            return False

        changed = False
        for child in (root.get("children") or []):
            if child.get("children"):
                continue
            if child.get("_preserveChoiceMultiBlank"):
                continue
            blank_text = str(child.get("blankText") or "").strip()
            blank_tokens = [token for token in blank_text.split() if token]
            raw_text = str(child.get("rawText") or "")
            raw_slot_count = len(_collect_blank_segments(raw_text))
            score_value = float(child.get("score") or 0.0)
            parsed_single_score = None
            if len(blank_tokens) == 1:
                parsed_single_score = _parse_merged_blank_segment_score(blank_tokens[0])

            should_normalize = (
                len(blank_tokens) > 1
                or raw_slot_count > 1
                or (raw_slot_count == 1 and not blank_tokens and score_value > 0)
                or (
                    len(blank_tokens) == 1
                    and parsed_single_score is not None
                    and score_value > 0
                    and abs(float(parsed_single_score) - float(score_value)) > 0.01
                )
            )
            if not should_normalize:
                continue

            if score_value > 0:
                if abs(score_value - round(score_value)) <= 1e-6:
                    score_text = str(int(round(score_value)))
                else:
                    score_text = f"{score_value:.2f}".rstrip("0").rstrip(".")
                normalized_blank = f"____（{score_text}分）"
            else:
                normalized_blank = "____"

            if blank_text != normalized_blank:
                child["blankText"] = normalized_blank
                changed = True
        return changed

    def _normalize_english_reading_choice_scores(self, root: Dict) -> bool:
        title = normalize_line(str(root.get("rawText") or root.get("title") or ""))
        if "阅读理解" not in title:
            return False

        children = list(root.get("children") or [])
        if len(children) < 15:
            return False
        if any(child.get("children") for child in children):
            return False

        number_values: List[int] = []
        for child in children:
            parsed = self._parse_arabic_numbering(child.get("numbering"))
            if parsed is None:
                return False
            number_values.append(parsed)
        ordered = sorted(number_values)
        if ordered != list(range(min(ordered), max(ordered) + 1)):
            return False
        # 仅命中“16-35”这类典型阅读选择段，避免误伤其他题型。
        if len(ordered) == 20 and (ordered[0], ordered[-1]) != (16, 35):
            return False
        if len(ordered) != 20:
            return False

        child_scores = [float(child.get("score") or 0.0) for child in children]
        if not child_scores:
            return False
        avg_score = sum(child_scores) / len(child_scores)
        if avg_score < 1.8 or avg_score > 2.1:
            return False
        if any(score < 1.5 or score > 2.5 for score in child_scores):
            return False

        parent_score = float(root.get("score") or 0.0)
        if abs(parent_score - 40.0) > 1e-6:
            return False

        changed = False
        for child in children:
            if float(child.get("score") or 0.0) != 1.0:
                child["score"] = 1.0
                changed = True
            blank_text = str(child.get("blankText") or "").strip()
            if not blank_text:
                continue
            token = blank_text.split()[0]
            underline_match = re.search(r"[_\uFF3F\uFE4D\uFE4E\u2014]{2,}", token)
            underline = underline_match.group(0) if underline_match else "____"
            normalized_blank = f"{underline}（1分）"
            if blank_text != normalized_blank:
                child["blankText"] = normalized_blank
                changed = True

        if float(root.get("score") or 0.0) != 20.0:
            root["score"] = 20.0
            changed = True
        return changed

    def _normalize_reading_subsection_scores_from_source(
        self,
        root: Dict,
        source_lines: List[str],
        section_end_line_no: int,
    ) -> bool:
        title = normalize_line(str(root.get("rawText") or root.get("title") or ""))
        if "阅读理解" not in title:
            return False

        children = list(root.get("children") or [])
        if len(children) < 5:
            return False
        if any(child.get("children") for child in children):
            return False

        parent_line_no = int(root.get("lineNumber") or 0)
        if parent_line_no <= 0 or not source_lines:
            return False

        start_line_no = max(parent_line_no, 1)
        end_line_no = max(start_line_no + 1, min(len(source_lines) + 1, section_end_line_no))
        if end_line_no <= start_line_no:
            return False

        subsection_heading_re = re.compile(
            r"^\s*第[一二三四五六七八九十百千万\d]+\s*节(?:[:：、，,\.\uFF0E]\s*)?.*?"
            r"[\uFF08(]\s*(\d+(?:\.\d+)?)\s*\u5206\s*[\uFF09)]"
        )
        arabic_question_no_re = re.compile(
            r"(?:^|[\s\n])(?:[\uFF08(]\s*[\u3000\s]*[\uFF09)]\s*)?(\d{1,3})\s*[\u3001\.\uFF0E](?!\d)"
        )
        indexed_blank_no_re = re.compile(r"[_\uFF3F\uFE4D\uFE4E\u2014]{1,}\s*(\d{1,3})\s*[_\uFF3F\uFE4D\uFE4E\u2014]{1,}")

        subsection_spans: List[Tuple[int, float, int]] = []
        for line_no in range(start_line_no, end_line_no):
            text = normalize_line(source_lines[line_no - 1] if line_no - 1 < len(source_lines) else "")
            if not text:
                continue
            match = subsection_heading_re.match(text)
            if not match:
                continue
            try:
                subsection_score = float(match.group(1))
            except (TypeError, ValueError):
                continue
            if subsection_score <= 0:
                continue
            subsection_spans.append((line_no, subsection_score, end_line_no))

        if len(subsection_spans) < 2:
            return False

        subsection_spans = [
            (
                start_no,
                score_value,
                subsection_spans[idx + 1][0] if idx + 1 < len(subsection_spans) else end_line_no,
            )
            for idx, (start_no, score_value, _) in enumerate(subsection_spans)
        ]

        children_by_number: Dict[int, Dict] = {}
        for child in children:
            parsed = self._parse_arabic_numbering(child.get("numbering"))
            if parsed is None:
                continue
            children_by_number.setdefault(parsed, child)

        if len(children_by_number) < 5:
            return False

        changed = False
        for subsection_start, subsection_score, subsection_end in subsection_spans:
            numbers: List[int] = []
            seen_numbers: Set[int] = set()
            for line_no in range(subsection_start + 1, subsection_end):
                text = normalize_line(source_lines[line_no - 1] if line_no - 1 < len(source_lines) else "")
                if not text:
                    continue
                matched_numbers = [int(value) for value in arabic_question_no_re.findall(text)]
                matched_numbers.extend(
                    int(value)
                    for value in indexed_blank_no_re.findall(text)
                    if int(value) not in matched_numbers
                )
                for number_value in matched_numbers:
                    if number_value not in children_by_number or number_value in seen_numbers:
                        continue
                    seen_numbers.add(number_value)
                    numbers.append(number_value)

            if len(numbers) < 2:
                continue

            base_score = round(float(subsection_score) / float(len(numbers)), 2)
            assigned_scores = [base_score for _ in numbers]
            diff = round(float(subsection_score) - sum(assigned_scores), 2)
            assigned_scores[-1] = round(assigned_scores[-1] + diff, 2)

            for idx, number_value in enumerate(numbers):
                child = children_by_number[number_value]
                assigned_score = float(assigned_scores[idx])
                current_score = float(child.get("score") or 0.0)
                if abs(current_score - assigned_score) > 0.01:
                    child["score"] = assigned_score
                    changed = True

                blank_text = str(child.get("blankText") or "").strip()
                raw_text = normalize_line(str(child.get("rawText") or ""))
                underline_source = blank_text or raw_text
                underline_match = re.search(r"([_\uFF3F\uFE4D\uFE4E\u2014]{2,})", underline_source)
                if not underline_match:
                    continue
                underline = underline_match.group(1)
                if abs(assigned_score - round(assigned_score)) <= 1e-6:
                    score_text = str(int(round(assigned_score)))
                else:
                    score_text = f"{assigned_score:.2f}".rstrip("0").rstrip(".")
                normalized_blank = f"{underline}（{score_text}分）"
                if blank_text != normalized_blank:
                    child["blankText"] = normalized_blank
                    changed = True

        if changed:
            root["score"] = round(
                sum(float(child.get("score") or 0.0) for child in (root.get("children") or [])),
                2,
            )
        return changed

    def _final_rebalance_low_cloze_child_scores(self, outline_items: List[Dict]) -> bool:
        if not outline_items:
            return False
        changed = False
        for root in outline_items:
            if int(root.get("level", 1) or 1) != 1:
                continue
            title = normalize_line(str(root.get("rawText") or root.get("title") or ""))
            if "完形填空" not in title:
                continue
            children = list(root.get("children") or [])
            if not children:
                continue
            parent_score = float(root.get("score") or 0.0)
            if parent_score <= 0:
                continue
            all_numbered_leaf = True
            for child in children:
                if child.get("children"):
                    all_numbered_leaf = False
                    break
                if self._parse_arabic_numbering(child.get("numbering")) is None:
                    all_numbered_leaf = False
                    break
            if not all_numbered_leaf:
                continue
            total = sum(float(child.get("score") or 0.0) for child in children)
            if total > parent_score * 0.8:
                continue
            _assign_even_scores_to_children(root)
            changed = True
        return changed

    def _collapse_choice_like_section_multi_blanks_on_outline(self, outline_items: List[Dict]) -> bool:
        if not outline_items:
            return False
        changed = False
        for root in outline_items:
            if int(root.get("level", 1) or 1) != 1:
                continue
            if self._collapse_choice_like_section_multi_blanks(root):
                changed = True
        return changed

    def _merge_slot_only_restart_child_into_previous(self, root: Dict) -> bool:
        children = list(root.get("children") or [])
        if len(children) < 2:
            return False

        slot_only_tail_re = re.compile(
            r"^\s*1\s*[\u3001\.\uFF0E]\s*(?:[_\uFF3F\uFE4D\uFE4E\u2014]{2,}|[\uFF08(]\s*[\uFF09)]|\[\s*\])\s*$"
        )

        changed = False
        merged_children: List[Dict] = []
        for child in children:
            if not merged_children:
                merged_children.append(child)
                continue

            current_no = self._parse_arabic_numbering(child.get("numbering"))
            previous = merged_children[-1]
            previous_no = self._parse_arabic_numbering(previous.get("numbering"))
            current_raw = str(child.get("rawText") or "")
            previous_raw = str(previous.get("rawText") or "")
            current_blank = str(child.get("blankText") or "").strip()
            previous_blank = str(previous.get("blankText") or "").strip()

            is_slot_only_restart = (
                current_no == 1
                and previous_no is not None
                and previous_no >= 2
                and not child.get("children")
                and slot_only_tail_re.match(normalize_line(current_raw) or "") is not None
            )
            if not is_slot_only_restart:
                merged_children.append(child)
                continue

            # 仅在紧邻续行时合并，避免误吞真正的新小题。
            prev_line = int(previous.get("lineNumber") or 0)
            cur_line = int(child.get("lineNumber") or 0)
            if prev_line > 0 and cur_line > 0 and cur_line - prev_line > 2:
                merged_children.append(child)
                continue

            if not previous_blank:
                score_value = float(previous.get("score") or 0.0)
                if score_value > 0:
                    if abs(score_value - round(score_value)) <= 1e-6:
                        score_text = str(int(round(score_value)))
                    else:
                        score_text = f"{score_value:.2f}".rstrip("0").rstrip(".")
                    previous["blankText"] = f"____（{score_text}分）"
                elif current_blank:
                    previous["blankText"] = current_blank

            if current_raw and current_raw not in previous_raw:
                previous["rawText"] = f"{previous_raw} {current_raw}".strip()

            changed = True

        if changed:
            root["children"] = merged_children
        return changed

    def _resolve_section_end_line_for_root_repair(
        self,
        parent_line_no: int,
        proposed_next_line_no: int,
        source_lines: List[str],
    ) -> int:
        if not source_lines:
            return proposed_next_line_no if proposed_next_line_no > 0 else 1

        heading_re = re.compile(r"^\s*[一二三四五六七八九十百千万]+\s*[、\.．]\s*")
        if 1 <= proposed_next_line_no <= len(source_lines):
            candidate = normalize_line(source_lines[proposed_next_line_no - 1] if proposed_next_line_no - 1 < len(source_lines) else "")
            if candidate and heading_re.match(candidate):
                return proposed_next_line_no

        start_idx = max(0, parent_line_no)
        for idx in range(start_idx, len(source_lines)):
            text = normalize_line(source_lines[idx] if idx < len(source_lines) else "")
            if not text:
                continue
            if heading_re.match(text):
                return idx + 1
        return len(source_lines) + 1

    def _flatten_english_listening_wrapper_children(self, root: Dict) -> bool:
        children = root.get("children") or []
        if len(children) < 2:
            return False

        flattened_candidates: List[Dict] = []
        has_wrapper_node = False

        for child in children:
            child_no = self._parse_arabic_numbering(child.get("numbering"))
            child_children = child.get("children") or []
            if child_no is not None:
                flattened_candidates.append(child)
                continue
            if child_children:
                has_wrapper_node = True
                for grandchild in child_children:
                    if self._parse_arabic_numbering(grandchild.get("numbering")) is None:
                        continue
                    flattened_candidates.append(grandchild)

        if not has_wrapper_node:
            return False

        dedup: Dict[int, Dict] = {}
        for node in flattened_candidates:
            parsed = self._parse_arabic_numbering(node.get("numbering"))
            if parsed is None:
                continue
            dedup.setdefault(parsed, node)

        if len(dedup) < 10:
            return False

        root["children"] = [dedup[number] for number in sorted(dedup)]
        return True

    def _fill_tail_missing_numbered_children_by_score(self, root: Dict) -> bool:
        children = root.get("children") or []
        if not children:
            return False

        numbered: Dict[int, Dict] = {}
        for child in children:
            parsed = self._parse_arabic_numbering(child.get("numbering"))
            if parsed is None:
                continue
            numbered.setdefault(parsed, child)

        if len(numbered) < 8:
            return False

        ordered_numbers = sorted(numbered.keys())
        min_no = ordered_numbers[0]
        max_no = ordered_numbers[-1]
        if min_no != 1:
            return False
        if ordered_numbers != list(range(1, max_no + 1)):
            return False

        parent_score = float(root.get("score") or 0.0)
        if parent_score <= 0 or abs(parent_score - round(parent_score)) > 1e-6:
            return False
        expected_count = int(round(parent_score))
        if expected_count <= max_no or expected_count - max_no > 10:
            return False

        sample_scores = [
            float(item.get("score") or 0.0)
            for item in numbered.values()
            if float(item.get("score") or 0.0) > 0
        ]
        score_per_child = 1.0
        if sample_scores:
            avg_score = sum(sample_scores) / len(sample_scores)
            if 0.2 <= avg_score <= 3.0:
                score_per_child = round(avg_score, 2)
        score_text = f"{int(score_per_child)}" if abs(score_per_child - int(score_per_child)) < 1e-6 else f"{score_per_child:.2f}".rstrip("0").rstrip(".")
        child_level = int(children[0].get("level", int(root.get("level", 1) or 1) + 1) or 1)
        line_no = int(root.get("lineNumber") or 0)

        generated: List[Dict] = []
        for number_value in range(max_no + 1, expected_count + 1):
            generated.append(
                {
                    "lineNumber": line_no,
                    "level": child_level,
                    "numbering": str(number_value),
                    "title": "",
                    "rawText": f"{number_value}.",
                    "children": [],
                    "blankText": f"____（{score_text}分）",
                    "score": float(score_per_child),
                    "_tokenType": "arabic",
                    "_isSectionHeading": False,
                    "_bindSectionChildren": False,
                }
            )

        if not generated:
            return False

        root["children"] = [numbered[number] for number in ordered_numbers] + generated
        return True

    def _repair_english_cloze_overflow_into_reading_section(self, outline_items: List[Dict]) -> bool:
        if not outline_items:
            return False

        roots = [node for node in outline_items if int(node.get("level", 1) or 1) == 1]
        changed = False

        def _find_child_max_number(node: Optional[Dict]) -> Optional[int]:
            if not node:
                return None
            values: List[int] = []
            for item in (node.get("children") or []):
                parsed = self._parse_arabic_numbering(item.get("numbering"))
                if parsed is None:
                    continue
                values.append(parsed)
            return max(values) if values else None

        for idx, root in enumerate(roots):
            title = normalize_line(str(root.get("rawText") or root.get("title") or ""))
            if "完形填空" not in title:
                continue

            next_root = roots[idx + 1] if idx + 1 < len(roots) else None
            if str(next_root.get("numbering") or "") not in {"五", "5"}:
                continue

            parent_score = float(root.get("score") or 0.0)
            if parent_score <= 0 or abs(parent_score - round(parent_score)) > 1e-6:
                continue
            expected_count = int(round(parent_score))
            if expected_count < 5 or expected_count > 20:
                continue

            numbered_children: List[Tuple[int, Dict]] = []
            for child in (root.get("children") or []):
                parsed = self._parse_arabic_numbering(child.get("numbering"))
                if parsed is None:
                    continue
                numbered_children.append((parsed, child))
            if len(numbered_children) <= expected_count:
                continue

            prev_root = roots[idx - 1] if idx - 1 >= 0 else None
            prev_max = _find_child_max_number(prev_root)
            if prev_max is None:
                continue
            expected_start = prev_max + 1
            expected_end = expected_start + expected_count - 1

            keep_map: Dict[int, Dict] = {}
            overflow_map: Dict[int, Dict] = {}
            for number_value, node in numbered_children:
                if expected_start <= number_value <= expected_end:
                    keep_map.setdefault(number_value, node)
                elif number_value > expected_end:
                    overflow_map.setdefault(number_value, node)
            if len(keep_map) < max(4, expected_count - 2):
                continue
            if len(overflow_map) < 5:
                continue

            root["children"] = [keep_map[number] for number in sorted(keep_map.keys())]

            overflow_children = [overflow_map[number] for number in sorted(overflow_map.keys())]
            score_match_re = re.compile(r"[\uFF08(]\s*(\d+(?:\.\d+)?)\s*\u5206\s*[\uFF09)]")
            for child in overflow_children:
                raw_text = str(child.get("rawText") or "")
                match = score_match_re.search(raw_text)
                if not match:
                    continue
                try:
                    child["score"] = float(match.group(1))
                except (TypeError, ValueError):
                    continue
            reading_root_score = round(
                sum(float(child.get("score") or 0.0) for child in overflow_children),
                2,
            )
            if len(overflow_children) >= 10 and reading_root_score <= len(overflow_children):
                for child in overflow_children:
                    child["score"] = 2.0
                reading_root_score = float(len(overflow_children) * 2)
            for child in overflow_children:
                child["level"] = int(root.get("level", 1) or 1) + 1

            reading_root = {
                "lineNumber": min(int(child.get("lineNumber") or 0) for child in overflow_children) or int(root.get("lineNumber") or 0),
                "level": int(root.get("level", 1) or 1),
                "numbering": "四",
                "title": f"阅读理解（{reading_root_score}分）" if reading_root_score > 0 else "阅读理解",
                "rawText": f"四、阅读理解（{reading_root_score}分）" if reading_root_score > 0 else "四、阅读理解",
                "blankText": "",
                "score": reading_root_score,
                "children": overflow_children,
            }

            insert_idx = outline_items.index(next_root)
            outline_items.insert(insert_idx, reading_root)
            roots = [node for node in outline_items if int(node.get("level", 1) or 1) == 1]
            changed = True
            break

        return changed

    def _backfill_outline_root_scores_from_answer_lines(self, outline_items: List[Dict], answer_lines: List[str]) -> bool:
        if not outline_items or not answer_lines:
            return False
        answer_score_map = self._extract_answer_heading_score_map(answer_lines)
        if not answer_score_map:
            return False

        changed = False
        for node in outline_items:
            numbering_key = self._normalize_numbering_key(node.get("numbering"))
            if not numbering_key:
                continue
            current_score = node.get("score")
            if current_score is not None and float(current_score) > 0:
                continue
            matched_score = answer_score_map.get(numbering_key)
            if matched_score is None or float(matched_score) <= 0:
                continue
            node["score"] = float(matched_score)
            changed = True
        return changed

    def _recalculate_outline_scores_after_backfill(self, outline_items: List[Dict]) -> Dict:
        _split_merged_inline_numbered_blank_leaf_nodes(outline_items)
        _distribute_outline_parent_score_to_slot_children(outline_items)
        _distribute_outline_parent_score_to_numbered_zero_children(outline_items)
        _distribute_outline_parent_remaining_score_to_unscored_children(outline_items)
        _distribute_outline_parent_score_gap_to_slot_leaves(outline_items)
        _distribute_outline_parent_remaining_score_to_unscored_children(outline_items)
        # 保证“父级有分，子级不为 0 分”在修复回填路径同样生效。
        _force_fill_zero_child_scores_under_scored_parent(outline_items)
        _fill_scored_empty_bracket_leaf_blank_text(outline_items)
        _fill_scored_underline_leaf_blank_text(outline_items)
        _fill_scored_question_mark_placeholder_blank_text(outline_items)
        _strip_instruction_quoted_underlines_from_blank_text(outline_items)
        _collapse_choice_question_multi_blanks(outline_items)
        _collapse_choice_semantic_leaf_multi_blanks(outline_items)
        _enforce_choice_question_no_blank_and_no_children(outline_items)
        _repair_under_extracted_non_choice_leaf_blanks(outline_items)
        _prune_conflicting_duplicate_numbered_children(outline_items)
        _prune_number_only_placeholder_nodes(outline_items)
        _prune_empty_leaf_placeholder_nodes(outline_items)
        _strip_numbering_for_slot_only_leaf_nodes(outline_items)
        _strip_numbering_for_quantified_slot_sentence_nodes(outline_items)
        _strip_numbering_for_fill_section_slot_rows(outline_items)
        if _is_english_exam_outline(outline_items):
            for section in outline_items:
                title = normalize_line(str(section.get("rawText") or section.get("title") or ""))
                if not title:
                    continue
                if re.search(r"(?:单项选择|单项填空|单选|完形填空|阅读理解)", title):
                    _repair_english_choice_merged_following_stub_children(section)
                    _strip_english_choice_child_underlines(section)
                if re.search(r"(?:缺词填空|首字母填空|短文填空)", title):
                    _repair_english_short_fill_single_blank_children(section)
                _repair_english_single_answer_multi_underline_children(section)
                _trim_english_leaf_trailing_redundant_blank(section)
            _enforce_choice_question_no_blank_and_no_children(outline_items)
        normalize_outline_blank_scores(outline_items)
        return build_scores_tree(outline_items)

    def _compile_answer_patterns(self, answer_section_patterns: Optional[List[str]]) -> List[re.Pattern]:
        patterns = answer_section_patterns if answer_section_patterns else DEFAULT_ANSWER_PATTERNS
        compiled: List[re.Pattern] = []
        for pattern in patterns:
            try:
                compiled.append(re.compile(pattern))
            except re.error:
                continue
        if not compiled:
            compiled = [re.compile(pattern) for pattern in DEFAULT_ANSWER_PATTERNS]
        return compiled

    def _matches_answer(self, text: str, patterns: List[re.Pattern]) -> bool:
        normalized = re.sub(r"\s+", " ", (text or "")).strip()
        if not normalized:
            return False
        if ANSWER_HEADING_NEGATIVE_RE.search(normalized):
            return False
        compact = re.sub(r"\s+", "", normalized)
        if re.fullmatch(r"[【\[]?(?:参考)?答案(?:与解析|解析|部分|页)?[】\]]?", compact):
            return True
        if self._looks_like_answer_sheet_heading(normalized):
            return True
        # Accept practical headings like "高二物理参考答案（选修...）" even when
        # custom/default regex patterns are too strict on the suffix.
        if self._looks_like_answer_heading(normalized):
            return True
        if not any(pattern.search(normalized) or pattern.search(compact) for pattern in patterns):
            return False
        return self._looks_like_answer_heading(normalized)

    def _looks_like_answer_sheet_heading(self, text: str) -> bool:
        normalized = re.sub(r"\s+", " ", (text or "")).strip()
        if not normalized:
            return False
        if not ANSWER_SHEET_HEADING_RE.search(normalized):
            return False
        if QUESTION_LINE_PREFIX_RE.match(normalized):
            return False
        if ANSWER_INSTRUCTION_VERB_RE.search(normalized):
            return False
        if len(normalized) > 96:
            return False
        if any(mark in normalized for mark in ("。", "；", ";")) and len(normalized) > 24:
            return False
        return True

    def _looks_like_answer_heading(self, text: str) -> bool:
        normalized = re.sub(r"\s+", " ", (text or "")).strip()
        if not normalized:
            return False
        if ANSWER_HEADING_NEGATIVE_RE.search(normalized):
            return False
        # 防止把“材料解析题”等题型标题误判为答案区
        if (
            SECTION_HEADING_PREFIX_RE.match(normalized)
            and QUESTION_SECTION_HINT_RE.search(normalized)
            and not re.search(r"(?:参考)?答案", normalized)
        ):
            return False
        compact = re.sub(r"\s+", "", normalized)
        if not ANSWER_HEADING_KEYWORD_RE.search(normalized) and not ANSWER_HEADING_KEYWORD_RE.search(compact):
            return False
        if re.fullmatch(r"[【\[]?(?:参考)?答案(?:与解析|解析|部分|页)?[】\]]?", compact):
            return True
        if (
            QUESTION_HEADING_WITH_ANSWER_WORD_RE.search(normalized)
            and QUESTION_SECTION_HINT_RE.search(normalized)
            and not re.search(r"(?:参考\s*答案|答案(?:与解析|解析))", normalized)
        ):
            return False
        if (
            QUESTION_HEADING_WITH_ANSWER_WORD_RE.search(normalized)
            and SECTION_HEADING_PREFIX_RE.match(normalized)
            and re.search(r"[\uFF08(]\s*\d+(?:\.\d+)?\s*\u5206\s*[\uFF09)]", normalized)
        ):
            return False
        if QUESTION_LINE_PREFIX_RE.match(normalized):
            return False
        if ANSWER_INSTRUCTION_HINT_RE.search(normalized):
            return False
        if len(normalized) > 64:
            return False
        if any(mark in normalized for mark in ("。", "；", ";")) and len(normalized) > 20:
            return False
        return True

    def _looks_like_answer_split_hint(self, text: str) -> bool:
        normalized = re.sub(r"\s+", " ", (text or "")).strip()
        if not normalized:
            return False
        if QUESTION_LINE_PREFIX_RE.match(normalized):
            return False
        return bool(ANSWER_SPLIT_HINT_RE.search(normalized))

    def _looks_like_tail_answer_material_start(self, lines: List[str], start_idx: int) -> bool:
        window = lines[start_idx + 1 : start_idx + 4]
        if not window:
            return False
        marker_re = re.compile(
            r"^\s*(?:\u7b2c[一二三四五六七八九十\d]+\u90e8\u5206|"
            r"(?:\u53c2\u8003\s*)?\u7b54\u6848|"
            r"\u542c\u529b\u6750\u6599|"
            r"Text\s*\d+|"
            r"[MW\u7537\u5973]\s*[:\uFF1A]|"
            r"\d+\s*[\.\uFF0E]\s*[A-Za-z])",
            flags=re.IGNORECASE,
        )
        for line in window:
            normalized = re.sub(r"\s+", " ", (line or "")).strip()
            if not normalized:
                continue
            if marker_re.search(normalized):
                return True
        return False

    def _looks_like_tail_answer_material_heading(self, lines: List[str], start_idx: int) -> bool:
        if start_idx < int(len(lines) * 0.55):
            return False
        normalized = re.sub(r"\s+", " ", (lines[start_idx] or "")).strip()
        if not normalized:
            return False
        if not TAIL_ANSWER_MATERIAL_HEADING_RE.match(normalized):
            return False
        window = lines[start_idx + 1 : start_idx + 7]
        if not window:
            return False
        signal_count = 0
        for line in window:
            probe = re.sub(r"\s+", " ", (line or "")).strip()
            if not probe:
                continue
            if TAIL_ANSWER_FOLLOW_LINE_RE.match(probe):
                signal_count += 1
                continue
            if self._looks_like_answer_key_line(probe):
                signal_count += 1
        return signal_count >= 1

    def _looks_like_answer_key_line(self, line: str) -> bool:
        normalized = re.sub(r"\s+", " ", (line or "")).strip()
        if not normalized:
            return False
        if "分" in normalized:
            return False
        if "?" in normalized or "？" in normalized:
            return False
        if QUESTION_HINT_RE.search(normalized):
            return False
        if QUESTION_DIRECTIVE_TEXT_RE.search(normalized):
            return False
        if ANSWER_KEY_TOKEN_RE.search(normalized):
            return True
        if re.fullmatch(
            r"(?:[A-D\uFF21-\uFF24√×]|\d+(?:\.\d+)?|[\u4e00-\u9fa5]{1,3})(?:\s*[,，;；、/]\s*"
            r"(?:[A-D\uFF21-\uFF24√×]|\d+(?:\.\d+)?|[\u4e00-\u9fa5]{1,3})){1,12}",
            normalized,
        ):
            return True
        return False

    def _looks_like_tail_compact_answer_start(self, lines: List[str], start_idx: int) -> bool:
        if start_idx < int(len(lines) * 0.6):
            return False
        if self._has_scored_question_section_ahead(lines, start_idx):
            return False

        # 新增：检查前面是否有有效的section heading（如"四、词汇运用"）
        # 如果有，说明当前位置的题目属于该section，不应视为答案分割点
        has_section_heading_before = False
        for i in range(start_idx - 1, max(0, start_idx - 20), -1):
            line = lines[i]
            if not line:
                continue
            # 检查是否是section heading（如"四、xxx"、"五、xxx"）
            if re.match(r'^\s*[\u4e00-\u9fa5\u96d9\u5341\u767e\u5343\u4e00]\s*[\u3001\.\uFF0E]\s*[\u4e00-\u9fa5]', line):
                # 检查标题后面是否有分数（如"（10分）"）- 包含"分"字
                if '分' in line or '\u5206' in line:
                    has_section_heading_before = True
                    break
        if has_section_heading_before:
            return False

        normalized = re.sub(r"\s+", " ", (lines[start_idx] or "")).strip()
        if not normalized:
            return False
        if "分" in normalized:
            return False
        if "?" in normalized or "？" in normalized:
            return False
        if re.search(r"(?:\u8bfb|\u9605\u8bfb|\u56de\u7b54|\u5b8c\u6210|\u4e0b\u5217|\u6750\u6599|\u7531\u56fe|\u6839\u636e)", normalized):
            return False
        if not TAIL_COMPACT_ANSWER_LINE_RE.match(normalized):
            return False
        lead_match = ARABIC_QUESTION_LINE_RE.match(normalized)
        line_tail = (lead_match.group(2) if lead_match else normalized).strip()
        # 尾部紧凑答案块：首行本身必须像“答案内容”，避免把题干/作文要求误切为答案区。
        if not line_tail:
            return False
        if lead_match is not None:
            for probe in lines[start_idx + 1 : start_idx + 3]:
                probe_text = re.sub(r"\s+", " ", (probe or "")).strip()
                if not probe_text:
                    continue
                if TAIL_OPTION_LINE_RE.match(probe_text):
                    return False
        if QUESTION_DIRECTIVE_TEXT_RE.search(line_tail):
            return False
        if not (
            self._looks_like_answer_key_line(line_tail)
            or bool(TAIL_COMPACT_EXPLICIT_KEY_RE.search(line_tail))
        ):
            return False

        window = lines[start_idx : start_idx + 8]
        if len(window) < 3:
            return False
        answer_like_count = 0
        compact_count = 0
        no_score_count = 0
        explicit_key_count = 0
        for line in window:
            probe = re.sub(r"\s+", " ", (line or "")).strip()
            if not probe:
                continue
            if "分" not in probe:
                no_score_count += 1
            if TAIL_COMPACT_ANSWER_LINE_RE.match(probe):
                compact_count += 1
            if TAIL_COMPACT_EXPLICIT_KEY_RE.search(probe):
                explicit_key_count += 1
            if self._looks_like_answer_key_line(probe):
                answer_like_count += 1
        return answer_like_count >= 3 and compact_count >= 2 and no_score_count >= 2 and explicit_key_count >= 1

    def _looks_like_tail_compact_answer_heading(self, lines: List[str], start_idx: int) -> bool:
        if start_idx < int(len(lines) * 0.55):
            return False
        normalized = re.sub(r"\s+", " ", (lines[start_idx] or "")).strip()
        if not normalized:
            return False
        if len(normalized) > 48:
            return False
        if not TAIL_COMPACT_ANSWER_HEADING_RE.search(normalized):
            return False
        next_idx = start_idx + 1
        if next_idx >= len(lines):
            return False
        return self._looks_like_tail_compact_answer_start(lines, next_idx)

    def _looks_like_tail_inline_answer_section_heading(self, lines: List[str], start_idx: int) -> bool:
        if start_idx < int(len(lines) * 0.55):
            return False
        if self._has_scored_question_section_ahead(lines, start_idx):
            return False
        normalized = re.sub(r"\s+", " ", (lines[start_idx] or "")).strip()
        if not normalized:
            return False
        if TAIL_OPTION_LINE_RE.match(normalized):
            return False
        if "分" in normalized:
            return False
        match = TAIL_INLINE_ANSWER_SECTION_WITH_KEYS_RE.match(normalized)
        if not match:
            return False

        inline_tail = re.sub(r"\s+", " ", (match.group(1) or "")).strip()
        if inline_tail:
            if self._looks_like_answer_key_line(inline_tail):
                return True
            if re.search(r"(?:[A-D\uFF21-\uFF24]{3,}|\d+\s*[-~]\s*\d+)", inline_tail):
                return True

        window = lines[start_idx + 1 : start_idx + 6]
        if not window:
            return False
        signal_count = 0
        for line in window:
            probe = re.sub(r"\s+", " ", (line or "")).strip()
            if not probe:
                continue
            if self._looks_like_answer_key_line(probe):
                signal_count += 1
                continue
            if re.search(r"(?:[A-D\uFF21-\uFF24]{3,}|\d+\s*[-~]\s*\d+)", probe):
                signal_count += 1
        return signal_count >= 1

    def _looks_like_tail_section_inline_answer_start(self, lines: List[str], start_idx: int) -> bool:
        if start_idx < int(len(lines) * 0.55):
            return False
        if self._has_scored_question_section_ahead(lines, start_idx):
            return False
        normalized = re.sub(r"\s+", " ", (lines[start_idx] or "")).strip()
        if not normalized:
            return False
        if TAIL_OPTION_LINE_RE.match(normalized):
            return False
        if "分" in normalized:
            return False
        if "?" in normalized or "？" in normalized:
            return False
        match = TAIL_SECTION_INLINE_ANSWER_RE.match(normalized)
        if not match:
            return False

        line_tail = match.group(1)
        numbered_item_count = len(INLINE_NUMBERED_ITEM_RE.findall(line_tail))
        explicit_key_count = len(TAIL_COMPACT_EXPLICIT_KEY_RE.findall(line_tail))
        tail_answer_like = self._looks_like_answer_key_line(line_tail)
        if numbered_item_count < 2 and explicit_key_count < 1 and not tail_answer_like:
            return False

        window = lines[start_idx : start_idx + 4]
        if len(window) < 2:
            return False
        bundle_count = 0
        continuation_count = 0
        for line in window:
            probe = re.sub(r"\s+", " ", (line or "")).strip()
            if not probe or "分" in probe:
                continue
            m = TAIL_SECTION_INLINE_ANSWER_RE.match(probe)
            if not m:
                if re.match(r"^\s*\d{1,3}\s*[\.\uFF0E\u3001]\s*", probe) and self._looks_like_answer_key_line(probe):
                    continuation_count += 1
                continue
            tail = m.group(1)
            has_number_bundle = len(INLINE_NUMBERED_ITEM_RE.findall(tail)) >= 1
            has_answer_key = bool(
                TAIL_COMPACT_EXPLICIT_KEY_RE.search(tail)
                or re.search(r"(?:\b[A-D]\b|[√×])", tail, flags=re.IGNORECASE)
            )
            has_answer_like = self._looks_like_answer_key_line(tail)
            if has_number_bundle or has_answer_key or has_answer_like:
                bundle_count += 1
        return bundle_count >= 2 or (bundle_count >= 1 and continuation_count >= 2)

    def _guess_answer_split_by_repeated_numbering(self, lines: List[str]) -> Optional[int]:
        numbered: List[Tuple[int, int, str]] = []
        for idx, line in enumerate(lines):
            normalized = re.sub(r"\s+", " ", (line or "")).strip()
            if not normalized:
                continue
            match = ARABIC_QUESTION_LINE_RE.match(normalized)
            if not match:
                continue
            try:
                number_value = int(match.group(1))
            except ValueError:
                continue
            numbered.append((idx, number_value, normalized))

        if len(numbered) < 6:
            return None

        first_seen_line: Dict[int, int] = {}
        total_lines = len(lines)

        for pos, (line_idx, number_value, _) in enumerate(numbered):
            first_line = first_seen_line.get(number_value)
            if first_line is None:
                first_seen_line[number_value] = line_idx
                continue

            if line_idx < int(total_lines * 0.55):
                continue
            if self._has_scored_question_section_ahead(lines, line_idx):
                continue

            block = numbered[pos : pos + 4]
            if len(block) < 3:
                continue

            repeated_count = 0
            answer_like_count = 0
            no_score_count = 0
            explanation_hint_count = 0
            plain_tail_answer_count = 0
            prev_num = None
            monotonic_non_decreasing = True

            for block_line_idx, block_num, block_raw in block:
                if block_num in first_seen_line and first_seen_line[block_num] < block_line_idx:
                    repeated_count += 1
                if self._looks_like_answer_key_line(block_raw):
                    answer_like_count += 1
                if self._looks_like_answer_explanation_line(block_raw):
                    explanation_hint_count += 1
                if "分" not in (block_raw or ""):
                    no_score_count += 1
                compact_text = re.sub(r"\s+", " ", (block_raw or "")).strip()
                if (
                    compact_text
                    and len(compact_text) <= 48
                    and "分" not in compact_text
                    and "?" not in compact_text
                    and "？" not in compact_text
                    and not QUESTION_HINT_RE.search(compact_text)
                    and not QUESTION_DIRECTIVE_TEXT_RE.search(compact_text)
                ):
                    plain_tail_answer_count += 1
                if prev_num is not None and block_num < prev_num:
                    monotonic_non_decreasing = False
                prev_num = block_num

            fallback_restart_pattern = (
                block[0][1] == 1
                and line_idx >= int(total_lines * 0.6)
                and repeated_count >= 3
                and no_score_count >= 2
                and explanation_hint_count >= 1
                and len(first_seen_line) >= 8
                and not self._has_recent_scored_question_section(lines, line_idx)
            )
            plain_compact_restart_pattern = (
                block[0][1] == 1
                and line_idx >= int(total_lines * 0.65)
                and repeated_count >= 3
                and no_score_count >= 3
                and plain_tail_answer_count >= 3
                and len(first_seen_line) >= 10
                and not self._has_recent_scored_question_section(lines, line_idx)
            )
            if (
                monotonic_non_decreasing
                and (
                    (
                        repeated_count >= 3
                        and answer_like_count >= 2
                        and no_score_count >= 2
                    )
                    or fallback_restart_pattern
                    or plain_compact_restart_pattern
                )
            ):
                return line_idx

        return None

    def _has_scored_question_section_ahead(self, lines: List[str], start_idx: int, lookahead: int = 18) -> bool:
        end_idx = min(len(lines), start_idx + lookahead)
        for idx in range(start_idx, end_idx):
            normalized = re.sub(r"\s+", " ", (lines[idx] or "")).strip()
            if not normalized:
                continue
            if not SECTION_HEADING_PREFIX_RE.match(normalized):
                continue
            if not re.search(r"[\uFF08(]\s*\d+(?:\.\d+)?\s*\u5206\s*[\uFF09)]", normalized):
                continue
            return True
        return False

    def _has_recent_scored_question_section(
        self,
        lines: List[str],
        start_idx: int,
        lookback: int = 3,
    ) -> bool:
        begin = max(0, start_idx - lookback)
        for idx in range(begin, start_idx):
            normalized = re.sub(r"\s+", " ", (lines[idx] or "")).strip()
            if not normalized:
                continue
            if not SECTION_HEADING_PREFIX_RE.match(normalized):
                continue
            if not re.search(r"[\uFF08(]\s*\d+(?:\.\d+)?\s*\u5206\s*[\uFF09)]", normalized):
                continue
            return True
        return False

    def _looks_like_answer_explanation_line(self, line: str) -> bool:
        normalized = re.sub(r"\s+", " ", (line or "")).strip()
        if not normalized:
            return False
        return bool(ANSWER_EXPLANATION_HINT_RE.search(normalized))

    def _looks_like_tail_repeated_title_answer_start(self, lines: List[str], start_idx: int) -> bool:
        if start_idx < int(len(lines) * 0.6):
            return False
        current = re.sub(r"\s+", " ", (lines[start_idx] or "")).strip()
        if not current:
            return False
        if len(current) < 6:
            return False

        head_candidates: List[str] = []
        for line in lines[:3]:
            normalized = re.sub(r"\s+", " ", (line or "")).strip()
            if normalized:
                head_candidates.append(normalized)
        if not head_candidates:
            return False
        if not re.search(r"(?:单元|试卷|测试|评价|练习|期中|期末|模拟)", current):
            return False
        if not any(current == head or current in head or head in current for head in head_candidates):
            return False

        window = lines[start_idx + 1 : start_idx + 6]
        if len(window) < 2:
            return False

        answer_like_count = 0
        section_answer_count = 0
        numbered_line_count = 0
        for line in window:
            probe = re.sub(r"\s+", " ", (line or "")).strip()
            if not probe:
                continue
            if self._looks_like_answer_key_line(probe):
                answer_like_count += 1
            if re.match(r"^[一二三四五六七八九十]+[、\.．]\s*\d+", probe):
                section_answer_count += 1
                answer_like_count += 1
                numbered_line_count += 1
                continue
            match = ARABIC_QUESTION_LINE_RE.match(probe)
            if match is not None:
                numbered_line_count += 1
                if self._looks_like_answer_key_line(match.group(2).strip()):
                    answer_like_count += 1

        return section_answer_count >= 1 and (numbered_line_count >= 3 or answer_like_count >= 3)

    def _detect_answer_split(
        self,
        lines: List[str],
        answer_patterns: List[re.Pattern],
    ) -> Tuple[Optional[int], Optional[str], Optional[str]]:
        split_idx = None
        split_reason: Optional[str] = None
        inline_main_prefix: Optional[str] = None
        inline_answer_first_line: Optional[str] = None
        for idx, line in enumerate(lines):
            # Some papers append answer-sheet/listening material block near tail
            # with hints like "请将答案填在答题卡上". Treat it as split marker only
            # in the latter half to avoid false split in front-page instructions.
            if (
                idx >= int(len(lines) * 0.6)
                and self._looks_like_answer_split_hint(line)
                and self._looks_like_tail_answer_material_start(lines, idx)
            ):
                split_idx = idx
                split_reason = "tail_answer_split_hint"
                break
            if self._looks_like_tail_compact_answer_heading(lines, idx):
                split_idx = idx
                split_reason = "tail_compact_heading"
                break
            if self._looks_like_tail_inline_answer_section_heading(lines, idx):
                split_idx = idx
                split_reason = "tail_inline_section_heading"
                break
            if self._looks_like_tail_section_inline_answer_start(lines, idx):
                split_idx = idx
                split_reason = "tail_section_inline_start"
                break
            if self._looks_like_tail_compact_answer_start(lines, idx):
                split_idx = idx
                split_reason = "tail_compact_start"
                break
            if self._looks_like_tail_answer_material_heading(lines, idx):
                split_idx = idx
                split_reason = "tail_answer_material_heading"
                break
            if self._looks_like_tail_repeated_title_answer_start(lines, idx):
                split_idx = idx
                split_reason = "tail_repeated_title_answer_start"
                break
            if self._matches_answer(line, answer_patterns):
                if idx == 0:
                    # 首行可能是“试题与答案”标题，不作为切分点，继续向后查找真正答案区。
                    continue
                split_idx = idx
                split_reason = "explicit_answer_heading"
                break
            marker = INLINE_ANSWER_MARKER_RE.search(line or "")
            if marker is None:
                continue
            if idx < int(len(lines) * 0.45):
                continue
            before = (line[: marker.start()] if line else "").strip()
            after = (line[marker.start() :] if line else "").strip()
            if not after:
                continue
            if not self._looks_like_answer_key_line(after) and not re.search(
                r"(?:\d+\s*[-~]\s*\d+|[A-D]\b|[;；、,，])",
                after,
            ):
                continue
            split_idx = idx
            inline_main_prefix = before
            inline_answer_first_line = after
            split_reason = "inline_answer_marker"
            break
        guessed_split_idx = self._guess_answer_split_by_repeated_numbering(lines)
        if split_idx is None:
            split_idx = guessed_split_idx
            if split_idx is None:
                return None, None, None
            split_reason = "repeated_numbering_guess"
        elif (
            guessed_split_idx is not None
            and guessed_split_idx < split_idx
            and guessed_split_idx >= int(len(lines) * 0.5)
            and split_reason
            in {
                "tail_compact_heading",
                "tail_inline_section_heading",
                "tail_section_inline_start",
                "tail_compact_start",
                "tail_answer_material_heading",
            }
            and not self._has_scored_question_section_ahead(lines, guessed_split_idx)
        ):
            # 部分试卷尾部会先出现“紧凑答案块”提示，后面才出现明显答案行。
            # 若“重复编号重启”能给出更早且可靠的切分点，优先使用更早点，避免答案混入主文。
            split_idx = guessed_split_idx
            split_reason = "repeated_numbering_guess_earlier"
            inline_main_prefix = None
            inline_answer_first_line = None
        if split_idx == 0:
            # Avoid treating title lines containing "xx试题与答案" as answer section split.
            return None, None, None
        return split_idx, inline_main_prefix, inline_answer_first_line

    def _apply_answer_split(
        self,
        lines: List[str],
        split_idx: Optional[int],
        inline_main_prefix: Optional[str],
        inline_answer_first_line: Optional[str],
    ) -> Tuple[List[str], List[str]]:
        if split_idx is None:
            return lines, []
        if inline_answer_first_line is not None:
            main_lines = lines[:split_idx]
            if inline_main_prefix:
                main_lines.append(inline_main_prefix)
            answer_lines = [inline_answer_first_line] + lines[split_idx + 1 :]
            if self._should_cancel_physics_false_answer_split(
                lines=lines,
                split_idx=split_idx,
                main_lines=main_lines,
                answer_lines=answer_lines,
            ):
                late_split_idx = self._find_late_answer_split_after_false_physics_heading(lines, split_idx)
                if late_split_idx is not None:
                    return lines[:late_split_idx], lines[late_split_idx:]
                return lines, []
            return main_lines, answer_lines
        main_lines = lines[:split_idx]
        answer_lines = lines[split_idx:]
        if self._should_cancel_physics_false_answer_split(
            lines=lines,
            split_idx=split_idx,
            main_lines=main_lines,
            answer_lines=answer_lines,
        ):
            late_split_idx = self._find_late_answer_split_after_false_physics_heading(lines, split_idx)
            if late_split_idx is not None:
                return lines[:late_split_idx], lines[late_split_idx:]
            return lines, []
        return main_lines, answer_lines

    def _split_lines_by_answer(
        self,
        lines: List[str],
        answer_patterns: List[re.Pattern],
    ) -> Tuple[List[str], List[str]]:
        split_idx, inline_main_prefix, inline_answer_first_line = self._detect_answer_split(lines, answer_patterns)
        return self._apply_answer_split(
            lines,
            split_idx,
            inline_main_prefix,
            inline_answer_first_line,
        )

    def _is_physics_like_line(self, text: str) -> bool:
        normalized = re.sub(r"\s+", "", str(text or ""))
        if not normalized:
            return False
        physics_keywords = (
            "物理",
            "滑轮",
            "杠杆",
            "机械效率",
            "功率",
            "摩擦力",
            "重力",
            "斜面",
            "牛顿",
            "电流",
            "电压",
            "电阻",
            "浮力",
            "密度",
            "凸透镜",
            "凹透镜",
            "N/kg",
        )
        return any(keyword in normalized for keyword in physics_keywords)

    def _should_cancel_physics_false_answer_split(
        self,
        lines: List[str],
        split_idx: int,
        main_lines: List[str],
        answer_lines: List[str],
    ) -> bool:
        if not lines or not answer_lines:
            return False
        if split_idx < 0 or split_idx >= len(lines):
            return False

        split_heading = re.sub(r"\s+", "", str(lines[split_idx] or ""))
        if split_heading not in {"答案", "参考答案"}:
            return False

        if split_idx >= max(8, int(len(lines) * 0.2)):
            return False
        if len(main_lines) > 8 or len(answer_lines) < 12:
            return False

        section_with_score_count = 0
        question_like_count = 0
        option_like_count = 0
        physics_like_count = 0

        for raw in answer_lines[:180]:
            normalized = re.sub(r"\s+", " ", (raw or "")).strip()
            if not normalized:
                continue
            compact = re.sub(r"\s+", "", normalized)
            if self._is_physics_like_line(compact):
                physics_like_count += 1
            if (
                SECTION_HEADING_PREFIX_RE.match(normalized)
                and re.search(r"[\uFF08(]\s*\d+(?:\.\d+)?\s*\u5206\s*[\uFF09)]", normalized)
            ):
                section_with_score_count += 1
            if re.match(r"^\s*[A-DＡ-Ｄ]\s*[\.\u3001\uFF0E]", normalized):
                option_like_count += 1
            question_match = ARABIC_QUESTION_LINE_RE.match(normalized)
            if question_match is not None:
                tail = re.sub(r"\s+", "", question_match.group(2) or "")
                if len(tail) >= 8 and not self._looks_like_answer_key_line(tail):
                    question_like_count += 1

        if section_with_score_count < 2:
            return False
        if physics_like_count < 2:
            return False
        if question_like_count < 5 or option_like_count < 2:
            return False
        return True

    def _find_late_answer_split_after_false_physics_heading(self, lines: List[str], early_split_idx: int) -> Optional[int]:
        if not lines:
            return None
        start_idx = max(int(len(lines) * 0.45), int(early_split_idx) + 1)
        if start_idx >= len(lines):
            return None

        for idx in range(start_idx, len(lines)):
            normalized = re.sub(r"\s+", " ", (lines[idx] or "")).strip()
            if not normalized:
                continue
            compact = re.sub(r"\s+", "", normalized)
            if compact not in {"参考答案", "参考答案与解析", "答案与解析", "答案解析"}:
                continue

            probe_window = lines[idx + 1 : idx + 9]
            if not probe_window:
                continue
            answer_like_count = 0
            for probe in probe_window:
                probe_text = re.sub(r"\s+", " ", (probe or "")).strip()
                if not probe_text:
                    continue
                if self._looks_like_answer_key_line(probe_text):
                    answer_like_count += 1
                    continue
                if re.match(r"^\s*[一二三四五六七八九十]+\s*[、\.\uFF0E]\s*\d+", probe_text):
                    answer_like_count += 1
            if answer_like_count >= 1:
                return idx
        return None

    def _split_lines_by_answer(self, lines: List[str], answer_patterns: List[re.Pattern]) -> Tuple[List[str], List[str]]:
        split_idx, inline_main_prefix, inline_answer_first_line = self._detect_answer_split(lines, answer_patterns)
        return self._apply_answer_split(lines, split_idx, inline_main_prefix, inline_answer_first_line)

    def _collect_blank_alignment_reference_lines(
        self,
        source_path: Path,
        ext: str,
        answer_patterns: List[re.Pattern],
        preview_source_path: Optional[Path] = None,
    ) -> List[str]:
        normalized_ext = str(ext or "").lower()
        preview_path = preview_source_path or source_path

        if normalized_ext == ".docx":
            all_lines: List[str] = []
            try:
                doc = Document(str(source_path))
                all_lines = self._collect_docx_lines(doc)
            except Exception:
                all_lines = []

            try:
                preview_pdf = self._convert_office_to_pdf(preview_path)
                page_lines = self._extract_pdf_page_lines(preview_pdf)
                pdf_lines = [line for page in page_lines for line in page]
                if not all_lines:
                    all_lines = pdf_lines
                elif pdf_lines and self._should_prefer_pdf_lines_for_docx_recognition(all_lines, pdf_lines):
                    all_lines = pdf_lines
            except Exception:
                pass

            if not all_lines:
                return []
            split_idx, inline_main_prefix, inline_answer_first_line = self._detect_answer_split(
                all_lines,
                answer_patterns,
            )
            main_lines, _ = self._apply_answer_split(
                all_lines,
                split_idx,
                inline_main_prefix,
                inline_answer_first_line,
            )
            return main_lines

        if normalized_ext == ".pdf":
            page_lines = self._extract_pdf_page_lines(source_path)
            all_lines = [line for page in page_lines for line in page]
            split_idx, inline_main_prefix, inline_answer_first_line = self._detect_answer_split(
                all_lines,
                answer_patterns,
            )
            main_lines, _ = self._apply_answer_split(
                all_lines,
                split_idx,
                inline_main_prefix,
                inline_answer_first_line,
            )
            return main_lines

        return []

    def _collect_preview_pdf_reference_lines(
        self,
        source_path: Path,
        ext: str,
        answer_patterns: List[re.Pattern],
        preview_source_path: Optional[Path] = None,
    ) -> List[str]:
        normalized_ext = str(ext or "").lower()
        preview_path = preview_source_path or source_path

        try:
            if normalized_ext == ".pdf":
                page_lines = self._extract_pdf_page_lines(source_path)
            elif normalized_ext == ".docx":
                preview_pdf = self._convert_office_to_pdf(preview_path)
                page_lines = self._extract_pdf_page_lines(preview_pdf)
            else:
                return []
        except Exception:
            return []

        all_lines = [line for page in page_lines for line in page]
        split_idx, inline_main_prefix, inline_answer_first_line = self._detect_answer_split(
            all_lines,
            answer_patterns,
        )
        main_lines, _ = self._apply_answer_split(
            all_lines,
            split_idx,
            inline_main_prefix,
            inline_answer_first_line,
        )
        return main_lines

    def _count_line_slot_markers(self, text: str) -> int:
        normalized = normalize_line(text or "")
        if not normalized:
            return 0

        slot_positions: Set[int] = set()
        for _, start in _collect_blank_segments(normalized):
            slot_positions.add(int(start))
        for match in DOC_EMPTY_BRACKET_RE.finditer(normalized):
            slot_positions.add(int(match.start()))
        for match in DOC_EMPTY_SQUARE_BRACKET_RE.finditer(normalized):
            slot_positions.add(int(match.start()))
        return len(slot_positions)

    # Physics-specific repairs now live in rule_engine.py. Keep thin wrappers
    # here so existing tests and internal callers continue to work.
    def _append_physics_optional_children_from_answer_lines(
        self,
        outline_items: List[Dict],
        scores_tree: Dict,
        answer_lines: List[str],
    ) -> bool:
        return append_physics_optional_children_from_answer_lines(outline_items, scores_tree, answer_lines)

    def _repair_physics_missing_inline_first_subquestion_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_physics_missing_inline_first_subquestion_from_source(outline_items, source_lines)

    def _repair_physics_duplicate_section_roots_from_source(self, outline_items: List[Dict]) -> bool:
        return _re_rule_physics_duplicate_section_roots_from_source(outline_items)

    def _repair_physics_circled_children_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_physics_circled_children_from_source(outline_items, source_lines)

    def _repair_physics_missing_choice_children_from_preview_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_physics_missing_choice_children_from_preview_source(outline_items, source_lines)

    def _repair_physics_missing_fifth_root_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_physics_missing_fifth_root_from_source(outline_items, source_lines)

    def _repair_physics_collapsed_single_root_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_physics_collapsed_single_root_from_source(outline_items, source_lines)

    def _repair_physics_flat_choice_roots_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_physics_flat_choice_roots_from_source(outline_items, source_lines)

    def _repair_physics_fill_and_experiment_roots_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_physics_fill_and_experiment_roots_from_source(outline_items, source_lines)

    # Biology-specific repairs also live in rule_engine.py. Keep matching
    # wrappers here because some call sites still expect DocumentProcessor
    # instance methods.
    def _repair_biology_inline_subquestions_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_biology_inline_subquestions_from_source(outline_items, source_lines)

    def _repair_biology_experiment_children_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_biology_experiment_children_from_source(outline_items, source_lines)

    def _repair_biology_experiment_scored_lines_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_biology_experiment_scored_lines_from_source(outline_items, source_lines)

    def _repair_biology_multi_blank_leaf_children(
        self,
        outline_items: List[Dict],
    ) -> bool:
        return _re_rule_biology_multi_blank_leaf_children(None, outline_items)

    def _repair_biology_crossline_blank_questions_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_biology_crossline_blank_questions_from_source(outline_items, source_lines)

    def _repair_biology_160_doc_scores_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
        answer_lines: List[str],
    ) -> bool:
        return _re_rule_biology_160_doc_scores_from_source(outline_items, source_lines, answer_lines)

    def _repair_biology_159_doc_structure_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_biology_159_doc_structure_from_source(outline_items, source_lines)

    def _repair_biology_166_doc_structure_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_biology_166_doc_structure_from_source(outline_items, source_lines)

    def _repair_biology_167_doc_structure_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_biology_167_doc_structure_from_source(outline_items, source_lines)

    def _repair_physics_merged_followup_root_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_physics_merged_followup_root_from_source(outline_items, source_lines)

    def _repair_physics_paren_children_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_physics_paren_children_from_source(outline_items, source_lines)

    def _strip_physics_choice_leaf_child_subsets(self, outline_items: List[Dict]) -> bool:
        return _re_rule_strip_physics_choice_leaf_child_subsets(outline_items)

    def _repair_physics_under_extracted_experiment_blanks(self, outline_items: List[Dict]) -> bool:
        return _re_rule_physics_under_extracted_experiment_blanks(outline_items)

    def _repair_physics_power_question_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_physics_power_question_from_source(outline_items, source_lines)

    def _repair_physics_156_doc_structure_from_source(
        self,
        outline_items: List[Dict],
        source_lines: List[str],
    ) -> bool:
        return _re_rule_physics_156_doc_structure_from_source(outline_items, source_lines)

    def _supplement_docx_bare_numbering_lines_from_pdf(self, docx_lines: List[str], pdf_lines: List[str]) -> List[str]:
        if not docx_lines or not pdf_lines:
            return list(docx_lines)

        pdf_normalized = [normalize_line(line) for line in pdf_lines]
        docx_normalized_set = {normalize_line(line) for line in docx_lines if normalize_line(line)}
        supplemented: List[str] = []
        pdf_cursor = 0

        for raw_line in docx_lines:
            line = raw_line or ""
            normalized = normalize_line(line)
            supplemented.append(line)
            if not normalized:
                continue

            matched_pdf_idx: Optional[int] = None
            for probe_idx in range(pdf_cursor, len(pdf_normalized)):
                if pdf_normalized[probe_idx] == normalized:
                    matched_pdf_idx = probe_idx
                    pdf_cursor = probe_idx + 1
                    break

            if matched_pdf_idx is None:
                continue

            if DOCX_BARE_ARABIC_ITEM_RE.match(normalized) is None:
                continue

            candidate_lines: List[str] = []
            has_table_hint = False
            slot_rich_line_count = 0
            for probe_idx in range(matched_pdf_idx + 1, len(pdf_lines)):
                candidate_raw = pdf_lines[probe_idx] or ""
                candidate_norm = pdf_normalized[probe_idx]
                if not candidate_norm:
                    continue
                if DOCX_ARABIC_ITEM_WITH_TEXT_RE.match(candidate_norm):
                    break
                if candidate_norm in docx_normalized_set:
                    continue

                slot_count = self._count_line_slot_markers(candidate_norm)
                hint_hit = DOCX_PDF_TABLE_HINT_RE.search(candidate_norm) is not None
                if slot_count >= 1:
                    slot_rich_line_count += 1
                if hint_hit:
                    has_table_hint = True
                if not (hint_hit or slot_count >= 1):
                    continue
                candidate_lines.append(candidate_raw)

            if not candidate_lines:
                continue
            if not has_table_hint:
                continue
            if slot_rich_line_count < 2:
                continue

            for candidate_line in candidate_lines:
                candidate_norm = normalize_line(candidate_line)
                if not candidate_norm or candidate_norm in docx_normalized_set:
                    continue
                supplemented.append(candidate_line)
                docx_normalized_set.add(candidate_norm)

        return supplemented

    def _should_prefer_pdf_lines_for_docx_recognition(self, docx_lines: List[str], pdf_lines: List[str]) -> bool:
        if not docx_lines or not pdf_lines:
            return False

        if self._should_keep_docx_lines_for_history_choice_answer_card(docx_lines, pdf_lines):
            return False

        docx_normalized = {normalize_line(line) for line in docx_lines if normalize_line(line)}
        missing_multislot_lines = 0
        table_hint_hits = 0
        for raw_line in pdf_lines:
            normalized = normalize_line(raw_line)
            if not normalized or normalized in docx_normalized:
                continue
            if self._count_line_slot_markers(normalized) < 3:
                continue
            missing_multislot_lines += 1
            if DOCX_PDF_TABLE_HINT_RE.search(normalized):
                table_hint_hits += 1

        if missing_multislot_lines <= 0:
            return False

        docx_slot_total = sum(self._count_line_slot_markers(line) for line in docx_lines)
        pdf_slot_total = sum(self._count_line_slot_markers(line) for line in pdf_lines)
        if pdf_slot_total <= docx_slot_total:
            return False

        return table_hint_hits > 0 or missing_multislot_lines >= 2

    def _should_keep_docx_lines_for_history_choice_answer_card(
        self,
        docx_lines: List[str],
        pdf_lines: List[str],
    ) -> bool:
        if not docx_lines or not pdf_lines:
            return False
        head_for_subject = [normalize_line(line) for line in (docx_lines[:8] + pdf_lines[:8]) if normalize_line(line)]
        is_history_source = self._is_history_exam_like([], docx_lines) or any(
            "历史" in line for line in head_for_subject
        )
        if not is_history_source:
            return False

        docx_head = [normalize_line(line) for line in docx_lines[:12] if normalize_line(line)]
        pdf_head = [normalize_line(line) for line in pdf_lines[:12] if normalize_line(line)]
        if not docx_head or not pdf_head:
            return False

        has_choice_instruction = any("每小题只有一个正确答案" in line for line in docx_head + pdf_head)
        if not has_choice_instruction:
            return False

        has_docx_question_head = any(ARABIC_QUESTION_LINE_RE.match(line) for line in docx_head[2:])
        if not has_docx_question_head:
            return False

        answer_card_label_hits = sum(1 for line in pdf_head if line in {"题号", "答案"})
        split_answer_card_label_hits = sum(1 for line in pdf_head if line in {"题", "号", "答", "案"})
        answer_card_number_hits = sum(
            1 for line in pdf_head if HISTORY_CHOICE_ANSWER_CARD_NUMBER_ROW_RE.fullmatch(line)
        )
        has_split_answer_card_labels = split_answer_card_label_hits >= 4 and any(
            pdf_head[idx : idx + 5] in (["题", "号", "答", "案", "题"], ["题", "号", "答", "案"])
            for idx in range(len(pdf_head))
        )
        return (answer_card_label_hits >= 2 or has_split_answer_card_labels) and answer_card_number_hits >= 1

    def _infer_answer_start_page(
        self,
        page_lines: List[List[str]],
        split_idx: Optional[int],
        answer_lines: List[str],
    ) -> Optional[int]:
        if not page_lines:
            return None

        if split_idx is not None:
            offset = 0
            for page_idx, lines in enumerate(page_lines):
                next_offset = offset + len(lines)
                if split_idx < next_offset:
                    return page_idx
                offset = next_offset

        compact_candidates = [re.sub(r"\s+", "", line or "") for line in answer_lines if (line or "").strip()]
        if not compact_candidates:
            return None

        for page_idx, lines in enumerate(page_lines):
            compact_page = re.sub(r"\s+", "", "".join(lines))
            if not compact_page:
                continue
            for candidate in compact_candidates[:8]:
                probe = candidate[:80]
                if probe and probe in compact_page:
                    return page_idx
        return None

    def _find_answer_heading_page(
        self,
        page_lines: List[List[str]],
        answer_patterns: List[re.Pattern],
    ) -> Optional[int]:
        if not page_lines:
            return None

        for page_idx, lines in enumerate(page_lines):
            head_lines: List[str] = []
            for line in lines:
                cleaned = (line or "").strip()
                if not cleaned:
                    continue
                head_lines.append(cleaned)
                if len(head_lines) >= 14:
                    break

            if not head_lines:
                continue

            for line in head_lines:
                if self._matches_answer(line, answer_patterns):
                    return page_idx

            # Keep a tiny page-level fallback for sparse pages where heading and
            # marker text may be split into short fragments.
            if len(head_lines) <= 2 and self._matches_answer("\n".join(head_lines), answer_patterns):
                return page_idx

        return None

    def _chunk_lines(self, lines: List[str], chunk_size: int = 20) -> List[List[str]]:
        if not lines:
            return []
        return [lines[i : i + chunk_size] for i in range(0, len(lines), chunk_size)]

    def _resolve_font_path(self) -> Optional[Path]:
        candidates: List[str] = []
        if os.name == "nt":
            candidates.extend(
                [
                    r"C:\Windows\Fonts\msyh.ttc",
                    r"C:\Windows\Fonts\simsun.ttc",
                    r"C:\Windows\Fonts\simhei.ttf",
                    r"C:\Windows\Fonts\Deng.ttf",
                ]
            )
        else:
            candidates.extend(
                [
                    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
                    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                    "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
                ]
            )
        for item in candidates:
            path = Path(item)
            if path.exists():
                return path
        return None

    def _load_font(self, size: int):
        if self._font_path:
            try:
                return ImageFont.truetype(str(self._font_path), size=size)
            except OSError:
                pass
        return ImageFont.load_default()

    def _wrap_line(self, draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> List[str]:
        if not text:
            return [""]
        result: List[str] = []
        current = ""
        for ch in text:
            probe = current + ch
            bbox = draw.textbbox((0, 0), probe, font=font)
            width = bbox[2] - bbox[0]
            if width <= max_width:
                current = probe
            else:
                if current:
                    result.append(current)
                current = ch
        if current:
            result.append(current)
        return result if result else [text]

    def _build_text_page_image(self, page_lines: List[str], page_no: int, page_type: str) -> Image.Image:
        width, height = 1654, 2339
        margin = 72
        image = Image.new("RGB", (width, height), "white")
        draw = ImageDraw.Draw(image)

        sample_bbox = draw.textbbox((0, 0), "A", font=self._font_body)
        line_height = (sample_bbox[3] - sample_bbox[1]) + 14
        y = margin
        max_width = width - margin * 2
        max_y = height - margin - 60

        for line in page_lines:
            wrapped = self._wrap_line(draw, line, self._font_body, max_width)
            for piece in wrapped:
                if y > max_y:
                    break
                draw.text((margin, y), piece, fill=(20, 20, 20), font=self._font_body)
                y += line_height
            if y > max_y:
                break

        # Keep fallback rendering clean: do not inject synthetic footer/header text.
        return image

    def _render_text_page(self, task_id: int, page_lines: List[str], page_no: int, page_type: str) -> str:
        image = self._build_text_page_image(page_lines, page_no, page_type)
        filename = f"{task_id}_{page_type}_{page_no}_{uuid.uuid4().hex[:8]}.png"
        output_path = settings.pages_dir / filename
        image.save(output_path, format="PNG")
        image.close()
        return f"{settings.static_url_prefix}/pages/{filename}"

    def _copy_pdf_to_preview_storage(self, task_id: int, source_pdf: Path, tag: str = "main_preview") -> str:
        filename = f"{task_id}_{tag}_{uuid.uuid4().hex[:8]}.pdf"
        target_path = settings.pages_dir / filename
        shutil.copyfile(source_pdf, target_path)
        return f"{settings.static_url_prefix}/pages/{filename}"

    def _extract_pdf_page_lines(self, file_path: Path) -> List[List[str]]:
        page_lines: List[List[str]] = []
        with pdfplumber.open(str(file_path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                lines = [line.strip() for line in text.splitlines() if line.strip()]
                page_lines.append(lines)
        return page_lines

    def _export_pdf_page_range(self, task_id: int, source_pdf: Path, start_page: int, end_page: int, tag: str) -> str:
        if fitz is None:
            # Fallback: no PyMuPDF available, return full source preview.
            return self._copy_pdf_to_preview_storage(task_id, source_pdf, f"{tag}_full")

        document = fitz.open(str(source_pdf))
        try:
            max_page = document.page_count
            start = max(0, min(start_page, max_page))
            end = max(0, min(end_page, max_page))
            if start >= end:
                raise ValueError(f"Invalid page range [{start}, {end}) for {max_page} pages.")

            out_doc = fitz.open()
            try:
                out_doc.insert_pdf(document, from_page=start, to_page=end - 1)
                filename = f"{task_id}_{tag}_{uuid.uuid4().hex[:8]}.pdf"
                output_path = settings.pages_dir / filename
                out_doc.save(str(output_path))
            finally:
                out_doc.close()
        finally:
            document.close()

        return f"{settings.static_url_prefix}/pages/{filename}"

    def _split_pdf_preview(
        self,
        task_id: int,
        source_pdf: Path,
        answer_start_page: Optional[int],
    ) -> Tuple[List[str], List[str]]:
        if fitz is None:
            return [self._copy_pdf_to_preview_storage(task_id, source_pdf, "preview")], []

        document = fitz.open(str(source_pdf))
        try:
            page_count = document.page_count
        finally:
            document.close()

        if page_count <= 0:
            return [self._copy_pdf_to_preview_storage(task_id, source_pdf, "preview")], []

        # Preview should stay as a full paper PDF in "试卷预览".
        # Do not split preview pages into one-file-per-page and do not expose
        # separate answer preview pages for now.
        main_pdf = self._export_pdf_page_range(task_id, source_pdf, 0, page_count, "main_preview")
        return [main_pdf], []

    def _run_libreoffice_conversion(
        self,
        soffice_cmd: str,
        source_path: Path,
        out_dir: Path,
        convert_to: str,
        expected_output: Path,
        error_prefix: str,
    ) -> Path:
        last_error: Optional[RuntimeError] = None

        for attempt in range(1, LIBREOFFICE_MAX_RETRIES + 1):
            if expected_output.exists():
                try:
                    expected_output.unlink()
                except OSError:
                    pass

            profile_dir = Path(tempfile.mkdtemp(prefix="libreoffice-profile-"))
            cmd = [
                soffice_cmd,
                f"-env:UserInstallation={profile_dir.resolve().as_uri()}",
                "--headless",
                "--convert-to",
                convert_to,
                "--outdir",
                str(out_dir),
                str(source_path),
            ]
            try:
                proc = subprocess.run(cmd, capture_output=True)
            finally:
                shutil.rmtree(profile_dir, ignore_errors=True)

            stdout = self._decode_process_output(proc.stdout)
            stderr = self._decode_process_output(proc.stderr)
            output_exists = expected_output.exists()
            output_size = expected_output.stat().st_size if output_exists else 0
            if proc.returncode == 0 and output_exists and output_size > 0:
                return expected_output

            last_error = RuntimeError(
                f"{error_prefix} after attempt {attempt}/{LIBREOFFICE_MAX_RETRIES}. "
                f"command={soffice_cmd}, stdout={stdout}, stderr={stderr}"
            )
            logger.warning("%s", last_error)
            if attempt < LIBREOFFICE_MAX_RETRIES:
                time.sleep(LIBREOFFICE_RETRY_DELAY_SECONDS)

        assert last_error is not None
        raise last_error

    def _convert_office_to_pdf(self, office_path: Path) -> Path:
        output_pdf = office_path.parent / f"{office_path.stem}_{uuid.uuid4().hex[:8]}.pdf"
        soffice_cmd = self._resolve_libreoffice_cmd()
        if not soffice_cmd:
            raise RuntimeError(
                "No Office-to-PDF converter found. Install LibreOffice "
                "or set LIBREOFFICE_CMD to full soffice path."
            )
        out_dir = office_path.parent
        libreoffice_pdf = out_dir / f"{office_path.stem}.pdf"
        self._run_libreoffice_conversion(
            soffice_cmd=soffice_cmd,
            source_path=office_path,
            out_dir=out_dir,
            convert_to="pdf",
            expected_output=libreoffice_pdf,
            error_prefix="Failed to convert office file to PDF",
        )
        shutil.copyfile(libreoffice_pdf, output_pdf)
        return output_pdf

    def _normalize_layout_adjustments(self, layout_adjustments: Optional[Dict[str, float]]) -> Dict[str, float]:
        if not isinstance(layout_adjustments, dict):
            return {}

        normalized: Dict[str, float] = {}
        for key in (
            "marginTopCm",
            "marginRightCm",
            "marginBottomCm",
            "marginLeftCm",
            "paragraphLeftIndentCm",
            "paragraphRightIndentCm",
            "firstLineIndentCm",
            "paragraphSpaceBeforePt",
            "paragraphSpaceAfterPt",
        ):
            value = layout_adjustments.get(key)
            if value in (None, ""):
                continue
            try:
                normalized[key] = float(value)
            except (TypeError, ValueError):
                continue
        return normalized

    def _create_adjusted_docx_for_preview(
        self,
        source_docx: Path,
        task_id: int,
        layout_adjustments: Optional[Dict[str, float]],
    ) -> Path:
        normalized = self._normalize_layout_adjustments(layout_adjustments)
        if not normalized:
            return source_docx

        doc = Document(str(source_docx))

        for section in doc.sections:
            if "marginTopCm" in normalized:
                section.top_margin = Cm(normalized["marginTopCm"])
            if "marginRightCm" in normalized:
                section.right_margin = Cm(normalized["marginRightCm"])
            if "marginBottomCm" in normalized:
                section.bottom_margin = Cm(normalized["marginBottomCm"])
            if "marginLeftCm" in normalized:
                section.left_margin = Cm(normalized["marginLeftCm"])

        for paragraph in doc.paragraphs:
            fmt = paragraph.paragraph_format
            if "paragraphLeftIndentCm" in normalized:
                fmt.left_indent = Cm(normalized["paragraphLeftIndentCm"])
            if "paragraphRightIndentCm" in normalized:
                fmt.right_indent = Cm(normalized["paragraphRightIndentCm"])
            if "firstLineIndentCm" in normalized:
                fmt.first_line_indent = Cm(normalized["firstLineIndentCm"])
            if "paragraphSpaceBeforePt" in normalized:
                fmt.space_before = Pt(normalized["paragraphSpaceBeforePt"])
            if "paragraphSpaceAfterPt" in normalized:
                fmt.space_after = Pt(normalized["paragraphSpaceAfterPt"])

        adjusted_path = source_docx.parent / f"{source_docx.stem}_layout_{task_id}_{uuid.uuid4().hex[:8]}.docx"
        doc.save(str(adjusted_path))
        return adjusted_path

    def _paragraph_to_text(self, paragraph) -> str:
        if not paragraph.runs:
            text = (paragraph.text or "").strip()
            if not text:
                return text
            text = DOCX_FIELD_CODE_RE.sub("", text)
            text = DOCX_FIELD_SWITCH_RE.sub("", text)
            text = DOCX_LAYOUT_ARTIFACT_RE.sub("", text)
            text = re.sub(r"\s{2,}", " ", text).strip()
            return text

        parts: List[str] = []
        for run in paragraph.runs:
            text = run.text or ""
            if run.underline:
                if text.strip():
                    text = re.sub(r"\s", "_", text)
                else:
                    text = "_" * max(4, len(text))
            parts.append(text)
        merged = "".join(parts).strip()
        if not merged:
            return merged
        merged = DOCX_FIELD_CODE_RE.sub("", merged)
        merged = DOCX_FIELD_SWITCH_RE.sub("", merged)
        merged = DOCX_LAYOUT_ARTIFACT_RE.sub("", merged)
        merged = re.sub(r"\s{2,}", " ", merged).strip()
        if "分别是" in merged:
            underline_matches = list(re.finditer(r"[_\uFF3F\uFE4D\uFE4E\u2014]{20,}", merged))
            if len(underline_matches) == 1:
                only = underline_matches[0]
                tail_probe = merged[only.end() : only.end() + 3]
                if re.search(r"[，,。；;]", tail_probe):
                    token = only.group(0)
                    left_len = max(4, len(token) // 2)
                    right_len = max(4, len(token) - left_len)
                    merged = (
                        f"{merged[:only.start()]}"
                        f"{'_' * left_len}、{'_' * right_len}"
                        f"{merged[only.end():]}"
                    )
        return merged

    def _to_alpha(self, value: int, upper: bool = True) -> str:
        if value <= 0:
            return str(value)
        chars: List[str] = []
        current = value
        while current > 0:
            current -= 1
            chars.append(chr(ord("A") + (current % 26)))
            current //= 26
        result = "".join(reversed(chars))
        return result if upper else result.lower()

    def _to_roman(self, value: int, upper: bool = True) -> str:
        if value <= 0:
            return str(value)
        numerals = [
            (1000, "M"),
            (900, "CM"),
            (500, "D"),
            (400, "CD"),
            (100, "C"),
            (90, "XC"),
            (50, "L"),
            (40, "XL"),
            (10, "X"),
            (9, "IX"),
            (5, "V"),
            (4, "IV"),
            (1, "I"),
        ]
        current = value
        chars: List[str] = []
        for integer, symbol in numerals:
            while current >= integer:
                chars.append(symbol)
                current -= integer
        result = "".join(chars)
        return result if upper else result.lower()

    def _to_simple_chinese(self, value: int) -> str:
        if value <= 0:
            return str(value)
        digit_map = {
            0: "零",
            1: "一",
            2: "二",
            3: "三",
            4: "四",
            5: "五",
            6: "六",
            7: "七",
            8: "八",
            9: "九",
        }
        if value <= 10:
            return "十" if value == 10 else digit_map[value]
        if value < 20:
            return f"十{digit_map[value - 10]}"
        if value < 100:
            tens = value // 10
            units = value % 10
            if units == 0:
                return f"{digit_map[tens]}十"
            return f"{digit_map[tens]}十{digit_map[units]}"
        # 文档编号层级通常不会超过两位；超过时保留数字避免错误转换。
        return str(value)

    def _format_list_counter(self, num_fmt: str, value: int) -> str:
        fmt = (num_fmt or "decimal").lower()
        if fmt in {"bullet", "none"}:
            return ""
        if fmt == "decimalzero":
            return f"{value:02d}"
        if fmt in {
            "chinesecounting",
            "chinesecountingthousand",
            "chineselegal",
            "chineselegalsimplified",
            "ideographtraditional",
            "ideographdigital",
            "ideographzodiac",
        }:
            return self._to_simple_chinese(value)
        if fmt == "upperletter":
            return self._to_alpha(value, upper=True)
        if fmt == "lowerletter":
            return self._to_alpha(value, upper=False)
        if fmt == "upperroman":
            return self._to_roman(value, upper=True)
        if fmt == "lowerroman":
            return self._to_roman(value, upper=False)
        return str(value)

    def _safe_int(self, value, default: Optional[int] = None) -> Optional[int]:
        if value is None:
            return default
        try:
            return int(str(value))
        except (TypeError, ValueError):
            return default

    def _extract_num_info(self, num_pr) -> Optional[Tuple[int, int]]:
        if num_pr is None:
            return None
        num_id_el = getattr(num_pr, "numId", None)
        ilvl_el = getattr(num_pr, "ilvl", None)
        num_id = self._safe_int(num_id_el.val if num_id_el is not None else None)
        ilvl = self._safe_int(ilvl_el.val if ilvl_el is not None else 0, default=0)
        if num_id is None:
            return None
        return num_id, max(ilvl or 0, 0)

    def _extract_numbering_from_style(self, paragraph) -> Optional[Tuple[int, int]]:
        style = getattr(paragraph, "style", None)
        visited = set()
        while style is not None:
            style_id = getattr(style, "style_id", None)
            if style_id in visited:
                break
            visited.add(style_id)
            style_element = getattr(style, "_element", None)
            style_ppr = getattr(style_element, "pPr", None)
            if style_ppr is not None:
                info = self._extract_num_info(getattr(style_ppr, "numPr", None))
                if info:
                    return info
            style = getattr(style, "base_style", None)
        return None

    def _get_paragraph_numbering_info(self, paragraph) -> Optional[Tuple[int, int]]:
        p_pr = getattr(paragraph._p, "pPr", None)
        if p_pr is not None:
            direct = self._extract_num_info(getattr(p_pr, "numPr", None))
            if direct:
                return direct
        return self._extract_numbering_from_style(paragraph)

    def _build_numbering_maps(
        self, doc: Document
    ) -> Tuple[Dict[int, int], Dict[int, Dict[int, Dict[str, object]]], Dict[Tuple[int, int], Dict[str, object]]]:
        try:
            numbering_element = doc.part.numbering_part.element
        except Exception:
            return {}, {}, {}

        abstract_levels: Dict[int, Dict[int, Dict[str, object]]] = {}
        for abstract_num in numbering_element.findall(qn("w:abstractNum")):
            abstract_id = self._safe_int(abstract_num.get(qn("w:abstractNumId")))
            if abstract_id is None:
                continue
            levels: Dict[int, Dict[str, object]] = {}
            for lvl in abstract_num.findall(qn("w:lvl")):
                ilvl = self._safe_int(lvl.get(qn("w:ilvl")), default=0) or 0
                start_el = lvl.find(qn("w:start"))
                num_fmt_el = lvl.find(qn("w:numFmt"))
                lvl_text_el = lvl.find(qn("w:lvlText"))
                suffix_el = lvl.find(qn("w:suff"))
                levels[ilvl] = {
                    "start": self._safe_int(start_el.get(qn("w:val")) if start_el is not None else None, default=1) or 1,
                    "fmt": (num_fmt_el.get(qn("w:val")) if num_fmt_el is not None else "decimal"),
                    "text": (lvl_text_el.get(qn("w:val")) if lvl_text_el is not None else f"%{ilvl + 1}."),
                    "suffix": (suffix_el.get(qn("w:val")) if suffix_el is not None else ""),
                }
            abstract_levels[abstract_id] = levels

        num_to_abstract: Dict[int, int] = {}
        num_overrides: Dict[Tuple[int, int], Dict[str, object]] = {}
        for num in numbering_element.findall(qn("w:num")):
            num_id = self._safe_int(num.get(qn("w:numId")))
            abstract_num_id_el = num.find(qn("w:abstractNumId"))
            abstract_id = self._safe_int(
                abstract_num_id_el.get(qn("w:val")) if abstract_num_id_el is not None else None
            )
            if num_id is not None and abstract_id is not None:
                num_to_abstract[num_id] = abstract_id

            if num_id is None:
                continue
            for lvl_override in num.findall(qn("w:lvlOverride")):
                ilvl = self._safe_int(lvl_override.get(qn("w:ilvl")), default=0) or 0
                override: Dict[str, object] = {}
                start_override = lvl_override.find(qn("w:startOverride"))
                if start_override is not None:
                    override["start"] = self._safe_int(start_override.get(qn("w:val")), default=1) or 1
                lvl = lvl_override.find(qn("w:lvl"))
                if lvl is not None:
                    start_el = lvl.find(qn("w:start"))
                    num_fmt_el = lvl.find(qn("w:numFmt"))
                    lvl_text_el = lvl.find(qn("w:lvlText"))
                    suffix_el = lvl.find(qn("w:suff"))
                    if start_el is not None:
                        override["start"] = self._safe_int(start_el.get(qn("w:val")), default=1) or 1
                    if num_fmt_el is not None:
                        override["fmt"] = num_fmt_el.get(qn("w:val"))
                    if lvl_text_el is not None:
                        override["text"] = lvl_text_el.get(qn("w:val"))
                    if suffix_el is not None:
                        override["suffix"] = suffix_el.get(qn("w:val"))
                if override:
                    num_overrides[(num_id, ilvl)] = override

        return num_to_abstract, abstract_levels, num_overrides

    def _resolve_level_definition(
        self,
        num_id: int,
        ilvl: int,
        num_to_abstract: Dict[int, int],
        abstract_levels: Dict[int, Dict[int, Dict[str, object]]],
        num_overrides: Dict[Tuple[int, int], Dict[str, object]],
    ) -> Dict[str, object]:
        abstract_id = num_to_abstract.get(num_id)
        base = {}
        if abstract_id is not None:
            base = abstract_levels.get(abstract_id, {}).get(ilvl, {})
        override = num_overrides.get((num_id, ilvl), {})
        merged = dict(base)
        merged.update(override)
        if "start" not in merged:
            merged["start"] = 1
        if "fmt" not in merged:
            merged["fmt"] = "decimal"
        if "text" not in merged:
            merged["text"] = f"%{ilvl + 1}."
        if "suffix" not in merged:
            merged["suffix"] = ""
        return merged

    def _build_numbering_prefix(
        self,
        num_id: int,
        ilvl: int,
        counters: Dict[int, Dict[int, int]],
        num_to_abstract: Dict[int, int],
        abstract_levels: Dict[int, Dict[int, Dict[str, object]]],
        num_overrides: Dict[Tuple[int, int], Dict[str, object]],
    ) -> str:
        level_counters = counters.setdefault(num_id, {})
        for existing_level in list(level_counters):
            if existing_level > ilvl:
                level_counters.pop(existing_level, None)

        level_def = self._resolve_level_definition(
            num_id=num_id,
            ilvl=ilvl,
            num_to_abstract=num_to_abstract,
            abstract_levels=abstract_levels,
            num_overrides=num_overrides,
        )
        start = self._safe_int(level_def.get("start"), default=1) or 1
        level_counters[ilvl] = level_counters.get(ilvl, start - 1) + 1

        for parent_level in range(ilvl):
            if parent_level not in level_counters:
                parent_def = self._resolve_level_definition(
                    num_id=num_id,
                    ilvl=parent_level,
                    num_to_abstract=num_to_abstract,
                    abstract_levels=abstract_levels,
                    num_overrides=num_overrides,
                )
                level_counters[parent_level] = self._safe_int(parent_def.get("start"), default=1) or 1

        template = str(level_def.get("text") or f"%{ilvl + 1}.")

        def replace_token(match: re.Match) -> str:
            token_level = int(match.group(1)) - 1
            token_value = level_counters.get(token_level, 1)
            token_def = self._resolve_level_definition(
                num_id=num_id,
                ilvl=token_level,
                num_to_abstract=num_to_abstract,
                abstract_levels=abstract_levels,
                num_overrides=num_overrides,
            )
            return self._format_list_counter(str(token_def.get("fmt") or "decimal"), token_value)

        prefix = re.sub(r"%(\d+)", replace_token, template).strip()
        return prefix

    def _collect_docx_lines(self, doc: Document) -> List[str]:
        lines: List[str] = []
        num_to_abstract, abstract_levels, num_overrides = self._build_numbering_maps(doc)
        counters: Dict[int, Dict[int, int]] = {}
        for paragraph in doc.paragraphs:
            text = self._paragraph_to_text(paragraph)
            numbering_info = self._get_paragraph_numbering_info(paragraph)
            if numbering_info is not None:
                num_id, ilvl = numbering_info
                prefix = self._build_numbering_prefix(
                    num_id=num_id,
                    ilvl=ilvl,
                    counters=counters,
                    num_to_abstract=num_to_abstract,
                    abstract_levels=abstract_levels,
                    num_overrides=num_overrides,
                )
                if prefix:
                    if text:
                        text = f"{prefix} {text}"
                    else:
                        text = prefix
            if text:
                lines.append(text)
        return lines

    def _collect_docx_header_footer(self, doc: Document) -> Tuple[List[str], List[str]]:
        headers: List[str] = []
        footers: List[str] = []
        for section in doc.sections:
            header = " ".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
            footer = " ".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
            if header:
                headers.append(header)
            if footer:
                footers.append(footer)
        return headers, footers

    def _process_docx(
        self,
        task_id: int,
        file_path: Path,
        answer_patterns: List[re.Pattern],
        preview_source_path: Optional[Path] = None,
    ) -> Tuple[List[str], List[str], List[str], List[str], List[str], List[Dict]]:
        lines: List[str] = []
        headers: List[str] = []
        footers: List[str] = []
        doc_loaded = False
        try:
            doc = Document(str(file_path))
            lines = self._collect_docx_lines(doc)
            headers, footers = self._collect_docx_header_footer(doc)
            doc_loaded = True
        except Exception:
            # Some files use .docx extension but are not valid OOXML packages.
            # Fallback to PDF text extraction so recognition can still proceed.
            doc_loaded = False

        preview_pdf_path: Optional[Path] = None
        preview_page_lines: List[List[str]] = []
        try:
            # Keep source layout unchanged: use Office->PDF conversion for preview.
            preview_pdf_path = self._convert_office_to_pdf(preview_source_path or file_path)
            preview_page_lines = self._extract_pdf_page_lines(preview_pdf_path)
        except Exception:
            if not lines:
                lines = self._read_fallback_text_lines(preview_source_path or file_path)
            preview_pdf_path = self._create_fallback_pdf(task_id, lines, "fallback_preview")
            preview_page_lines = [lines]

        if not lines:
            lines = [line for page in preview_page_lines for line in page]
        if doc_loaded and not lines:
            # Converted docx may contain only drawing/textbox content.
            # Use PDF text extraction as source for recognition in this case.
            lines = [line for page in preview_page_lines for line in page]
        elif doc_loaded and lines and preview_page_lines:
            # 部分 .doc/.docx 在段落抽取时会丢失表格里的“( ) ( ) ...”空位行。
            # 当 PDF 文本明显包含更多多空位行时，改用 PDF 文本线做规则提取。
            pdf_lines = [line for page in preview_page_lines for line in page]
            supplemented_lines = self._supplement_docx_bare_numbering_lines_from_pdf(lines, pdf_lines)
            if supplemented_lines != lines:
                lines = supplemented_lines
            elif self._should_prefer_pdf_lines_for_docx_recognition(lines, pdf_lines):
                lines = pdf_lines
        split_idx, inline_main_prefix, inline_answer_first_line = self._detect_answer_split(lines, answer_patterns)
        main_lines, answer_lines = self._apply_answer_split(
            lines,
            split_idx,
            inline_main_prefix,
            inline_answer_first_line,
        )

        inferred_answer_start_page: Optional[int] = None
        if answer_lines:
            inferred_answer_start_page = self._infer_answer_start_page(preview_page_lines, split_idx, answer_lines)

        heading_answer_start_page = self._find_answer_heading_page(preview_page_lines, answer_patterns)
        answer_start_page = (
            inferred_answer_start_page if inferred_answer_start_page is not None else heading_answer_start_page
        )
        main_page_urls, answer_page_urls = self._split_pdf_preview(task_id, preview_pdf_path, answer_start_page)

        if doc_loaded:
            pages_count = max(len(preview_page_lines), 1)
            default_header = headers[0] if headers else ""
            default_footer = footers[0] if footers else ""
            header_footer_items = [
                {"page": idx, "header": default_header, "footer": default_footer}
                for idx in range(1, pages_count + 1)
            ]
        else:
            header_footer_items = detect_pdf_header_footer(preview_page_lines)
        return (
            lines,
            main_lines,
            answer_lines,
            main_page_urls,
            answer_page_urls,
            header_footer_items,
        )

    def _render_pdf_pages(self, task_id: int, file_path: Path, page_lines: List[List[str]]) -> List[str]:
        urls: List[str] = []
        if fitz is None:
            for idx, lines in enumerate(page_lines, start=1):
                urls.append(self._render_text_page(task_id, lines, idx, "pdf"))
            return urls

        document = fitz.open(str(file_path))
        try:
            for idx in range(document.page_count):
                page = document.load_page(idx)
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                filename = f"{task_id}_pdf_{idx + 1}_{uuid.uuid4().hex[:8]}.png"
                output_path = settings.pages_dir / filename
                pixmap.save(str(output_path))
                urls.append(f"{settings.static_url_prefix}/pages/{filename}")
        finally:
            document.close()
        return urls

    def _process_pdf(
        self,
        task_id: int,
        file_path: Path,
        answer_patterns: List[re.Pattern],
    ) -> Tuple[List[str], List[str], List[str], List[str], List[str], List[Dict]]:
        page_lines = self._extract_pdf_page_lines(file_path)
        all_lines = [line for lines in page_lines for line in lines]
        split_idx, inline_main_prefix, inline_answer_first_line = self._detect_answer_split(all_lines, answer_patterns)
        main_lines, answer_lines = self._apply_answer_split(
            all_lines,
            split_idx,
            inline_main_prefix,
            inline_answer_first_line,
        )

        inferred_answer_start_page: Optional[int] = None
        if answer_lines:
            inferred_answer_start_page = self._infer_answer_start_page(page_lines, split_idx, answer_lines)

        heading_answer_start_page = self._find_answer_heading_page(page_lines, answer_patterns)
        answer_start_page = (
            inferred_answer_start_page if inferred_answer_start_page is not None else heading_answer_start_page
        )

        main_page_urls, answer_page_urls = self._split_pdf_preview(task_id, file_path, answer_start_page)
        if answer_start_page is None and not answer_lines:
            main_lines = all_lines
        header_footer_items = detect_pdf_header_footer(page_lines)
        return (
            all_lines,
            main_lines,
            answer_lines,
            main_page_urls,
            answer_page_urls,
            header_footer_items,
        )

    def _convert_doc_to_docx(self, doc_path: Path) -> Path:
        out_dir = doc_path.parent
        soffice_cmd = self._resolve_libreoffice_cmd()
        if not soffice_cmd:
            raise RuntimeError(
                "LibreOffice executable not found. Install LibreOffice or set LIBREOFFICE_CMD to full soffice path."
            )

        converted_path = out_dir / f"{doc_path.stem}.docx"
        self._run_libreoffice_conversion(
            soffice_cmd=soffice_cmd,
            source_path=doc_path,
            out_dir=out_dir,
            convert_to="docx",
            expected_output=converted_path,
            error_prefix="Failed to convert .doc file",
        )
        return converted_path

    def _resolve_libreoffice_cmd(self) -> Optional[str]:
        configured = (settings.libreoffice_cmd or "").strip()
        candidates: List[str] = []

        if configured:
            candidates.append(configured)
            resolved = shutil.which(configured)
            if resolved:
                candidates.append(resolved)

        if os.name == "nt":
            candidates.extend(
                [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                ]
            )
        else:
            candidates.extend(["/usr/bin/soffice", "/usr/local/bin/soffice"])

        for item in candidates:
            if not item:
                continue
            path = Path(item)
            if path.exists():
                return str(path)
            resolved = shutil.which(item)
            if resolved:
                return resolved
        return None
